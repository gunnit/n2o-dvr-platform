import asyncio
import io
import json
import os
import re
import shutil
import uuid
from collections import OrderedDict
from collections.abc import Collection
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.billing.entitlements import Entitlements, get_entitlements
from app.billing.gates import (
    ensure_company_slot,
    ensure_doc_type_allowed,
    ensure_subscription_active,
)
from app.billing.metering import (
    count_active_companies,
    is_company_active,
    record_activation_for_azienda,
)
from app.core.exceptions import BadRequestError, NotFoundError
from app.db.session import get_db
from app.core.permissions import DOCUMENTS_GENERATE
from app.dependencies import get_current_org, require_capability
from app.models.ambiente import Ambiente
from app.models.azienda import Azienda
from app.models.documento_generato import DocumentoGenerato
from app.models.persona import Persona
from app.models.user import User
from app.schemas.document import (
    DocumentBatchRequest,
    DocumentEditLinkResponse,
    DocumentGenerateRequest,
    DocumentPreviewResponse,
    DocumentResponse,
    DocumentSnapshotResponse,
    OverridesPatchRequest,
    OverridesResponse,
)
from app.services.ambiente_photo import (
    PhotoBackfillResult,
    backfill_document_images_for_dvr,
)


def _doc_to_response(doc: DocumentoGenerato, generated_by_name: str | None) -> DocumentResponse:
    """Serialise a DocumentoGenerato row adding the resolved user name.

    The `generated_by_name` is resolved via a join on users.full_name in
    the caller (US-2.9) since SQLAlchemy relationships for `generated_by`
    aren't wired on the model.
    """
    gdoc_file_id = getattr(doc, "gdoc_file_id", None)
    gdoc_edit_url = (
        f"https://docs.google.com/document/d/{gdoc_file_id}/edit"
        if gdoc_file_id else None
    )
    options = getattr(doc, "options", None) or {}
    edited_in_gdocs = bool(options.get("edited_in_gdocs"))
    # Mirrors edited_in_gdocs: set by save-edited-version so the frontend
    # can render the "Modificato" badge without re-parsing options JSON.
    edited_inline = bool(options.get("edited_inline"))
    return DocumentResponse(
        id=doc.id,
        azienda_id=doc.azienda_id,
        tipo_documento=doc.tipo_documento,
        versione=doc.versione,
        status=doc.status,
        file_path=doc.file_path,
        gdrive_file_id=doc.gdrive_file_id,
        gdoc_file_id=gdoc_file_id,
        gdoc_edit_url=gdoc_edit_url,
        edited_in_gdocs=edited_in_gdocs,
        edited_inline=edited_inline,
        error_message=doc.error_message,
        created_at=doc.created_at,
        generated_by_name=generated_by_name,
        # US-5.2 AC2: pass the worker-set drift flag to the documents
        # page. Defaults to False on legacy rows where the column was
        # NULL before the d3e4f5a6b7c8 migration applied a default.
        stale_snapshot=bool(getattr(doc, "stale_snapshot", False)),
    )


async def _resolve_user_name(user_id: uuid.UUID | None, db: AsyncSession) -> str | None:
    if user_id is None:
        return None
    result = await db.execute(select(User.full_name).where(User.id == user_id))
    return result.scalar_one_or_none()

router = APIRouter(prefix="/aziende/{azienda_id}/documents", tags=["documents"])

# Global (non-nested) router for download-by-id endpoint
download_router = APIRouter(prefix="/documenti", tags=["documents"])


# US-4.1: document types that require the DVR Master to already exist before
# they can be generated. The DVR carries the anagrafica + environments that
# these dependent documents reuse, so generating them first would produce
# incomplete output.
_DVR_DEPENDENT_TYPES: set[str] = {"pee_azienda", "pee_comune"}


async def _ensure_anagrafica_complete_for_dvr(
    azienda: Azienda, tipo_documento: str
) -> None:
    """Block DVR generation when ALL legally-required contact fields are NULL.

    Audit F-004 (2026-04-29 rerun): the DVR Anagrafica section requires
    Codice Fiscale, Telefono, Email and PEC. When all four are NULL the
    document renders four "Non comunicato" rows that an inspector will
    reject. Allowing generation with at least one field populated keeps
    the door open for small artigiani without (e.g.) a PEC yet.
    """
    if tipo_documento != "dvr_master":
        return
    fields = (
        azienda.codice_fiscale,
        azienda.telefono,
        azienda.email,
        azienda.pec,
    )
    if not any((f or "").strip() for f in fields):
        raise BadRequestError(
            "Anagrafica incompleta: inserisci almeno uno tra "
            "Codice Fiscale, Telefono, Email o PEC sull'Azienda "
            "prima di generare il DVR."
        )


async def _ensure_dvr_exists_for_dependent(
    azienda_id: uuid.UUID, tipo_documento: str, db: AsyncSession
) -> None:
    """If tipo_documento depends on the DVR Master, raise 400 when none exists.

    Matches US-4.1 AC2: "Given no DVR exists yet, When I attempt to generate the
    PEE, Then the action is blocked with the message 'Genera prima il DVR Master'."
    A DVR counts as "existing" when at least one DocumentoGenerato row exists
    with tipo_documento == 'dvr_master' and a successful status
    (completed / ready). Bozza / failed / pending rows do not unblock.
    """
    if tipo_documento not in _DVR_DEPENDENT_TYPES:
        return
    result = await db.execute(
        select(DocumentoGenerato.id)
        .where(
            DocumentoGenerato.azienda_id == azienda_id,
            DocumentoGenerato.tipo_documento == "dvr_master",
            DocumentoGenerato.status.in_(("completed", "ready")),
        )
        .limit(1)
    )
    if result.scalar_one_or_none() is None:
        raise BadRequestError("Genera prima il DVR Master")


@download_router.get("/{document_id}/download")
async def download_document(
    document_id: uuid.UUID,
    raw: bool = False,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Stream back the generated file (.docx or .zip).

    Prefers the ``file_content`` bytes stored in Postgres (works across
    Render's separate API / Worker disks). Falls back to ``file_path``
    on the local filesystem for backwards compatibility with documents
    generated before the DB-storage migration.

    When the row carries pending inline-edit ``content_overrides`` and the
    payload is a .docx, the overrides are applied to the bytes in memory
    before streaming (same filename). ``?raw=true`` skips application and
    returns the pristine generated bytes. Zip bundles never apply.
    """
    result = await db.execute(
        select(DocumentoGenerato)
        .join(Azienda, Azienda.id == DocumentoGenerato.azienda_id)
        .where(DocumentoGenerato.id == document_id, Azienda.organization_id == org_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundError("Document not found")
    if doc.status != "completed":
        raise NotFoundError("Document not ready yet")

    # Determine filename and MIME type from whichever source is available
    filename = doc.file_name or (os.path.basename(doc.file_path) if doc.file_path else None)
    if not filename:
        raise NotFoundError("Document not ready yet")
    media_type = "application/zip" if filename.endswith(".zip") else (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Apply pending inline edits to the .docx bytes before streaming.
    # Best-effort: if application blows up (corrupt payload, unexpected
    # XML) we log and fall through to the pristine bytes — a stale
    # override must never break the download.
    if not raw and doc.content_overrides and filename.endswith(".docx"):
        content = doc.file_content
        if not content and doc.file_path and os.path.exists(doc.file_path):
            with open(doc.file_path, "rb") as f:
                content = f.read()
        if content:
            try:
                from app.services.document_preview import apply_overrides_to_docx

                # CPU-bound docx rewrite — off the event loop (B5).
                edited = await asyncio.to_thread(
                    apply_overrides_to_docx, content, doc.content_overrides
                )
                return StreamingResponse(
                    io.BytesIO(edited),
                    media_type=media_type,
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'},
                )
            except Exception:
                import logging
                logging.getLogger(__name__).exception(
                    "Override application failed for %s — streaming original", doc.id
                )

    # Prefer DB content (works cross-service on Render)
    if doc.file_content:
        return StreamingResponse(
            io.BytesIO(doc.file_content),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # Fallback: serve from local disk (pre-migration documents or local dev)
    if doc.file_path and os.path.exists(doc.file_path):
        return FileResponse(doc.file_path, media_type=media_type, filename=filename)

    raise NotFoundError("File non disponibile. Rigenera il documento.")


async def _get_azienda(azienda_id: uuid.UUID, org_id: uuid.UUID, db: AsyncSession) -> Azienda:
    result = await db.execute(
        select(Azienda).where(Azienda.id == azienda_id, Azienda.organization_id == org_id)
    )
    azienda = result.scalar_one_or_none()
    if not azienda:
        raise NotFoundError("Azienda not found")
    return azienda


async def _ensure_company_slot_available(
    ent: Entitlements, org_id: uuid.UUID, azienda_id: uuid.UUID, db: AsyncSession
) -> None:
    """Gate the Model A active-company meter, synchronously (MB-2.3).

    The worker records the activation, but the ceiling is enforced *here* so the
    user gets a 402 on the request instead of a background job that quietly
    produces nothing. A company already counted this period passes for free —
    otherwise a consultant sitting at 15/15 could never finish the fifteenth
    company's documents.
    """
    if ent.max_companies is None:
        return
    already = await is_company_active(org_id, azienda_id, db, ent)
    if already:
        return
    active = await count_active_companies(org_id, db, ent)
    ensure_company_slot(ent, active, already_active=False, org_id=org_id)


async def _ensure_new_version_allowed(
    ent: Entitlements,
    org_id: uuid.UUID,
    azienda_id: uuid.UUID,
    tipo_documento: str | None,
    db: AsyncSession,
) -> None:
    """The three gates every *new completed version* must pass (MB-6.2).

    ``/generate`` reaches the worker through ``_enqueue_generation``, which
    carries its own gates. Three endpoints — restore, sync-from-gdoc and
    save-edited-version — skip the worker entirely, mint a completed row
    themselves and then call ``record_activation_for_azienda``. Until this
    helper existed they were an unguarded side door: a canceled tenant could
    keep producing versions, and an activation could push a consultant past the
    active-company ceiling without ever consulting it.

    Ordering matches ``/generate``: subscription first (the most actionable
    message), then doc type, then the company slot.
    """
    ensure_subscription_active(ent, org_id)
    ensure_doc_type_allowed(ent, tipo_documento, org_id)
    await _ensure_company_slot_available(ent, org_id, azienda_id, db)


def _enqueue_generation(doc: DocumentoGenerato, ent: Entitlements, org_id: uuid.UUID) -> None:
    """The ONLY place a generation task is dispatched (MB-2.2).

    Both the single and batch endpoints funnel through here so no future code
    path can reach the worker without passing the doc-type gate first (INV-5).
    ``tests/test_billing_enforcement.py`` fails the build if a bare
    ``generate_document_task.delay(`` appears anywhere else.

    The gate is re-checked at dispatch rather than trusted from the caller: the
    endpoint's earlier check and this one bracket the row creation, so a type
    that slipped through a future refactor still cannot reach a worker.

    MB-4.5 lives here too: a lapsed subscription may not *generate*. Read and
    download paths deliberately never call this — a canceled tenant keeps access
    to documents it already produced, which D.Lgs. 81/2008 retention requires.
    """
    ensure_subscription_active(ent, org_id)
    ensure_doc_type_allowed(ent, doc.tipo_documento, org_id)
    try:
        from app.tasks.document_tasks import generate_document_task

        generate_document_task.delay(str(doc.id))
    except HTTPException:
        # A 402 from the gate must reach the client, not be swallowed as a
        # broker failure.
        raise
    except Exception:
        # If the broker is unavailable, log and leave the pending record in
        # place — the task can be retried manually.
        import logging

        logging.getLogger(__name__).exception("Celery dispatch failed for %s", doc.id)


async def _preflight_dvr_photo_transport(
    azienda_id: uuid.UUID,
    requested_types: Collection[str],
    db: AsyncSession,
) -> PhotoBackfillResult | None:
    if "dvr_master" not in requested_types:
        return None
    return await backfill_document_images_for_dvr(azienda_id, db)


@router.post("/generate", response_model=DocumentResponse, status_code=202)
async def generate_document(
    azienda_id: uuid.UUID,
    body: DocumentGenerateRequest,
    org_id: uuid.UUID = Depends(get_current_org),
    user: User = Depends(require_capability(DOCUMENTS_GENERATE)),
    ent: Entitlements = Depends(get_entitlements),
    db: AsyncSession = Depends(get_db),
):
    """Trigger async document generation for a single document type.

    Creates a DocumentoGenerato record with status='pending' and returns
    immediately. The actual generation will be handled by a Celery worker.
    """
    azienda = await _get_azienda(azienda_id, org_id, db)
    # MB-4.5 first: a lapsed tenant should be told their subscription is
    # inactive, not sent off to fix anagrafica for a document they cannot
    # generate either way. Re-checked at dispatch, which is the guarantee.
    ensure_subscription_active(ent, org_id)
    # MB-2.1: is this document type in the plan? Checked before any work — no
    # row is created for a document the tenant cannot have.
    ensure_doc_type_allowed(ent, body.tipo_documento, org_id)
    # MB-2.3: would this consume a new client-company slot?
    await _ensure_company_slot_available(ent, org_id, azienda_id, db)
    # Audit F-004: refuse DVR Master with all anagrafica contact fields NULL.
    await _ensure_anagrafica_complete_for_dvr(azienda, body.tipo_documento)
    # US-4.1 AC2: block dependent documents (PEE) until the DVR Master exists.
    await _ensure_dvr_exists_for_dependent(azienda_id, body.tipo_documento, db)

    await _preflight_dvr_photo_transport(azienda_id, [body.tipo_documento], db)

    # Determine the next version number for this document type
    result = await db.execute(
        select(DocumentoGenerato)
        .where(
            DocumentoGenerato.azienda_id == azienda_id,
            DocumentoGenerato.tipo_documento == body.tipo_documento,
        )
        .order_by(DocumentoGenerato.versione.desc())
        .limit(1)
    )
    latest = result.scalar_one_or_none()
    next_version = (latest.versione + 1) if latest else 1

    doc = DocumentoGenerato(
        azienda_id=azienda_id,
        tipo_documento=body.tipo_documento,
        versione=next_version,
        status="pending",
        generated_by=user.id,
        generation_started_at=datetime.utcnow(),
        # US-4.4: persist the dialog-supplied options (e.g. HACCP forms
        # selected_codes) so the async worker can read them back.
        options=body.options,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    _enqueue_generation(doc, ent, org_id)

    return _doc_to_response(doc, await _resolve_user_name(doc.generated_by, db))


@router.get("", response_model=list[DocumentResponse])
async def list_documents(
    azienda_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """List all generated documents for an azienda.

    US-5.2 AC2 — drift detection runs on every list call: we compute the
    current survey hash once and update any completed rows whose stored
    hash no longer matches. This catches the case where the operator
    edits the survey *after* a job completes (the worker can't see those
    edits since it's already done). Cost: one extra SHA256 + small UPDATE
    per page load.
    """
    await _get_azienda(azienda_id, org_id, db)

    # Recompute the live hash + flip any completed docs whose snapshot is
    # now stale. Done in a single helper so the survey-edit endpoints can
    # call the same code path proactively.
    try:
        from app.services.survey_snapshot import mark_documents_stale_for

        await mark_documents_stale_for(azienda_id, db)
        await db.commit()
    except Exception:  # pragma: no cover — never fail the list call on this
        import logging

        logging.getLogger(__name__).exception(
            "Stale-snapshot recompute failed for %s", azienda_id
        )

    # Left-join on users so rows with a NULL generated_by (legacy records)
    # still appear, just without an author name.
    result = await db.execute(
        select(DocumentoGenerato, User.full_name)
        .outerjoin(User, User.id == DocumentoGenerato.generated_by)
        .where(DocumentoGenerato.azienda_id == azienda_id)
        .order_by(DocumentoGenerato.created_at.desc())
    )
    return [_doc_to_response(doc, name) for doc, name in result.all()]


@router.get("/{document_id}/status", response_model=DocumentResponse)
async def get_document_status(
    azienda_id: uuid.UUID,
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Check generation status for a specific document."""
    await _get_azienda(azienda_id, org_id, db)
    result = await db.execute(
        select(DocumentoGenerato, User.full_name)
        .outerjoin(User, User.id == DocumentoGenerato.generated_by)
        .where(
            DocumentoGenerato.id == document_id,
            DocumentoGenerato.azienda_id == azienda_id,
        )
    )
    row = result.one_or_none()
    if not row:
        raise NotFoundError("Document not found")
    doc, name = row
    return _doc_to_response(doc, name)


# B1 — survey statuses that count as "submitted" for the Genera-Documenti
# precondition. ``draft`` (the survey was never started) and ``in_progress``
# (mid-wizard) both fail the gate; once /survey/complete or /survey/sign
# has run we trust the operator decided the survey is fit to generate
# against. Mirrors how survey.py advances the status (step_N -> completed
# -> firmato -> in_revisione).
_SURVEY_SUBMITTED_STATUSES: set[str] = {"completed", "firmato", "in_revisione"}


async def _check_batch_preconditions(
    azienda_id: uuid.UUID, db: AsyncSession
) -> list[str]:
    """Return a list of Italian descriptions for any missing prerequisites.

    The batch generator is a foot-gun on an empty azienda — every dependent
    document silently produces a placeholder file with "Nessun ambiente
    registrato" boilerplate. We require the survey to be at least submitted
    once and to carry the minimal entities the DVR Master needs (>=1
    ambiente, >=1 persona, >=1 RSPP).
    """
    missing: list[str] = []

    # Resolve the survey_status alongside the entity counts in three small
    # queries instead of one mega-join — clearer and fast on the row counts
    # we actually deal with (tens, not thousands).
    az_status = (
        await db.execute(
            select(Azienda.survey_status).where(Azienda.id == azienda_id)
        )
    ).scalar_one_or_none()
    if az_status is None or az_status not in _SURVEY_SUBMITTED_STATUSES:
        missing.append("Sopralluogo non completato o non firmato")

    ambienti_count = (
        await db.execute(
            select(func.count(Ambiente.id)).where(Ambiente.azienda_id == azienda_id)
        )
    ).scalar_one()
    if not ambienti_count:
        missing.append("Nessun ambiente di lavoro registrato")

    persone_count = (
        await db.execute(
            select(func.count(Persona.id)).where(Persona.azienda_id == azienda_id)
        )
    ).scalar_one()
    if not persone_count:
        missing.append("Nessuna persona registrata")

    # RSPP gate matches the survey/sign flow's expectation that the survey
    # carries a designated safety manager before any document is produced.
    rspp_count = (
        await db.execute(
            select(func.count(Persona.id)).where(
                Persona.azienda_id == azienda_id,
                Persona.ruolo_rspp.is_(True),
            )
        )
    ).scalar_one()
    if not rspp_count:
        missing.append("Nessun RSPP designato")

    return missing


@router.post("/batch", response_model=list[DocumentResponse], status_code=202)
async def batch_generate_documents(
    azienda_id: uuid.UUID,
    body: DocumentBatchRequest,
    org_id: uuid.UUID = Depends(get_current_org),
    user: User = Depends(require_capability(DOCUMENTS_GENERATE)),
    ent: Entitlements = Depends(get_entitlements),
    db: AsyncSession = Depends(get_db),
):
    """Trigger async generation for multiple document types at once."""
    await _get_azienda(azienda_id, org_id, db)

    # MB-4.5 — see the note in `generate_document`: the subscription answer
    # comes before any completeness complaint.
    ensure_subscription_active(ent, org_id)

    # B1 — refuse to enqueue tasks against an incomplete sopralluogo. The
    # frontend disables the button when survey_status == draft, but we
    # double-check server-side because curl + stale tabs both bypass that.
    missing = await _check_batch_preconditions(azienda_id, db)
    if missing:
        raise HTTPException(
            status_code=409,
            detail=f"Sopralluogo incompleto: {', '.join(missing)}",
        )

    # MB-2.1: gate every requested type up front. All-or-nothing on purpose —
    # partially fulfilling a batch would leave the user guessing which of the
    # documents they asked for actually exist.
    for tipo in body.tipi_documento:
        ensure_doc_type_allowed(ent, tipo, org_id)
    # MB-2.3: a batch activates the company once, not once per document.
    await _ensure_company_slot_available(ent, org_id, azienda_id, db)

    await _preflight_dvr_photo_transport(azienda_id, body.tipi_documento, db)

    created_docs: list[DocumentoGenerato] = []

    for tipo in body.tipi_documento:
        # Determine the next version number for each document type
        result = await db.execute(
            select(DocumentoGenerato)
            .where(
                DocumentoGenerato.azienda_id == azienda_id,
                DocumentoGenerato.tipo_documento == tipo,
            )
            .order_by(DocumentoGenerato.versione.desc())
            .limit(1)
        )
        latest = result.scalar_one_or_none()
        next_version = (latest.versione + 1) if latest else 1

        doc = DocumentoGenerato(
            azienda_id=azienda_id,
            tipo_documento=tipo,
            versione=next_version,
            status="pending",
            generated_by=user.id,
            generation_started_at=datetime.utcnow(),
        )
        db.add(doc)
        created_docs.append(doc)

    await db.commit()

    responses: list[DocumentResponse] = []
    for doc in created_docs:
        await db.refresh(doc)
        _enqueue_generation(doc, ent, org_id)
        responses.append(
            _doc_to_response(doc, await _resolve_user_name(doc.generated_by, db))
        )

    return responses


@router.get("/{document_id}/snapshot", response_model=DocumentSnapshotResponse)
async def get_document_snapshot(
    azienda_id: uuid.UUID,
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Return a structured text snapshot of a generated .docx (US-2.9).

    Used by the frontend diff viewer. Since we don't persist snapshots,
    we parse the .docx on demand. If the file is missing (e.g. bozza
    rollback per US-2.8), we 404 — there's nothing to diff.
    """
    await _get_azienda(azienda_id, org_id, db)
    result = await db.execute(
        select(DocumentoGenerato).where(
            DocumentoGenerato.id == document_id,
            DocumentoGenerato.azienda_id == azienda_id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundError("Document not found")

    # Resolve the document source — DB content or disk file
    file_source: io.BytesIO | str | None = None
    filename = doc.file_name or (os.path.basename(doc.file_path) if doc.file_path else "")
    if doc.file_content:
        file_source = io.BytesIO(doc.file_content)
    elif doc.file_path and os.path.exists(doc.file_path):
        file_source = doc.file_path
    else:
        raise NotFoundError("Snapshot non disponibile per questo documento")

    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    # Only parse .docx files; .zip bundles (e.g. haccp_forms) are not
    # structurally diffable — fall back to empty lists so the frontend
    # can still show the metadata header.
    if filename.endswith(".docx"):
        try:
            from docx import Document

            document = Document(file_source)
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                rows: list[list[str]] = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append(rows)
        except Exception:
            # Swallow parse errors — return what we have so the UI can
            # still render the version metadata. Diff will just be empty.
            import logging
            logging.getLogger(__name__).exception(
                "Failed to parse .docx for snapshot %s", doc.id
            )

    generated_by_name = await _resolve_user_name(doc.generated_by, db)
    return DocumentSnapshotResponse(
        id=doc.id,
        versione=doc.versione,
        generated_at=doc.generation_completed_at or doc.created_at,
        generated_by_name=generated_by_name,
        paragraphs=paragraphs,
        tables=tables,
    )


@router.post("/{document_id}/restore", response_model=DocumentResponse, status_code=201)
async def restore_document(
    azienda_id: uuid.UUID,
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    user: User = Depends(require_capability(DOCUMENTS_GENERATE)),
    ent: Entitlements = Depends(get_entitlements),
    db: AsyncSession = Depends(get_db),
):
    """Restore a historical version as a new version (US-2.9).

    MVP approach: rather than re-running the generator (templates and
    live data may have drifted), we copy the source .docx on disk and
    register it as a new `DocumentoGenerato` row with status=completed.
    This preserves the exact bytes the user is trying to "go back to".
    """
    await _get_azienda(azienda_id, org_id, db)
    result = await db.execute(
        select(DocumentoGenerato).where(
            DocumentoGenerato.id == document_id,
            DocumentoGenerato.azienda_id == azienda_id,
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise NotFoundError("Document not found")
    if not source.file_content and (not source.file_path or not os.path.exists(source.file_path)):
        raise BadRequestError("Impossibile ripristinare una bozza")

    # A restore mints a *new completed version* and records an activation, so
    # it is document production, not retrieval — it takes the same three gates
    # as /generate. Downloading an existing version stays ungated on purpose
    # (D.Lgs. 81/2008 retention); only minting v+1 is charged for.
    await _ensure_new_version_allowed(ent, org_id, azienda_id, source.tipo_documento, db)

    # Next version number for this document type
    result = await db.execute(
        select(DocumentoGenerato)
        .where(
            DocumentoGenerato.azienda_id == azienda_id,
            DocumentoGenerato.tipo_documento == source.tipo_documento,
        )
        .order_by(DocumentoGenerato.versione.desc())
        .limit(1)
    )
    latest = result.scalar_one_or_none()
    next_version = (latest.versione + 1) if latest else 1

    # Build restored filename
    src_name = source.file_name or (os.path.basename(source.file_path) if source.file_path else f"{source.tipo_documento}_v{source.versione}")
    stem, ext = os.path.splitext(src_name)
    new_name = f"{stem}_v{next_version}_restored{ext}"

    # Copy the file content (prefer DB, fall back to disk)
    restored_content: bytes | None = source.file_content
    new_path: str | None = None
    if not restored_content and source.file_path and os.path.exists(source.file_path):
        try:
            with open(source.file_path, "rb") as f:
                restored_content = f.read()
        except OSError as exc:
            raise BadRequestError(f"Copia del file fallita: {exc}") from exc

    # Also write to disk if local filesystem is available (backwards compat)
    if source.file_path:
        src_dir = os.path.dirname(source.file_path)
        new_path = os.path.join(src_dir, new_name)
        try:
            if restored_content:
                os.makedirs(src_dir, exist_ok=True)
                with open(new_path, "wb") as f:
                    f.write(restored_content)
            elif os.path.exists(source.file_path):
                shutil.copy2(source.file_path, new_path)
        except OSError:
            new_path = None  # disk write failed, DB content will suffice

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    new_doc = DocumentoGenerato(
        azienda_id=azienda_id,
        tipo_documento=source.tipo_documento,
        versione=next_version,
        status="completed",
        file_path=new_path,
        file_content=restored_content,
        file_name=new_name,
        error_message=None,
        generated_by=user.id,
        generation_started_at=now,
        generation_completed_at=now,
    )
    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)
    # MB-2.3 — a restore mints a completed document, so it activates the
    # company just like a fresh generation. ON CONFLICT keeps it to one row.
    await record_activation_for_azienda(new_doc.azienda_id, db)

    return _doc_to_response(new_doc, await _resolve_user_name(new_doc.generated_by, db))


# ---------------------------------------------------------------------------
# Google Docs round-trip: open-for-editing + sync-from-gdoc (DVR Master only)
# ---------------------------------------------------------------------------

# Document types eligible for the in-browser Google Docs editing flow.
# Start with DVR Master; add attachments once the round-trip is proven.
_GDOC_EDITABLE_TYPES: set[str] = {"dvr_master"}


@download_router.post(
    "/{document_id}/open-for-editing",
    response_model=DocumentEditLinkResponse,
    # Editing the finished document is the office operator's review job, not the
    # field operator's. Reading and downloading it stay open to everyone.
    dependencies=[Depends(require_capability(DOCUMENTS_GENERATE))],
)
async def open_document_for_editing(
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Return an editable Google Docs URL for the document.

    First call: upload the .docx bytes to Drive with conversion to Google Doc,
    grant "anyone with link can edit" permission, persist the new Google Doc
    file ID on the row, and return the edit URL.
    Subsequent calls: return the existing edit URL (idempotent).
    """
    from fastapi import HTTPException, status

    result = await db.execute(
        select(DocumentoGenerato)
        .join(Azienda, Azienda.id == DocumentoGenerato.azienda_id)
        .where(DocumentoGenerato.id == document_id, Azienda.organization_id == org_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundError("Document not found")

    if doc.tipo_documento not in _GDOC_EDITABLE_TYPES:
        raise BadRequestError(
            "La modifica in Google Docs è disponibile solo per il DVR Master"
        )
    if doc.status != "completed":
        raise NotFoundError("Document not ready yet")

    # Idempotent reopen: if we already have the Google Doc, just return its URL.
    if doc.gdoc_file_id:
        return DocumentEditLinkResponse(
            gdoc_file_id=doc.gdoc_file_id,
            edit_url=f"https://docs.google.com/document/d/{doc.gdoc_file_id}/edit",
        )

    if not doc.file_content:
        raise NotFoundError("File non disponibile. Rigenera il documento.")

    # Resolve the azienda name for the Drive folder
    azienda_result = await db.execute(select(Azienda).where(Azienda.id == doc.azienda_id))
    azienda = azienda_result.scalar_one_or_none()
    if azienda is None:
        # Orphaned document (parent azienda deleted) — 404, not an unhandled 500.
        raise NotFoundError("Azienda not found")

    from app.services.gdrive_service import (
        create_gdoc_from_docx_bytes,
        share_anyone_with_link,
    )

    filename = doc.file_name or f"{doc.tipo_documento}_v{doc.versione}.docx"
    gdoc_id = await create_gdoc_from_docx_bytes(
        doc.file_content, filename, azienda.ragione_sociale[:100]
    )
    if not gdoc_id:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Google Drive non configurato o errore di conversione",
        )

    await share_anyone_with_link(gdoc_id)

    doc.gdoc_file_id = gdoc_id
    await db.commit()
    await db.refresh(doc)

    return DocumentEditLinkResponse(
        gdoc_file_id=gdoc_id,
        edit_url=f"https://docs.google.com/document/d/{gdoc_id}/edit",
    )


@download_router.post(
    "/{document_id}/sync-from-gdoc", response_model=DocumentResponse, status_code=201
)
async def sync_document_from_gdoc(
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    user: User = Depends(require_capability(DOCUMENTS_GENERATE)),
    ent: Entitlements = Depends(get_entitlements),
    db: AsyncSession = Depends(get_db),
):
    """Pull the latest Google Doc content back into a new version row.

    Exports the live Google Doc as .docx via Drive API, inserts a new
    DocumentoGenerato row with incremented `versione` and status=completed,
    and tags `options.edited_in_gdocs=True` for the version history UI.
    """
    from fastapi import HTTPException, status

    result = await db.execute(
        select(DocumentoGenerato)
        .join(Azienda, Azienda.id == DocumentoGenerato.azienda_id)
        .where(DocumentoGenerato.id == document_id, Azienda.organization_id == org_id)
    )
    source = result.scalar_one_or_none()
    if not source:
        raise NotFoundError("Document not found")

    if not source.gdoc_file_id:
        raise BadRequestError("Nessuna modifica in Google Docs da sincronizzare")

    # Pulling edited content back mints a new completed version (MB-6.2).
    await _ensure_new_version_allowed(
        ent, org_id, source.azienda_id, source.tipo_documento, db
    )

    from app.services.gdrive_service import (
        delete_gdoc,
        export_gdoc_as_docx,
        get_gdoc_times,
    )

    # Dirty-check: if the Google Doc's modifiedTime is within a few seconds of
    # its createdTime, the user never actually edited. Reject so double-clicks
    # or stale sync attempts don't produce a spurious v+1 identical to v.
    times = await get_gdoc_times(source.gdoc_file_id)
    if times is not None:
        from datetime import datetime as _dt
        try:
            created_dt = _dt.fromisoformat(times[0].replace("Z", "+00:00"))
            modified_dt = _dt.fromisoformat(times[1].replace("Z", "+00:00"))
            # 5s tolerance covers Drive's own post-conversion writes.
            if (modified_dt - created_dt).total_seconds() < 5:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Nessuna modifica rilevata in Google Docs",
                )
        except ValueError:
            # Fall through if Drive returned something we couldn't parse.
            pass

    exported = await export_gdoc_as_docx(source.gdoc_file_id)
    if not exported:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Impossibile esportare il documento da Google Docs",
        )

    # Next version for this document type
    result = await db.execute(
        select(DocumentoGenerato)
        .where(
            DocumentoGenerato.azienda_id == source.azienda_id,
            DocumentoGenerato.tipo_documento == source.tipo_documento,
        )
        .order_by(DocumentoGenerato.versione.desc())
        .limit(1)
    )
    latest = result.scalar_one_or_none()
    next_version = (latest.versione + 1) if latest else 1

    # Build filename: append -edited to stem
    src_name = source.file_name or f"{source.tipo_documento}_v{source.versione}.docx"
    stem, ext = os.path.splitext(src_name)
    new_name = f"{stem}_v{next_version}_edited{ext or '.docx'}"

    merged_options = dict(source.options or {})
    merged_options["edited_in_gdocs"] = True
    merged_options["source_version_id"] = str(source.id)

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    new_doc = DocumentoGenerato(
        azienda_id=source.azienda_id,
        tipo_documento=source.tipo_documento,
        versione=next_version,
        status="completed",
        # Carry the source's file_path (parity with restore): frontend
        # surfaces gate their download/preview buttons on a non-NULL
        # file_path, while the download endpoint prefers file_content —
        # so the NEW bytes are always the ones served.
        file_path=source.file_path,
        file_content=exported,
        file_name=new_name,
        options=merged_options,
        generated_by=user.id,
        generation_started_at=now,
        generation_completed_at=now,
    )
    db.add(new_doc)
    # Self-cleanup: once we've captured the edits as a new DB-backed version,
    # the Google Doc is no longer authoritative. Clear the source row's
    # gdoc_file_id and delete the Drive file so the UI stops offering sync on
    # a stale link. Best-effort — a Drive delete failure is logged and the
    # commit still proceeds (the user can manually clean up from Drive).
    stale_gdoc_id = source.gdoc_file_id
    source.gdoc_file_id = None
    await db.commit()
    await db.refresh(new_doc)
    # MB-2.3 — a Google-Doc sync mints a completed document.
    await record_activation_for_azienda(new_doc.azienda_id, db)
    try:
        await delete_gdoc(stale_gdoc_id)
    except Exception:
        import logging
        logging.getLogger(__name__).warning(
            "Post-sync Drive cleanup failed for %s", stale_gdoc_id
        )

    return _doc_to_response(new_doc, await _resolve_user_name(new_doc.generated_by, db))


@download_router.delete(
    "/{document_id}/gdoc",
    response_model=DocumentResponse,
    dependencies=[Depends(require_capability(DOCUMENTS_GENERATE))],
)
async def discard_gdoc_edits(
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Delete the editable Google Doc for this row without importing its content.

    Used when the user realises they don't want to keep the in-browser edits —
    removes the Drive file and clears ``gdoc_file_id`` so the UI hides the
    sync/discard buttons. Idempotent: if the Doc is already gone on Drive,
    still clears the DB state.
    """
    result = await db.execute(
        select(DocumentoGenerato)
        .join(Azienda, Azienda.id == DocumentoGenerato.azienda_id)
        .where(DocumentoGenerato.id == document_id, Azienda.organization_id == org_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundError("Document not found")
    if not doc.gdoc_file_id:
        # Nothing to discard — return current state without error so the
        # frontend treats a double-click as a no-op.
        return _doc_to_response(doc, await _resolve_user_name(doc.generated_by, db))

    from app.services.gdrive_service import delete_gdoc as _delete_gdoc

    stale_id = doc.gdoc_file_id
    doc.gdoc_file_id = None
    await db.commit()
    await db.refresh(doc)
    try:
        await _delete_gdoc(stale_id)
    except Exception:
        import logging
        logging.getLogger(__name__).warning(
            "Drive delete failed on discard for %s", stale_id
        )

    return _doc_to_response(doc, await _resolve_user_name(doc.generated_by, db))


# ---------------------------------------------------------------------------
# In-browser preview + inline text editing (per-paragraph overrides)
# ---------------------------------------------------------------------------

# Address grammar accepted by PATCH /overrides: "12" for a top-level block
# or up to "table:row:cell:para" for a cell paragraph. Validation on save is
# purely syntactic (no .docx re-parse per keystroke) — addresses that no
# longer resolve are skipped silently at apply time.
_OVERRIDE_ADDR_RE = re.compile(r"^\d+(:\d+){0,3}$")

# Hard cap (bytes-ish: json.dumps length) on the MERGED override map stored
# on a row. Each request is already schema-limited (500 entries / 20k chars
# per value) but repeated PATCHes could still grow the JSONB unboundedly.
_OVERRIDES_MAX_SERIALIZED = 2_000_000

# Parsed-blocks LRU for the preview endpoint. Keyed by document_id: a
# completed row's bytes are IMMUTABLE — regenerate, restore, sync-from-gdoc
# and save-edited-version all mint NEW rows with new ids — so a cached
# parse can never go stale and needs no invalidation. Entries are treated
# as read-only by consumers. Thread-safety is best-effort under the GIL:
# worst case two concurrent misses parse the same document twice and the
# second insert wins.
_PREVIEW_BLOCKS_CACHE: OrderedDict[uuid.UUID, list[dict]] = OrderedDict()
_PREVIEW_BLOCKS_CACHE_MAX = 4


async def _get_org_document(
    document_id: uuid.UUID,
    org_id: uuid.UUID,
    db: AsyncSession,
    *,
    for_update: bool = False,
) -> tuple[DocumentoGenerato, str]:
    """Org-scoped fetch by document id, joined with the owning Azienda's name.

    ``for_update=True`` row-locks the DocumentoGenerato (Postgres
    ``FOR UPDATE OF``) so concurrent mutating handlers serialise —
    save-edited-version uses it to keep two racing saves from minting
    duplicate version rows. SQLite (unit tests) ignores the clause.
    """
    stmt = (
        select(DocumentoGenerato, Azienda.ragione_sociale)
        .join(Azienda, Azienda.id == DocumentoGenerato.azienda_id)
        .where(DocumentoGenerato.id == document_id, Azienda.organization_id == org_id)
    )
    if for_update:
        stmt = stmt.with_for_update(of=DocumentoGenerato)
    result = await db.execute(stmt)
    row = result.one_or_none()
    if row is None:
        raise NotFoundError("Documento non trovato")
    return row[0], row[1]


def _resolve_docx_filename(doc: DocumentoGenerato) -> str:
    """Return the filename for a completed .docx row; raise 404/400 otherwise.

    The preview/override endpoints only operate on .docx payloads — the
    haccp_forms .zip bundle has no paragraph model to preview or edit.
    Accepts the same ("completed", "ready") status pair the generation
    pipeline treats as successful (mirrors _ensure_dvr_exists_for_dependent).
    """
    if doc.status not in ("completed", "ready"):
        raise NotFoundError("Documento non ancora pronto")
    filename = doc.file_name or (os.path.basename(doc.file_path) if doc.file_path else None)
    if not filename:
        raise NotFoundError("Documento non ancora pronto")
    if not filename.endswith(".docx"):
        raise BadRequestError(
            "L'anteprima e la modifica inline sono disponibili solo per documenti .docx"
        )
    return filename


def _load_docx_bytes(doc: DocumentoGenerato) -> bytes:
    """Document bytes — DB content preferred, disk fallback (mirrors download)."""
    if doc.file_content:
        return doc.file_content
    if doc.file_path and os.path.exists(doc.file_path):
        with open(doc.file_path, "rb") as f:
            return f.read()
    raise NotFoundError("File non disponibile. Rigenera il documento.")


@download_router.get("/{document_id}/preview", response_model=DocumentPreviewResponse)
async def preview_document(
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Parse the stored .docx into the addressed block model for the editor.

    The .docx stays the source of truth — we re-parse on every call (no
    persisted snapshot), so a regenerated file immediately yields fresh
    blocks. Saved overrides ride along so the frontend can rehydrate
    pending edits onto the matching addresses.
    """
    doc, azienda_nome = await _get_org_document(document_id, org_id, db)
    filename = _resolve_docx_filename(doc)

    # LRU hit: the row's bytes are immutable after completion (see cache
    # comment above) so the parsed blocks can be reused as-is.
    blocks = _PREVIEW_BLOCKS_CACHE.get(doc.id)
    if blocks is not None:
        _PREVIEW_BLOCKS_CACHE.move_to_end(doc.id)
    else:
        docx_bytes = _load_docx_bytes(doc)

        from app.services.document_preview import parse_docx_to_blocks

        try:
            # CPU-bound parse (~seconds on the DVR) — off the event loop (B5).
            blocks = await asyncio.to_thread(parse_docx_to_blocks, docx_bytes)
        except Exception as exc:
            import logging
            logging.getLogger(__name__).exception("Preview parse failed for %s", doc.id)
            raise BadRequestError(
                "Impossibile generare l'anteprima. Rigenera il documento."
            ) from exc
        _PREVIEW_BLOCKS_CACHE[doc.id] = blocks
        while len(_PREVIEW_BLOCKS_CACHE) > _PREVIEW_BLOCKS_CACHE_MAX:
            _PREVIEW_BLOCKS_CACHE.popitem(last=False)

    return DocumentPreviewResponse(
        id=doc.id,
        azienda_id=doc.azienda_id,
        azienda_nome=azienda_nome,
        tipo_documento=doc.tipo_documento,
        versione=doc.versione,
        file_name=filename,
        stale_snapshot=bool(getattr(doc, "stale_snapshot", False)),
        generated_at=(
            doc.generation_completed_at.isoformat()
            if doc.generation_completed_at
            else None
        ),
        blocks=blocks,
        overrides=doc.content_overrides or {},
    )


@download_router.get("/{document_id}/preview/images/{image_id}")
async def preview_document_image(
    document_id: uuid.UUID,
    image_id: str,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Stream one inline image part (by relationship id) for the preview.

    ``image_id`` is the rId the parse output emitted in ``images[].image_id``;
    unknown ids (or ids pointing at non-image parts) 404.
    """
    doc, _ = await _get_org_document(document_id, org_id, db)
    _resolve_docx_filename(doc)
    docx_bytes = _load_docx_bytes(doc)

    from app.services.document_preview import extract_image_fast

    # Zip-level lookup (no full Document parse), off the event loop (B5/B6).
    extracted = await asyncio.to_thread(extract_image_fast, docx_bytes, image_id)
    if extracted is None:
        raise NotFoundError("Immagine non trovata nel documento")
    blob, content_type = extracted
    return StreamingResponse(
        io.BytesIO(blob),
        media_type=content_type,
        # A row's bytes never change (edits mint new rows), so the browser
        # may cache each image blob for as long as it likes.
        headers={"Cache-Control": "private, max-age=31536000, immutable"},
    )


@download_router.patch(
    "/{document_id}/overrides",
    response_model=OverridesResponse,
    dependencies=[Depends(require_capability(DOCUMENTS_GENERATE))],
)
async def patch_document_overrides(
    document_id: uuid.UUID,
    body: OverridesPatchRequest,
    org_id: uuid.UUID = Depends(get_current_org),
    db: AsyncSession = Depends(get_db),
):
    """Merge inline-edit overrides into the row: ``{addr: text}``, null deletes.

    Returns the full current map so the frontend can replace its copy
    wholesale after each save.
    """
    doc, _ = await _get_org_document(document_id, org_id, db)
    _resolve_docx_filename(doc)

    merged = dict(doc.content_overrides or {})
    for addr, text in body.set.items():
        if not _OVERRIDE_ADDR_RE.fullmatch(addr):
            raise BadRequestError(f"Indirizzo di blocco non valido: {addr}")
        if text is None:
            merged.pop(addr, None)
        else:
            # Normalize Windows/mac line endings before persisting so the
            # stored values stay clean and apply emits exactly one <w:br/>
            # per visual line break.
            merged[addr] = text.replace("\r\n", "\n").replace("\r", "\n")

    # Cap the MERGED map, not just the request — repeated PATCHes must not
    # bloat the JSONB row (or the apply path's memory) without bound.
    if len(json.dumps(merged, ensure_ascii=False)) > _OVERRIDES_MAX_SERIALIZED:
        raise BadRequestError(
            "Le modifiche superano la dimensione massima consentita."
        )

    # Reassign the whole dict — in-place mutation of a JSONB value is
    # invisible to SQLAlchemy's change tracking. Store NULL when the map
    # empties out so "no pending edits" stays queryable as IS NULL.
    doc.content_overrides = merged or None
    await db.commit()
    await db.refresh(doc)
    return OverridesResponse(overrides=doc.content_overrides or {})


@download_router.post(
    "/{document_id}/save-edited-version", response_model=DocumentResponse, status_code=201
)
async def save_edited_version(
    document_id: uuid.UUID,
    org_id: uuid.UUID = Depends(get_current_org),
    user: User = Depends(require_capability(DOCUMENTS_GENERATE)),
    ent: Entitlements = Depends(get_entitlements),
    db: AsyncSession = Depends(get_db),
):
    """Fold pending overrides into a NEW completed version row.

    Mirrors sync-from-gdoc: applies the overrides to the source bytes,
    inserts a v+1 row tagged ``options.edited_inline``, then clears the
    overrides on the source so its working copy is pristine again.
    """
    from fastapi import status

    # Row-lock the source (B8): two concurrent saves would otherwise both
    # read the same overrides and mint duplicate version rows. The second
    # transaction blocks here, re-reads the cleared content_overrides after
    # the first commits, and hits the 409 below.
    source, _ = await _get_org_document(document_id, org_id, db, for_update=True)
    _resolve_docx_filename(source)

    if not source.content_overrides:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Nessuna modifica da salvare",
        )

    # Folding overrides into v+1 mints a new completed version (MB-6.2).
    await _ensure_new_version_allowed(
        ent, org_id, source.azienda_id, source.tipo_documento, db
    )

    docx_bytes = _load_docx_bytes(source)

    from app.services.document_preview import apply_overrides_to_docx

    # CPU-bound docx rewrite — off the event loop (B5).
    new_bytes = await asyncio.to_thread(
        apply_overrides_to_docx, docx_bytes, source.content_overrides
    )

    # Next version for this document type
    result = await db.execute(
        select(DocumentoGenerato)
        .where(
            DocumentoGenerato.azienda_id == source.azienda_id,
            DocumentoGenerato.tipo_documento == source.tipo_documento,
        )
        .order_by(DocumentoGenerato.versione.desc())
        .limit(1)
    )
    latest = result.scalar_one_or_none()
    next_version = (latest.versione + 1) if latest else 1

    # Build filename: append -edited to stem
    src_name = source.file_name or f"{source.tipo_documento}_v{source.versione}.docx"
    stem, ext = os.path.splitext(src_name)
    new_name = f"{stem}_v{next_version}_edited{ext or '.docx'}"

    merged_options = dict(source.options or {})
    merged_options["edited_inline"] = True
    merged_options["source_version_id"] = str(source.id)

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    new_doc = DocumentoGenerato(
        azienda_id=source.azienda_id,
        tipo_documento=source.tipo_documento,
        versione=next_version,
        status="completed",
        # Carry the source's file_path (parity with restore): frontend
        # surfaces gate their download/preview buttons on a non-NULL
        # file_path, while the download endpoint prefers file_content —
        # so the NEW bytes are always the ones served.
        file_path=source.file_path,
        file_content=new_bytes,
        file_name=new_name,
        options=merged_options,
        generated_by=user.id,
        generation_started_at=now,
        generation_completed_at=now,
        # The edited version is the source bytes restyled — it inherits the
        # source's snapshot freshness verbatim (US-5.2 banner carries over).
        survey_snapshot_hash=source.survey_snapshot_hash,
        stale_snapshot=bool(source.stale_snapshot),
    )
    db.add(new_doc)
    # The edits are captured in the new row — clear the working overrides.
    source.content_overrides = None
    await db.commit()
    await db.refresh(new_doc)
    # MB-2.3 — saving an edited version mints a completed document.
    await record_activation_for_azienda(new_doc.azienda_id, db)

    return _doc_to_response(new_doc, await _resolve_user_name(new_doc.generated_by, db))
