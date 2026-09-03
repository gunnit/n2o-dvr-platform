"""DUVRI - Documento Unico Valutazione Rischi Interferenze (art. 26 D.Lgs. 81/2008)."""

import os
import unicodedata

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_duvri
from app.services.document_generator.design import finish_document
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    page_break,
    replace_placeholders,
    scrub_body,
    slugify,
)

TEMPLATE = TEMPLATES_DIR / "DUVRI.docx"
TIPO_DOC = "duvri"
UNKNOWN_ENVIRONMENT = "Ambiente non disponibile"


def _normalize_display_text(value) -> str:
    return " ".join(unicodedata.normalize("NFKC", str(value or "")).split())


def _company_equipment_rows(data: dict) -> list[list[str]]:
    environment_names = {
        str(ambiente.id): _normalize_display_text(ambiente.nome)
        or UNKNOWN_ENVIRONMENT
        for ambiente in data["ambienti"]
        if getattr(ambiente, "id", None) is not None
    }
    records = []
    for equipment in data["attrezzature"]:
        description = _normalize_display_text(
            getattr(equipment, "descrizione", None)
        )
        if not description:
            continue
        environment = environment_names.get(
            str(getattr(equipment, "ambiente_id", None)),
            UNKNOWN_ENVIRONMENT,
        )
        records.append((
            environment,
            description,
            str(getattr(equipment, "id", "")),
        ))
    records.sort(
        key=lambda record: (
            record[0].casefold(),
            record[1].casefold(),
            record[2],
        )
    )
    return [[description, environment] for environment, description, _ in records]


def _body_paragraph_with_marker(doc: Document, marker: str):
    marker = marker.casefold()
    return next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if marker in paragraph.text.casefold()
        ),
        None,
    )


def _add_company_equipment(doc: Document, rows: list[list[str]]) -> list:
    blocks = [add_heading(doc, "Attrezzature del committente", level=2)._p]
    if rows:
        blocks.append(
            add_data_table(
                doc,
                ["Attrezzatura del committente", "Ambiente"],
                rows,
            )._element
        )
    else:
        blocks.append(
            add_paragraph(
                doc,
                "Nessuna attrezzatura registrata nel Rischio Master.",
                italic=True,
            )._p
        )
    return blocks


def _insert_company_equipment_at_legacy_anchor(
    doc: Document,
    rows: list[list[str]],
) -> bool:
    anchor = _body_paragraph_with_marker(
        doc,
        "Attrezzature e mezzi messi a disposizione dal Committente",
    )
    if anchor is None:
        return False
    for block in _add_company_equipment(doc, rows):
        anchor._p.addprevious(block)
    return True


def _remove_body_paragraph_block(
    doc: Document,
    start_marker: str,
    end_marker: str,
    *,
    replacement: str | None = None,
) -> bool:
    """Remove a legacy body block delimited by stable text anchors."""
    paragraphs = list(doc.paragraphs)
    start = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if start_marker.casefold() in paragraph.text.casefold()
        ),
        None,
    )
    if start is None:
        return False
    end = next(
        (
            index
            for index, paragraph in enumerate(paragraphs[start:], start)
            if end_marker.casefold() in paragraph.text.casefold()
        ),
        None,
    )
    if end is None:
        return False
    if replacement:
        replacement_paragraph = add_paragraph(doc, replacement)
        paragraphs[start]._p.addprevious(replacement_paragraph._p)
    for paragraph in paragraphs[start : end + 1]:
        paragraph._p.getparent().remove(paragraph._p)
    return True


def _replace_body_tables_with_anchor(
    doc: Document,
    marker: str,
    replacement: str,
) -> None:
    """Replace donor tables with a narrow reference to current appalto data."""
    marker = marker.casefold()
    for table in list(doc.tables):
        table_text = "\n".join(
            cell.text for row in table.rows for cell in row.cells
        ).casefold()
        if marker in table_text:
            replacement_paragraph = add_paragraph(doc, replacement)
            table._element.addprevious(replacement_paragraph._p)
            table._element.getparent().remove(table._element)


def _remove_legacy_donor_equipment(doc: Document) -> None:
    """Drop template-only equipment examples without touching generic DUVRI text."""
    _remove_body_paragraph_block(
        doc,
        "Attrezzature e mezzi messi a disposizione dal Committente",
        "L’utilizzo di tali attrezzature avviene all’interno degli spazi aziendali condivisi",
    )
    _remove_body_paragraph_block(
        doc,
        "Movimentazione materiali – Muletto / Transpallet",
        "In caso di utilizzo contemporaneo, le attività sono coordinate tra preposti",
        replacement=(
            "Le misure di coordinamento applicabili sono riportate, per ciascun "
            "appalto, nella sezione Interferenze identificate in calce al "
            "presente documento."
        ),
    )
    _remove_body_paragraph_block(
        doc,
        "ATTIVITÀ DI SALDATURA – SVOLTA DA",
        "È vietato depositare materiali combustibili nelle aree limitrofe durante le operazioni",
    )
    _replace_body_tables_with_anchor(
        doc,
        "Attività operative RECOM",
        "Le date e le attività di ciascun appalto sono riportate nelle relative "
        "sezioni in calce al presente documento.",
    )
    _replace_body_tables_with_anchor(
        doc,
        "Possibili rischi/interferenze per RECOM",
        "Le attrezzature e le attività dell’appaltatore, con le relative "
        "interferenze, sono riportate nelle sezioni dedicate ai singoli appalti "
        "in calce al presente documento.",
    )


class DuvriGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        duvri_rows = await load_duvri(self.db, self.azienda_id)
        company_equipment_rows = _company_equipment_rows(data)
        company_equipment_inserted = False

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
            # Blank the sample dates left in the template body; the generator
            # supplies the real appalto/sottoscrizione dates below.
            scrub_body(doc, {"01.11.2025": "__/__/____", "15/01/2026": "__/__/____"})
            company_equipment_inserted = _insert_company_equipment_at_legacy_anchor(
                doc,
                company_equipment_rows,
            )
            _remove_legacy_donor_equipment(doc)
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"DUVRI - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Committente", azienda.ragione_sociale or ""),
            ("Sede", format_sede(azienda, "legale")),
            ("P.IVA committente", azienda.partita_iva or ""),
            ("Data emissione", generated_at.strftime("%d/%m/%Y")),
            ("Riferimento normativo", "Art. 26 D.Lgs. 81/2008 e D.Lgs. 106/2009"),
        ])

        add_heading(doc, "Oggetto del documento", level=2)
        add_paragraph(doc, "Il presente DUVRI individua le misure di prevenzione e protezione necessarie ad eliminare o ridurre al minimo i rischi derivanti dalle interferenze tra le attività del committente e quelle dell'impresa appaltatrice.")
        add_paragraph(doc, "Ai sensi dell'art. 26 comma 3-bis del D.Lgs. 81/2008 (introdotto dal D.Lgs. 106/2009), l'obbligo di redazione del DUVRI non si applica ai servizi di natura intellettuale, alle mere forniture di materiali o attrezzature, nonché ai lavori o servizi la cui durata non superi i cinque uomini-giorno, salvo che comportino rischi derivanti da agenti cancerogeni, biologici, atmosfere esplosive o dai rischi particolari di cui all'Allegato XI.")

        if not company_equipment_inserted:
            _add_company_equipment(doc, company_equipment_rows)

        if not duvri_rows:
            add_paragraph(doc, "Non risultano appalti attivi al momento della valutazione.", italic=True)
        for idx, d in enumerate(duvri_rows, 1):
            page_break(doc)
            add_heading(doc, f"{idx}. Appalto: {d.oggetto_appalto}", level=2)
            add_kv_table(doc, [
                ("Appaltatore", d.appaltatore_ragione_sociale or ""),
                ("P.IVA appaltatore", d.appaltatore_partita_iva or ""),
                ("Referente", d.appaltatore_referente or ""),
                ("Oggetto appalto", d.oggetto_appalto or ""),
                ("Data inizio", d.data_inizio.strftime("%d/%m/%Y") if d.data_inizio else "—"),
                ("Data fine", d.data_fine.strftime("%d/%m/%Y") if d.data_fine else "—"),
            ])

            attrezz = d.attrezzature_appaltatore or []
            if attrezz:
                add_heading(doc, "Attrezzature / attività appaltatore", level=3)
                rows = [
                    [a.get("tipo", ""), a.get("descrizione", "") or ""]
                    for a in attrezz
                    if isinstance(a, dict)
                ]
                add_data_table(doc, ["Tipo", "Descrizione"], rows)

            add_heading(doc, "Interferenze identificate", level=3)
            interfs = d.interferenze or []
            if interfs:
                rows = []
                for i in interfs:
                    dpi = ", ".join(i.get("dpi", [])) if isinstance(i.get("dpi"), list) else (i.get("dpi") or "")
                    rows.append([i.get("rischio", ""), i.get("misure", ""), dpi])
                add_data_table(doc, ["Rischio da interferenza", "Misure di coordinamento", "DPI"], rows)
            else:
                add_paragraph(doc, "Nessuna interferenza rilevante identificata.", italic=True)

            add_heading(doc, "Sottoscrizione", level=3)
            add_data_table(doc, ["Ruolo", "Firma"], [
                ["Committente (Datore di Lavoro)", "________________________"],
                ["Appaltatore", "________________________"],
                ["Data", generated_at.strftime("%d/%m/%Y")],
            ])

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        # Audit 2026-09-03: the donor template's header/footer carried the
        # legacy N2O letterhead and literal placeholders; rewrite them from
        # the organization's branding and set honest file properties.
        finish_document(
            doc,
            title='DUVRI - Valutazione dei Rischi da Interferenze',
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
            fill_cover=True,
            cover_values={
                "Oggetto dell'appalto": (duvri_rows[0].oggetto_appalto or '') if duvri_rows else '',
                'Azienda Committente': azienda.ragione_sociale or '',
                "Committente dell'opera": azienda.ragione_sociale or '',
                'Descrizione delle attività appaltate': (duvri_rows[0].oggetto_appalto or '') if duvri_rows else '',
                'Datore di Lavoro Committente': next(
                    ((getattr(p, 'nominativo', None) or '') for p in (data.get('persone') or []) if getattr(p, 'ruolo_datore_lavoro', False)),
                    '',
                ),
                "Indirizzo presso cui si svolgerà l'appalto": format_sede(azienda, 'operativa') if format_sede(azienda, 'operativa') != '—' else format_sede(azienda, 'legale'),
            },
        )
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "DUVRI"])
