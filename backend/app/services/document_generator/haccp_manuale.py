"""HACCP Manuale di Autocontrollo - Reg. CE 852/2004 + Reg. CE 178/2002.

The HACCP.docx template carries the comprehensive regulatory corpus
(definizioni, riferimenti legislativi, attrezzature, SSOP, SOP 01..N,
botulino procedures, allergeni) — ~1350 paragraphs, ~290 headings,
20 tables. This generator's job is azienda-specific *customization*
appended after the template body: cover header, attivita/CCPs/forms
table, supplier protocol references, monitoring schedule, sign-off.
"""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_haccp
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    page_break,
    replace_placeholders,
    slugify,
)

TEMPLATE = TEMPLATES_DIR / "HACCP.docx"
TIPO_DOC = "haccp"

# Standard food-safety SOPs that the customization section reminds the
# operator are documented in the template body. Helps an inspector see at
# a glance which SOPs apply to this azienda's tipologia.
SOP_INDEX: list[tuple[str, str]] = [
    ("SOP 01", "Manutenzione dei locali e delle attrezzature"),
    ("SOP 02", "Sanificazione (pulizia e disinfezione)"),
    ("SOP 03", "Lotta agli infestanti"),
    ("SOP 04", "Approvvigionamento idrico"),
    ("SOP 05", "Gestione rifiuti"),
    ("SOP 06", "Ricevimento materie prime e qualifica fornitori"),
    ("SOP 07", "Stoccaggio (temperatura, separazione crudo/cotto, FIFO)"),
    ("SOP 08", "Preparazione e lavorazione alimenti"),
    ("SOP 09", "Cottura e trattamenti termici"),
    ("SOP 10", "Raffreddamento rapido (blast-chiller)"),
    ("SOP 11", "Conservazione a caldo / a freddo"),
    ("SOP 12", "Rigenerazione e somministrazione"),
    ("SOP 13", "Trasporto degli alimenti"),
    ("SOP 14", "Igiene del personale"),
    ("SOP 15", "Formazione del personale"),
    ("SOP 16", "Gestione allergeni (Reg. UE 1169/2011)"),
    ("SOP 17", "Tracciabilita e rintracciabilita (Reg. CE 178/2002)"),
    ("SOP 18", "Gestione non conformita e azioni correttive"),
    ("SOP 19", "Verifica e revisione del piano HACCP"),
    ("SOP 20", "Botulino - controllo conserve e sottovuoto"),
    ("SOP 21", "Listeria - controllo prodotti pronti al consumo (RTE)"),
    ("SOP 22", "Salmonella - uova, pollame, prodotti a base di carne"),
    ("SOP 23", "Anisakis - prodotti ittici crudi/marinati"),
    ("SOP 24", "Acrilamide (Reg. UE 2017/2158)"),
    ("SOP 25", "Tarature termometri e strumenti di misura"),
]


class HaccpManualeGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        config, forms = await load_haccp(self.db, self.azienda_id)

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"MANUALE HACCP - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Sede operativa", format_sede(azienda, "legale")),
            ("P.IVA", azienda.partita_iva or "—"),
            ("Tipologia attivita", (config.tipologia_attivita if config else "—") or "—"),
            ("Numero pasti/giorno", str(config.numero_pasti_giorno) if (config and config.numero_pasti_giorno) else "—"),
            ("Responsabile HACCP", (config.responsabile_haccp if config else "—") or "—"),
            ("Data emissione", generated_at.strftime("%d/%m/%Y")),
            ("Riferimenti normativi", "Reg. CE 852/2004, Reg. CE 178/2002, Reg. CE 2073/2005, Reg. UE 1169/2011"),
        ])

        add_heading(doc, "Campo di applicazione", level=2)
        add_paragraph(doc, "Il presente manuale si applica a tutte le fasi di preparazione, cottura, conservazione e somministrazione degli alimenti svolte presso l'azienda. Le procedure operative standard (SOP) descritte nel corpo del documento sono adottate integralmente con eventuali personalizzazioni indicate di seguito.")

        add_heading(doc, "Principi HACCP applicati", level=2)
        add_paragraph(doc, "Il sistema si fonda sui 7 principi Codex Alimentarius: 1) analisi dei pericoli, 2) identificazione CCP, 3) limiti critici, 4) monitoraggio, 5) azioni correttive, 6) verifica, 7) documentazione. Per ciascun principio le evidenze sono registrate sulle schede di autocontrollo SA-01 ÷ SA-16.")

        if config:
            add_heading(doc, "Alimenti trattati", level=2)
            tipi = config.tipi_alimenti_trattati or []
            if tipi:
                add_paragraph(doc, ", ".join(tipi))
            else:
                add_paragraph(doc, "Da definire — completare al primo audit del Responsabile HACCP.", italic=True)

            add_heading(doc, "Punti critici di controllo (CCP) - personalizzazione azienda", level=2)
            ccps = config.ccps or []
            if ccps:
                rows = [[
                    c.get("codice", ""),
                    c.get("nome", ""),
                    c.get("limite_critico", ""),
                    c.get("monitoraggio", "Verifica giornaliera + registrazione su scheda"),
                ] for c in ccps]
                add_data_table(doc, ["Codice", "CCP", "Limite critico", "Monitoraggio"], rows)
            else:
                add_paragraph(doc, "CCP da definire — il piano di base prevede CCP su Ricevimento, Stoccaggio, Cottura, Raffreddamento, Conservazione. Confermare al primo audit.", italic=True)

            # Feedback #65 — equipment list, highlighting which items are
            # subject to HACCP control (cleaning/maintenance/temperature plan).
            add_heading(doc, "Attrezzature e controllo HACCP", level=2)
            attrezzature = config.attrezzature or []
            if attrezzature:
                rows = [[
                    a.get("nome", ""),
                    "Sì" if a.get("sotto_controllo_haccp") else "No",
                ] for a in attrezzature]
                add_data_table(doc, ["Attrezzatura", "Sottoposta a controllo HACCP"], rows)
            else:
                add_paragraph(doc, "Elenco attrezzature da definire — censire le attrezzature e indicare quelle sottoposte a controllo HACCP al primo audit.", italic=True)

        add_heading(doc, "Indice delle Procedure Operative Standard (SOP)", level=2)
        add_paragraph(doc, "Le seguenti SOP sono documentate per esteso nel corpo del manuale (D.Lgs. 193/2007, Reg. CE 852/2004 Allegato II). Il responsabile HACCP verifica che ciascuna SOP applicabile sia stata trasmessa al personale.", italic=True, size=9)
        add_data_table(doc, ["Codice", "Titolo"], [[code, title] for code, title in SOP_INDEX])

        add_heading(doc, "Procedure di monitoraggio", level=2)
        add_paragraph(doc, "Ogni CCP e monitorato mediante le schede di autocontrollo SA-01 ÷ SA-16 (allegate), compilate secondo la frequenza indicata in ciascuna scheda. Le registrazioni sono archiviate per almeno 12 mesi e disponibili alle Autorita di controllo (ASL, NAS, USMAF).")

        add_heading(doc, "Gestione delle non conformita", level=2)
        add_paragraph(doc, "In caso di superamento dei limiti critici, l'alimento viene isolato e identificato. Se non recuperabile, viene smaltito con registrazione sulla scheda SA-13 (gestione non conformita). L'azione correttiva viene comunicata al responsabile HACCP e, se l'alimento e gia stato distribuito, si attiva la procedura di richiamo/ritiro (Reg. CE 178/2002 art. 19).")

        add_heading(doc, "Formazione del personale", level=2)
        add_paragraph(doc, "Tutti gli operatori del settore alimentare (O.S.A.) ricevono formazione HACCP ai sensi dell'art. 4 Reg. CE 852/2004 all'assunzione, con aggiornamento biennale. Per gli addetti alla manipolazione di alimenti deperibili e prevista formazione specifica annuale (Accordo Stato-Regioni 27/01/2010).")

        add_heading(doc, "Verifica e revisione del piano", level=2)
        add_paragraph(doc, "Il piano HACCP e oggetto di verifica almeno annuale da parte del Responsabile HACCP, e di revisione straordinaria in caso di: variazione attivita/menu, nuova attrezzatura, non conformita ripetute, modifiche normative, segnalazioni delle Autorita di controllo.")

        add_heading(doc, "Schede di autocontrollo allegate", level=2)
        if forms:
            add_data_table(doc, ["Codice", "Titolo"], [[f.form_code, f.form_title] for f in forms])
        else:
            add_paragraph(doc, "Schede in fase di configurazione — generare tramite il modulo HACCP_FORMS.", italic=True)

        add_heading(doc, "Sottoscrizione", level=2)
        add_data_table(doc, ["Ruolo", "Nominativo", "Firma"], [
            ["Datore di lavoro / O.S.A.", azienda.ragione_sociale or "", "________________________"],
            ["Responsabile HACCP", (config.responsabile_haccp if config else "") or "________________________", "________________________"],
            ["Data emissione", generated_at.strftime("%d/%m/%Y"), ""],
        ])

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento.in_([TIPO_DOC, "HACCP"]))
        )
        r = await self.db.execute(stmt)
        return (r.scalar() or 0) + 1
