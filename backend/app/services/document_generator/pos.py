"""POS - Piano Operativo di Sicurezza (D.Lgs. 81/2008 Titolo IV All. XV)."""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_pos
from app.services.document_generator.design import finish_document
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    page_break,
    replace_placeholders,
    slugify,
)
from app.services.dpi_rules import (
    DPI_CATALOG,
    PHASES_CONSTRUCTION,
    ROLES_CONSTRUCTION,
    build_default_matrix,
)

# Italian labels for the standard phases — used when the operator hasn't
# defined custom fasi yet so the document still ships with a defensible
# 8-phase skeleton (D.Lgs. 81/2008 All. XV punto 3.2.1 d).
_PHASE_LABELS_IT: dict[str, str] = {
    "allestimento_cantiere": "Allestimento del cantiere",
    "scavi": "Scavi e movimento terra",
    "fondazioni": "Fondazioni",
    "getto_calcestruzzo": "Getto del calcestruzzo",
    "montaggio_ponteggi": "Montaggio dei ponteggi",
    "opere_murarie": "Opere murarie / strutturali",
    "finiture": "Finiture (intonaco, copertura, facciata)",
    "smobilizzo_cantiere": "Smobilizzo del cantiere",
}

# Italian labels for the standard construction roles.
_ROLE_LABELS_IT: dict[str, str] = {
    "carpentiere": "Carpentiere",
    "manovale": "Manovale",
    "gruista": "Gruista",
    "operatore_escavatore": "Operatore escavatore",
    "ponteggiatore": "Ponteggiatore",
    "saldatore": "Saldatore",
    "elettricista": "Elettricista",
    "muratore": "Muratore",
    "capo_cantiere": "Capo cantiere",
    "autista_mezzi": "Autista mezzi",
}

# ---------------------------------------------------------------------------
# Figure di sicurezza sul cantiere (client request 2026-08-13)
# ---------------------------------------------------------------------------

# Italian labels for FIGURE_SICUREZZA_RUOLI (schemas/pos.py). Same key split
# as the DPI matrix: stable keys in the DB, labels here and on the frontend.
_FIGURE_LABELS_IT: dict[str, str] = {
    "datore_lavoro": "Datore di Lavoro",
    "direttore_tecnico_cantiere": "Direttore Tecnico di Cantiere",
    "capocantiere_preposto": "Capocantiere / Preposto",
    "rspp": "Responsabile del Servizio di Prevenzione e Protezione (RSPP)",
    "rls": "Rappresentante dei Lavoratori per la Sicurezza (RLS)",
    "medico_competente": "Medico Competente",
    "addetto_primo_soccorso": "Addetto al Primo Soccorso",
    "addetto_antincendio": "Addetto Antincendio",
}

# Mansionario boilerplate — abridged from the N2O original POS ("MANSIONARIO"
# chapter). Rendered under the figure table for every figure present so the
# printed POS carries the default diciture the operator expects. The original
# texts run several pages; these are the lead paragraphs, verbatim where
# possible. Full-length texts to be confirmed with the client.
_MANSIONARIO_IT: dict[str, list[str]] = {
    "datore_lavoro": [
        "Al Datore di Lavoro competono i compiti individuati dalla vigente "
        "normativa, ed in particolare quelli sanciti dal D.Lgs. 81/2008 e "
        "s.m.i. in quanto titolare dei poteri illimitati di gestione e di "
        "spesa in materia di sicurezza ed igiene del lavoro.",
        "Ogni squadra di lavoro dovrà essere formata almeno da un preposto, "
        "due addetti al servizio primo soccorso ed un addetto al servizio "
        "emergenze ed antincendio (le figure possono essere "
        "contemporaneamente coperte da un medesimo lavoratore).",
        "Il POS può essere sottoscritto esclusivamente dal datore di lavoro "
        "in quanto la valutazione del rischio, ai sensi e per gli effetti di "
        "cui al D.Lgs. 81/2008 e s.m.i., non è fra le attività delegabili.",
    ],
    "direttore_tecnico_cantiere": [
        "Il direttore tecnico di cantiere è una figura apicale obbligatoria "
        "dell'affidataria, incaricata dell'organizzazione, della gestione e "
        "della conduzione del cantiere. Mantiene i rapporti con la direzione "
        "dei lavori, coordina e segue l'esecuzione delle prestazioni in "
        "contratto e sovrintende all'adattamento, all'applicazione e "
        "all'osservanza dei piani di sicurezza.",
        "Verifica, con l'ausilio del RSPP, la rispondenza alle norme delle "
        "attrezzature di lavoro (macchine, impianti, DPI) messe a "
        "disposizione dei propri lavoratori e mette a disposizione dei "
        "lavoratori i necessari DPI.",
    ],
    "capocantiere_preposto": [
        "Il capocantiere rappresenta il raccordo tra la direzione generale "
        "di cantiere e le maestranze dell'impresa esecutrice; cura "
        "l'attuazione delle misure di sicurezza previste dalle norme in "
        "vigore e secondo le disposizioni impartite dalla Direzione del "
        "cantiere anche tramite il piano di sicurezza, e sospende il lavoro "
        "qualora, a suo giudizio, la prosecuzione dello stesso si rivelasse "
        "pericolosa per l'incolumità dei lavoratori o di terzi.",
        "I preposti organizzano e coordinano il lavoro nelle aree di propria "
        "competenza, rendendo edotti i lavoratori dei rischi cui sono "
        "sottoposti ed esigendo l'osservanza delle norme antinfortunistiche; "
        "obbligano i lavoratori ad indossare i dispositivi di protezione "
        "individuale e sono destinatari degli obblighi di cui all'art. 19 "
        "del D.Lgs. 81/2008 e s.m.i.",
    ],
    "rspp": [
        "Il Servizio di Prevenzione e Protezione provvede "
        "all'individuazione dei fattori di rischio, alla valutazione dei "
        "rischi e all'individuazione delle misure per la sicurezza e la "
        "salubrità degli ambienti di lavoro nel rispetto della normativa "
        "vigente; elabora le procedure di sicurezza per le varie attività "
        "aziendali, propone i programmi di informazione e formazione dei "
        "lavoratori e partecipa alla riunione periodica di prevenzione e "
        "protezione dai rischi.",
    ],
    "rls": [
        "Il Rappresentante dei lavoratori per la sicurezza collabora con il "
        "datore di lavoro e con il servizio di prevenzione e protezione, "
        "sulla base della specifica conoscenza dell'organizzazione "
        "dell'azienda e delle situazioni di rischio, alla predisposizione e "
        "attuazione delle misure per la tutela della salute e "
        "dell'integrità psico-fisica dei lavoratori.",
    ],
    "medico_competente": [
        "Il Medico Competente collabora con il servizio di prevenzione e "
        "protezione alla predisposizione ed all'attuazione delle misure per "
        "la tutela della salute e dell'integrità psicofisica dei "
        "lavoratori; effettua gli accertamenti sanitari preventivi e "
        "periodici previsti dalla legge ed esprime i giudizi di idoneità "
        "alla mansione specifica.",
    ],
    "addetto_primo_soccorso": [
        "L'addetto al primo soccorso interviene in caso di necessità prima "
        "dell'arrivo dei soccorsi specializzati: riconosce l'emergenza "
        "sanitaria, allerta il sistema di soccorso (112), attua gli "
        "interventi di primo soccorso e tiene sotto controllo la cassetta "
        "di primo soccorso, rifornendola quando occorre.",
    ],
    "addetto_antincendio": [
        "L'addetto antincendio è un lavoratore preposto ad attuare le "
        "misure di prevenzione incendi e di lotta antincendio, nonché a "
        "gestire le relative emergenze che si possono presentare nel "
        "cantiere: valuta tempestivamente l'entità dell'emergenza, agisce "
        "anche attraverso l'utilizzo di estintori se l'emergenza si "
        "manifesta di lieve entità, attiva le procedure di emergenza e i "
        "contatti con il soccorso esterno qualora l'emergenza non sia "
        "facilmente controllabile e agevola l'evacuazione dei presenti "
        "garantendo a tutti la possibilità di raggiungere il punto di "
        "ritrovo.",
    ],
}

# Persona.ruolo_* flag → figure key, used to prefill the figure table from
# the organigramma when the operator hasn't picked anyone yet (platform
# principle: prefill what the operator can correct).
_FIGURE_PERSONA_FLAGS: list[tuple[str, str]] = [
    ("datore_lavoro", "ruolo_datore_lavoro"),
    ("capocantiere_preposto", "ruolo_preposto"),
    ("rspp", "ruolo_rspp"),
    ("rls", "ruolo_rls"),
    ("medico_competente", "ruolo_medico_competente"),
    ("addetto_primo_soccorso", "ruolo_primo_soccorso"),
    ("addetto_antincendio", "ruolo_antincendio"),
]

# ---------------------------------------------------------------------------
# Default diciture ported from the N2O original POS (Drive templates folder,
# "POS.docx"). Verbatim except where marked. See MIGRATION-NOTES-pos.md and
# the implementation report for which texts still need client confirmation.
# ---------------------------------------------------------------------------

_DOCUMENTAZIONE_CANTIERE_INTRO = (
    "Si riporta di seguito un elenco indicativo e non esaustivo della "
    "documentazione che deve essere conservata in cantiere a cura "
    "dell'Impresa:"
)

_DOCUMENTAZIONE_CANTIERE: list[str] = [
    "Piano operativo di sicurezza",
    "Certificato di iscrizione C.C.I.A.A.",
    "Modelli UNILAV",
    "Denuncia INAIL inizio attività e variazioni",
    "Documento Unico di Regolarità Contributiva (D.U.R.C.)",
    "Dichiarazione organico medio annuo e tipo di contratto applicato con i dipendenti",
    "Documento di valutazione dei rischi ex art. 17 D.Lgs. 81/2008 e s.m.i.",
    "Designazione del Responsabile del Servizio di Prevenzione e Protezione",
    "Adempimento dell'obbligo formativo/informativo, ex D.Lgs. 81/2008 e s.m.i.",
    "Nomina dei coordinatori dell'emergenza ed elenco dei componenti",
    "Registro degli infortuni debitamente vidimato",
    "Nomina Medico Competente e registro visite mediche dipendenti ed elenco accertamenti sanitari periodici",
    "Copia della dichiarazione di conformità dell'impianto di messa a terra e di protezione contro le scariche atmosferiche e copia della richiesta delle verifiche periodiche",
    "Denuncia degli apparecchi di sollevamento di portata superiore a kg 200 e relative verifiche",
    "Libretti degli apparecchi di sollevamento con portata superiore a 200 kg",
    "Schede delle verifiche trimestrali alle funi e catene, anche per gli apparecchi di portata inferiore a kg 200",
    "Copia dell'autorizzazione ministeriale del ponteggio metallico, ovvero disegno esecutivo e relazione di calcolo ove richiesti",
    "Libretto rilasciato dal costruttore del ponteggio, indicante i limiti di carico e le modalità di impiego",
    "Denuncia annuale concernente produzione, trasporto, stoccaggio dei rifiuti",
    "Registro di carico e scarico, vidimato dall'Ufficio del Registro",
]

_MISURE_LAVORATORI_DEVONO: list[str] = [
    "osservare con attenzione le misure di sicurezza predisposte dall'impresa;",
    "usare con cura e costantemente i dispositivi di protezione individuale "
    "e gli altri mezzi messi a loro disposizione;",
    "segnalare al responsabile di cantiere gli eventuali guasti sopravvenuti "
    "ai dispositivi di protezione individuale o la loro intollerabilità, "
    "chiedendone la sostituzione;",
    "avvertire immediatamente il responsabile di cantiere qualora "
    "individuino o sospettino situazioni di pericolo;",
    "adoperarsi direttamente, in caso di urgenza e nell'ambito delle proprie "
    "competenze e possibilità, per eliminare o ridurre i pericoli.",
]

_MISURE_LAVORATORI_NON_DEVONO: list[str] = [
    "rimuovere o modificare i dispositivi e gli altri mezzi di sicurezza e "
    "di protezione senza averne ottenuta l'autorizzazione;",
    "compiere, di propria iniziativa, operazioni o manovre che non siano di "
    "loro competenza e che possano compromettere la sicurezza propria o di "
    "altre persone.",
]

_MISURE_DPI_PARAGRAFI: list[str] = [
    "Il lavoratore deve indossare i dispositivi di protezione individuale "
    "messi a sua disposizione e rispettare le istruzioni di impiego.",
    "Tali dispositivi devono essere custoditi in luogo adatto e "
    "accessibile, mantenuti in condizioni di perfetta efficienza e "
    "contrassegnati col nome dell'assegnatario.",
    "Il lavoratore che noti un qualsiasi difetto o la mancata tolleranza "
    "del proprio dispositivo di protezione individuale deve chiederne la "
    "sostituzione.",
    "Il rifiuto ad indossare il dispositivo di protezione individuale, "
    "previsto per l'attività in atto, comporta la mancanza di idoneità al "
    "lavoro stesso.",
]

# Adapted from the original's art. 97 passage (Direttore Tecnico chapter).
_SUBAPPALTI_VIGILANZA = (
    "L'impresa vigila sulla sicurezza dei lavori affidati in subappalto e "
    "sull'applicazione delle disposizioni e delle prescrizioni del piano di "
    "sicurezza e coordinamento, verificando la congruenza dei piani "
    "operativi di sicurezza delle imprese subappaltatrici rispetto al "
    "proprio, prima della trasmissione al coordinatore per l'esecuzione "
    "(art. 97 D.Lgs. 81/2008 e s.m.i.)."
)

# Standard wording, NOT in the original (which has no dedicated subappalti
# section) — to be confirmed with the client.
_SUBAPPALTI_ASSENTI = (
    "Non è previsto l'affidamento di lavorazioni in subappalto per il "
    "presente cantiere. Qualora nel corso dei lavori si rendesse necessario "
    "ricorrere al subappalto, il presente POS sarà aggiornato e i piani "
    "operativi di sicurezza delle imprese subappaltatrici saranno "
    "verificati ai sensi dell'art. 97 del D.Lgs. 81/2008 e s.m.i."
)

_DICHIARAZIONE_TESTO = (
    "che il procedimento sulla valutazione dei rischi ex art. 17 del "
    "D.Lgs. 81/2008 e s.m.i., è stato attuato in collaborazione con il "
    "Servizio di Prevenzione e Protezione dai rischi, con il Medico "
    "Competente previa consultazione del Rappresentante dei lavoratori per "
    "la sicurezza."
)
from app.services.pos_phases import dependency_violations_after_ordering
from app.schemas.pos_phase import PosPhase
from pydantic import ValidationError

TEMPLATE = TEMPLATES_DIR / "POS.docx"
TIPO_DOC = "pos"


class PosGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        persone = data.get("persone") or []
        generated_at = data["generated_at"]
        pos_rows = await load_pos(self.db, self.azienda_id)

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"POS - {azienda.ragione_sociale}", level=1)

        # Default dicitura from the N2O original: elenco documentazione da
        # conservare in cantiere (Introduzione chapter).
        _render_documentazione_cantiere(doc)

        if not pos_rows:
            add_paragraph(doc, "Nessun cantiere registrato per questa azienda.", italic=True)
        for idx, p in enumerate(pos_rows, 1):
            page_break(doc)
            add_heading(doc, f"{idx}. Cantiere: {p.cantiere_indirizzo}", level=2)

            # Anagrafica + dati cantiere (header)
            add_kv_table(doc, [
                ("Impresa esecutrice", azienda.ragione_sociale or ""),
                ("Indirizzo cantiere", p.cantiere_indirizzo or ""),
                ("Descrizione", p.cantiere_descrizione or ""),
                ("Data inizio", p.data_inizio.strftime("%d/%m/%Y") if p.data_inizio else "—"),
                ("Data fine", p.data_fine.strftime("%d/%m/%Y") if p.data_fine else "—"),
                ("Importo lavori", f"{float(p.importo_lavori):,.2f} EUR" if p.importo_lavori else "—"),
                ("Numero massimo lavoratori", str(p.numero_massimo_lavoratori) if p.numero_massimo_lavoratori else "—"),
            ])

            # Soggetti di riferimento (All. XV punto 3.2.1 b)
            add_heading(doc, "Soggetti di riferimento", level=3)
            add_kv_table(doc, [
                ("Committente", p.committente or "—"),
                ("Progettista responsabile", p.progettista_responsabile or "—"),
                ("Direttore dei lavori", p.direttore_lavori or "—"),
                ("Direttore operativo edilizia / strutture", p.direttore_operativo_edilizia or "—"),
                ("Direttore operativo impianti", p.direttore_operativo_impianti or "—"),
                ("Responsabile dei lavori", p.responsabile_lavori or "—"),
                ("Coordinatore per la sicurezza in fase di progettazione (CSP)", p.coordinatore_progettazione or "—"),
                ("Coordinatore per la sicurezza in fase di esecuzione (CSE)", p.coordinatore_sicurezza or "—"),
            ])

            # Figure di sicurezza sul cantiere (client request 2026-08-13).
            # Prefilled from the organigramma when the operator hasn't
            # assigned them yet.
            _render_figure_sicurezza(doc, p, persone)

            # Dipendenti dell'azienda — tabella ruoli operativi. When the
            # operator selected the dipendenti presenti in cantiere, only
            # those rows are printed.
            _render_dipendenti_table(
                doc, persone, getattr(p, "dipendenti_cantiere", None) or []
            )

            # Modalità organizzative (All. XV punto 3.2.1 c)
            _render_modalita_organizzative(doc, p)

            # Organizzazione logistica
            _render_organizzazione_logistica(doc, p)

            # Subappalti (client request 2026-08-13)
            _render_subappalti(doc, p)

            add_heading(doc, "Fasi lavorative", level=3)
            fasi = p.fasi_lavorative or []
            if fasi:
                _render_phase_sections(doc, fasi)
            else:
                _render_default_phase_skeleton(doc)

            add_heading(doc, "Valutazioni specifiche", level=3)
            rum = p.valutazione_rumore or {}
            vib = p.valutazione_vibrazioni or {}
            add_kv_table(doc, [
                ("Lex 8h (dB(A))", str(rum.get("lex_8h_dba", "—"))),
                ("Fascia rumore", rum.get("fascia", "—")),
                ("DPI uditivi obbligatori", "SI" if rum.get("dpi_obbligatori") else "NO"),
                ("a8 mano-braccio (m/s^2)", str(vib.get("a8_mano_braccio", "—"))),
                ("a8 corpo intero (m/s^2)", str(vib.get("a8_corpo_intero", "—"))),
                ("Entro i limiti di legge", "SI" if vib.get("entro_limiti") else "NO"),
            ])

            add_heading(doc, "Mezzi e attrezzature", level=3)
            mezzi = p.mezzi_attrezzature or []
            add_data_table(doc, ["Tipo"], [[m.get("tipo", "")] for m in mezzi] or [["—"]])

            # Sostanze pericolose (All. XV punto 3.2.1 e) — flag-driven:
            # renders the manual list or the original's default dicitura.
            _render_sostanze_pericolose(doc, p, azienda.ragione_sociale or "L'impresa")

            # US-4.8: DPI matrix (role x phase). Only emit when the operator
            # has actually built one — we never auto-seed at generation time
            # because the matrix is a per-client override surface.
            _render_dpi_matrix(doc, p)

        # Default diciture from the N2O original: principali misure di
        # prevenzione + DPI, then the closing Dichiarazione with the firma
        # block (DdL / RSPP / RLS).
        _render_misure_prevenzione(doc)
        _render_dichiarazione(doc, azienda, persone, generated_at)

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        # Audit 2026-09-03: the donor template's header/footer carried the
        # legacy N2O letterhead and literal placeholders; rewrite them from
        # the organization's branding and set honest file properties.
        finish_document(
            doc,
            title='Piano Operativo di Sicurezza',
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
            fill_cover=True,
            cover_values={
                "Oggetto dell'appalto": (pos_rows[0].cantiere_descrizione or '') if pos_rows else '',
                'Indirizzo del cantiere': (pos_rows[0].cantiere_indirizzo or '') if pos_rows else '',
            },
        )
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "POS"])


# ---------------------------------------------------------------------------
# Dipendenti / modalità organizzative / organizzazione logistica
# ---------------------------------------------------------------------------


def _bullet(doc, text: str) -> None:
    """Bullet paragraph with a graceful fallback when the docx template
    doesn't carry the "List Bullet" style (safety documents must never fail
    to render over a missing style)."""
    try:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    except Exception:
        add_paragraph(doc, f"• {text}")


def _render_documentazione_cantiere(doc) -> None:
    """Default dicitura from the N2O original: elenco della documentazione
    che l'impresa deve conservare in cantiere."""
    add_heading(doc, "Documentazione da conservare in cantiere", level=2)
    add_paragraph(doc, _DOCUMENTAZIONE_CANTIERE_INTRO)
    for item in _DOCUMENTAZIONE_CANTIERE:
        _bullet(doc, item)


def _figura_entries(pos, persone: list) -> tuple[list[tuple[str, str]], bool]:
    """Resolve the printable (label, nominativo) rows for the figure table.

    Operator-selected entries (``pos.figure_sicurezza``) win. When the list
    is empty we derive assignees from the Persona ``ruolo_*`` flags so the
    document is prefilled from the organigramma (platform principle) —
    the second return value tells the caller the rows are auto-derived.
    """
    persona_by_id = {str(getattr(pe, "id", "")): pe for pe in persone}

    selected = getattr(pos, "figure_sicurezza", None) or []
    if selected:
        rows: list[tuple[str, str]] = []
        for f in selected:
            if not isinstance(f, dict):
                continue
            ruolo = f.get("ruolo") or ""
            label = _FIGURE_LABELS_IT.get(ruolo, ruolo.replace("_", " ").capitalize())
            nominativo = f.get("nominativo") or ""
            if not nominativo and f.get("persona_id"):
                pe = persona_by_id.get(str(f["persona_id"]))
                nominativo = getattr(pe, "nominativo", "") if pe else ""
            rows.append((label, nominativo or "—"))
        return rows, False

    # Derive from the organigramma flags. Roles with no flagged persona
    # still appear with "—" so the reviewer sees what is missing.
    derived: dict[str, list[str]] = {}
    for figura_key, flag_attr in _FIGURE_PERSONA_FLAGS:
        names = [
            pe.nominativo
            for pe in persone
            if getattr(pe, flag_attr, False) and getattr(pe, "nominativo", None)
        ]
        if names:
            derived[figura_key] = names
    rows = [
        (label, ", ".join(derived.get(key, [])) or "—")
        for key, label in _FIGURE_LABELS_IT.items()
    ]
    return rows, True


def _render_figure_sicurezza(doc, pos, persone: list) -> None:
    """Emit the "Figure di sicurezza sul cantiere" table + mansionario.

    The mansionario paragraphs (default diciture from the N2O original) are
    rendered for every figure that appears in the table, so the printed POS
    always explains each role's duties.
    """
    rows, derived = _figura_entries(pos, persone)
    if not rows:
        return
    add_heading(doc, "Figure di sicurezza sul cantiere", level=3)
    if derived:
        add_paragraph(
            doc,
            "Figure precompilate automaticamente dall'organigramma "
            "aziendale. Il Datore di Lavoro verifica e integra le nomine "
            "specifiche per il cantiere.",
            italic=True,
            size=9,
        )
    add_data_table(doc, ["Figura", "Nominativo"], [[label, nome] for label, nome in rows])

    add_heading(doc, "Mansionario delle figure di sicurezza", level=4)
    rendered_labels = {label for label, _ in rows}
    for key, label in _FIGURE_LABELS_IT.items():
        if label not in rendered_labels:
            continue
        paragraphs = _MANSIONARIO_IT.get(key)
        if not paragraphs:
            continue
        add_paragraph(doc, label, bold=True)
        for text in paragraphs:
            add_paragraph(doc, text)


def _render_subappalti(doc, pos) -> None:
    """Emit the "Subappalti" section (client request 2026-08-13).

    Flag off → default "nessun subappalto" dicitura. Flag on → table of
    subappaltatori (ragione sociale + lavori affidati) plus the art. 97
    vigilanza dicitura adapted from the N2O original.
    """
    add_heading(doc, "Subappalti", level=3)
    presenti = bool(getattr(pos, "subappalti_presenti", False))
    subappaltatori = getattr(pos, "subappaltatori", None) or []
    if not presenti or not subappaltatori:
        add_paragraph(doc, _SUBAPPALTI_ASSENTI)
        return
    rows = []
    for s in subappaltatori:
        if not isinstance(s, dict):
            continue
        rows.append([s.get("ragione_sociale", "") or "—", s.get("lavori", "") or "—"])
    add_data_table(doc, ["Impresa subappaltatrice", "Lavori affidati"], rows or [["—", "—"]])
    add_paragraph(doc, _SUBAPPALTI_VIGILANZA)


def _render_sostanze_pericolose(doc, pos, ragione_sociale: str) -> None:
    """Emit the "Sostanze pericolose" section (All. XV punto 3.2.1 e).

    Flag-driven: when the operator declared no hazardous substances, the
    default dicitura from the N2O original is printed instead of an empty
    table. A populated legacy list keeps rendering even if the flag was
    never set (pre-2026-08 rows).
    """
    add_heading(doc, "Sostanze pericolose utilizzate in cantiere", level=3)
    sostanze = getattr(pos, "sostanze_pericolose", None) or []
    presenti = bool(getattr(pos, "sostanze_pericolose_presenti", False)) or bool(sostanze)
    if not presenti:
        add_paragraph(
            doc,
            f"{ragione_sociale} non utilizzerà sostanze chimiche nelle "
            "lavorazioni che effettuerà presso il cantiere.",
        )
        return
    rows = [
        [s.get("nome", ""), s.get("uso", "") or "—"]
        for s in sostanze
        if isinstance(s, dict)
    ]
    add_data_table(doc, ["Sostanza", "Uso"], rows or [["—", "—"]])
    add_paragraph(
        doc,
        "Le schede di sicurezza (SDS) delle sostanze elencate sono "
        "conservate in cantiere a disposizione dei lavoratori e degli "
        "organi di vigilanza (All. XV punto 3.2.1 lettera e, D.Lgs. "
        "81/2008 e s.m.i.).",
    )


def _render_misure_prevenzione(doc) -> None:
    """Default diciture from the N2O original: principali misure di
    prevenzione (obblighi/divieti dei lavoratori) + dispositivi di
    protezione individuale."""
    add_heading(doc, "Principali misure di prevenzione", level=2)
    add_paragraph(doc, "I lavoratori devono:", bold=True)
    for item in _MISURE_LAVORATORI_DEVONO:
        _bullet(doc, item)
    add_paragraph(doc, "I lavoratori non devono:", bold=True)
    for item in _MISURE_LAVORATORI_NON_DEVONO:
        _bullet(doc, item)
    add_paragraph(doc, "Dispositivi di protezione individuale", bold=True)
    for text in _MISURE_DPI_PARAGRAFI:
        add_paragraph(doc, text)


def _render_dichiarazione(doc, azienda, persone: list, generated_at) -> None:
    """Closing "Dichiarazione" ported from the N2O original: DdL statement
    + firma block (Datore di Lavoro / RSPP / RLS)."""
    ddl = next(
        (
            pe.nominativo
            for pe in persone
            if getattr(pe, "ruolo_datore_lavoro", False)
            and getattr(pe, "nominativo", None)
        ),
        None,
    )
    sede = ", ".join(
        str(part)
        for part in (
            getattr(azienda, "sede_legale_via", None),
            getattr(azienda, "sede_legale_citta", None),
        )
        if part
    )
    add_heading(doc, "Dichiarazione", level=2)
    add_paragraph(
        doc,
        f"Il sottoscritto, {ddl or '________________'}, in qualità di "
        f"datore di lavoro della Impresa {azienda.ragione_sociale or '—'} "
        f"con sede a {sede or '________________'}",
    )
    add_paragraph(doc, "D I C H I A R A", bold=True)
    add_paragraph(doc, _DICHIARAZIONE_TESTO)
    add_data_table(doc, ["Ruolo", "Firma"], [
        ["Il Datore di Lavoro", "________________________"],
        ["Il Responsabile del S.P.P.", "________________________"],
        ["Il Rappresentante dei lavoratori per la sicurezza", "________________________"],
        ["Data", generated_at.strftime("%d/%m/%Y")],
    ])


def _render_dipendenti_table(doc, persone: list, dipendenti_cantiere: list | None = None) -> None:
    """Render the "Dipendenti dell'azienda" table required by Luca's
    2026-05-25 annotated template — Nominativo / Mansione / Primo Soccorso /
    Antincendio / Preposto. Pulled live from the azienda's Persona rows.

    ``dipendenti_cantiere`` (client request 2026-08-13) is the list of
    Persona ids the operator marked as working on this cantiere. Non-empty →
    only those rows are printed under a cantiere-specific heading; empty →
    legacy behaviour, every registered dipendente.
    """
    selected_ids = {str(x) for x in (dipendenti_cantiere or []) if x}
    if selected_ids:
        filtered = [pe for pe in persone if str(getattr(pe, "id", "")) in selected_ids]
        # Stale ids (persona deleted after selection) fall back to the full
        # roster rather than printing an empty table in a safety document.
        if filtered:
            add_heading(doc, "Dipendenti impegnati in cantiere", level=3)
            _render_dipendenti_rows(doc, filtered)
            return
    add_heading(doc, "Dipendenti dell'azienda", level=3)
    _render_dipendenti_rows(doc, persone)


def _render_dipendenti_rows(doc, persone: list) -> None:
    if not persone:
        add_paragraph(doc, "Nessun dipendente registrato.", italic=True)
        return
    rows = []
    for pe in persone:
        rows.append([
            pe.nominativo or "—",
            pe.mansione or "—",
            "SI" if getattr(pe, "ruolo_primo_soccorso", False) else "NO",
            "SI" if getattr(pe, "ruolo_antincendio", False) else "NO",
            "SI" if getattr(pe, "ruolo_preposto", False) else "NO",
        ])
    add_data_table(
        doc,
        ["Nominativo", "Mansione", "Addetto Primo Soccorso", "Addetto Antincendio", "Preposto"],
        rows,
    )


def _render_modalita_organizzative(doc, pos) -> None:
    """Render the "Modalità organizzative" section (All. XV punto 3.2.1 c).

    All three fields are free-text. Skip the section entirely if none of
    them is populated, to avoid printing an empty section on small POS.
    """
    items = [
        ("Orario di lavoro", pos.orario_lavoro_cantiere),
        ("Turni", pos.turni_descrizione),
        # Feedback #57 (2026-05-26): label "Riunioni di coordinamento"
        # rinominato in "Descrizione del cantiere". DB column unchanged
        # (riunioni_coordinamento) — solo etichetta utente / docx.
        ("Descrizione del cantiere", pos.riunioni_coordinamento),
    ]
    if not any(v for _, v in items):
        return
    add_heading(doc, "Modalità organizzative", level=3)
    for label, value in items:
        if not value:
            continue
        add_paragraph(doc, label + ":", bold=True)
        add_paragraph(doc, value)


def _render_organizzazione_logistica(doc, pos) -> None:
    """Render the "Organizzazione logistica" section.

    `monoblocchi_installati` drives the boilerplate line ("Non saranno
    installati monoblocchi" vs the dettagli text). `modalita_pasti` is a
    free-text paragraph.
    """
    if not (
        pos.monoblocchi_installati
        or pos.monoblocchi_dettagli
        or pos.modalita_pasti
    ):
        return
    add_heading(doc, "Organizzazione logistica", level=3)
    if pos.monoblocchi_installati:
        add_paragraph(doc, "Monoblocchi installati in cantiere:", bold=True)
        add_paragraph(doc, pos.monoblocchi_dettagli or "Sì — dettagli da specificare.")
    else:
        add_paragraph(doc, "Non saranno installati monoblocchi in cantiere.")
    if pos.modalita_pasti:
        add_paragraph(doc, "Modalità consumazione pasti:", bold=True)
        add_paragraph(doc, pos.modalita_pasti)


# ---------------------------------------------------------------------------
# DPI matrix (US-4.8)
# ---------------------------------------------------------------------------


def _dpi_labels(codes: list[str]) -> str:
    """Render a list of DPI codes as comma-separated Italian labels.

    Feedback #61 (2026-05-26): the frontend stores the literal string
    ``__non_effettua__`` in a cell when the operator declared that a
    role does not perform a given phase. We collapse it to a single
    Italian-readable label so the printed POS doesn't leak the sentinel.
    """
    if not codes:
        return "—"
    if "__non_effettua__" in codes:
        return "Non effettua questa operazione"
    return ", ".join(DPI_CATALOG.get(c, c) for c in codes)


def _parse_phases(fasi_raw: list) -> list[PosPhase]:
    """Tolerantly parse the JSONB column into structured PosPhase rows.

    Older POS rows (pre-US-4.7) store loose ``{"fase": "...", "descrizione": "..."}``
    dicts without an ``id`` or ``ordine``. We promote them lazily so the
    generator can render either shape. Anything that still fails validation
    is dropped — safety documents must never fail to render over a stray
    row.
    """
    out: list[PosPhase] = []
    for i, raw in enumerate(fasi_raw or []):
        if not isinstance(raw, dict):
            continue
        # Back-compat: legacy rows used "fase" for the name.
        promoted = dict(raw)
        if "nome" not in promoted and "fase" in promoted:
            promoted["nome"] = promoted.pop("fase")
        promoted.setdefault("id", f"legacy-{i}")
        promoted.setdefault("ordine", i)
        try:
            out.append(PosPhase(**promoted))
        except ValidationError:
            continue
    return out


def _render_default_phase_skeleton(doc) -> None:
    """Emit the 8 standard construction phases as a review-and-customize
    skeleton when the operator hasn't created custom fasi yet.

    Per D.Lgs. 81/2008 All. XV punto 3.2.1 d the POS must list le fasi
    lavorative previste in cantiere. The 8 default phases are the most
    common phases for an Italian cantiere edile; the operator is expected
    to remove/modify in fase di audit so the matrix matches the actual
    project. The Italian heading explains this so an inspector reading
    the doc understands why the rows look generic.
    """
    add_paragraph(
        doc,
        "Le fasi lavorative non sono ancora state personalizzate per "
        "questo cantiere. Di seguito le 8 fasi standard per un cantiere "
        "edile da rivedere con il Coordinatore CSE prima dell'inizio "
        "dei lavori.",
        italic=True,
        size=9,
    )
    rows = []
    for i, phase_key in enumerate(PHASES_CONSTRUCTION, 1):
        rows.append([str(i), _PHASE_LABELS_IT.get(phase_key, phase_key)])
    add_data_table(doc, ["#", "Fase standard"], rows)


def _render_phase_sections(doc, fasi_raw: list) -> None:
    """Emit the per-phase sections for one POS (US-4.7).

    Structure:
      1. "Quadro sinottico" summary table — phases in drag-drop order with
         their dependencies, for the Gantt-like overview that the AC calls for.
      2. Per-phase narrative with rischi / DPI / mezzi and any NIOSH / rumore
         / vibrazioni snapshots.
      3. Footnote listing dependency violations (a phase declared after one
         of its declared predecessors — the generator does not refuse to
         render, per the endpoint comment).
    """
    phases = _parse_phases(fasi_raw)
    if not phases:
        # Fall back to the pre-US-4.7 tabular shape.
        rows = []
        for f in fasi_raw:
            if not isinstance(f, dict):
                continue
            rischi = ", ".join(f.get("rischi", [])) if isinstance(f.get("rischi"), list) else (f.get("rischi") or "")
            dpi = ", ".join(f.get("dpi", [])) if isinstance(f.get("dpi"), list) else (f.get("dpi") or "")
            mezzi = ", ".join(f.get("mezzi", [])) if isinstance(f.get("mezzi"), list) else (f.get("mezzi") or "")
            rows.append([f.get("fase", f.get("nome", "")), f.get("descrizione", ""), rischi, dpi, mezzi])
        if rows:
            add_data_table(doc, ["Fase", "Descrizione", "Rischi", "DPI", "Mezzi"], rows)
        return

    phases.sort(key=lambda p: p.ordine)
    name_by_id = {p.id: p.nome for p in phases}

    # --- 1. Quadro sinottico (fasi + precedenze) -----------------------
    add_paragraph(
        doc,
        "Quadro sinottico delle fasi lavorative in ordine di esecuzione. "
        "La colonna 'Dipende da' esplicita le fasi che devono essere "
        "completate prima dell'avvio della fase in riga (Gantt logico).",
        italic=True,
    )
    synoptic_rows = []
    for i, ph in enumerate(phases, 1):
        deps = ", ".join(name_by_id.get(d, d) for d in ph.dipende_da) or "—"
        synoptic_rows.append([str(i), ph.nome, deps])
    add_data_table(doc, ["#", "Fase", "Dipende da"], synoptic_rows)

    # --- 2. Per-phase detail ------------------------------------------
    for i, ph in enumerate(phases, 1):
        add_heading(doc, f"{i}. {ph.nome}", level=4)
        if ph.descrizione:
            add_paragraph(doc, ph.descrizione)

        detail_rows: list[tuple[str, str]] = [
            ("Rischi", ", ".join(ph.rischi) or "—"),
            ("DPI", ", ".join(ph.dpi) or "—"),
            ("Mezzi / attrezzature", ", ".join(ph.mezzi) or "—"),
        ]
        if ph.dipende_da:
            detail_rows.append(
                ("Precedenze", ", ".join(name_by_id.get(d, d) for d in ph.dipende_da))
            )
        add_kv_table(doc, detail_rows)

        if ph.niosh is not None:
            add_paragraph(doc, "Valutazione NIOSH (movimentazione manuale dei carichi):", bold=True)
            add_kv_table(
                doc,
                [
                    ("Peso sollevato (kg)", f"{ph.niosh.peso_sollevato:.2f}"),
                    ("Costante di peso CP (kg)", f"{ph.niosh.cp:.2f}"),
                    ("Fattori A·B·C·D·E·F", (
                        f"{ph.niosh.fattore_a:.2f} · {ph.niosh.fattore_b:.2f} · "
                        f"{ph.niosh.fattore_c:.2f} · {ph.niosh.fattore_d:.2f} · "
                        f"{ph.niosh.fattore_e:.2f} · {ph.niosh.fattore_f:.2f}"
                    )),
                    ("PLR (kg)", f"{ph.niosh.plr:.2f}" if ph.niosh.plr is not None else "—"),
                    ("IR", f"{ph.niosh.ir:.2f}" if ph.niosh.ir is not None else "—"),
                    ("Zona di rischio", ph.niosh.livello or "—"),
                ],
            )

        if ph.rumore is not None:
            add_paragraph(doc, "Esposizione al rumore:", bold=True)
            add_kv_table(
                doc,
                [
                    ("LEX,8h (dB(A))", f"{ph.rumore.lex_8h_dba:.1f}"),
                    ("Fascia", ph.rumore.fascia or "—"),
                    ("DPI uditivi obbligatori", "SI" if ph.rumore.dpi_obbligatori else "NO"),
                    ("Note", ph.rumore.note or "—"),
                ],
            )

        if ph.vibrazioni is not None:
            add_paragraph(doc, "Esposizione a vibrazioni meccaniche:", bold=True)
            add_kv_table(
                doc,
                [
                    ("A(8) mano-braccio (m/s²)",
                     f"{ph.vibrazioni.a8_mano_braccio:.2f}" if ph.vibrazioni.a8_mano_braccio is not None else "—"),
                    ("A(8) corpo intero (m/s²)",
                     f"{ph.vibrazioni.a8_corpo_intero:.2f}" if ph.vibrazioni.a8_corpo_intero is not None else "—"),
                    ("Entro i limiti di legge", "SI" if ph.vibrazioni.entro_limiti else "NO"),
                    ("Note", ph.vibrazioni.note or "—"),
                ],
            )

    # --- 3. Dependency-order footnote ---------------------------------
    violations = dependency_violations_after_ordering(phases)
    if violations:
        add_paragraph(
            doc,
            "Nota: le seguenti fasi risultano ordinate prima di una loro "
            "dichiarata precedenza. Verificare la programmazione del cantiere "
            "prima dell'avvio dei lavori.",
            italic=True,
            bold=True,
        )
        for dependent, missing in violations:
            add_paragraph(doc, f"  • '{dependent}' dipende da '{missing}' ma la precede nell'ordine.")


def _render_dpi_matrix(doc, pos) -> None:
    """Emit the role x phase DPI matrix for one POS.

    Layout: rows = roles, columns = phases (first column is the role
    label). Cells list the DPI as Italian labels from DPI_CATALOG. Where
    an operator edited a cell (differs from the rules-engine default for
    this role/phase), we append " (personalizzato)" so the reviewer can
    spot customisations.

    Merge rule: if two adjacent rows have identical DPI across every
    phase, their phase cells are merged vertically (their role labels
    stay separate since they're semantically different).
    """
    matrix = pos.dpi_matrix or {}
    roles = pos.dpi_matrix_roles or []
    phases = pos.dpi_matrix_phases or []

    # If the operator hasn't configured a matrix yet, fall back to the
    # standard 10 roles × 8 phases default produced by build_default_matrix.
    # Mark the whole table as "default" so the reviewer knows it's a
    # rules-engine suggestion, not a per-client choice.
    using_default = False
    if not matrix or not roles or not phases:
        roles = ROLES_CONSTRUCTION
        phases = PHASES_CONSTRUCTION
        matrix = build_default_matrix(roles, phases)
        using_default = True

    add_heading(doc, "Matrice DPI per ruolo e fase", level=3)
    if using_default:
        add_paragraph(
            doc,
            "La matrice seguente e generata automaticamente dalle regole "
            "predefinite (D.Lgs. 81/2008 art. 77 e Tit. III). Il Datore di "
            "Lavoro verifica e personalizza ogni cella in fase di audit.",
            italic=True,
            size=9,
        )

    # Pre-compute defaults once so we can detect operator overrides.
    defaults = build_default_matrix(roles, phases)

    header = ["Ruolo"] + [_PHASE_LABELS_IT.get(p, p) for p in phases]
    data_rows: list[list[str]] = []
    for role in roles:
        row: list[str] = [_ROLE_LABELS_IT.get(role, role)]
        for phase in phases:
            cell_codes = (matrix.get(phase, {}) or {}).get(role, []) or []
            default_codes = (defaults.get(phase, {}) or {}).get(role, []) or []
            label = _dpi_labels(cell_codes)
            if sorted(cell_codes) != sorted(default_codes):
                label = f"{label} (personalizzato)" if label != "—" else "— (personalizzato)"
            row.append(label)
        data_rows.append(row)

    table = add_data_table(doc, header, data_rows)

    # Vertical merge for adjacent rows with identical DPI across every
    # phase. python-docx merges are additive — cell.merge(other) merges
    # the rectangle they span, so we call it per column for each run of
    # identical rows. Column 0 (role) is left untouched so the labels
    # remain distinct.
    if len(data_rows) < 2:
        return
    run_start = 0
    for i in range(1, len(data_rows) + 1):
        same = (
            i < len(data_rows)
            and data_rows[i][1:] == data_rows[run_start][1:]
        )
        if not same:
            if i - run_start > 1:
                # +1 because the header occupies row 0 in the docx table.
                top_row = table.rows[run_start + 1]
                bot_row = table.rows[i - 1 + 1]
                for col in range(1, len(header)):
                    top_row.cells[col].merge(bot_row.cells[col])
            run_start = i
