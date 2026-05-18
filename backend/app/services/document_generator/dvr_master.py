"""
DVR Master document generator.

Generates the Documento di Valutazione dei Rischi (DVR) — the master risk
assessment document required by D.Lgs. 81/2008 for every Italian workplace.

Output: A .docx file with professional formatting including:
- Cover page with company name, date, and logo placeholder
- Table of contents placeholder
- Part I: Company data tables
- Part III: Risk assessment tables per environment with P/D/I color coding
- Part IV: Improvement measures placeholder
"""

import os
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

from app.data.regional_regulations import get_regulations_for_comune
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.reference_data import (
    HAZARD_LIBRARY,
    RISK_CATEGORIES,
    normalize_categoria_to_long,
)
from app.services.risk_calculator import calculate_risk_index


# ---------------------------------------------------------------------------
# Parte II — Definizioni (Template Table 19, MIXED). Core D.Lgs. 81/2008
# terms from art. 2 — kept compact so the emitted table mirrors the
# template's 27-row shape without exploding the .docx size.
# ---------------------------------------------------------------------------

_DEFINIZIONI_ROWS: list[tuple[str, str]] = [
    ("LAVORATORE (LAV)",
     "Persona che, indipendentemente dalla tipologia contrattuale, svolge "
     "un'attivita lavorativa nell'ambito dell'organizzazione di un datore "
     "di lavoro pubblico o privato, con o senza retribuzione."),
    ("DATORE DI LAVORO (DL)",
     "Soggetto titolare del rapporto di lavoro con il lavoratore o, "
     "comunque, il soggetto che ha la responsabilita dell'organizzazione "
     "stessa o dell'unita produttiva."),
    ("AZIENDA",
     "Il complesso della struttura organizzata dal datore di lavoro "
     "pubblico o privato."),
    ("DIRIGENTE",
     "Persona che attua le direttive del datore di lavoro organizzando "
     "l'attivita lavorativa e vigilando su di essa."),
    ("PREPOSTO",
     "Persona che, in ragione delle competenze professionali e nei limiti "
     "di poteri gerarchici e funzionali adeguati alla natura dell'incarico "
     "conferitogli, sovrintende alla attivita lavorativa."),
    ("RSPP",
     "Responsabile del Servizio di Prevenzione e Protezione — persona "
     "designata dal datore di lavoro, in possesso di attitudini e "
     "capacita adeguate, a cui il datore di lavoro si avvale per "
     "organizzare il servizio di prevenzione e protezione."),
    ("ASPP",
     "Addetto del Servizio di Prevenzione e Protezione — persona in "
     "possesso di attitudini e capacita adeguate che supporta il RSPP "
     "nell'organizzazione del servizio."),
    ("RLS",
     "Rappresentante dei Lavoratori per la Sicurezza — persona eletta o "
     "designata per rappresentare i lavoratori per quanto concerne gli "
     "aspetti della salute e della sicurezza durante il lavoro."),
    ("MEDICO COMPETENTE (MC)",
     "Medico in possesso di uno dei titoli e requisiti formativi e "
     "professionali richiesti dalla normativa, che collabora con il "
     "datore di lavoro ai fini della valutazione dei rischi e della "
     "sorveglianza sanitaria."),
    ("VALUTAZIONE DEI RISCHI",
     "Valutazione globale e documentata di tutti i rischi per la salute "
     "e la sicurezza dei lavoratori presenti nell'ambito "
     "dell'organizzazione in cui essi prestano la propria attivita."),
    ("PERICOLO",
     "Proprieta o qualita intrinseca di un determinato fattore avente "
     "il potenziale di causare danni."),
    ("RISCHIO",
     "Probabilita di raggiungimento del livello potenziale di danno "
     "nelle condizioni di impiego o di esposizione a un determinato "
     "fattore o agente oppure alla loro combinazione."),
    ("PROBABILITA (P)",
     "Frequenza con cui un evento dannoso puo verificarsi, valutata "
     "su scala 1-4 (Bassa, Medio-Bassa, Medio-Alta, Elevata)."),
    ("DANNO (D)",
     "Entita del danno atteso per il lavoratore esposto, valutata su "
     "scala 1-4 (Trascurabile, Modesto, Notevole, Ingente)."),
    ("INDICE DI RISCHIO (I)",
     "Calcolato come I = 2*D + P; range 3-12; livelli ACCETTABILE, "
     "MODESTO, GRAVE, GRAVISSIMO."),
    ("UNITA PRODUTTIVA",
     "Stabilimento o struttura finalizzati alla produzione di beni "
     "o all'erogazione di servizi, dotati di autonomia finanziaria e "
     "tecnico-funzionale."),
    ("DPI",
     "Dispositivo di Protezione Individuale — qualsiasi attrezzatura "
     "destinata ad essere indossata e tenuta dal lavoratore allo scopo "
     "di proteggerlo contro uno o piu rischi."),
    ("SORVEGLIANZA SANITARIA",
     "Insieme degli atti medici finalizzati alla tutela dello stato di "
     "salute e sicurezza dei lavoratori, in relazione all'ambiente di "
     "lavoro, ai fattori di rischio professionali e alle modalita di "
     "svolgimento dell'attivita lavorativa."),
    ("FORMAZIONE",
     "Processo educativo attraverso il quale trasferire ai lavoratori "
     "conoscenze e procedure utili alla acquisizione di competenze "
     "per lo svolgimento in sicurezza dei rispettivi compiti."),
    ("INFORMAZIONE",
     "Complesso delle attivita dirette a fornire conoscenze utili alla "
     "identificazione, alla riduzione e alla gestione dei rischi "
     "nell'ambiente di lavoro."),
    ("ADDESTRAMENTO",
     "Complesso delle attivita dirette a fare apprendere ai lavoratori "
     "l'uso corretto di attrezzature, macchine, impianti, sostanze, "
     "dispositivi, anche di protezione individuale, e le procedure di lavoro."),
    ("AGENTE",
     "L'agente chimico, fisico o biologico presente durante il lavoro "
     "e potenzialmente dannoso per la salute."),
    ("NORMA TECNICA",
     "Specifica tecnica, approvata e pubblicata da un'organizzazione "
     "internazionale, da un organismo europeo o nazionale di "
     "normalizzazione, la cui osservanza non e obbligatoria."),
    ("BUONA PRASSI",
     "Soluzioni organizzative o procedurali coerenti con la normativa "
     "vigente e con le norme di buona tecnica, adottate volontariamente e "
     "finalizzate a promuovere la salute e sicurezza sui luoghi di lavoro."),
    ("LINEE GUIDA",
     "Atti di indirizzo e coordinamento per l'applicazione della "
     "normativa in materia di salute e sicurezza."),
    ("MODELLO DI ORGANIZZAZIONE E GESTIONE",
     "Modello organizzativo e gestionale per la definizione e "
     "attuazione di una politica aziendale per la salute e sicurezza."),
    ("RESPONSABILITA SOCIALE",
     "Integrazione volontaria delle preoccupazioni sociali ed "
     "ecologiche delle aziende nelle loro operazioni commerciali e nei "
     "loro rapporti con le parti interessate."),
]

# Parte II — P/D criteria lookup tables (Template Tables 21 and 22).
_PROBABILITA_CRITERI_ROWS = [
    ("4", "ELEVATA",
     "Esiste correlazione diretta tra mancanza rilevata e possibilita "
     "che l'evento dannoso si verifichi; si sono gia verificati casi "
     "analoghi in azienda o in realta simili."),
    ("3", "MEDIO ALTA",
     "La mancanza rilevata puo provocare un danno, anche se non "
     "direttamente, seppur in modo automatico; sono noti rari episodi "
     "in azienda o in realta simili."),
    ("2", "MEDIO BASSA",
     "La mancanza rilevata puo provocare un danno in circostanze "
     "particolari; non sono noti episodi in azienda."),
    ("1", "BASSA",
     "La mancanza rilevata puo provocare un danno solo in circostanze "
     "eccezionali e in concomitanza con piu eventi sfavorevoli; non "
     "sono noti episodi nel settore."),
]

_DANNO_CRITERI_ROWS = [
    ("4", "INGENTE",
     "Infortunio o episodio di esposizione con effetti letali o "
     "invalidita totale permanente."),
    ("3", "NOTEVOLE",
     "Infortunio o episodio di esposizione acuta con effetti di "
     "invalidita parziale permanente; patologie gravi a effetti "
     "progressivi."),
    ("2", "MODESTO",
     "Infortunio o episodio di esposizione acuta con inabilita "
     "temporanea reversibile."),
    ("1", "TRASCURABILE",
     "Infortunio o episodio di esposizione acuta con inabilita "
     "reversibile di rapida guarigione."),
]


# Ordered list of the 11 canonical risk categories grouped by macro-area.
# Drives both the SI/NO checklist (Table 26 per env) and the per-category
# 5-col risk tables (Tables 27+). Order matches the DVR template.
_CATEGORY_ORDER: list[tuple[str, str]] = [
    (rc["macro_categoria"], rc["categoria"]) for rc in RISK_CATEGORIES
]

# Macro-area row labels interleaved into the checklist so it mirrors the
# template's 3 section headers + 11 data rows = 14 rows layout.
_MACRO_LABELS: list[str] = [
    "Rischi per la Sicurezza",
    "Rischi per la Salute",
    "Rischi Trasversali",
]


# ---------------------------------------------------------------------------
# Parte IV §4.2–4.12 — static procedural sections.
#
# Each entry mirrors the master template (DVR_TEMPLATE_MAPPING.md §10, paragraphs
# #1800–#2425) and is condensed to the procedural intent + responsibilities,
# not the verbatim 650-paragraph copy. Verbatim ingestion of the full template
# text is tracked separately; the contents below are sufficient to make the
# DVR a legally signable draft (audit verdict 2026-04-28: SHELL-ONLY → DRAFT).
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class _ProceduralSection:
    heading: str
    body: tuple[str, ...] = ()
    bullets: tuple[str, ...] = ()
    docs: tuple[str, ...] = ()


# ---------------------------------------------------------------------------
# Misura di miglioramento templates — keyed by the canonical risk categoria
# (long form, matches RISK_CATEGORIES). Each entry differentiates the
# Procedura and Risorse columns of T109 so the auto-seed pass produces
# operationally distinct rows instead of one repeated template string
# (audit F-002, 2026-04-28).
# ---------------------------------------------------------------------------
_MISURA_TEMPLATE_DEFAULT: dict[str, str] = {
    "procedura": (
        "Verifica dell'efficacia delle misure di prevenzione gia in atto, "
        "definizione di procedura operativa dedicata e formazione "
        "specifica dei lavoratori esposti."
    ),
    "risorse": (
        "Tempo del RSPP / Preposto per la stesura della procedura; "
        "eventuale aggiornamento DPI; ore di formazione."
    ),
}

_MISURA_TEMPLATE_BY_CATEGORIA: dict[str, dict[str, str]] = {
    "Strutture": {
        "procedura": (
            "Sopralluogo tecnico per verifica conformita ai requisiti "
            "di sicurezza dei luoghi di lavoro (Allegato IV D.Lgs. 81/08); "
            "interventi edili o impiantistici sui non-conformi; "
            "verifica dell'efficacia tramite controllo programmato."
        ),
        "risorse": (
            "Tecnico abilitato per il sopralluogo; ditta esterna per "
            "interventi edili / impiantistici; budget per interventi "
            "strutturali."
        ),
    },
    "Macchine": {
        "procedura": (
            "Verifica della marcatura CE e dei manuali d'uso; controllo "
            "della funzionalita di protezioni fisse, mobili e di emergenza; "
            "manutenzione periodica programmata e registrata."
        ),
        "risorse": (
            "Manutentore qualificato; budget ricambi; tempo del Preposto "
            "per la verifica e la registrazione delle manutenzioni."
        ),
    },
    "Impianti Elettrici": {
        "procedura": (
            "Verifica periodica obbligatoria degli impianti elettrici e di "
            "messa a terra (D.P.R. 462/01); manutenzione programmata; "
            "informazione ai lavoratori sui rischi residui di natura "
            "elettrica."
        ),
        "risorse": (
            "Verificatore abilitato (ASL/ARPA o organismo notificato); "
            "elettricista qualificato; eventuale aggiornamento "
            "dell'impianto."
        ),
    },
    "Incendio-Esplosioni": {
        "procedura": (
            "Aggiornamento della valutazione del rischio incendio (D.M. "
            "03/09/2021); verifica periodica di estintori, idranti, "
            "porte REI e sistemi di rilevazione; esercitazioni di "
            "evacuazione almeno annuali."
        ),
        "risorse": (
            "Tecnico abilitato per il rinnovo del CPI; ditta manutentrice "
            "antincendio; tempo dei lavoratori per l'esercitazione."
        ),
    },
    "Agenti Chimici": {
        "procedura": (
            "Aggiornamento dell'inventario delle sostanze e delle SDS; "
            "valutazione del rischio chimico ai sensi degli artt. 222-232 "
            "D.Lgs. 81/08; verifica DPI specifici e adeguatezza della "
            "ventilazione."
        ),
        "risorse": (
            "Igienista industriale o consulente chimico; SDS aggiornate; "
            "eventuali campagne di campionamento ambientale."
        ),
    },
    "Agenti Fisici": {
        "procedura": (
            "Misurazione strumentale dell'agente fisico (rumore, "
            "vibrazioni, microclima, illuminamento, campi EM) ai sensi "
            "del Titolo VIII; confronto con i valori limite e azione "
            "se VLE/VAM superati."
        ),
        "risorse": (
            "Tecnico competente in acustica / vibrazioni; strumentazione "
            "di misura; eventuale Allegato dedicato (Microclima, Rumore)."
        ),
    },
    "Agenti Biologici": {
        "procedura": (
            "Valutazione del rischio biologico ai sensi del Titolo X "
            "D.Lgs. 81/08; classificazione degli agenti per gruppo di "
            "pericolo; protocollo di pulizia e disinfezione; sorveglianza "
            "sanitaria mirata."
        ),
        "risorse": (
            "Medico Competente; DPI specifici; protocollo di disinfezione "
            "(prodotti, frequenze, registro)."
        ),
    },
    "Agenti Cancerogeni": {
        "procedura": (
            "Valutazione del rischio cancerogeno ai sensi del Titolo IX "
            "Capo II; sostituzione dell'agente ove tecnicamente possibile; "
            "monitoraggio biologico e ambientale; registro degli esposti."
        ),
        "risorse": (
            "Medico Competente; tecnico per il monitoraggio; eventuale "
            "investimento in sistemi di confinamento."
        ),
    },
    "Organizzazione del Lavoro": {
        "procedura": (
            "Revisione dei carichi di lavoro, dei turni e delle pause; "
            "valutazione del rischio stress lavoro-correlato secondo "
            "metodologia INAIL; monitoraggio degli indicatori sentinella."
        ),
        "risorse": (
            "Tempo del SPP per la revisione dei processi; eventuale "
            "Allegato Stress; possibile coinvolgimento di consulente "
            "del lavoro."
        ),
    },
    "Fattori Psicologici": {
        "procedura": (
            "Indagine sui rischi psicosociali (carichi mentali, conflitti, "
            "molestie); definizione di canali di segnalazione e di un "
            "comitato per i casi piu complessi; formazione su gestione "
            "dello stress."
        ),
        "risorse": (
            "Psicologo del lavoro / consulente HR; formazione dedicata; "
            "eventuali sportelli di ascolto."
        ),
    },
    "Fattori Ergonomici": {
        "procedura": (
            "Valutazione dei rischi ergonomici (MMC, posture, VDT, "
            "movimenti ripetitivi); riprogettazione delle postazioni; "
            "addestramento sulle corrette tecniche di sollevamento e "
            "lavoro."
        ),
        "risorse": (
            "Allegati MMC e VDT; ergonomo o fisioterapista occupazionale; "
            "budget per ausili meccanici e arredi conformi UNI EN 1335."
        ),
    },
}


# ---------------------------------------------------------------------------
# Formazione contenuti differenziati per famiglia di mansione (audit F-008).
# Maps a coarse-grained family classifier to (durata, contenuti) tuples.
# Defaults to "rischio medio" for unknown families.
# ---------------------------------------------------------------------------
_FORMAZIONE_BY_FAMILY: dict[str, tuple[str, str]] = {
    "operaio_industriale": (
        "12 ore (rischio alto)",
        "Macchine utensili, organi in movimento, marcatura CE, rumore, "
        "vibrazioni, MMC, agenti chimici industriali, DPI specifici "
        "(udito, mani, occhi, vie respiratorie).",
    ),
    "saldatura_termica": (
        "12 ore (rischio alto)",
        "Fumi di saldatura, raggi UV, rischio incendio/esplosione, "
        "spazi confinati, DPI specifici (maschera, visiera, guanti "
        "saldatura, abbigliamento ignifugo).",
    ),
    "magazziniere_carrellista": (
        "12 ore (rischio alto)",
        "Movimentazione manuale e meccanica dei carichi (NIOSH), "
        "stoccaggio merci, urti contro ostacoli, abilitazione "
        "carrelli elevatori.",
    ),
    "ristorazione": (
        "12 ore (rischio alto)",
        "Rischio biologico alimentare (HACCP), tagli e ustioni, "
        "rischio chimico (detergenti, sgrassanti), MMC, posture, "
        "DPI specifici (calzature antiscivolo, grembiuli, guanti).",
    ),
    "ufficio_vdt": (
        "8 ore (rischio medio)",
        "Esposizione a videoterminali (VDT) ai sensi dell'art. 173, "
        "ergonomia della postazione (Allegato XXXIV), illuminazione, "
        "stress correlato all'uso del PC, microclima.",
    ),
    "edile_cantiere": (
        "16 ore (rischio alto)",
        "Lavori in quota, opere provvisionali, lavori in spazi "
        "confinati, rischio elettrico, scavi, utilizzo macchine "
        "movimento terra, DPI III categoria.",
    ),
    "default": (
        "8 ore (rischio medio)",
        "Rischi specifici della mansione e dell'ambiente, misure di "
        "prevenzione, comportamenti di sicurezza, DPI assegnati.",
    ),
}


def _classify_mansione_family(nome: str) -> str:
    """Coarse-grained mansione classifier for the formazione grid."""
    s = (nome or "").lower()
    if any(k in s for k in ("salda",)):
        return "saldatura_termica"
    if any(k in s for k in ("magazz", "carrell", "logist")):
        return "magazziniere_carrellista"
    if any(k in s for k in ("torni", "fresat", "monta", "operaio", "officin")):
        return "operaio_industriale"
    if any(k in s for k in ("cuoc", "mensa", "cucina", "alimenta")):
        return "ristorazione"
    if any(k in s for k in (
        "impieg", "ufficio", "amministrat", "commerc", "tecnic", "videoterminal"
    )):
        return "ufficio_vdt"
    if any(k in s for k in ("cantier", "edile", "muratore", "carpentiere")):
        return "edile_cantiere"
    return "default"


_PART_IV_PROCEDURAL_SECTIONS: tuple[_ProceduralSection, ...] = (
    _ProceduralSection(
        heading="4.2 Gestione Leggi e Regolamenti",
        body=(
            "Il RSPP, in collaborazione con il Datore di Lavoro, ha la "
            "responsabilita di mantenere aggiornata la mappa delle leggi e "
            "dei regolamenti applicabili all'attivita aziendale e di "
            "garantirne la diffusione presso i soggetti interessati.",
        ),
        bullets=(
            "ricerca delle leggi e regolamenti applicabili e identificazione "
            "di quelli relativi alle attivita aziendali",
            "valutazione delle ricadute operative e organizzative delle "
            "novita normative",
            "diffusione interna e archiviazione delle versioni vigenti",
        ),
    ),
    _ProceduralSection(
        heading="4.3 Gestione Sorveglianza Sanitaria",
        body=(
            "Il Datore di Lavoro, tramite il Servizio di Prevenzione e "
            "Protezione, verifica la necessita di sottoporre i lavoratori "
            "a sorveglianza sanitaria nei casi previsti dalla normativa "
            "(art. 41 D.Lgs. 81/2008) e in funzione degli esiti della "
            "valutazione dei rischi.",
            "La sorveglianza ricorre nei casi previsti dalla legge — "
            "esposizione ad agenti chimici, fisici, biologici, "
            "movimentazione manuale dei carichi, videoterminali oltre la "
            "soglia di esposizione, lavoro notturno, gestanti e altri "
            "gruppi sensibili — e in tutti i casi richiesti dal lavoratore "
            "qualora il MC ne riconosca la correlazione con i rischi "
            "professionali.",
            "La nomina e l'eventuale revoca del Medico Competente sono "
            "formalizzate per iscritto e comunicate ai lavoratori e al RLS.",
        ),
        docs=(
            "Lettera di nomina del MC",
            "Protocollo di sorveglianza sanitaria",
            "Giudizi di idoneita alla mansione",
        ),
    ),
    _ProceduralSection(
        heading="4.4 Gestione Informazione, Formazione ed Addestramento",
        body=(
            "Il Datore di Lavoro, in collaborazione con il RSPP e previa "
            "consultazione del RLS, programma le attivita di informazione, "
            "formazione ed addestramento ai sensi degli artt. 36 e 37 del "
            "D.Lgs. 81/2008 e dell'Accordo Stato-Regioni 21/12/2011 "
            "(rep. atti 221/CSR) e successivi aggiornamenti.",
            "La programmazione tiene conto della valutazione dei rischi, "
            "delle mansioni svolte, dell'introduzione di nuove "
            "attrezzature o sostanze e delle segnalazioni provenienti dai "
            "preposti, dai lavoratori e dal Medico Competente.",
        ),
        bullets=(
            "formazione generale e specifica per ogni lavoratore in "
            "funzione del livello di rischio della mansione",
            "formazione aggiuntiva per preposti, dirigenti, RLS, addetti "
            "primo soccorso e antincendio",
            "addestramento sull'uso corretto di attrezzature, DPI e "
            "sostanze pericolose",
            "aggiornamento periodico nei termini previsti dalla normativa",
        ),
        docs=(
            "Registro formazione",
            "Attestati di partecipazione",
            "Verbali di addestramento",
        ),
    ),
    _ProceduralSection(
        heading="4.5 Riunione Periodica",
        body=(
            "Il Datore di Lavoro, direttamente o tramite il RSPP, indice "
            "la riunione periodica ai sensi dell'art. 35 del D.Lgs. "
            "81/2008. La riunione si tiene almeno una volta all'anno per "
            "le aziende con piu di 15 lavoratori, ed in occasione di "
            "significative variazioni delle condizioni di esposizione al "
            "rischio.",
            "Della riunione e redatto verbale tenuto a disposizione dei "
            "partecipanti per la consultazione.",
        ),
        docs=(
            "Convocazione della riunione",
            "Verbale di riunione periodica",
        ),
    ),
    _ProceduralSection(
        heading="4.6 Gestione degli Infortuni",
        body=(
            "Ogni infortunio e ogni mancato infortunio (near miss) sono "
            "segnalati al preposto e, per il tramite del RSPP, al Datore "
            "di Lavoro. La gestione comprende la segnalazione tempestiva, "
            "l'indagine sulle cause, l'eventuale costituzione di una "
            "commissione di indagine e l'aggiornamento della valutazione "
            "dei rischi qualora emergano fattori non precedentemente "
            "identificati.",
        ),
        bullets=(
            "segnalazione immediata al preposto e ai soccorsi",
            "denuncia all'INAIL nei termini di legge",
            "indagine interna sulle cause dirette e di radice",
            "aggiornamento del DVR e delle procedure operative ove "
            "necessario",
        ),
        docs=(
            "Registro infortuni",
            "Verbali di indagine",
            "Comunicazioni INAIL",
        ),
    ),
    _ProceduralSection(
        heading="4.7 Gestione Comportamenti Scorretti dei Lavoratori",
        body=(
            "Il preposto vigila sull'osservanza da parte dei lavoratori "
            "degli obblighi di legge e delle disposizioni aziendali in "
            "materia di salute e sicurezza (art. 19 D.Lgs. 81/2008). I "
            "comportamenti scorretti sono richiamati e gestiti attraverso "
            "il sistema disciplinare aziendale, in coerenza con il CCNL "
            "applicato.",
        ),
    ),
    _ProceduralSection(
        heading="4.8 Gestione DPI",
        body=(
            "I Dispositivi di Protezione Individuale sono forniti ai "
            "lavoratori a titolo personale, ai sensi degli artt. 74-79 "
            "del D.Lgs. 81/2008, dopo una valutazione dei rischi residui "
            "che non possono essere eliminati alla fonte o ridotti con "
            "misure tecniche o organizzative.",
        ),
        bullets=(
            "scelta dei DPI in funzione dei rischi specifici della "
            "mansione e in conformita al Reg. UE 2016/425",
            "consegna registrata su modulo nominativo (tipologia, codice, "
            "data di consegna, scadenza)",
            "addestramento all'uso corretto, alla manutenzione e alla "
            "conservazione dei DPI di III categoria e di tutti i DPI per "
            "udito",
            "gestione dei casi di inadeguatezza, intolleranza o "
            "deterioramento del DPI",
        ),
        docs=(
            "Modulo di consegna DPI",
            "Comunicazione di inadeguatezza DPI",
            "Schede tecniche dei DPI in dotazione",
        ),
    ),
    _ProceduralSection(
        heading="4.9 Gestione Infrastrutture",
        body=(
            "Per tutte le macchine, attrezzature, impianti e infrastrutture "
            "presenti in azienda sono definite le responsabilita di "
            "manutenzione ordinaria, straordinaria e di verifica "
            "periodica obbligatoria (D.M. 11/04/2011, all. VII D.Lgs. "
            "81/2008).",
        ),
        docs=(
            "Registro delle attrezzature soggette a verifica periodica",
            "Verbali di manutenzione",
            "Libretti d'uso e manutenzione",
        ),
    ),
    _ProceduralSection(
        heading="4.10 Gestione dei Lavoratori particolarmente sensibili al rischio",
        body=(
            "L'azienda individua e tutela i lavoratori appartenenti a "
            "gruppi particolarmente sensibili al rischio: lavoratrici "
            "gestanti, puerpere o in periodo di allattamento (D.Lgs. "
            "151/2001), minori (D.Lgs. 345/1999), lavoratori con disabilita, "
            "lavoratori con specifiche prescrizioni del Medico Competente.",
            "Per ciascuna categoria sono adottate misure di prevenzione e "
            "protezione differenziate, riportate negli specifici allegati "
            "del presente DVR (Allegato Gestanti, Allegato Minori, valutazioni "
            "ergonomiche e di stress lavoro-correlato).",
        ),
    ),
    _ProceduralSection(
        heading="4.11 Gestione Acquisti",
        body=(
            "Ogni acquisto di macchine, attrezzature, sostanze chimiche, "
            "DPI e arredi e preceduto dalla verifica della conformita ai "
            "requisiti di sicurezza e salute previsti dalla normativa "
            "vigente. Il RSPP e coinvolto nella definizione dei capitolati "
            "tecnici e nella valutazione preventiva dei rischi introdotti.",
        ),
    ),
    _ProceduralSection(
        heading="4.12 Gestione delle Lavorazioni affidate in Appalto",
        body=(
            "Per le lavorazioni affidate in appalto o a lavoratori "
            "autonomi, il Datore di Lavoro committente verifica i "
            "requisiti tecnico-professionali delle ditte appaltatrici "
            "(art. 26 D.Lgs. 81/2008) e promuove la cooperazione e il "
            "coordinamento, redigendo il D.U.V.R.I. nei casi previsti "
            "dalla normativa.",
            "N.B.: nel caso si rientri nel campo di applicazione del "
            "Titolo IV del D.Lgs. 81/2008 (cantieri temporanei o mobili) "
            "si applicano le disposizioni specifiche relative al "
            "coordinamento per la progettazione e per l'esecuzione "
            "dell'opera.",
        ),
        docs=(
            "Verifica idoneita tecnico-professionale",
            "D.U.V.R.I.",
            "Contratti d'appalto",
        ),
    ),
)


# ---------------------------------------------------------------------------
# Logo asset path (embedded on cover page when present)
# ---------------------------------------------------------------------------

_LOGO_PATH = Path(__file__).resolve().parents[3] / "assets" / "logo.png"


# ---------------------------------------------------------------------------
# Part II static boilerplate content (extracted to keep the file lean)
# ---------------------------------------------------------------------------

_METODOLOGIA_INTRO_1 = (
    "La presente valutazione dei rischi e redatta ai sensi dell'art. 28 del "
    "D.Lgs. 9 aprile 2008, n. 81 e s.m.i., che impone al Datore di Lavoro la "
    "valutazione di tutti i rischi per la sicurezza e la salute dei lavoratori, "
    "tenendo conto della specificita delle mansioni, delle attrezzature e degli "
    "ambienti di lavoro."
)

_METODOLOGIA_INTRO_2 = (
    "Il metodo adottato si fonda sulla stima dell'Indice di Rischio (I) "
    "calcolato attraverso la formula I = 2*D + P, dove P rappresenta la "
    "Probabilita di accadimento dell'evento dannoso (scala 1-4) e D il Danno "
    "atteso per il lavoratore esposto (scala 1-4). L'indice risultante, "
    "compreso nell'intervallo 3-12, e associato a un livello di rischio e a "
    "una relativa priorita di intervento."
)

_RISK_LEVEL_TABLE_ROWS = [
    ("3-4", "ACCETTABILE", "Monitoraggio", "Continuo"),
    ("5-6", "MODESTO", "Strumenti di minimizzazione", "1 anno"),
    ("7-8", "GRAVE", "Sensibilizzazione + controllo", "6 mesi"),
    ("9-12", "GRAVISSIMO", "Ricerca urgente misure", "Immediatamente"),
]

# ---------------------------------------------------------------------------
# Color palette for risk levels
# ---------------------------------------------------------------------------

_RISK_COLORS = {
    "ACCETTABILE": RGBColor(0x4C, 0xAF, 0x50),   # Green
    "MODESTO": RGBColor(0xFF, 0xC1, 0x07),        # Amber
    "GRAVE": RGBColor(0xFF, 0x98, 0x00),           # Orange
    "GRAVISSIMO": RGBColor(0xF4, 0x43, 0x36),      # Red
}

_HEADER_BG = RGBColor(0x1A, 0x23, 0x7E)           # Dark blue for table headers
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)          # White text on headers
_LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)          # Alternating row background


class DVRMasterGenerator(BaseDocumentGenerator):
    """Generates the DVR Master document (.docx)."""

    async def generate(self) -> str:
        """Generate the DVR Master document.

        Returns:
            Absolute path to the generated .docx file.
        """
        data = await self.load_data()
        azienda = data["azienda"]
        # DVR-specific extras: foto per ambiente, sorveglianza-sanitaria
        # rows per mansione, allegato presence flags, VDT exposure flags,
        # and the misure di miglioramento program. Loaded once here so the
        # Part III/IV renderers don't fan out into N+1 queries.
        extras = await self._load_dvr_extras(data)

        # Look up the version BEFORE building the doc so the cover page
        # can show it. Same value is used in the filename below — single
        # source of truth.
        version = await self._get_next_version()

        doc = Document()
        self._setup_styles(doc)

        # Build document sections
        self._add_cover_page(doc, azienda, data["generated_at"], version)
        toc_anchors = self._add_table_of_contents(doc)
        self._add_premessa(doc)
        self._add_pre_parte_i(doc, azienda, data["generated_at"])
        self._add_part_i(
            doc,
            azienda,
            data["persone"],
            data["attrezzature"],
            data["sostanze_chimiche"],
            data["ambienti"],
        )
        self._add_part_ii(doc, azienda)
        self._add_part_iii(
            doc,
            azienda,
            data["persone"],
            data["ambienti"],
            data["attrezzature"],
            extras,
        )
        self._add_part_iv(doc, azienda, data["persone"], extras)

        # Replace the TOC's cached body with the real outline that the
        # rest of the doc just emitted. Without this Word users who
        # decline the "update fields?" prompt on open see only the
        # placeholder line — the user feedback ("indice deve essere
        # esploso con paginazione e piu dettagliato"). After F9 the
        # field re-resolves with page numbers from the real headings.
        self._finalize_table_of_contents(doc, *toc_anchors)

        # Save with the filename pattern required by US-2.8 AC2:
        # DVR_<ragione_sociale>_<YYYYMMDD>_v<N>.docx.
        # The <ragione_sociale> segment is slugified (lowercase,
        # alphanumeric + underscore) so the filename stays safe on both
        # POSIX and Windows checkouts. The date is the generation day
        # (UTC) so regenerations on the same day keep the same stamp.
        output_dir = self._get_output_dir()
        slug = self._slugify(azienda.ragione_sociale or "azienda")
        date_stamp = datetime.now(timezone.utc).strftime("%Y%m%d")
        filename = f"DVR_{slug}_{date_stamp}_v{version}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)
        return filepath

    @staticmethod
    def _slugify(text: str, max_length: int = 40) -> str:
        """Produce a filesystem-safe slug from a free-form label.

        Lowercases the input, replaces any non-alphanumeric character with
        an underscore, collapses repeated underscores and trims them from
        the edges, then truncates to ``max_length`` characters.
        """
        lowered = (text or "").lower()
        replaced = re.sub(r"[^a-z0-9]+", "_", lowered)
        collapsed = re.sub(r"_+", "_", replaced).strip("_")
        if not collapsed:
            collapsed = "azienda"
        return collapsed[:max_length].rstrip("_") or "azienda"

    async def _load_dvr_extras(self, data: dict) -> dict:
        """Hydrate DVR-only data not loaded by the shared base loader.

        Returns a dict with:
          - ``foto_by_ambiente``: ambiente_id -> list[AmbienteFoto]
          - ``vdt_esposti_persona_ids``: set of persona_ids with VDT esposto=True
          - ``allegati_presenti``: ordered list of (slug, label) for applicable
            allegati, derived from MMC/VDT/Stress/Incendio/Gestanti/
            Microclima/Biologico assessment rows on this azienda
          - ``misure_miglioramento``: list[MisuraMiglioramento] ordered for
            T109 emission (auto-seeded from pericoli I>=7 when empty so the
            document is signable on day 1)

        DPI + rischi specifici flags now live on Persona itself
        (persona.dpi_codes / persona.rischi_specifici_codes), so the
        renderer reads them directly from data["persone"] — no extras
        lookup needed.

        Failure to load any single sub-source is non-fatal — the renderer
        falls back to the empty case rather than crashing the doc job.
        """
        from sqlalchemy import select
        from app.models.ambiente_foto import AmbienteFoto
        from app.models.biologico_valutazione import BiologicoValutazione
        from app.models.gestanti_valutazione import GestantiValutazione
        from app.models.incendio_valutazione import IncendioValutazione
        from app.models.microclima_valutazione import MicroclimaValutazione
        from app.models.misura_miglioramento import MisuraMiglioramento
        from app.models.mmc_valutazione import MmcValutazione
        from app.models.stress_valutazione import StressValutazione
        from app.models.vdt_valutazione import VdtValutazione

        ambienti = data.get("ambienti") or []
        ambiente_ids = [a.id for a in ambienti]

        # Foto per ambiente
        foto_by_ambiente: dict = {}
        if ambiente_ids:
            r = await self.db.execute(
                select(AmbienteFoto)
                .where(AmbienteFoto.ambiente_id.in_(ambiente_ids))
                .order_by(AmbienteFoto.created_at)
            )
            for foto in r.scalars().all():
                foto_by_ambiente.setdefault(foto.ambiente_id, []).append(foto)

        # VDT esposti persona ids
        vdt_esposti_persona_ids: set = set()
        r = await self.db.execute(
            select(VdtValutazione).where(
                VdtValutazione.azienda_id == self.azienda_id,
                VdtValutazione.esposto.is_(True),
            )
        )
        for v in r.scalars().all():
            if v.persona_id is not None:
                vdt_esposti_persona_ids.add(v.persona_id)

        # Allegato presence detection — one row in the corresponding table
        # is enough for the allegato to be referenced by name.
        allegati_presenti: list[tuple[str, str]] = []
        incendio_present = False
        for model, slug, label in (
            (MmcValutazione, "mmc", "Allegato Movimentazione Manuale dei Carichi (MMC)"),
            (VdtValutazione, "vdt", "Allegato Videoterminali (VDT)"),
            (StressValutazione, "stress", "Allegato Stress Lavoro-Correlato"),
            (IncendioValutazione, "incendio", "Allegato Rischio Incendio"),
            (GestantiValutazione, "gestanti", "Allegato Lavoratrici Gestanti"),
            (MicroclimaValutazione, "microclima", "Allegato Microclima"),
            (BiologicoValutazione, "biologico", "Allegato Rischio Biologico"),
        ):
            r = await self.db.execute(
                select(model.id).where(model.azienda_id == self.azienda_id).limit(1)
            )
            if r.first() is not None:
                allegati_presenti.append((slug, label))
                if model is IncendioValutazione:
                    incendio_present = True

        # PEE (Piano di Emergenza ed Evacuazione) is mandatory under
        # D.M. 03/09/2021 art. 5 + D.Lgs. 81/2008 art. 43 when the workplace
        # has > 10 lavoratori OR an applicable incendio assessment. Audit
        # F-010 flagged the missing cross-reference; we add it here so the
        # allegati list correctly cites the document.
        n_persone = len(data.get("persone") or [])
        if n_persone > 10 or incendio_present:
            allegati_presenti.append((
                "pee", "Piano di Emergenza ed Evacuazione (PEE)"
            ))

        # Misure di miglioramento — auto-seed from pericoli I>=7 when empty
        # so the document is signable from day 1. Operators edit/extend the
        # rows via the API; subsequent generations preserve them.
        r = await self.db.execute(
            select(MisuraMiglioramento)
            .where(MisuraMiglioramento.azienda_id == self.azienda_id)
            .order_by(MisuraMiglioramento.ordine, MisuraMiglioramento.created_at)
        )
        misure = list(r.scalars().all())
        if not misure:
            misure = await self._auto_seed_misure(data)

        return {
            "foto_by_ambiente": foto_by_ambiente,
            "vdt_esposti_persona_ids": vdt_esposti_persona_ids,
            "allegati_presenti": allegati_presenti,
            "misure_miglioramento": misure,
        }

    async def _auto_seed_misure(self, data: dict) -> list:
        """Derive a starter set of misure di miglioramento from pericoli with
        I >= 5 (MODESTO and above).

        Persists the rows so subsequent generations skip the seeding pass and
        operators can edit them via the API. Tempi:
          I >= 9 GRAVISSIMO → "Immediatamente (entro 30 giorni)"
          I = 7-8 GRAVE     → "Entro 6 mesi"
          I = 5-6 MODESTO   → "Entro 12 mesi (ciclo annuale)"

        Procedura/risorse text is differentiated by the parent valutazione's
        ``categoria_rischio`` rather than emitted as one identical template
        string — addresses audit F-002.
        """
        from app.models.misura_miglioramento import MisuraMiglioramento

        ambienti = data.get("ambienti") or []
        persone = data.get("persone") or []
        ddl = next((p for p in persone if p.ruolo_datore_lavoro), None)
        rspp = next((p for p in persone if p.ruolo_rspp), None)
        ddl_name = (ddl.nominativo if ddl else "DATORE DI LAVORO").upper()
        rspp_name = (rspp.nominativo if rspp else "RSPP").upper()

        # Per-ambiente preposto lookup (ambiente.preposto_id -> persona),
        # falling back to any ambiente persona with ruolo_preposto, else
        # any azienda-level preposto. Mirrors the resolution order used
        # in _add_env_identity_table so the same name flows through the
        # measures table.
        persone_by_id = {p.id: p for p in persone}
        global_preposto = next((p for p in persone if p.ruolo_preposto), None)

        def _resolve_preposto(amb) -> str:
            pid = getattr(amb, "preposto_id", None)
            if pid and pid in persone_by_id:
                return persone_by_id[pid].nominativo.upper()
            for pers in getattr(amb, "persone", None) or []:
                if getattr(pers, "ruolo_preposto", False):
                    return pers.nominativo.upper()
            if global_preposto:
                return global_preposto.nominativo.upper()
            return "PREPOSTO DA NOMINARE"

        # Categories whose interventions are typically structural -> DdL.
        # Categories that are operational/behavioural -> Preposto.
        # Everything else (procedure, formazione, chemical, biological,
        # psychosocial) -> RSPP.
        ddl_categories = {
            "Strutture",
            "Impianti Elettrici",
            "Incendio-Esplosioni",
        }
        preposto_categories = {
            "Macchine",
            "Agenti Fisici",
            "Fattori Ergonomici",
        }

        def _classify_responsabile(amb, categoria: str, indice: int) -> str:
            # Severity escalation: GRAVISSIMO always escalates to DdL.
            if indice >= 9:
                return ddl_name
            if categoria in ddl_categories:
                return ddl_name
            if categoria in preposto_categories:
                return _resolve_preposto(amb)
            return rspp_name

        seeded: list[MisuraMiglioramento] = []
        ordine = 0
        for amb in ambienti:
            for v in getattr(amb, "valutazioni_rischio", None) or []:
                cat_long = normalize_categoria_to_long(
                    getattr(v, "categoria_rischio", "") or ""
                ) or "Generico"
                template = _MISURA_TEMPLATE_BY_CATEGORIA.get(
                    cat_long, _MISURA_TEMPLATE_DEFAULT
                )
                for per in getattr(v, "pericoli", None) or []:
                    indice = getattr(per, "indice_i", None)
                    if indice is None or indice < 5:
                        continue
                    if not getattr(per, "applicabile", True):
                        continue
                    livello = getattr(per, "livello_rischio", None) or ""
                    if indice >= 9:
                        tempi = "Immediatamente (entro 30 giorni)"
                    elif indice >= 7:
                        tempi = "Entro 6 mesi"
                    else:
                        tempi = "Entro 12 mesi (ciclo annuale)"

                    pericolo_text = (per.pericolo or "").strip()
                    misura_text = (
                        f"Riduzione del rischio in {(amb.nome or '').upper()}: "
                        f"{pericolo_text or cat_long}"
                    )
                    # Procedura/risorse from the categoria template are a
                    # starting point, not a finished plan. Mark them with a
                    # DA COMPILARE prefix so an inspector immediately sees
                    # which rows still need the RSPP to write the specific
                    # procedure for the actual pericolo (audit F-002,
                    # 2026-04-29 rerun). The marker disappears the moment
                    # the operator edits the row via the API.
                    da_compilare = "[DA COMPILARE — RSPP] "
                    misura = MisuraMiglioramento(
                        azienda_id=self.azienda_id,
                        pericolo_valutazione_id=per.id,
                        misura=misura_text,
                        procedura=da_compilare + template["procedura"],
                        risorse=da_compilare + template["risorse"],
                        responsabile=_classify_responsabile(amb, cat_long, indice),
                        scadenza=tempi,
                        priorita=livello,
                        ordine=ordine,
                    )
                    self.db.add(misura)
                    seeded.append(misura)
                    ordine += 1

        if seeded:
            await self.db.flush()
            # We do not commit here — leave the transaction open so the
            # surrounding generation request controls the boundary. The rows
            # are visible to subsequent reads in this session.
        return seeded

    async def _get_next_version(self) -> int:
        """Determine the next version number for this azienda's DVR."""
        from sqlalchemy import func, select
        from app.models.documento_generato import DocumentoGenerato

        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento == "dvr_master")
        )
        result = await self.db.execute(stmt)
        current_max = result.scalar()
        return current_max + 1

    # ------------------------------------------------------------------
    # Document styles setup
    # ------------------------------------------------------------------

    def _setup_styles(self, doc: Document) -> None:
        """Configure document-wide styles and defaults.

        Pagination hardening (user feedback 2026-05-12 — "Impaginazzioe
        dovrebbe rispettare piu le tabelle ei capitoli"):
          - All heading styles get `keepNext` and `keepLines` set on the
            style itself so an H1/H2/H3 never sits orphaned at the bottom
            of a page without at least one following content line and
            never splits mid-heading on long titles.
          - `Normal` keeps lines together so a single paragraph doesn't
            split awkwardly across pages.
        Both attributes are applied via the underlying XML because the
        python-docx ParagraphFormat helpers only act on per-paragraph
        overrides, not the base style.
        """
        from docx.oxml import OxmlElement

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # Heading styles — bind pagination behavior to the style so every
        # heading inherits it without per-call boilerplate.
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.color.rgb = _HEADER_BG

            # `style.paragraph_format.keep_with_next = True` would set the
            # value on a per-paragraph override; for a style we have to
            # touch <w:pPr> directly via the element API.
            pPr = heading_style.element.get_or_add_pPr()
            for tag in ("w:keepNext", "w:keepLines"):
                # Drop any existing entries so toggling is idempotent.
                for existing in pPr.findall(qn(tag)):
                    pPr.remove(existing)
                pPr.append(OxmlElement(tag))

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(
        self, doc: Document, azienda, generated_at: datetime, version: int
    ) -> None:
        """Add a professional cover page.

        Layout (top → bottom): logo, title block, company identity block
        (name + address + P.IVA + ATECO), date+version footer.
        Every field falls back gracefully when missing so generation
        never crashes on a sparse survey.
        """
        # Top spacer — kept tight (3 lines) so the title sits above the
        # vertical center, leaving room for the identity block below.
        for _ in range(3):
            doc.add_paragraph("")

        # Logo: embed from assets/logo.png if available, otherwise fall back
        # to an italic gray text placeholder so generation never breaks.
        if _LOGO_PATH.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(str(_LOGO_PATH), width=Inches(2.0))
            except Exception:
                # Any image-loading issue degrades gracefully to the text
                # placeholder below (e.g. corrupt file).
                run.text = "[LOGO AZIENDALE]"
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.italic = True
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[LOGO AZIENDALE]")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.italic = True

        doc.add_paragraph("")

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DOCUMENTO DI VALUTAZIONE DEI RISCHI")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = _HEADER_BG

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ai sensi degli artt. 17 e 28 del D.Lgs. 81/2008 e s.m.i.")
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")
        doc.add_paragraph("")

        # Company name — guard against None ragione_sociale (sparse surveys
        # were crashing here previously with AttributeError on .upper()).
        ragione = (azienda.ragione_sociale or "—").upper()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ragione)
        run.bold = True
        run.font.size = Pt(18)

        # Address
        address_parts = []
        if azienda.sede_legale_via:
            address_parts.append(azienda.sede_legale_via)
        if azienda.sede_legale_citta:
            address_parts.append(azienda.sede_legale_citta)
        if address_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" — ".join(address_parts))
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Identity line: P.IVA + ATECO when available
        identity_bits: list[str] = []
        partita_iva = getattr(azienda, "partita_iva", None)
        if partita_iva:
            identity_bits.append(f"P.IVA {partita_iva}")
        codice_ateco = getattr(azienda, "codice_ateco", None)
        if codice_ateco:
            identity_bits.append(f"ATECO {codice_ateco}")
        if identity_bits:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" · ".join(identity_bits))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Spacer before footer
        for _ in range(3):
            doc.add_paragraph("")

        # Date and version block — single centered paragraph with both bits.
        # Format the version with a 2-digit pad (Rev. 01) to mirror the
        # convention Luca uses on the master template.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Revisione {version:02d} — {generated_at.strftime('%d/%m/%Y')}"
        )
        run.font.size = Pt(12)
        run.bold = True

        # Page break
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Table of contents placeholder
    # ------------------------------------------------------------------

    def _add_table_of_contents(self, doc: Document):
        """Add a real Word TOC field that resolves to a clickable indice on
        first open in Word/LibreOffice (Ctrl+A, F9 to refresh).

        We emit a `<w:fldChar fldCharType="begin">` + instrText `TOC \\o "1-3" \\h \\z \\u`
        + `<w:fldChar fldCharType="end">` triplet wrapped in a paragraph. Word
        auto-prompts to update fields on open via the `<w:updateFields/>`
        setting below; after refresh the field resolves to the real heading
        list with page numbers.

        The cached body emitted here is a single placeholder line. It is
        replaced post-generation by `_finalize_table_of_contents`, which
        walks all real H1/H2/H3 headings emitted after the TOC and rebuilds
        the body so the unrefreshed view also shows the full structure.

        Returns ``(field_start_p, end_p)`` — the paragraph containing the
        `fldChar` begin/separate runs and the paragraph containing the
        `fldChar` end run. The finalizer uses these as anchors to identify
        the cached-body span.
        """
        from docx.oxml import OxmlElement

        doc.add_heading("INDICE", level=1)

        # The TOC field paragraph
        field_start_p = doc.add_paragraph()
        run = field_start_p.add_run()
        # 1) fldChar begin
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        # 2) instrText with the TOC instruction
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr)

        # 3) fldChar separate — anything between separate and end is the
        # cached TOC body that Word displays before the user refreshes.
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        # Cached body — one placeholder line so the field is well-formed
        # even if the finalizer never runs (e.g. exception during a later
        # part). The finalizer replaces this with the real heading list.
        placeholder_p = doc.add_paragraph()
        placeholder_run = placeholder_p.add_run(
            "Indice in fase di aggiornamento — premere Ctrl+A e poi F9 per "
            "ricaricare la tabella dei contenuti."
        )
        placeholder_run.font.size = Pt(11)
        placeholder_run.font.italic = True

        # 4) fldChar end on a fresh run so the cached body sits between
        end_p = doc.add_paragraph()
        end_run = end_p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end_run._r.append(fld_end)

        # Tell Word that the TOC needs an update on next open. The setting
        # is on the document-level settings part; insert the element if it
        # is not already present so we don't fight existing config.
        try:
            settings_root = doc.settings.element
            update_fields = settings_root.find(qn("w:updateFields"))
            if update_fields is None:
                update_fields = OxmlElement("w:updateFields")
                update_fields.set(qn("w:val"), "true")
                settings_root.append(update_fields)
        except Exception:
            # Non-fatal — the TOC field still works, only the auto-update
            # prompt is lost. Cached body covers the gap.
            pass

        doc.add_page_break()

        return field_start_p, end_p

    def _finalize_table_of_contents(self, doc: Document, field_start_p, end_p) -> None:
        """Replace the TOC field's cached body with the real outline.

        Walks all paragraphs emitted after ``end_p`` (i.e. Premessa onward),
        keeps Heading 1/2/3 styles, then rewrites the cached body between
        ``field_start_p`` and ``end_p`` so the pre-F9 view of the TOC shows
        the actual document structure instead of a generic placeholder.

        We can't emit real page numbers here — those come from Word's layout
        engine when the field is refreshed — so cached entries are
        page-number-free. The `<w:updateFields/>` setting written by
        `_add_table_of_contents` still prompts Word to refresh on first
        open, restoring page numbers automatically.
        """
        body = doc.element.body
        children = list(body)
        try:
            start_idx = children.index(field_start_p._p)
            end_idx = children.index(end_p._p)
        except ValueError:
            # Anchors not found — be defensive and leave the placeholder.
            return

        # Collect H1/H2/H3 headings emitted after the TOC. Anything before
        # ``end_p`` belongs to the cover page or the TOC itself.
        entries: list[tuple[int, str]] = []
        for para in doc.paragraphs:
            style_name = (para.style.name if para.style else "") or ""
            if not style_name.startswith("Heading "):
                continue
            try:
                level = int(style_name.split(" ")[1])
            except (IndexError, ValueError):
                continue
            if level not in (1, 2, 3):
                continue
            # Only headings after the TOC end fldChar contribute.
            try:
                para_idx = children.index(para._p)
            except ValueError:
                continue
            if para_idx <= end_idx:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            entries.append((level, text))

        if not entries:
            return

        # Remove the existing cached body paragraphs (between
        # field_start_p exclusive and end_p exclusive).
        for el in children[start_idx + 1 : end_idx]:
            body.remove(el)

        # Re-locate end_p after removals.
        end_idx = list(body).index(end_p._p)

        # Insert new cached body paragraphs in order before end_p. We
        # create each via doc.add_paragraph (appended at body tail), then
        # move the element to the correct slot. python-docx doesn't
        # expose a public "insert paragraph before" helper.
        for level, text in entries:
            new_p = doc.add_paragraph()
            # Indent by level so H2/H3 read as a sub-outline pre-refresh.
            # 0.0/0.4/0.8cm matches what Word's default "TOC 1/2/3"
            # styles render after F9 refresh — close enough that the
            # cached view doesn't visually jump when the field updates.
            indent_cm = (level - 1) * 0.4
            if indent_cm > 0:
                new_p.paragraph_format.left_indent = Cm(indent_cm)
            new_p.paragraph_format.space_after = Pt(2)
            run = new_p.add_run(text)
            run.font.size = Pt(11) if level == 1 else Pt(10)
            run.bold = level == 1

            # Move new_p from body tail to just before end_p.
            body.remove(new_p._p)
            end_idx = list(body).index(end_p._p)
            body.insert(end_idx, new_p._p)

    def _add_premessa(self, doc: Document) -> None:
        """Premessa — the legal preamble required by D.Lgs. 81/2008 art. 28.

        The master template carries ~20 paragraphs of foreword text under a
        Heading 1 "Premessa". The condensed version below preserves the
        binding obligations, the rielaborazione triggers (art. 29 c. 3) and
        the conservation requirement, and is sufficient for the document to
        read as a complete DVR rather than a numbered shell.
        """
        doc.add_heading("Premessa", level=1)

        for body in (
            "Il presente documento rappresenta l'attuazione dell'obbligo "
            "previsto per il datore di lavoro dall'art. 17, comma 1, "
            "lettera a) del D.Lgs. 9 aprile 2008, n. 81 e s.m.i., relativo "
            "alla valutazione di tutti i rischi per la salute e la "
            "sicurezza dei lavoratori, e dall'art. 28 dello stesso decreto, "
            "che ne disciplina l'oggetto e le modalita di redazione.",
            "La valutazione dei rischi e stata effettuata in collaborazione "
            "con il Responsabile del Servizio di Prevenzione e Protezione "
            "(RSPP), con il Medico Competente — ove la sorveglianza "
            "sanitaria sia richiesta dalla normativa o dalla valutazione "
            "stessa — e previa consultazione del Rappresentante dei "
            "Lavoratori per la Sicurezza (RLS).",
            "L'analisi prende in considerazione la natura dell'attivita, "
            "i luoghi di lavoro, le mansioni svolte, le attrezzature "
            "impiegate, le sostanze e i preparati pericolosi utilizzati, "
            "nonche le caratteristiche dei lavoratori, con particolare "
            "attenzione ai gruppi piu sensibili al rischio quali "
            "lavoratrici gestanti, minori, lavoratori con prescrizioni "
            "del Medico Competente e lavoratori notturni.",
            "Il documento e conservato presso l'unita produttiva alla "
            "quale si riferisce e reso disponibile per la consultazione "
            "ai soggetti previsti dall'art. 18 e dagli artt. 50 e 47 del "
            "D.Lgs. 81/2008. La sua rielaborazione e prevista nei casi "
            "stabiliti dall'art. 29, comma 3 — modifiche significative "
            "del processo produttivo o dell'organizzazione del lavoro, "
            "infortuni significativi, evoluzione della tecnica, esiti "
            "della sorveglianza sanitaria — e comunque ad ogni successiva "
            "modifica delle condizioni operative.",
            "Le misure di prevenzione e protezione attuate sono descritte, "
            "per ogni ambiente di lavoro, nella Parte III. Il programma "
            "delle misure di miglioramento di cui all'art. 28, comma 2, "
            "lettera c) e riportato in Parte IV §4.1.",
        ):
            p = doc.add_paragraph()
            run = p.add_run(body)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Pre-Parte I — Frontispiece (Template Tables 0, 1, 2)
    # ------------------------------------------------------------------

    def _add_pre_parte_i(
        self, doc: Document, azienda, generated_at: datetime
    ) -> None:
        """Render the front-matter block that appears before Parte I.

        Tables 0, 1, 2 from DVR_TEMPLATE_MAPPING.md — azienda identity,
        revision history, and a stamped-signature placeholder. The revision
        row uses the azienda's own DVR version (next-version - 1 because
        ``_get_next_version`` is called later; for the front page we reflect
        the current emission).
        """
        p = doc.add_paragraph()
        run = p.add_run(
            "ex art. 17, comma 1, lettera a) ed art. 28 del "
            "D.Lgs. 81/2008 e s.m.i."
        )
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph("")
        self._add_azienda_header_table(doc, azienda)
        doc.add_paragraph("")

        doc.add_heading("Storico Revisioni", level=3)
        self._add_revision_history_table(doc, generated_at)
        doc.add_paragraph("")

        self._add_timbro_firma_table(doc)
        doc.add_page_break()

    def _add_revision_history_table(
        self, doc: Document, generated_at: datetime
    ) -> None:
        """Template Table 1 — Rev./Motivazione/Data (7×3 DYNAMIC).

        Emits a single row for the current emission. Real clients will have
        this backed by the ``DocumentoGenerato`` version log in a later
        iteration; for now it's a truthful single-entry record.
        """
        headers = ["Rev.", "Motivazione", "Data"]
        rows = [[
            "00",
            "Emissione",
            generated_at.strftime("%d/%m/%Y"),
        ]]
        self._add_data_table(doc, headers, rows)

    def _add_timbro_firma_table(self, doc: Document) -> None:
        """Template Table 2 — single-cell 'Timbro e Firma' placeholder (2×1)."""
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        header_cell = table.rows[0].cells[0]
        header_cell.text = ""
        p = header_cell.paragraphs[0]
        run = p.add_run("Timbro e Firma")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _HEADER_TEXT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_cell_bg(header_cell, _HEADER_BG)

        body_cell = table.rows[1].cells[0]
        body_cell.text = ""
        body_cell.paragraphs[0].add_run("\n\n\n")

    # ------------------------------------------------------------------
    # Part I — Company data (Template Tables 3–17)
    # ------------------------------------------------------------------

    def _add_part_i(
        self,
        doc: Document,
        azienda,
        persone: list,
        attrezzature: list,
        sostanze_chimiche: list,
        ambienti: list,
    ) -> None:
        """Full Parte I with 15-table parity against the master template.

        Layout (Tables 3–17 in DVR_TEMPLATE_MAPPING.md):
          3   Azienda header (5×2)
          4   Anagrafica Aziendale (12×2)
          5   Dati occupazionali grid (N×5)
          6–9 Single-role titles (Datore di Lavoro / RSPP / RLS / Medico)
          10  Addetti Primo Soccorso
          11  Addetti Antincendio
          12  Ambienti + N.Lavoratori
          13  Attrezzature / Marcatura CE / Verifiche
          14  Sostanze chimiche / Produttore / Attivita
          15–17 Static hazard library (Sicurezza / Salute / Trasversali)
        """
        # Each Parte starts on a fresh page so the section break is
        # visually unambiguous regardless of where the previous part
        # ended. We do this on the heading itself rather than appending
        # an extra `add_page_break()` so the cached TOC outline (which
        # walks paragraphs in order) stays compact.
        h = doc.add_heading("PARTE I — DATI GENERALI DELL'AZIENDA", level=1)
        h.paragraph_format.page_break_before = True

        # Table 3 — Presentazione dell'azienda
        doc.add_heading("1. Presentazione dell'Azienda", level=2)
        self._add_azienda_header_table(doc, azienda)
        doc.add_paragraph("")

        # Table 4 — Anagrafica Aziendale (key-value)
        # Required rows (Ragione Sociale, Datore di Lavoro, Sede Legale,
        # Sede Operativa) always render. Optional rows are SKIPPED entirely
        # when null/empty — empty placeholder rows like "—" or
        # "Non comunicato" are noise that an inspector flags. The audit
        # gate in api/v1/documents._ensure_anagrafica_complete_for_dvr
        # already blocks generation when ALL of CF/Tel/Email/PEC are
        # missing, so at least one of those rows will always render.
        doc.add_heading("2. Anagrafica Aziendale", level=2)

        ddl = next((p for p in persone if p.ruolo_datore_lavoro), None)
        ddl_label = (
            f"{(ddl.nominativo or '').upper()}"
            f"{f' — {ddl.mansione}' if ddl and ddl.mansione else ''}"
            if ddl else "—"
        )
        # Count of effective workers prefers the actual roster (Persone) over
        # the operator-declared figure on the Azienda. The declared count is
        # surfaced in parentheses when it diverges, so reviewers can spot
        # discrepancies between visura camerale and the live survey.
        actual_count = len(persone or [])
        declared = getattr(azienda, "numero_dipendenti_dichiarati", None)
        if declared and declared != actual_count:
            dipendenti_label = f"{actual_count} (dichiarati: {declared})"
        else:
            dipendenti_label = str(actual_count)

        # Required rows — always rendered, even when the underlying field
        # is blank, so the table structure is recognisable.
        anagrafica_rows: list[tuple[str, str]] = [
            ("Ragione Sociale", azienda.ragione_sociale or "—"),
            ("Datore di Lavoro", ddl_label),
            ("Sede Legale", self._format_address(
                azienda.sede_legale_via, azienda.sede_legale_citta
            )),
            ("Sede Operativa", self._format_address(
                azienda.sede_operativa_via, azienda.sede_operativa_citta
            )),
        ]

        # Optional rows — skip entirely when null/empty. Order preserved
        # to match the prior layout (Attivita / ATECO above contacts,
        # operativi at the bottom).
        def _add_optional(label: str, value, formatter=lambda v: str(v)) -> None:
            if value is None:
                return
            if isinstance(value, str) and not value.strip():
                return
            anagrafica_rows.append((label, formatter(value)))

        _add_optional("Attivita", getattr(azienda, "attivita", None))
        _add_optional("Codice ATECO", getattr(azienda, "codice_ateco", None))
        _add_optional("Partita IVA", getattr(azienda, "partita_iva", None))
        _add_optional("Codice Fiscale", getattr(azienda, "codice_fiscale", None))
        _add_optional("Telefono", getattr(azienda, "telefono", None))
        _add_optional("Email", getattr(azienda, "email", None))
        _add_optional("PEC", getattr(azienda, "pec", None))
        _add_optional("Orario di Lavoro", getattr(azienda, "orario_lavoro", None))
        # Numero dipendenti always rendered — derived from the Persone roster
        # and the declared figure, both of which the survey forces.
        anagrafica_rows.append(("Numero Totale Dipendenti", dipendenti_label))
        _add_optional(
            "Metratura Totale",
            getattr(azienda, "metratura_totale", None),
            formatter=lambda v: f"{v} mq",
        )
        _add_optional(
            "Zona Sismica",
            getattr(azienda, "zona_sismica", None),
            formatter=lambda v: str(v),
        )

        self._add_key_value_table(doc, anagrafica_rows)
        doc.add_paragraph("")

        # Table 5 — Dati occupazionali (Nominativo | Mansione | Ambiente | Note | Contratto)
        doc.add_heading("3. Dati Occupazionali", level=2)
        self._add_dati_occupazionali_table(doc, persone)
        doc.add_paragraph("")

        # Tables 6–9 — Organizzazione Aziendale della Sicurezza
        doc.add_heading("4. Organizzazione Aziendale della Sicurezza", level=2)
        role_tables = [
            ("Datore di Lavoro",
             [p for p in persone if p.ruolo_datore_lavoro]),
            ("Responsabile del Servizio di Prevenzione e Protezione",
             [p for p in persone if p.ruolo_rspp]),
            ("Rappresentante dei Lavoratori per la Sicurezza",
             [p for p in persone if p.ruolo_rls]),
            ("Medico Competente",
             [p for p in persone if getattr(p, "ruolo_medico_competente", False)]),
        ]
        for title, role_persone in role_tables:
            self._add_single_role_title_table(doc, title, role_persone)
            doc.add_paragraph("")

        # Tables 10, 11 — Addetti Primo Soccorso / Antincendio
        self._add_addetti_role_table(
            doc,
            "Addetti al Primo Soccorso",
            [p for p in persone if p.ruolo_primo_soccorso],
        )
        doc.add_paragraph("")
        self._add_addetti_role_table(
            doc,
            "Addetti alla Prevenzione Incendi e Lotta Antincendio",
            [p for p in persone if p.ruolo_antincendio],
        )
        doc.add_paragraph("")

        # Table 12 — Ambienti di Lavoro + N. Lavoratori
        doc.add_heading("5. Ambienti di Lavoro", level=2)
        self._add_ambienti_summary_table(doc, ambienti)
        doc.add_paragraph("")

        # §1.6 Servizi Igienico-Assistenziali — required under D.Lgs. 81/2008
        # Allegato IV §1.10-1.16. Mandatory for >10 lavoratori (this client
        # has 18). Boilerplate references the law and surfaces the existing
        # azienda field if populated, otherwise emits a default paragraph
        # referencing the legal requirement (audit F-003).
        self._add_servizi_igienico_assistenziali_section(doc, azienda, persone)

        # Table 13 — Macchine, attrezzature ed impianti.
        # Pass ambienti so the renderer can resolve ambiente_id → nome
        # without lazy-loading `attrezzatura.ambiente` (the base loader
        # doesn't selectinload that relationship, and lazy access in an
        # async session raises MissingGreenlet).
        doc.add_heading("7. Macchine, Attrezzature ed Impianti", level=2)
        self._add_attrezzature_table(doc, attrezzature, ambienti)
        doc.add_paragraph("")

        # Table 14 — Sostanze, prodotti e preparati chimici
        doc.add_heading("8. Sostanze, Prodotti e Preparati Chimici", level=2)
        self._add_sostanze_table(doc, sostanze_chimiche)
        doc.add_paragraph("")

        # Tables 15, 16, 17 — Static hazard library
        doc.add_heading("9. Elenco Fattori di Pericolo (Riferimento)", level=2)
        p = doc.add_paragraph()
        run = p.add_run(
            "N.B. Gli elenchi seguenti sono da intendersi indicativi e non "
            "esaustivi. Sono valutati in dettaglio per ogni ambiente di "
            "lavoro nella Parte III."
        )
        run.font.size = Pt(9)
        run.font.italic = True
        doc.add_paragraph("")

        self._add_hazard_library_group(doc, "Rischi per la Sicurezza", [
            "Strutture", "Macchine", "Impianti Elettrici", "Incendio-Esplosioni",
        ])
        doc.add_paragraph("")
        self._add_hazard_library_group(doc, "Rischi per la Salute", [
            "Agenti Chimici", "Agenti Fisici", "Agenti Biologici", "Agenti Cancerogeni",
        ])
        doc.add_paragraph("")
        self._add_hazard_library_group(doc, "Rischi Trasversali", [
            "Organizzazione del Lavoro", "Fattori Psicologici", "Fattori Ergonomici",
        ])

        doc.add_page_break()

    def _add_dati_occupazionali_table(self, doc: Document, persone: list) -> None:
        """Template Table 5 — 5-col lavoratori grid including ambiente assignments.

        Cross-site roles (DdL / RSPP / RLS / MC) typically operate across
        the whole company and have no explicit ambiente assignment. We
        surface that with an explicit "Tutta l'azienda" label rather than
        an em-dash so reviewers don't read it as missing data
        (audit F-011).
        """
        # Column 4 is the persona's codice fiscale, not a free-form Note —
        # rename the header to match the actual content so an em-dash reads
        # as "CF not entered" rather than "broken Note column" (audit F-007,
        # 2026-04-29 rerun).
        headers = ["Nominativo", "Mansione", "Ambiente di Lavoro", "Codice Fiscale", "Tipologia contrattuale"]
        if not persone:
            self._add_data_table(doc, headers, [["—", "—", "—", "—", "—"]])
            return

        rows = []
        for p in persone:
            ambienti_names = ", ".join(
                (a.nome or "")
                for a in (getattr(p, "ambienti", None) or [])
                if getattr(a, "nome", None)
            )
            if not ambienti_names:
                if (
                    p.ruolo_datore_lavoro
                    or p.ruolo_rspp
                    or p.ruolo_rls
                    or getattr(p, "ruolo_medico_competente", False)
                ):
                    ambienti_names = "Tutta l'azienda (ruolo trasversale)"
                else:
                    ambienti_names = "—"
            cf = getattr(p, "codice_fiscale", None) or "—"
            rows.append([
                (p.nominativo or "—").upper(),
                (p.mansione or "—").upper(),
                ambienti_names.upper(),
                cf,
                (p.tipologia_contrattuale or "—").upper(),
            ])
        self._add_data_table(doc, headers, rows)

    def _add_single_role_title_table(
        self, doc: Document, title: str, role_persone: list
    ) -> None:
        """Template Tables 6–9 — single-column title table with the role-holder's name.

        Feedback #10 (2026-04-29): MC and RSPP are frequently external
        consultants. We append "(ESTERNO)" to the name when `is_esterno` is
        set. The flag is only writable on MC/RSPP in the survey UI, so DdL
        and RLS rows naturally carry it false and render unchanged.
        """
        def _format(p) -> str:
            base = (p.nominativo or "").upper()
            return f"{base} (ESTERNO)" if getattr(p, "is_esterno", False) else base

        names = [_format(p) for p in role_persone] or ["—"]
        table = doc.add_table(rows=1 + len(names), cols=1)
        table.style = "Table Grid"

        header_cell = table.rows[0].cells[0]
        header_cell.text = ""
        hp = header_cell.paragraphs[0]
        hrun = hp.add_run(title)
        hrun.bold = True
        hrun.font.size = Pt(9)
        hrun.font.color.rgb = _HEADER_TEXT
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_cell_bg(header_cell, _HEADER_BG)

        for i, name in enumerate(names, start=1):
            c = table.rows[i].cells[0]
            c.text = ""
            cp = c.paragraphs[0]
            crun = cp.add_run(name)
            crun.font.size = Pt(9)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_addetti_role_table(
        self, doc: Document, title: str, role_persone: list
    ) -> None:
        """Template Tables 10, 11 — Nominativo/Mansione grid with a spanning header."""
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        merged = table.rows[0].cells[0].merge(table.rows[0].cells[1])
        merged.text = ""
        p = merged.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _HEADER_TEXT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_cell_bg(merged, _HEADER_BG)

        sub_row = table.rows[1]
        for i, text in enumerate(["Nominativo", "Mansione"]):
            c = sub_row.cells[i]
            c.text = ""
            cp = c.paragraphs[0]
            crun = cp.add_run(text)
            crun.bold = True
            crun.font.size = Pt(9)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(c, _LIGHT_GRAY)

        if not role_persone:
            row = table.add_row()
            row.cells[0].text = "—"
            row.cells[1].text = "—"
            return

        for p in role_persone:
            row = table.add_row()
            for i, text in enumerate([
                (p.nominativo or "—").upper(),
                (p.mansione or "—").upper(),
            ]):
                c = row.cells[i]
                c.text = ""
                cp = c.paragraphs[0]
                crun = cp.add_run(text)
                crun.font.size = Pt(9)

    def _add_servizi_igienico_assistenziali_section(
        self, doc: Document, azienda, persone: list
    ) -> None:
        """§1.6 — Servizi Igienico-Assistenziali (D.Lgs. 81/2008 Allegato IV
        §1.10-1.16). Mandatory section for any workplace; the M/F separation
        requirement triggers above 10 lavoratori.
        """
        doc.add_heading("6. Servizi Igienico-Assistenziali", level=2)

        n_dipendenti = len(persone or [])
        sex_distribution = {p.sesso for p in persone if getattr(p, "sesso", None)}
        mixed = sex_distribution >= {"M", "F"}

        intro = (
            "Le dotazioni igienico-assistenziali sono conformi ai requisiti "
            "minimi previsti dall'Allegato IV del D.Lgs. 81/2008 (§1.10 - "
            "1.16) per quanto riguarda servizi igienici, lavabi, spogliatoi "
            "e locali di riposo / refezione."
        )
        p = doc.add_paragraph()
        run = p.add_run(intro)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        if n_dipendenti > 10 and mixed:
            extra = (
                f"Con un numero di lavoratori pari a {n_dipendenti} e in "
                "presenza di personale di entrambi i sessi, l'azienda "
                "garantisce servizi igienici e spogliatoi distinti per "
                "uomini e donne (Allegato IV §1.13)."
            )
        elif n_dipendenti > 10:
            extra = (
                f"Con un numero di lavoratori pari a {n_dipendenti}, "
                "l'azienda garantisce servizi igienici e spogliatoi "
                "adeguati al numero degli addetti (Allegato IV §1.13)."
            )
        else:
            extra = (
                "Sono garantiti servizi igienici e spogliatoi adeguati "
                "al numero degli addetti (Allegato IV §1.13)."
            )
        p = doc.add_paragraph()
        run = p.add_run(extra)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        servizi_descr = (
            getattr(azienda, "servizi_igienici_descrizione", None) or ""
        ).strip()
        if servizi_descr:
            p = doc.add_paragraph()
            run = p.add_run(servizi_descr)
            run.font.size = Pt(10)
        else:
            for bullet in (
                "Servizi igienici aerati e illuminati naturalmente o "
                "con ricambio d'aria forzato.",
                "Lavabi con acqua corrente e dispenser per il lavaggio "
                "delle mani.",
                "Spogliatoi con armadietti individuali a doppio scomparto "
                "ove l'attivita comporti l'uso di indumenti da lavoro.",
                "Locale di riposo/refezione adeguato quando i lavoratori "
                "consumano i pasti in azienda.",
            ):
                bp = doc.add_paragraph(style="List Bullet")
                run = bp.add_run(bullet)
                run.font.size = Pt(10)
        doc.add_paragraph("")

    def _add_ambienti_summary_table(self, doc: Document, ambienti: list) -> None:
        """Template Table 12 — Ambiente | Tipo | Metratura | N. Lavoratori.

        Note on N. Lavoratori: a worker assigned to multiple ambienti is
        counted once per ambiente, so the column sum can exceed the total
        dipendenti reported in T004 (e.g. preposti + addetti emergenza
        operate across several aree). A footnote below the table makes the
        convention explicit so an inspector does not flag the discrepancy.
        """
        headers = ["Ambiente", "Tipo", "Metratura", "N. Lavoratori"]
        if not ambienti:
            rows = [["—", "—", "—", "0"]]
        else:
            rows = []
            for a in ambienti:
                mq = getattr(a, "superficie_mq", None)
                rows.append([
                    (a.nome or "—").upper(),
                    (a.tipo or "—"),
                    f"{mq} mq" if mq else "—",
                    str(len(getattr(a, "persone", None) or [])),
                ])
        self._add_data_table(doc, headers, rows)

        note = doc.add_paragraph()
        run = note.add_run(
            "Nota: la colonna \"N. Lavoratori\" conteggia ciascun addetto "
            "una volta per ogni ambiente in cui opera; preposti, addetti "
            "alle emergenze e personale con mansioni trasversali "
            "compaiono pertanto in piu righe. Il totale dei dipendenti "
            "in forza all'azienda e riportato in §1.1."
        )
        run.font.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_attrezzature_table(
        self,
        doc: Document,
        attrezzature: list,
        ambienti: list | None = None,
    ) -> None:
        """Template Table 13 — global equipment inventory (Parte I §7).

        Columns: Descrizione | Ambiente | Marcata CE | Verifiche periodiche.

        Rationale for the Ambiente column (feedback 2026-05-12: "macchine
        attrezzature impianti non mette le cose giuste"): the previous
        3-column variant rendered 1 row per piece of equipment with no
        location context, so identical descriptions (e.g. two "TRAPANO A
        COLONNA" rows from two different ambienti) looked like
        duplicates. Surfacing the ambiente name keeps the per-row signal
        and lets the reader cross-reference the per-ambiente blocks in
        Parte III.

        We resolve `ambiente_id → nome` via the in-memory ambienti list
        rather than lazy-loading `attrezzatura.ambiente`: the base
        loader doesn't selectinload that backref, and lazy access in an
        async session raises MissingGreenlet. ``ambienti=None`` keeps
        the old "Ambiente: —" behavior for any caller that doesn't
        thread the list through (defensive — no current callsite hits
        this branch).

        Rows are grouped by ambiente name so identical ambienti cluster
        together — matches operator expectation when scanning the list.
        """
        amb_name_by_id: dict = {}
        if ambienti:
            for amb in ambienti:
                amb_name_by_id[amb.id] = (amb.nome or "—").upper()

        headers = [
            "Macchine, Attrezzature ed Impianti",
            "Ambiente",
            "Marcata CE",
            "Verifiche Periodiche",
        ]
        if not attrezzature:
            rows = [["Nessuna attrezzatura registrata.", "—", "—", "—"]]
        else:
            def _amb_label(a) -> str:
                amb_id = getattr(a, "ambiente_id", None)
                return amb_name_by_id.get(amb_id, "—")

            sorted_att = sorted(
                attrezzature,
                key=lambda a: (_amb_label(a), (a.descrizione or "").lower()),
            )
            rows = [
                [
                    (a.descrizione or "—").upper(),
                    _amb_label(a),
                    "SI" if a.marcatura_ce else "NO",
                    "SI" if a.verifiche_periodiche else "NO",
                ]
                for a in sorted_att
            ]
        self._add_data_table(doc, headers, rows)

    def _add_sostanze_table(self, doc: Document, sostanze: list) -> None:
        """Template Table 14 — chemical inventory + SDS hazard detail.

        Two-block layout:
          1. Inventory table (4 cols) — name, manufacturer, state, GHS
             pictogram codes (joined). Uses SDS-extracted pittogrammi when
             available; falls back to "—" for manually entered rows.
          2. Per-sostanza H/P phrase detail block, emitted only for
             sostanze that have at least one hazard phrase. Skipped
             entirely when no SDS data exists, so manually-entered rows
             stay compact.
        """
        headers = [
            "Sostanza / Prodotto",
            "Produttore / Distributore",
            "Attivita / Uso",
            "Destinazione d'uso",
            "Stato",
            "Pittogrammi GHS",
        ]
        if not sostanze:
            rows = [[
                "Nessuna sostanza chimica registrata.",
                "—",
                "—",
                "—",
                "—",
                "—",
            ]]
            self._add_data_table(doc, headers, rows)
            return

        rows = []
        for s in sostanze:
            pittogrammi = getattr(s, "pittogrammi", None) or []
            pittogrammi_text = ", ".join(pittogrammi) if pittogrammi else "—"
            rows.append([
                (s.nome_prodotto or "—").upper(),
                (s.produttore or "—").upper(),
                (getattr(s, "attivita_uso", None) or "—"),
                (getattr(s, "destinazione_uso", None) or "—"),
                (getattr(s, "stato_miscela", None) or "—").upper(),
                pittogrammi_text,
            ])
        self._add_data_table(doc, headers, rows)

        # Phase 8.5 — emit per-sostanza H/P detail only when SDS data is
        # present. We don't add a heading when nothing has SDS data, so the
        # absence is invisible (operator-friendly).
        sostanze_with_sds = [
            s for s in sostanze
            if (getattr(s, "frasi_h", None) or [])
            or (getattr(s, "frasi_p", None) or [])
        ]
        if not sostanze_with_sds:
            return

        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Dettaglio frasi di pericolo (H) e consigli di prudenza (P)")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _HEADER_BG

        for s in sostanze_with_sds:
            frasi_h = getattr(s, "frasi_h", None) or []
            frasi_p = getattr(s, "frasi_p", None) or []

            p = doc.add_paragraph()
            run = p.add_run((s.nome_prodotto or "—").upper())
            run.bold = True
            run.font.size = Pt(9)

            if frasi_h:
                p = doc.add_paragraph()
                run = p.add_run("Frasi H: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run("; ".join(frasi_h))
                run.font.size = Pt(9)

            if frasi_p:
                p = doc.add_paragraph()
                run = p.add_run("Frasi P: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run("; ".join(frasi_p))
                run.font.size = Pt(9)

            doc.add_paragraph("")

    def _add_hazard_library_group(
        self, doc: Document, macro_label: str, categorie: list[str]
    ) -> None:
        """Template Tables 15/16/17 — 2-col static hazard catalog per macro-area."""
        rows: list[list[str]] = []
        for categoria in categorie:
            items = HAZARD_LIBRARY.get(categoria, [])
            for item in items:
                rows.append([categoria, item])

        if not rows:
            rows = [["—", "—"]]

        self._add_data_table(doc, headers=["Categoria", macro_label], rows=rows)

    # ------------------------------------------------------------------
    # Part II — Activity description and risk assessment methodology
    # ------------------------------------------------------------------

    def _add_part_ii(self, doc: Document, azienda) -> None:
        """Add Part II: activity description and risk methodology.

        Renders four sub-sections:
          2.1 Descrizione dell'Attivita (from azienda fields, with placeholder
              fallback when the survey field is empty).
          2.2 Metodologia di Valutazione dei Rischi (static boilerplate with
              a color-coded risk-level lookup table).
          2.3 Scala di Probabilita (P).
          2.4 Scala del Danno (D).
        """
        h = doc.add_heading(
            "PARTE II — DESCRIZIONE DELL'ATTIVITA E METODOLOGIA DI VALUTAZIONE",
            level=1,
        )
        h.paragraph_format.page_break_before = True

        # Template Table 18 — Azienda identity block at the top of Parte II
        self._add_azienda_header_table(doc, azienda)
        doc.add_paragraph("")

        # 2.1 — Activity description
        doc.add_heading("2.1 Descrizione dell'Attivita", level=2)

        descrizione = (azienda.descrizione_attivita or "").strip()
        if descrizione:
            # Split into multiple paragraphs on blank-line separators or on
            # sentence-group boundaries so the section reads as the
            # multi-paragraph narrative the spec expects (audit F-004).
            paragraphs = self._split_descrizione_paragraphs(descrizione)
            for body in paragraphs:
                p = doc.add_paragraph()
                run = p.add_run(body)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            run = p.add_run(
                "[Descrizione dell'attivita da compilare durante la revisione]"
            )
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Synthesized "ciclo produttivo" paragraph from ATECO + attivita +
        # zona sismica + orario lavoro — these structured fields exist
        # on the Azienda but aren't surfaced in the descrizione narrative.
        ciclo_bits: list[str] = []
        ateco = getattr(azienda, "codice_ateco", None)
        attivita = getattr(azienda, "attivita", None)
        if ateco and attivita:
            ciclo_bits.append(
                f"Ciclo produttivo: l'attivita e classificata sotto il "
                f"codice ATECO {ateco} ({attivita})."
            )
        zona = getattr(azienda, "zona_sismica", None)
        if zona:
            ciclo_bits.append(
                f"Zona sismica di appartenenza: {zona} "
                "(O.P.C.M. 3274/2003 e successive deliberazioni regionali)."
            )
        orario = getattr(azienda, "orario_lavoro", None)
        if orario:
            ciclo_bits.append(
                f"Articolazione dell'orario di lavoro: {orario}."
            )
        if ciclo_bits:
            p = doc.add_paragraph()
            run = p.add_run(" ".join(ciclo_bits))
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

        contesto = (getattr(azienda, "contesto_territoriale", None) or "").strip()
        if contesto:
            p = doc.add_paragraph()
            run = p.add_run(f"Contesto territoriale: {contesto}")
            run.font.size = Pt(10)

        # US-2.2 AC1: inject regione + applicable regional regulations into
        # Parte II Contesto Territoriale. Driven off sede_legale_citta (falls
        # back to sede_operativa_citta when absent). Silent no-op for comuni
        # not in the 158-entry lookup so the generator never fails because of
        # a missing regulation match — the operator fills it in during review.
        sede_citta = (
            getattr(azienda, "sede_legale_citta", None)
            or getattr(azienda, "sede_operativa_citta", None)
            or ""
        )
        regione, regulations = get_regulations_for_comune(sede_citta.strip())
        if regione and regulations:
            p = doc.add_paragraph()
            run = p.add_run(f"Regione di riferimento: {regione}")
            run.font.size = Pt(10)
            run.font.bold = True

            p = doc.add_paragraph()
            run = p.add_run(
                "Regolamenti regionali applicabili (in aggiunta al D.Lgs. 81/2008):"
            )
            run.font.size = Pt(10)
            run.font.italic = True

            for reg in regulations:
                bullet = doc.add_paragraph(style="List Bullet")
                run = bullet.add_run(f"{reg['titolo']} — ")
                run.font.size = Pt(10)
                run.font.bold = True
                run = bullet.add_run(reg["riferimento"])
                run.font.size = Pt(10)
                run = bullet.add_run(f" ({reg['ambito']})")
                run.font.size = Pt(10)
                run.font.italic = True

        doc.add_paragraph("")

        # Template Table 19 — Definizioni (glossary)
        doc.add_heading("2.2 Definizioni", level=2)
        self._add_data_table(
            doc,
            headers=["Termine", "Definizione"],
            rows=[list(r) for r in _DEFINIZIONI_ROWS],
        )
        doc.add_paragraph("")

        # 2.3 — Risk assessment methodology
        doc.add_heading("2.3 Metodologia di Valutazione dei Rischi", level=2)

        p = doc.add_paragraph()
        run = p.add_run(_METODOLOGIA_INTRO_1)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(_METODOLOGIA_INTRO_2)
        run.font.size = Pt(10)

        doc.add_paragraph("")
        self._add_risk_level_table(doc)
        doc.add_paragraph("")

        # Template Table 21 — Scala di Probabilita (P) with full criteria column
        doc.add_heading("2.4 Scala di Probabilita (P)", level=2)
        self._add_data_table(
            doc,
            headers=["P", "Livello", "Criteri"],
            rows=[list(row) for row in _PROBABILITA_CRITERI_ROWS],
        )
        doc.add_paragraph("")

        # Template Table 22 — Scala del Danno (D) with full criteria column
        doc.add_heading("2.5 Scala del Danno (D)", level=2)
        self._add_data_table(
            doc,
            headers=["D", "Livello", "Criteri"],
            rows=[list(row) for row in _DANNO_CRITERI_ROWS],
        )

        doc.add_page_break()

    def _add_risk_level_table(self, doc: Document) -> None:
        """Render the I-range -> Livello/Azione/Tempistica lookup table.

        The ``Livello`` column is shaded with the same ``_RISK_COLORS``
        palette used by Part III so the reader sees a consistent color
        language across the document.
        """
        headers = ["I", "Livello", "Azione", "Tempistica"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _HEADER_TEXT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(cell, _HEADER_BG)

        # Data rows
        for i_range, livello, azione, tempistica in _RISK_LEVEL_TABLE_ROWS:
            row = table.add_row()
            values = [i_range, livello, azione, tempistica]
            for col_idx, text in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9)

                if col_idx in (0, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Color the Livello cell (white bold text on palette color)
                if col_idx == 1 and livello in _RISK_COLORS:
                    self._set_cell_bg(cell, _RISK_COLORS[livello])
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Part III — Risk assessment per environment
    # ------------------------------------------------------------------

    def _add_part_iii(
        self,
        doc: Document,
        azienda,
        persone: list,
        ambienti: list,
        attrezzature: list,
        extras: dict,
    ) -> None:
        """Add Part III: per-environment risk assessment block + trailing
        sections (mansioni-rischi-specifici, DPI per mansione, segnaletica,
        programma di formazione).

        Emits the template-shaped env block (tables 23–33 in DVR_TEMPLATE_MAPPING.md):
          - Table 23 (once): azienda identity header — Ragione Sociale + Sede.
          - Per environment: identity, addetti, attrezzature present, foto
            (when AmbienteFoto rows exist), risk-category checklist (SI/NO),
            and one 5-col risk table per applicable macro-category.
          - Trailing: 4 sections covering mansioni con rischi specifici,
            DPI per mansione, segnaletica di sicurezza, programma formazione.
        """
        h = doc.add_heading(
            "PARTE III — VALUTAZIONE DEI RISCHI PER AMBIENTE DI LAVORO",
            level=1,
        )
        h.paragraph_format.page_break_before = True

        self._add_azienda_header_table(doc, azienda)
        doc.add_paragraph("")

        if not ambienti:
            p = doc.add_paragraph("Nessun ambiente di lavoro registrato.")
            p.runs[0].font.italic = True
            doc.add_page_break()
            return

        # Phase 8.2 — bucket attrezzature by ambiente_id once so each env
        # section gets its own slice. Anything without an ambiente_id is
        # silently skipped here (it still appears in the global Part I
        # inventory table, so it isn't lost).
        attrezzature_by_ambiente: dict = {}
        for att in attrezzature:
            amb_id = getattr(att, "ambiente_id", None)
            if amb_id is None:
                continue
            attrezzature_by_ambiente.setdefault(amb_id, []).append(att)

        foto_by_ambiente = extras.get("foto_by_ambiente", {})
        vdt_ids = extras.get("vdt_esposti_persona_ids", set())
        persone_by_id = {getattr(p, "id", None): p for p in persone}
        for ambiente in ambienti:
            env_attrezzature = attrezzature_by_ambiente.get(ambiente.id, [])
            env_foto = foto_by_ambiente.get(ambiente.id, [])
            # VDT-esposti workers in this ambiente trigger a synthetic row
            # in the Agenti Fisici risk table referencing art. 174 + the
            # Allegato VDT (audit F-007).
            env_persone = list(getattr(ambiente, "persone", None) or [])
            vdt_in_env = sum(
                1 for p in env_persone
                if getattr(p, "id", None) in vdt_ids
            )
            self._add_environment_section(
                doc,
                ambiente,
                persone_by_id,
                env_attrezzature,
                env_foto,
                vdt_in_env,
            )

        # Trailing Part III sections — driven by persone + extras.
        self._add_mansioni_rischi_specifici_section(doc, persone, extras)
        self._add_dpi_per_mansione_section(doc, persone, extras)
        self._add_segnaletica_section(doc)
        self._add_formazione_programma_section(doc, persone)

    def _add_azienda_header_table(self, doc: Document, azienda) -> None:
        """Template Table 23 — Azienda / Sede identity block (once, at top of Parte III)."""
        rows = [
            ("Azienda", (azienda.ragione_sociale or "—").upper()),
            ("Sede Legale", azienda.sede_legale_via or "—"),
            ("Sede Legale", azienda.sede_legale_citta or "—"),
        ]
        if azienda.sede_operativa_via or azienda.sede_operativa_citta:
            rows.append(
                ("Sede Operativa", self._format_address(
                    azienda.sede_operativa_via, azienda.sede_operativa_citta
                ))
            )
        # Issue #11 (2026-05-14): when the client has more than one
        # operating address, list the extras as additional rows under the
        # "Altre sedi operative" label. The list is whatever the operator
        # entered in the aziende form — empty rows have already been
        # filtered out at submit time, but defend against missing-attr
        # for older Azienda rows that pre-date the JSONB column.
        extras = getattr(azienda, "sedi_operative_extra", None) or []
        for sede in extras:
            if not isinstance(sede, dict):
                continue
            line = self._format_address(
                sede.get("via"),
                sede.get("citta") or sede.get("comune"),
            )
            if line and line != "—":
                rows.append(("Altre sedi operative", line))
        self._add_key_value_table(doc, rows)

    def _add_environment_section(
        self,
        doc: Document,
        ambiente,
        persone_by_id: dict,
        attrezzature: list,
        foto: list,
        vdt_count: int = 0,
    ) -> None:
        """Render the env section for a single environment.

        Phase 8.2 — DVR esploso per ambiente: now also lists the
        attrezzature present in this ambiente (the global Part I inventory
        stays in place; this is the per-env slice the operator asked for).

        Order mirrors tables 24–33 in the template:
          1. Identity (Table 24) — ambiente / preposto / descrizione.
          2. Addetti (Table 25) — nominativo / mansione.
          3. Attrezzature presenti in questo ambiente (Phase 8.2).
          4. Foto/planimetrie (when AmbienteFoto rows exist).
          5. Risk-category checklist (Table 26) — SI/NO per macro-area.
          6. One 5-col risk table per applicable macro-category.
        """
        nome_ambiente = (ambiente.nome or "—").upper()
        doc.add_heading(
            f"Identificazione dell'Ambiente di Lavoro e degli Addetti — {nome_ambiente}",
            level=2,
        )

        self._add_env_identity_table(doc, ambiente, persone_by_id)
        doc.add_paragraph("")
        self._add_env_addetti_table(doc, ambiente)
        doc.add_paragraph("")

        # Phase 8.2 — Attrezzature per ambiente. Reuses the same column
        # shape as the global Part I inventory so the visual pattern is
        # consistent. Empty slice → single placeholder row.
        doc.add_heading(
            f"Macchine, Attrezzature ed Impianti — {nome_ambiente}",
            level=3,
        )
        self._add_env_attrezzature_table(doc, attrezzature)
        doc.add_paragraph("")

        if foto:
            self._add_env_foto_block(doc, nome_ambiente, foto)

        doc.add_heading(
            f"Identificazione dei Fattori di Rischio — {nome_ambiente}",
            level=3,
        )
        self._add_env_risk_checklist(doc, ambiente)
        doc.add_paragraph("")

        self._add_env_risk_tables(doc, ambiente, vdt_count=vdt_count)
        doc.add_page_break()

    def _add_env_foto_block(
        self, doc: Document, nome_ambiente: str, foto: list
    ) -> None:
        """Embed up to 3 foto/planimetrie inline at consistent width.

        Skips photos whose filesystem file is missing rather than failing the
        whole document — happens during dev/test resets when the DB row
        survives but the storage path was wiped.
        """
        doc.add_heading(
            f"Foto e Planimetrie — {nome_ambiente}", level=3
        )
        embedded = 0
        for f in foto[:3]:
            path = getattr(f, "file_path", None)
            if not path or not os.path.exists(path):
                continue
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(path, width=Cm(12))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # "Fig. N — filename" prefix prevents downstream text-only
                # readers (audit tools, accessibility readers) from
                # misreading the styled italic caption as a stray bare
                # filename (audit F-003, 2026-04-29 rerun).
                cap_run = caption.add_run(
                    f"Fig. {embedded + 1} — {f.filename or ''}"
                )
                cap_run.font.italic = True
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                embedded += 1
            except Exception:
                continue
        if embedded == 0:
            p = doc.add_paragraph()
            run = p.add_run(
                "[Foto/planimetrie non disponibili al momento della "
                "generazione]"
            )
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Part III tail — mansioni rischi specifici, DPI, segnaletica, formazione
    # ------------------------------------------------------------------

    def _add_mansioni_rischi_specifici_section(
        self, doc: Document, persone: list, extras: dict
    ) -> None:
        """Emit the "Mansioni che espongono i lavoratori a rischi specifici"
        table required by the template tail (DOCUMENT_STRUCTURE.md §3.N+1).

        Driven by per-persona signals, then aggregated up to mansione for
        the legal table layout. When two persone with the same mansione
        diverge, the row shows the union of risks.

        Sources, per persona:
          - VDT esposti (>= 20h/week → "Videoterminali" rischio specifico)
          - persona.attrezzature_speciali codes (each maps to a rischio)
          - persona.rischi_specifici_codes ticked in the wizard
        """
        from app.services.reference_data import RISCHI_SPECIFICI_CATALOG

        doc.add_heading(
            "Mansioni che espongono i lavoratori a rischi specifici", level=2
        )

        # Build mansione → set[rischio_label] index.
        # Track the per-persona contribution count so we can flag mansioni
        # where the flags vary across persone (audit hint for the operator).
        mansioni_idx: dict[str, set[str]] = {}
        mansione_persona_codes: dict[str, list[frozenset[str]]] = {}
        vdt_ids = extras.get("vdt_esposti_persona_ids", set())

        attrezzatura_to_rischio = {
            "lavori_in_quota": "Lavori in quota",
            "trabattelli": "Utilizzo di trabattelli",
            "ponteggi": "Utilizzo di ponteggi",
            "carrello_elevatore": "Utilizzo di carrelli elevatori",
            "ple": "Utilizzo di piattaforme di lavoro elevabili (PLE)",
            "gru": "Utilizzo di gru",
            "ruspa_escavatore": "Utilizzo di ruspe ed escavatori",
            "patente_cde": "Guida professionale (patente C/D/E)",
            "adr": "Trasporto merci pericolose (ADR)",
        }

        for p in persone:
            mansione = (p.mansione or "").strip()
            if not mansione:
                continue
            bucket = mansioni_idx.setdefault(mansione, set())
            persona_codes: set[str] = set()
            if p.id in vdt_ids:
                label = "Videoterminali (VDT) — esposizione >= 20h/sett."
                bucket.add(label)
                persona_codes.add(label)
            for code in (p.attrezzature_speciali or []):
                label = attrezzatura_to_rischio.get(code)
                if label:
                    bucket.add(label)
                    persona_codes.add(label)
            for rs_code in (getattr(p, "rischi_specifici_codes", None) or []):
                meta = RISCHI_SPECIFICI_CATALOG.get(rs_code) or {}
                label = meta.get("etichetta")
                if label:
                    bucket.add(label)
                    persona_codes.add(label)
            mansione_persona_codes.setdefault(mansione, []).append(
                frozenset(persona_codes)
            )

        # Filter out mansioni without specific risks — they belong to the
        # generic risk assessment, not to this table.
        rows = []
        for mansione in sorted(mansioni_idx.keys()):
            rischi = sorted(mansioni_idx[mansione])
            if not rischi:
                continue
            count = sum(
                1
                for p in persone
                if (p.mansione or "").strip() == mansione
            )
            # If the persone with this mansione don't all share the same
            # set of rischi, flag the row so the Medico del Lavoro knows
            # to consult the per-persona detail (kept in the CRM).
            persona_sets = mansione_persona_codes.get(mansione, [])
            distinct = {s for s in persona_sets if s}
            varia = len(distinct) > 1
            rischi_cell = "; ".join(rischi)
            if varia:
                rischi_cell += " (varia per lavoratore — vedere scheda persona)"
            rows.append([
                mansione.upper(),
                str(count),
                rischi_cell,
            ])

        if not rows:
            p = doc.add_paragraph()
            run = p.add_run(
                "Nessuna mansione presenta esposizioni a rischi specifici "
                "che richiedano sorveglianza sanitaria mirata in aggiunta "
                "alla valutazione generale."
            )
            run.font.size = Pt(10)
            run.font.italic = True
            doc.add_paragraph("")
            return

        self._add_data_table(
            doc,
            headers=[
                "Mansione",
                "N. Lavoratori",
                "Rischi specifici / Sorveglianza Sanitaria",
            ],
            rows=rows,
        )
        doc.add_paragraph("")

    def _add_dpi_per_mansione_section(
        self, doc: Document, persone: list, extras: dict
    ) -> None:
        """Emit one DPI grid per mansione, listing the codes the operator
        has ticked on each persona in the wizard.

        Aggregates per-persona ``dpi_codes`` up to mansione (union). When
        the persone with the same mansione have divergent DPI selections,
        a "varia per lavoratore" note is appended so the operator and
        the Medico del Lavoro know to consult the per-persona detail.

        Mansioni without any persona-level DPI flags are omitted — they
        roll up into the generic Gestione DPI procedure (§4.8).
        """
        from app.services.reference_data import DPI_CATALOG

        doc.add_heading("DPI in dotazione per Mansione", level=2)

        # Aggregate per-persona DPI codes up to mansione (union).
        mansione_codes: dict[str, set[str]] = {}
        mansione_persona_codes: dict[str, list[frozenset[str]]] = {}
        for p in persone:
            nome = (p.mansione or "").strip()
            if not nome:
                continue
            codes = list(getattr(p, "dpi_codes", None) or [])
            mansione_codes.setdefault(nome, set()).update(codes)
            mansione_persona_codes.setdefault(nome, []).append(frozenset(codes))

        # Drop mansioni whose persone all have empty selections.
        mansioni: dict[str, list[str]] = {
            nome: sorted(codes)
            for nome, codes in mansione_codes.items()
            if codes
        }

        if not mansioni:
            p = doc.add_paragraph()
            run = p.add_run(
                "Le matrici DPI per mansione sono in fase di compilazione "
                "tramite la procedura di Sorveglianza Sanitaria. Nel "
                "frattempo, i DPI sono assegnati in base alla valutazione "
                "dei rischi descritta in Parte III e secondo quanto "
                "previsto al §4.8."
            )
            run.font.size = Pt(10)
            run.font.italic = True
            doc.add_paragraph("")
            return

        for mansione, codes in sorted(mansioni.items()):
            doc.add_heading(mansione.upper(), level=3)

            persona_sets = mansione_persona_codes.get(mansione, [])
            distinct_nonempty = {s for s in persona_sets if s}
            varia = len(distinct_nonempty) > 1
            if varia:
                p = doc.add_paragraph()
                run = p.add_run(
                    "Nota: la dotazione DPI varia tra i lavoratori con "
                    "questa mansione — la tabella riporta l'unione delle "
                    "selezioni. Per l'assegnazione individuale fare "
                    "riferimento alla scheda persona."
                )
                run.font.size = Pt(9)
                run.font.italic = True

            rows = []
            # Group by area so the table mirrors the DPI catalog layout.
            # Spec §3.N+2 requires Descrizione | Marca/Modello | Note;
            # Marca/Modello and Note are operator-completed at first audit
            # so we emit a defensible placeholder rather than blanks.
            by_area: dict[str, list[str]] = {}
            for code in codes:
                meta = DPI_CATALOG.get(code) or {}
                area = meta.get("area", "Altro")
                etichetta = meta.get("etichetta", code)
                by_area.setdefault(area, []).append(etichetta)
            for area in sorted(by_area.keys()):
                rows.append([
                    f"{area}: {'; '.join(sorted(by_area[area]))}",
                    "Da definire al primo audit DPI",
                    "Conforme a Reg. UE 2016/425; verifica annuale",
                ])
            self._add_data_table(
                doc,
                headers=["Descrizione DPI", "Marca / Modello", "Note"],
                rows=rows,
            )
            doc.add_paragraph("")

    def _add_segnaletica_section(self, doc: Document) -> None:
        """Static reference to the segnaletica di sicurezza obligation
        (D.Lgs. 81/2008 art. 161-166, Allegato XXIV-XXXII).

        Lists the four canonical sign categories with color/shape mapping.
        """
        doc.add_heading("Segnaletica di Sicurezza", level=2)

        p = doc.add_paragraph()
        run = p.add_run(
            "L'azienda adotta segnaletica di sicurezza conforme al D.Lgs. "
            "81/2008 (artt. 161-166, allegati XXIV-XXXII) per richiamare "
            "l'attenzione su rischi residui, divieti, prescrizioni e "
            "indicazioni di emergenza."
        )
        run.font.size = Pt(10)

        # Tabella 105 — Codifica cromatica (Allegato XXV D.Lgs. 81/2008).
        # Mappa colore -> significato e indicazioni d'uso.
        doc.add_heading("Codifica cromatica della segnaletica", level=3)
        color_rows = [
            ["Rosso", "Divieto", "Atteggiamenti pericolosi", "Stop, dispositivi di interruzione, sgombero"],
            ["Rosso", "Pericolo / allarme", "Identificazione attrezzature antincendio", "Estintori, idranti, punto di raccolta"],
            ["Giallo o giallo-arancio", "Avvertimento", "Attenzione, cautela, verifica", "Pericolo elettrico, infiammabile, biologico"],
            ["Azzurro", "Prescrizione", "Comportamento o azione specifica, obbligo DPI", "Obbligo casco, guanti, occhiali, otoprotettori"],
            ["Verde", "Salvataggio / soccorso", "Porte, uscite, percorsi, materiali, postazioni, locali", "Uscita di emergenza, cassetta primo soccorso, defibrillatore"],
            ["Verde", "Situazione di sicurezza", "Ritorno alla normalita", "Punto di raccolta evacuazione"],
        ]
        self._add_data_table(
            doc,
            headers=["Colore", "Significato", "Indicazioni e precisazioni", "Esempi"],
            rows=color_rows,
        )
        doc.add_paragraph("")

        # Tabella 106 — Forma dei cartelli (Allegato XXIV D.Lgs. 81/2008).
        # Mappa forma + colore -> categoria di cartello.
        doc.add_heading("Forma e categorie dei cartelli", level=3)
        shape_rows = [
            ["Cartelli di divieto", "Tonda", "Pittogramma nero su fondo bianco; bordo e banda diagonale rossi", "Vietato fumare; vietato l'accesso ai non addetti; vietato spegnere con acqua"],
            ["Cartelli di avvertimento", "Triangolare", "Pittogramma nero su fondo giallo; bordo nero", "Pericolo elettrico; materiale infiammabile; carichi sospesi; rischio biologico"],
            ["Cartelli di prescrizione", "Tonda", "Pittogramma bianco su fondo azzurro", "Obbligo casco; obbligo guanti; obbligo occhiali; obbligo otoprotettori"],
            ["Cartelli di salvataggio", "Rettangolare o quadrata", "Pittogramma bianco su fondo verde", "Uscita di emergenza; direzione vie di fuga; punto di raccolta; defibrillatore"],
            ["Cartelli antincendio", "Rettangolare o quadrata", "Pittogramma bianco su fondo rosso", "Estintore; idrante; pulsante di allarme; scala antincendio"],
            ["Cartelli supplementari", "Rettangolare", "Testo nero su fondo bianco o coordinato al cartello principale", "Distanza, altezza, indicazioni testuali integrative"],
        ]
        self._add_data_table(
            doc,
            headers=["Categoria", "Forma", "Colori e caratteristiche", "Esempi"],
            rows=shape_rows,
        )
        doc.add_paragraph("")

    def _add_formazione_programma_section(
        self, doc: Document, persone: list
    ) -> None:
        """Programma di Formazione, Informazione ed Addestramento (template
        §3.N+4). Driven by mansione/ruolo so the operator gets a starter
        grid covering the Accordo Stato-Regioni 21/12/2011 obligations.
        """
        doc.add_heading(
            "Programma di Formazione, Informazione ed Addestramento", level=2
        )

        rows: list[list[str]] = []

        # Generic worker formation — applies to every mansione.
        rows.append([
            "TUTTI I LAVORATORI",
            "Formazione generale",
            "4 ore",
            "Concetti generali in tema di prevenzione e sicurezza sul lavoro",
            "All'assunzione + aggiornamento quinquennale",
        ])

        # Risk-class-specific formation — durata + contenuti differentiated
        # by mansione family (operai industriali, ufficio, ristorazione,
        # logistica) so the operator gets a meaningful starter grid instead
        # of identical rows (audit F-008).
        seen_mansioni: set[str] = set()
        for p in persone:
            nome = (p.mansione or "").strip()
            if not nome or nome in seen_mansioni:
                continue
            seen_mansioni.add(nome)
            family = _classify_mansione_family(nome)
            durata, contenuti = _FORMAZIONE_BY_FAMILY[family]
            rows.append([
                nome.upper(),
                "Formazione specifica",
                durata,
                contenuti,
                "All'assunzione + aggiornamento quinquennale",
            ])

        # Role-specific formation
        if any(p.ruolo_rspp for p in persone):
            rows.append([
                "RSPP",
                "Modulo A + B + C",
                "Min. 100 ore",
                "Modulo A (28h), Modulo B (variabile per ATECO), Modulo C (24h)",
                "Aggiornamento 40 ore / 5 anni",
            ])
        if any(p.ruolo_rls for p in persone):
            rows.append([
                "RLS",
                "Formazione iniziale + aggiornamento",
                "32 ore + 8 h/anno",
                "Diritti e doveri del RLS, normativa, valutazione rischi",
                "Annuale (8 ore per >50 lavoratori, 4 ore per 15-50)",
            ])
        if any(p.ruolo_preposto for p in persone):
            rows.append([
                "PREPOSTO",
                "Formazione aggiuntiva",
                "8 ore",
                "Soggetti del SPP, vigilanza sull'applicazione delle misure",
                "Aggiornamento biennale (DL 146/2021)",
            ])
        if any(p.ruolo_primo_soccorso for p in persone):
            rows.append([
                "ADDETTI PRIMO SOCCORSO",
                "Formazione D.M. 388/2003",
                "12 o 16 ore (gruppo A/B/C)",
                "BLS, gestione emergenze sanitarie, contenuti cassetta P.S.",
                "Aggiornamento triennale",
            ])
        if any(p.ruolo_antincendio for p in persone):
            rows.append([
                "ADDETTI ANTINCENDIO",
                "Formazione D.M. 02/09/2021",
                "4 / 8 / 16 ore (livello 1/2/3)",
                "Misure preventive, attrezzature antincendio, esercitazioni",
                "Aggiornamento quinquennale",
            ])
        if any(p.attrezzature_speciali for p in persone):
            rows.append([
                "ADDETTI ATTREZZATURE SPECIFICHE",
                "Abilitazione (Acc. SR 22/02/2012)",
                "8-16 ore + pratica",
                "Carrelli elevatori, PLE, gru, ruspe, ADR, lavori in quota",
                "Aggiornamento quinquennale",
            ])

        self._add_data_table(
            doc,
            headers=["Destinatari", "Tipologia", "Durata", "Contenuti", "Cadenza"],
            rows=rows,
        )
        doc.add_paragraph("")

    def _add_env_attrezzature_table(
        self, doc: Document, attrezzature: list
    ) -> None:
        """Per-ambiente equipment table (Phase 8.2)."""
        headers = [
            "Macchine, Attrezzature ed Impianti",
            "Marcata CE",
            "Verifiche Periodiche",
        ]
        if not attrezzature:
            rows = [["Nessuna attrezzatura associata a questo ambiente.", "—", "—"]]
        else:
            rows = [
                [
                    (a.descrizione or "—").upper(),
                    "SI" if a.marcatura_ce else "NO",
                    "SI" if a.verifiche_periodiche else "NO",
                ]
                for a in attrezzature
            ]
        self._add_data_table(doc, headers, rows)

    def _add_env_identity_table(
        self, doc: Document, ambiente, persone_by_id: dict
    ) -> None:
        """Template Table 24 — DYNAMIC key-value block for the environment.

        Preposto resolution order:
          1. ``ambiente.preposto_id`` (explicit assignment)
          2. Any persona assigned to this ambiente with ``ruolo_preposto=True``
          3. Any azienda-level persona with ``ruolo_preposto=True``
          4. "Da nominare" — legally defensible placeholder per art. 19.
        """
        preposto_name: str | None = None
        preposto_id = getattr(ambiente, "preposto_id", None)
        if preposto_id and preposto_id in persone_by_id:
            preposto_name = persone_by_id[preposto_id].nominativo
        if not preposto_name:
            for pers in getattr(ambiente, "persone", None) or []:
                if getattr(pers, "ruolo_preposto", False):
                    preposto_name = pers.nominativo
                    break
        if not preposto_name:
            for pers in persone_by_id.values():
                if getattr(pers, "ruolo_preposto", False):
                    preposto_name = pers.nominativo
                    break
        preposto_name = (preposto_name or "Da nominare").upper()

        descrizione = (
            (ambiente.descrizione_attivita or "").strip()
            or (ambiente.tipo or "—")
        ).upper()

        rows: list[tuple[str, str]] = [
            ("Ambiente di lavoro", (ambiente.nome or "—").upper()),
            ("Tipologia", (ambiente.tipo or "—")),
        ]
        mq = getattr(ambiente, "superficie_mq", None)
        if mq:
            rows.append(("Metratura", f"{mq} mq"))
        rows.extend([
            ("Preposto per la sicurezza", preposto_name),
            ("Descrizione Attività", descrizione),
        ])
        self._add_key_value_table(doc, rows)

    def _add_env_addetti_table(self, doc: Document, ambiente) -> None:
        """Template Table 25 — Nominativo / Mansione for addetti assigned to this env.

        Always emits the table shell so the per-env layout matches the
        template even when no persone_ambienti mapping exists yet; a single
        placeholder row signals the missing assignment to the operator.
        """
        addetti = list(getattr(ambiente, "persone", []) or [])
        if addetti:
            rows = [
                [(a.nominativo or "—").upper(), (a.mansione or "—").upper()]
                for a in addetti
            ]
        else:
            rows = [["—", "—"]]
        self._add_data_table(doc, headers=["Nominativo Addetti", "Mansione"], rows=rows)

    def _add_env_risk_checklist(self, doc: Document, ambiente) -> None:
        """Template Table 26 — 14-row SI/NO checklist for the 11 risk categories.

        Row layout: macro-area label row, then its categories with SI/NO
        derived from whether at least one applicable valutazione_rischio
        exists with that categoria in this ambiente.
        """
        # Normalize short DB names ("Elettrici") to canonical long names
        # ("Impianti Elettrici") so the lookup against _CATEGORY_ORDER keys
        # actually matches. Without this every row silently shows NO.
        applicable_by_category = {
            normalize_categoria_to_long(r.categoria_rischio): True
            for r in ambiente.valutazioni_rischio
            if getattr(r, "applicabile", False)
        }

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header_row = table.rows[0]
        for i, text in enumerate(["Categoria di Rischio", "Applicabile"]):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _HEADER_TEXT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(cell, _HEADER_BG)

        current_macro = None
        for macro, categoria in _CATEGORY_ORDER:
            if macro != current_macro:
                macro_row = table.add_row()
                merged = macro_row.cells[0].merge(macro_row.cells[1])
                merged.text = ""
                p = merged.paragraphs[0]
                run = p.add_run(macro)
                run.bold = True
                run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._set_cell_bg(merged, _LIGHT_GRAY)
                current_macro = macro

            row = table.add_row()
            cell_label = row.cells[0]
            cell_label.text = ""
            p = cell_label.paragraphs[0]
            run = p.add_run(categoria)
            run.font.size = Pt(9)

            cell_flag = row.cells[1]
            cell_flag.text = ""
            p = cell_flag.paragraphs[0]
            flag = "SI" if applicable_by_category.get(categoria) else "NO"
            run = p.add_run(flag)
            run.font.size = Pt(9)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_env_risk_tables(
        self, doc: Document, ambiente, vdt_count: int = 0
    ) -> None:
        """Template Tables 27+ — one 5-col (PERICOLO/CONDIZIONI/RISCHIO/MISURE/I)
        table per applicable macro-category, emitted in the canonical order.

        ``vdt_count``: when > 0, the Agenti Fisici table is force-emitted
        (even if no other agente fisico is applicable) and a synthetic VDT
        pericolo row is appended that references art. 174 D.Lgs. 81/2008
        and the Allegato VDT (audit F-007).
        """
        # Normalize short DB names to canonical long names so the per-category
        # tables actually emit (see _add_env_risk_checklist comment).
        by_category: dict[str, list] = {}
        for r in ambiente.valutazioni_rischio:
            if not getattr(r, "applicabile", False):
                continue
            key = normalize_categoria_to_long(r.categoria_rischio)
            if not key:
                continue
            by_category.setdefault(key, []).append(r)

        # Force the Agenti Fisici section when VDT addetti are present in
        # the ambiente — even if no other agente fisico has been ticked.
        if vdt_count > 0 and "Agenti Fisici" not in by_category:
            by_category["Agenti Fisici"] = []

        ordered_keys = [cat for _, cat in _CATEGORY_ORDER if cat in by_category]
        trailing = [k for k in by_category.keys() if k not in ordered_keys]

        if not ordered_keys and not trailing:
            p = doc.add_paragraph(
                "Nessun rischio applicabile identificato per questo ambiente."
            )
            p.runs[0].font.italic = True
            return

        for cat_name in ordered_keys + trailing:
            inject_vdt = cat_name == "Agenti Fisici" and vdt_count > 0
            self._add_single_category_risk_table(
                doc,
                cat_name,
                by_category[cat_name],
                inject_vdt=inject_vdt,
                vdt_count=vdt_count,
            )
            doc.add_paragraph("")

    def _add_single_category_risk_table(
        self,
        doc: Document,
        categoria: str,
        risks: list,
        inject_vdt: bool = False,
        vdt_count: int = 0,
    ) -> None:
        """5-col risk table for a single category (Template Tables 27–33 shape).

        When ``inject_vdt`` is True a synthetic videoterminali pericolo row
        is appended at the end. The synthetic row references art. 174
        D.Lgs. 81/2008 and the dedicated Allegato VDT and is scored
        P=2 D=2 I=6 MODESTO by default — operators tune the score in the
        Allegato VDT and the next regeneration picks up the real values.
        """
        p = doc.add_paragraph()
        run = p.add_run(categoria.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _HEADER_BG

        headers = [
            "PERICOLO",
            "CONDIZIONI DI IMPIEGO O DI ESPOSIZIONE",
            "RISCHIO",
            "MISURE DI PREVENZIONE E PROTEZIONE ATTUATE E DPI",
            "I = 2*D + P",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            hp = cell.paragraphs[0]
            hrun = hp.add_run(text)
            hrun.bold = True
            hrun.font.size = Pt(8)
            hrun.font.color.rgb = _HEADER_TEXT
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(cell, _HEADER_BG)

        # Phase 3 (1:N): when the parent valutazione_rischio has children
        # in pericoli_valutazione, emit one row per child — that's the
        # template-faithful layout. When no children exist (legacy data),
        # fall back to the parent's single pericolo/condizioni/misure
        # block so older DVRs still render.
        rows_to_emit: list = []
        for risk in risks:
            children = [
                c for c in (getattr(risk, "pericoli", []) or [])
                if getattr(c, "applicabile", True)
            ]
            if children:
                rows_to_emit.extend(children)
            else:
                rows_to_emit.append(risk)

        # Synthetic VDT row when none of the existing children already
        # cover videoterminali — keeps the row dedupe-safe.
        if inject_vdt:
            already_covers_vdt = any(
                "videoterminal" in (getattr(r, "pericolo", "") or "").lower()
                for r in rows_to_emit
            )
            if not already_covers_vdt:
                from types import SimpleNamespace
                rows_to_emit.append(SimpleNamespace(
                    pericolo="Esposizione a videoterminali (VDT)",
                    condizioni_esposizione=(
                        f"{vdt_count} lavoratore/i con uso del VDT >= 20 "
                        "ore/settimana ai sensi dell'art. 173 D.Lgs. 81/08."
                    ),
                    rischio=(
                        "Affaticamento visivo, disturbi muscoloscheletrici, "
                        "stress correlato all'uso prolungato del VDT."
                    ),
                    misure_prevenzione=(
                        "Postazioni conformi all'allegato XXXIV; pause di "
                        "15 min ogni 120 min di applicazione continuativa; "
                        "sorveglianza sanitaria mirata. Vedi Allegato "
                        "Videoterminali (VDT)."
                    ),
                    probabilita_p=2,
                    danno_d=2,
                    valutazione_riferimento=None,
                ))

        for source in rows_to_emit:
            p_val = source.probabilita_p
            d_val = source.danno_d
            riferimento = getattr(source, "valutazione_riferimento", None)
            if p_val is not None and d_val is not None:
                result = calculate_risk_index(p_val, d_val)
                indice = result["indice_i"]
                livello = result["livello_rischio"]
                indice_text = f"P = {p_val}; D = {d_val}; I = {indice}; {livello}"
            elif riferimento:
                livello = None
                indice_text = riferimento
            else:
                livello = None
                indice_text = "—"

            row = table.add_row()
            values = [
                source.pericolo or "—",
                source.condizioni_esposizione or "—",
                source.rischio or "—",
                source.misure_prevenzione or "—",
                indice_text,
            ]
            for i, text in enumerate(values):
                cell = row.cells[i]
                cell.text = ""
                cp = cell.paragraphs[0]
                crun = cp.add_run(text)
                crun.font.size = Pt(8)
                if i == 4:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if livello and livello in _RISK_COLORS:
                        self._set_cell_bg(cell, _RISK_COLORS[livello])
                        crun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        crun.bold = True

        widths = [Cm(3.5), Cm(4.0), Cm(3.0), Cm(4.5), Cm(3.5)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    # ------------------------------------------------------------------
    # Part IV — Improvement measures (Template Tables 108, 109, 110)
    # ------------------------------------------------------------------

    def _add_part_iv(
        self, doc: Document, azienda, persone: list, extras: dict
    ) -> None:
        """Parte IV — programma di miglioramento, procedure operative §4.2–4.13,
        cross-references agli allegati applicabili, Dichiarazione del Datore
        di Lavoro, blocco firme.

        Structure mirrors the master template (DVR_TEMPLATE_MAPPING.md §10):
          §4.1  Programma di Miglioramento (Table 108 + Table 109 — real rows)
          §4.2  Gestione Leggi e Regolamenti
          §4.3  Gestione Sorveglianza Sanitaria
          §4.4  Gestione Informazione, Formazione ed Addestramento
          §4.5  Riunione Periodica
          §4.6  Gestione degli Infortuni
          §4.7  Gestione Comportamenti Scorretti
          §4.8  Gestione DPI
          §4.9  Gestione Infrastrutture
          §4.10 Lavoratori particolarmente sensibili al rischio
          §4.11 Gestione Acquisti
          §4.12 Gestione delle lavorazioni affidate in appalto
          (Documenti correlati al presente DVR — list of applicable allegati)
          §4.13 Dichiarazione del Datore di Lavoro
          (signature grid Table 110)
        """
        h = doc.add_heading("PARTE IV — PROGRAMMA DI MIGLIORAMENTO", level=1)
        h.paragraph_format.page_break_before = True

        # Template Table 108 — Azienda header
        self._add_azienda_header_table(doc, azienda)
        doc.add_paragraph("")

        doc.add_heading(
            "4.1 Programma e Procedure di attuazione delle Misure di Miglioramento",
            level=2,
        )
        p = doc.add_paragraph()
        run = p.add_run(
            "Il programma di miglioramento e definito sulla base delle "
            "criticita emerse dalla valutazione dei rischi. Le misure sono "
            "ordinate per priorita in funzione del livello di rischio: "
            "rischi GRAVISSIMI (I >= 9) richiedono intervento immediato, "
            "rischi GRAVI (I = 7-8) entro sei mesi, rischi MODESTI o "
            "ACCETTABILI sono pianificati nel ciclo annuale di revisione."
        )
        run.font.size = Pt(10)
        doc.add_paragraph("")

        # Template Table 109 — Misure di miglioramento (5-col grid)
        self._add_improvement_program_table(
            doc, extras.get("misure_miglioramento") or []
        )
        doc.add_paragraph("")

        # §4.2–4.12 — static procedural sections. After §4.3 we inject a
        # mansione-level Sorveglianza Sanitaria protocol table aggregated
        # from per-persona DPI + rischi flags (audit recommendation,
        # post-fix 2026-04-29; persona-level data 2026-04-30).
        for spec in _PART_IV_PROCEDURAL_SECTIONS:
            self._add_procedural_section(doc, spec)
            if spec.heading.startswith("4.3"):
                self._add_sorveglianza_protocol_table(doc, persone)

        # Cross-reference applicable allegati by name (audit F-016).
        self._add_allegati_cross_references(
            doc, extras.get("allegati_presenti") or []
        )

        # §4.13 — Dichiarazione del Datore di Lavoro (MIXED — names + date)
        self._add_dichiarazione_ddl(doc, azienda, persone)

        # Template Table 110 — Signature block (2×3)
        self._add_signature_table(doc, persone)

    def _add_sorveglianza_protocol_table(
        self, doc: Document, persone: list
    ) -> None:
        """Render the per-mansione Sorveglianza Sanitaria protocol table,
        aggregating per-persona DPI + rischi flags up to mansione (union).

        4 columns: Mansione | Rischi specifici | DPI assegnati | Periodicita'
        (compilata dal Medico Competente). Periodicita' is left as a
        placeholder because the actual cadence is a medical decision the
        MC must make per mansione (art. 41 D.Lgs. 81/2008) — we render the
        scaffold so the MC has a structured form to fill, not invented
        values that look authoritative.

        Falls back to a transparent "no data" paragraph when no persona
        carries any DPI/rischi flags for the azienda.
        """
        from app.services.reference_data import (
            DPI_CATALOG,
            RISCHI_SPECIFICI_CATALOG,
        )

        doc.add_heading(
            "Protocollo di Sorveglianza Sanitaria per Mansione",
            level=3,
        )

        # Aggregate per-persona codes up to mansione (union).
        mansione_dpi: dict[str, set[str]] = {}
        mansione_rischi: dict[str, set[str]] = {}
        for p in persone:
            nome = (p.mansione or "").strip()
            if not nome:
                continue
            dpi_codes = list(getattr(p, "dpi_codes", None) or [])
            rs_codes = list(getattr(p, "rischi_specifici_codes", None) or [])
            if dpi_codes:
                mansione_dpi.setdefault(nome, set()).update(dpi_codes)
            if rs_codes:
                mansione_rischi.setdefault(nome, set()).update(rs_codes)

        all_mansioni = sorted(set(mansione_dpi.keys()) | set(mansione_rischi.keys()))

        if not all_mansioni:
            p = doc.add_paragraph()
            run = p.add_run(
                "Nessuna mansione con sorveglianza sanitaria configurata. "
                "Il Medico Competente, in fase di nomina, definira il "
                "protocollo specifico per le mansioni esposte ai rischi "
                "previsti dalla normativa."
            )
            run.font.size = Pt(10)
            run.font.italic = True
            doc.add_paragraph("")
            return

        headers = [
            "Mansione",
            "Rischi specifici",
            "DPI assegnati",
            "Periodicita' (compilata dal MC)",
        ]

        rows: list[list[str]] = []
        for nome in all_mansioni:
            rischi_labels = sorted(
                RISCHI_SPECIFICI_CATALOG.get(c, {}).get("etichetta", c)
                for c in mansione_rischi.get(nome, set())
            )
            dpi_labels = sorted(
                DPI_CATALOG.get(c, {}).get("etichetta", c)
                for c in mansione_dpi.get(nome, set())
            )
            rows.append([
                nome.upper(),
                "; ".join(rischi_labels) if rischi_labels else "—",
                "; ".join(dpi_labels) if dpi_labels else "—",
                "[DA COMPILARE — MC]",
            ])

        self._add_data_table(doc, headers, rows)
        # Footnote clarifying authority for the periodicita' column.
        note = doc.add_paragraph()
        nrun = note.add_run(
            "Nota: la periodicita' della sorveglianza sanitaria e i "
            "protocolli specifici sono definiti dal Medico Competente in "
            "sede di nomina e di prima visita, ai sensi dell'art. 41 "
            "D.Lgs. 81/2008."
        )
        nrun.font.size = Pt(9)
        nrun.font.italic = True
        nrun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph("")

    def _add_allegati_cross_references(
        self, doc: Document, allegati: list[tuple[str, str]]
    ) -> None:
        """Emit a "Documenti correlati al presente DVR" section that names
        each applicable allegato verbatim — fixes audit F-016 / H16.

        The allegati list is computed in ``_load_dvr_extras`` from the actual
        assessment rows on this azienda, so the section only references
        documents that genuinely exist for this client.
        """
        doc.add_heading("Documenti correlati al presente DVR", level=2)

        if not allegati:
            p = doc.add_paragraph()
            run = p.add_run(
                "Nessun allegato di approfondimento applicabile sulla base "
                "delle valutazioni effettuate. La presente valutazione "
                "generale copre la totalita dei rischi rilevati."
            )
            run.font.size = Pt(10)
            run.font.italic = True
            doc.add_paragraph("")
            return

        p = doc.add_paragraph()
        run = p.add_run(
            "Costituiscono parte integrante del presente DVR i seguenti "
            "allegati di approfondimento, redatti secondo le metodologie "
            "previste dalla normativa di riferimento:"
        )
        run.font.size = Pt(10)

        for _slug, label in allegati:
            bullet = doc.add_paragraph(style="List Bullet")
            run = bullet.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
        doc.add_paragraph("")

    def _add_procedural_section(
        self, doc: Document, spec: "_ProceduralSection"
    ) -> None:
        """Emit one §4.x procedural section: heading, intro paragraph(s),
        optional bullet list, optional documentazione collegata footer.
        """
        doc.add_heading(spec.heading, level=2)
        for para in spec.body:
            p = doc.add_paragraph()
            run = p.add_run(para)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        if spec.bullets:
            for bullet in spec.bullets:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(bullet)
                run.font.size = Pt(10)
        if spec.docs:
            p = doc.add_paragraph()
            run = p.add_run("Documentazione collegata: ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.italic = True
            run2 = p.add_run("; ".join(spec.docs) + ".")
            run2.font.size = Pt(9)
            run2.font.italic = True
        doc.add_paragraph("")

    def _add_dichiarazione_ddl(
        self, doc: Document, azienda, persone: list
    ) -> None:
        """§4.13 — Binding declaration signed by the Datore di Lavoro.

        Names the actual DdL and ragione sociale and stamps luogo + data
        of redazione, mirroring the template's #2429 paragraph block.
        """
        doc.add_heading("4.13 Dichiarazione del Datore di Lavoro", level=2)

        ddl = next((p for p in persone if p.ruolo_datore_lavoro), None)
        ddl_name = (ddl.nominativo if ddl and ddl.nominativo else "________________________").upper()
        rs = (azienda.ragione_sociale or "________________________").upper()

        p = doc.add_paragraph()
        run = p.add_run(
            f"Il/la sottoscritto/a, {ddl_name}, in qualita di Datore di "
            f"Lavoro della {rs}, ai sensi e per gli effetti dell'art. 28 "
            "del D.Lgs. 81/2008 e s.m.i.,"
        )
        run.font.size = Pt(10)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sub.add_run("DICHIARA")
        srun.bold = True
        srun.font.size = Pt(11)

        for clause in (
            "che il procedimento di valutazione dei rischi ex art. 17, "
            "comma 1, lettera a) del D.Lgs. 81/2008 e s.m.i. e stato "
            "effettuato in collaborazione con il RSPP, il Medico "
            "Competente (ove nominato) e previa consultazione del RLS;",
            "che la valutazione e tutti i suoi aggiornamenti sono "
            "documentati nel presente DVR, conservato presso l'unita "
            "produttiva e reso disponibile per la consultazione ai "
            "soggetti previsti dalla legge;",
            "di aver individuato e descritto le misure di prevenzione e "
            "protezione attuate e di aver definito il programma delle "
            "misure di miglioramento di cui al §4.1;",
            "di impegnarsi a rielaborare la valutazione nei casi previsti "
            "dall'art. 29, comma 3 del D.Lgs. 81/2008 e s.m.i. "
            "(modifiche significative del processo produttivo, "
            "infortuni, evoluzione normativa, esiti della sorveglianza "
            "sanitaria).",
        ):
            cp = doc.add_paragraph()
            run = cp.add_run(clause)
            run.font.size = Pt(10)
            cp.paragraph_format.space_after = Pt(4)

        # Luogo, data — derived from sede operativa (or legale) + today
        citta = (
            azienda.sede_operativa_citta
            or azienda.sede_legale_citta
            or "_____________"
        )
        oggi = datetime.now(timezone.utc).strftime("%d/%m/%Y")
        luogo_data = doc.add_paragraph()
        luogo_data.paragraph_format.space_before = Pt(12)
        run = luogo_data.add_run(f"{citta}, li {oggi}")
        run.font.size = Pt(10)
        run.bold = True
        doc.add_paragraph("")

    def _add_improvement_program_table(
        self, doc: Document, misure: list
    ) -> None:
        """Template Table 109 — 5-col measures grid populated from the
        ``misure_miglioramento`` table.

        When ``misure`` is empty the operator has actively cleared all
        rows — we still emit a single italic placeholder row so the table
        renders, but flag it as not legally complete.
        """
        headers = [
            "Misure di miglioramento",
            "Procedure per l'attuazione delle misure di miglioramento",
            "Risorse necessarie per l'attuazione",
            "Responsabile",
            "Tempi di attuazione",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            hp = cell.paragraphs[0]
            hrun = hp.add_run(text)
            hrun.bold = True
            hrun.font.size = Pt(9)
            hrun.font.color.rgb = _HEADER_TEXT
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(cell, _HEADER_BG)

        if not misure:
            row = table.add_row()
            for i, text in enumerate([
                "[Nessuna misura inserita]",
                "—",
                "—",
                "—",
                "—",
            ]):
                c = row.cells[i]
                c.text = ""
                cp = c.paragraphs[0]
                crun = cp.add_run(text)
                crun.font.size = Pt(9)
                crun.font.italic = True
                crun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        for m in misure:
            row = table.add_row()
            values = [
                (m.misura or "—"),
                (m.procedura or "—"),
                (m.risorse or "—"),
                (m.responsabile or "—"),
                (m.scadenza or "—"),
            ]
            priorita = (m.priorita or "").upper()
            for i, text in enumerate(values):
                c = row.cells[i]
                c.text = ""
                cp = c.paragraphs[0]
                crun = cp.add_run(str(text))
                crun.font.size = Pt(9)
                # Color-band the misura cell when priorita maps to a
                # known livello — mirrors the risk-table convention.
                if i == 0 and priorita in _RISK_COLORS:
                    self._set_cell_bg(c, _RISK_COLORS[priorita])
                    crun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    crun.bold = True

    def _add_signature_table(self, doc: Document, persone: list) -> None:
        """Template Table 110 — 2×3 signature grid with DL / RSPP / Medico
        (row 1) and RLS (row 2, merged center cell)."""
        def _first(pred) -> str:
            match = next((p for p in persone if pred(p)), None)
            return (match.nominativo if match else "").upper()

        dl = _first(lambda p: p.ruolo_datore_lavoro) or "—"
        rspp = _first(lambda p: p.ruolo_rspp) or "—"
        medico = _first(lambda p: getattr(p, "ruolo_medico_competente", False)) or "—"
        rls = _first(lambda p: p.ruolo_rls) or "—"

        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        def _fill(cell, title_line: str, name_line: str) -> None:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(title_line)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(name_line)
            run2.font.size = Pt(9)
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run("___________________________")
            run3.font.size = Pt(9)

        _fill(table.rows[0].cells[0], "Il Datore di Lavoro", f"({dl})")
        _fill(table.rows[0].cells[1], "", "")
        _fill(table.rows[0].cells[2], "Il Responsabile del S.P.P.", f"({rspp})")

        _fill(table.rows[1].cells[0], "Il Medico Competente", f"({medico})")
        _fill(table.rows[1].cells[1], "", "")
        _fill(
            table.rows[1].cells[2],
            "Per consultazione\nIl Rappresentante dei Lavoratori",
            f"({rls})",
        )

    # ------------------------------------------------------------------
    # Helper methods
    # ------------------------------------------------------------------

    @staticmethod
    def _split_descrizione_paragraphs(text: str) -> list[str]:
        """Split a free-form descrizione_attivita into 1-N paragraphs.

        Splits on (a) blank-line separators if present, otherwise (b) into
        chunks of ~3 sentences each so a long single-paragraph string
        becomes a paragraph-formatted narrative. Empty inputs return [].
        """
        text = (text or "").strip()
        if not text:
            return []
        # Honor explicit blank-line splits first.
        blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
        if len(blocks) > 1:
            return blocks
        # Otherwise group sentences into chunks of 3.
        sentences = re.split(r"(?<=[\.!\?])\s+", text)
        chunks: list[str] = []
        buf: list[str] = []
        for s in sentences:
            buf.append(s)
            if len(buf) >= 3:
                chunks.append(" ".join(buf).strip())
                buf = []
        if buf:
            chunks.append(" ".join(buf).strip())
        return chunks or [text]

    def _format_address(self, via: str | None, citta: str | None) -> str:
        """Format an address from its components."""
        parts = [p for p in [via, citta] if p]
        return ", ".join(parts) if parts else "—"

    def _add_key_value_table(
        self, doc: Document, rows: list[tuple[str, str]]
    ) -> None:
        """Add a simple two-column key-value table.

        Each row gets `cantSplit` so a key/value pair never breaks across
        a page boundary — readability win for anagrafica / identity
        blocks where the label dangling alone on the previous page was
        the most jarring case in operator feedback (2026-05-12).
        """
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"

        for i, (key, value) in enumerate(rows):
            # Key cell
            cell_key = table.rows[i].cells[0]
            cell_key.text = ""
            p = cell_key.paragraphs[0]
            run = p.add_run(key)
            run.bold = True
            run.font.size = Pt(9)
            cell_key.width = Cm(5)

            # Value cell
            cell_val = table.rows[i].cells[1]
            cell_val.text = ""
            p = cell_val.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(9)
            cell_val.width = Cm(12)

            # Alternating row colors
            if i % 2 == 0:
                self._set_cell_bg(cell_key, _LIGHT_GRAY)
                self._set_cell_bg(cell_val, _LIGHT_GRAY)

            # Don't split this row across pages.
            self._set_row_cant_split(table.rows[i])

    def _add_data_table(
        self, doc: Document, headers: list[str], rows: list[list[str]]
    ) -> None:
        """Add a multi-column data table with styled header.

        Pagination behaviour:
          - The header row carries `tblHeader` (so Word repeats it on
            every page when the table breaks) plus `cantSplit` so the
            header itself never dangles split.
          - Data rows get `cantSplit` so a single row never breaks
            across pages. Whole tables are still allowed to span pages —
            blanket-locking them would push huge attrezzatura/pericoli
            tables onto a fresh page and waste vertical space.
        """
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Header row
        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _HEADER_TEXT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_bg(cell, _HEADER_BG)

        self._set_row_repeat_as_header(header_row)
        self._set_row_cant_split(header_row)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.add_row()
            for i, text in enumerate(row_data):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9)

                # Center-align the first column (row number)
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Alternating row colors
            if row_idx % 2 == 0:
                for cell in row.cells:
                    self._set_cell_bg(cell, _LIGHT_GRAY)

            self._set_row_cant_split(row)

    @staticmethod
    def _set_cell_bg(cell, color: RGBColor) -> None:
        """Set a table cell's background (shading) color.

        Uses the low-level XML API since python-docx does not expose
        cell shading directly.
        """
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.find(qn("w:shd"))
        if shading is None:
            shading = shading_elm.makeelement(qn("w:shd"), {})
            shading_elm.append(shading)
        shading.set(qn("w:fill"), f"{color}")
        shading.set(qn("w:val"), "clear")

    @staticmethod
    def _set_row_cant_split(row) -> None:
        """Mark a table row so Word never splits it across pages.

        Emits `<w:trPr><w:cantSplit/></w:trPr>` on the row's <w:tr>.
        Idempotent — re-applying just leaves the existing flag in place.
        """
        from docx.oxml import OxmlElement

        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _set_row_repeat_as_header(row) -> None:
        """Mark a row as a repeating table header.

        When Word reflows the table across pages it re-emits this row at
        the top of each subsequent page. Combined with `cantSplit` it's
        the standard pattern for long data tables.
        """
        from docx.oxml import OxmlElement

        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))
