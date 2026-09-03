"""Shared helpers for document generators.

Keeps each generator small and focused by extracting common docx
manipulation patterns (placeholder substitution, table helpers,
heading styling, color palettes, etc.).
"""

from __future__ import annotations

import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Design tokens — one palette for all 17 documents, mirroring the web app
# (frontend/src/app/globals.css) so product and paperwork read as one family.
# Document audit 2026-09-03: generators had drifted to Material indigo #1A237E
# for headers and stock Word blue for headings; the app is N2O navy #003d74.
# Print wants firmer rules than the screen border token, hence BRAND_RULE.
# ---------------------------------------------------------------------------

BRAND_NAVY = RGBColor(0x00, 0x3D, 0x74)      # --color-primary: headings, table headers
BRAND_NAVY_HEX = "003D74"
BRAND_DEEP = RGBColor(0x06, 0x1B, 0x31)      # --color-heading: body ink
BRAND_SLATE = RGBColor(0x64, 0x74, 0x8D)     # --color-body: captions, secondary lines
BRAND_SURFACE_HEX = "F6F9FC"                 # --color-surface-low: zebra rows, label cells
BRAND_RULE_HEX = "C9D3DF"                    # hairlines and table borders (print-firm)
FONT_FAMILY = "Calibri"                      # universal in Word; Carlito on LibreOffice

# Type scale in points. Body 10.5 keeps ~13 words per line on A4 text width.
TYPE_SCALE = {
    "cover_title": 26,
    "cover_subtitle": 12,
    "cover_client": 18,
    "h1": 16,
    "h2": 13,
    "h3": 11,
    "body": 10.5,
    "table": 9.5,
    "small": 8.5,
}

# Risk scale = the app's risk chips (--color-risk-*), so a GRAVE cell in the
# document is the same orange as the GRAVE chip on screen.
RISK_COLORS = {
    "ACCETTABILE": RGBColor(0x15, 0xBE, 0x53),
    "MODESTO": RGBColor(0xF5, 0x9E, 0x0B),
    "GRAVE": RGBColor(0xF9, 0x73, 0x16),
    "GRAVISSIMO": RGBColor(0xEF, 0x44, 0x44),
    "BASSO": RGBColor(0x15, 0xBE, 0x53),
    "MEDIO": RGBColor(0xF5, 0x9E, 0x0B),
    "ALTO": RGBColor(0xEF, 0x44, 0x44),
    "VERDE": RGBColor(0x15, 0xBE, 0x53),
    "GIALLO": RGBColor(0xF5, 0x9E, 0x0B),
    "ROSSO": RGBColor(0xEF, 0x44, 0x44),
}


# ---------------------------------------------------------------------------
# Address formatting (shared across all generators)
# ---------------------------------------------------------------------------

def format_comune(cap, citta, provincia) -> str:
    """Build the Italian comune segment ``CAP Comune (PROV)``.

    Any component may be missing; a blank comune never yields an orphan
    ``(PROV)``. Returns ``"—"`` when nothing is available.
    """
    seg = (citta or "").strip()
    cap = (cap or "").strip()
    provincia = (provincia or "").strip()
    if cap and seg:
        seg = f"{cap} {seg}"
    elif cap:
        seg = cap
    if provincia and seg:
        seg = f"{seg} ({provincia})"
    return seg or "—"


def format_sede(azienda, which: str = "legale") -> str:
    """Format a full Italian seat address as ``Via, CAP Comune (PROV)``.

    ``which`` selects the field family: ``"legale"`` reads
    ``sede_legale_via`` / ``sede_legale_citta`` / ``cap_legale`` /
    ``provincia_legale``; ``"operativa"`` the operative equivalents.
    Audit F-301 (2026-05-31): generators previously emitted only
    ``via, comune``, silently dropping the CAP and province held on the row.
    """
    via = (getattr(azienda, f"sede_{which}_via", None) or "").strip()
    seg = format_comune(
        getattr(azienda, f"cap_{which}", None),
        getattr(azienda, f"sede_{which}_citta", None),
        getattr(azienda, f"provincia_{which}", None),
    )
    parts = [p for p in [via, seg if seg != "—" else ""] if p]
    return ", ".join(parts) if parts else "—"

# Legacy names kept for the generators that import them; they now resolve to
# the brand tokens above so no call site has to change to pick up the palette.
HEADER_BG = BRAND_NAVY
HEADER_BG_HEX = BRAND_NAVY_HEX
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF6, 0xF9, 0xFC)


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

BACKEND_ROOT = Path(__file__).resolve().parents[3]
REPO_ROOT = BACKEND_ROOT.parent


def _resolve_templates_dir() -> Path:
    # On Render the service root is `backend/`, so the sibling `templates/`
    # folder only exists if the buildCommand copied it in. Prefer that
    # in-backend copy; fall back to the repo-root location for local dev.
    for candidate in (BACKEND_ROOT / "templates", REPO_ROOT / "templates"):
        if candidate.is_dir():
            return candidate
    return REPO_ROOT / "templates"


TEMPLATES_DIR = _resolve_templates_dir()
LOGO_PATH = BACKEND_ROOT / "assets" / "logo.png"


def slugify(text: str, max_length: int = 40) -> str:
    """Produce a filesystem-safe slug from free text."""
    lowered = (text or "").lower()
    replaced = re.sub(r"[^a-z0-9]+", "_", lowered)
    collapsed = re.sub(r"_+", "_", replaced).strip("_")
    if not collapsed:
        collapsed = "azienda"
    return collapsed[:max_length].rstrip("_") or "azienda"


# ---------------------------------------------------------------------------
# Cell shading (table header background)
# ---------------------------------------------------------------------------

def insert_in_order(parent, child, successors: tuple[str, ...]) -> None:
    """Insert ``child`` into ``parent`` before the first existing element
    whose tag is in ``successors`` (schema order), else append. Word
    validates child order inside property elements; appending blindly can
    make it report the file as corrupt."""
    tags = {qn(f"w:{t}") for t in successors}
    anchor = next((el for el in parent if el.tag in tags), None)
    if anchor is not None:
        anchor.addprevious(child)
    else:
        parent.append(child)


def shade_cell(cell, color_hex: str) -> None:
    """Set background shading on a table cell via raw w:shd XML.

    Replaces any shading already on the cell: zebra striping is applied by
    :func:`add_data_table` first and a generator may then colour a risk cell,
    and the last call must win deterministically (stacked ``w:shd`` elements
    render differently in Word and LibreOffice).
    """
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    insert_in_order(tc_pr, shd, ("noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "headers", "cellIns", "cellDel", "cellMerge", "tcPrChange"))


def style_header_row(row, bg_hex: str = BRAND_NAVY_HEX, text_color: RGBColor = HEADER_TEXT) -> None:
    """Bold white text on the brand-navy header background; repeats on every
    page the table spans, so a long table never prints headless."""
    from docx.oxml import OxmlElement
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        insert_in_order(tr_pr, OxmlElement("w:tblHeader"), ("tblCellSpacing", "jc", "hidden", "ins", "del", "trPrChange"))
    for cell in row.cells:
        shade_cell(cell, bg_hex)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(TYPE_SCALE["table"])


# ---------------------------------------------------------------------------
# Placeholder replacement — visit every paragraph in the document (including
# paragraphs in tables) and do straight string substitution across runs.
# ---------------------------------------------------------------------------

def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace keys with values inside a single paragraph.

    Runs may split a placeholder across multiple parts, so first join and
    then redistribute. We rewrite the paragraph text in the first run and
    clear subsequent runs — this loses per-run formatting on replaced
    paragraphs, which is acceptable for simple placeholder fields.
    """
    text = paragraph.text
    changed = False
    for k, v in replacements.items():
        if k in text:
            text = text.replace(k, str(v))
            changed = True
    if changed:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)


def replace_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    """Walk every paragraph (body + tables + headers/footers) and replace."""
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                replace_in_paragraph(p, replacements)


def scrub_body(
    doc: Document,
    replacements: dict[str, str],
    *,
    drop_paragraph_markers: Iterable[str] = (),
) -> int:
    """Body-only scrub for legacy templates built from a real completed client
    document that still carry that *origin* company's identity in the body.

    Several attachment templates (Stress, Gestanti, …) were authored by filling
    a real assessment, so the donor company's name/address/declaration print as
    the *assessed* subject. The page header/footer carries the consultancy's own
    letterhead (intentional branding) — this scrub therefore touches ONLY the
    body (``doc.paragraphs`` + table cells), never ``section.header/footer``.

    - ``drop_paragraph_markers``: body paragraphs whose text contains any of
      these substrings (case-insensitive) are removed entirely — use for donor
      free-prose (e.g. a company self-description) that can't be safely
      string-substituted to the client.
    - ``replacements``: case-insensitive literal swaps applied to surviving body
      paragraphs + table cells — use for structured donor identity → client data.

    Returns the number of paragraphs dropped (for logging/verification).
    """
    dropped = 0
    if drop_paragraph_markers:
        markers = [m.lower() for m in drop_paragraph_markers]
        for p in list(doc.paragraphs):
            low = (p.text or "").lower()
            if any(m in low for m in markers):
                p._p.getparent().remove(p._p)
                dropped += 1

    compiled = [
        (re.compile(re.escape(k), re.IGNORECASE), str(v)) for k, v in replacements.items()
    ]

    def _scrub(paragraph) -> None:
        text = paragraph.text
        if not text:
            return
        new = text
        for rx, val in compiled:
            new = rx.sub(lambda _m, _v=val: _v, new)
        if new != text:
            if paragraph.runs:
                paragraph.runs[0].text = new
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(new)

    for p in doc.paragraphs:
        _scrub(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _scrub(p)
    return dropped


def scrub_n2o_legacy_donor(
    doc: Document, azienda, *, drop_paragraph_markers: Iterable[str] = ()
) -> int:
    """Body-only scrub of the N2O origin-company identity baked into the legacy
    attachment templates (Stress, Gestanti) that were authored from a real N2O
    self-assessment. Swaps the donor's name / declarant / seat addresses for the
    client's, leaving the consultancy letterhead (header/footer) intact.

    The donor constants below are the literal strings found in those two
    templates (verified 2026-05-31). ``drop_paragraph_markers`` is forwarded to
    :func:`scrub_body` for donor free-prose that can't be string-substituted.
    """
    rs = azienda.ragione_sociale or ""
    legale = format_comune(
        getattr(azienda, "cap_legale", None),
        getattr(azienda, "sede_legale_citta", None),
        getattr(azienda, "provincia_legale", None),
    )
    oper = format_comune(
        getattr(azienda, "cap_operativa", None),
        getattr(azienda, "sede_operativa_citta", None),
        getattr(azienda, "provincia_operativa", None),
    )
    legale = "" if legale == "—" else legale
    oper = "" if oper == "—" else oper
    via_legale = getattr(azienda, "sede_legale_via", None) or ""
    via_oper = getattr(azienda, "sede_operativa_via", None) or via_legale
    return scrub_body(
        doc,
        {
            "N2O SRL": rs,
            "N2O S.R.L.": rs,
            "CIARAMITARO AMALIA": "",
            "VIA DEI CHIOSI 4": via_legale,
            "VIA MONZA 107/30": via_oper,
            "GORGONZOLA (MI)": legale,
            "GESSATE (MI)": oper,
        },
        drop_paragraph_markers=drop_paragraph_markers,
    )


# ---------------------------------------------------------------------------
# Simple heading/paragraph helpers
# ---------------------------------------------------------------------------

def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading; fall back to Normal style if custom heading missing."""
    try:
        h = doc.add_heading(text, level=level)
    except Exception:
        h = doc.add_paragraph(text)
        for r in h.runs:
            r.bold = True
            r.font.size = Pt(14 if level == 1 else 12)
    return h


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_consultancy_letterhead(doc: Document, branding, *, center: bool = True) -> None:
    """Render the consultancy's letterhead text block from ``branding``.

    Prints the firm name plus whatever optional letterhead detail is present
    (address, P.IVA / C.F., contacts, RSPP). Renders nothing beyond the firm
    name when the org hasn't filled the rest in. ``branding`` is a
    :class:`~app.services.document_generator.branding.Branding`.
    """
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    gray = BRAND_SLATE

    name_p = doc.add_paragraph()
    name_p.alignment = align
    name_run = name_p.add_run((branding.firm_name or "").upper())
    name_run.bold = True
    name_run.font.size = Pt(TYPE_SCALE["body"])
    name_run.font.color.rgb = BRAND_NAVY

    lines: list[str] = []
    addr = branding.address_line()
    if addr:
        lines.append(addr)
    tax_bits = []
    if branding.partita_iva:
        tax_bits.append(f"P.IVA {branding.partita_iva}")
    if branding.codice_fiscale and branding.codice_fiscale != branding.partita_iva:
        tax_bits.append(f"C.F. {branding.codice_fiscale}")
    if tax_bits:
        lines.append(" · ".join(tax_bits))
    contact = branding.contact_line()
    if contact:
        lines.append(contact)
    if branding.rspp_nome:
        lines.append(f"RSPP: {branding.rspp_nome}")

    for text in lines:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = gray


def _safe_table_style(table) -> None:
    """Apply Table Grid if available, otherwise apply manual cell borders."""
    try:
        table.style = "Table Grid"
    except KeyError:
        _apply_cell_borders_all(table)


def _apply_cell_borders_all(table) -> None:
    """Draw thin black borders on every cell (fallback when 'Table Grid' absent)."""
    from docx.oxml import OxmlElement
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), "808080")
                tc_borders.append(b)


def set_table_borders(table, color_hex: str = BRAND_RULE_HEX, size: int = 4) -> None:
    """Thin borders in the brand rule colour on every edge, inside and out.

    Overrides the black ``Table Grid`` borders so tables sit quietly on the
    page instead of shouting; ``size`` is in eighths of a point.
    """
    from docx.oxml import OxmlElement
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    # Word validates child order inside tblPr: tblBorders must precede shd,
    # tblLayout, tblCellMar, tblLook. Insert before the first of those.
    successors = [qn(f"w:{t}") for t in ("shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription")]
    anchor = next((child for child in tbl_pr if child.tag in successors), None)
    if anchor is not None:
        anchor.addprevious(borders)
    else:
        tbl_pr.append(borders)


def keep_row_together(row) -> None:
    """Forbid a row from splitting across a page break."""
    from docx.oxml import OxmlElement
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        insert_in_order(tr_pr, OxmlElement("w:cantSplit"), ("trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden", "ins", "del", "trPrChange"))


def add_kv_table(doc: Document, rows: Iterable[tuple[str, str]], *, width_label_cm: float = 5.0, width_value_cm: float = 11.0) -> None:
    """2-column key/value table: bold label on a tinted cell, plain value."""
    table = doc.add_table(rows=0, cols=2)
    style_applied = _try_set_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(k)
        cells[1].text = str(v) if v is not None else ""
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(TYPE_SCALE["table"])
                r.font.color.rgb = BRAND_DEEP
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(TYPE_SCALE["table"])
        shade_cell(cells[0], BRAND_SURFACE_HEX)
    if not style_applied:
        _apply_cell_borders_all(table)
    set_table_borders(table)
    return table


def _try_set_table_style(table) -> bool:
    """Try Table Grid; return True if applied."""
    try:
        table.style = "Table Grid"
        return True
    except KeyError:
        return False


def add_data_table(
    doc: Document,
    headers: list[str],
    data_rows: list[list[str]],
    *,
    column_widths_cm: list[float] | None = None,
):
    """Table with styled header + rows."""
    if column_widths_cm is not None and len(column_widths_cm) != len(headers):
        raise ValueError("column_widths_cm must contain one width per header")

    table = doc.add_table(rows=1, cols=len(headers))
    style_applied = _try_set_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)

    # No cantSplit on data rows: a keep-together row taller than a page makes
    # LibreOffice loop forever (the DUVRI interferenze cells hit this), and
    # Word splits such rows regardless.
    for n, row_data in enumerate(data_rows):
        row = table.add_row()
        zebra = n % 2 == 1
        for i, cell_val in enumerate(row_data):
            row.cells[i].text = "" if cell_val is None else str(cell_val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(TYPE_SCALE["table"])
            if zebra:
                shade_cell(row.cells[i], BRAND_SURFACE_HEX)

    if column_widths_cm is not None:
        table.autofit = False
        for column, width_cm in zip(table.columns, column_widths_cm):
            width = Cm(width_cm)
            column.width = width
            for cell in column.cells:
                cell.width = width

    if not style_applied:
        _apply_cell_borders_all(table)
    set_table_borders(table)
    return table


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Donor-template surgery (audit 2026-09-03)
#
# Several attachments open a real completed client document and append the
# generated assessment after it. The donor half carried personal data — a
# staff roster with codici fiscali, scanned signatures, Street View photos —
# that no string substitution can remove. These helpers cut it out by
# structure, before the generator appends anything of its own.
# ---------------------------------------------------------------------------

def _body_children(doc: Document) -> list:
    return list(doc.element.body)


def _table_header_text(tbl_el, doc: Document) -> str:
    from docx.table import Table
    table = Table(tbl_el, doc)
    if not table.rows:
        return ""
    return " | ".join((c.text or "").strip() for c in table.rows[0].cells)


def _norm_text(text: str) -> str:
    """Case-, whitespace- and apostrophe-insensitive comparison key."""
    text = (text or "").replace(chr(0x2019), "'").replace(chr(0x2018), "'").replace(chr(0xA0), " ")
    return " ".join(text.split()).strip().lower()


def find_table_index(doc: Document, header_prefix: str, *, occurrence: int = 1) -> int | None:
    """Body index of the n-th table whose first row starts with ``header_prefix``."""
    seen = 0
    for i, el in enumerate(_body_children(doc)):
        if el.tag != qn("w:tbl"):
            continue
        if _norm_text(_table_header_text(el, doc)).startswith(_norm_text(header_prefix)):
            seen += 1
            if seen == occurrence:
                return i
    return None


def remove_body_after(doc: Document, index: int) -> int:
    """Delete every body element after ``index`` (the final ``sectPr`` is kept).

    Returns the number of elements removed. Use it to drop a donor document's
    tail — its results, company description, roster and signatures — while
    keeping the front matter and method chapters that are genuinely generic.
    """
    body = doc.element.body
    removed = 0
    for el in _body_children(doc)[index + 1 :]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)
        removed += 1
    return removed


def remove_tables_with_header(doc: Document, needles: Iterable[str]) -> int:
    """Remove body tables whose first row contains ALL of ``needles``."""
    wanted = [n.lower() for n in needles]
    body = doc.element.body
    removed = 0
    for el in _body_children(doc):
        if el.tag != qn("w:tbl"):
            continue
        head = _table_header_text(el, doc).lower()
        if all(n in head for n in wanted):
            body.remove(el)
            removed += 1
    return removed


def strip_body_images(doc: Document) -> int:
    """Remove every inline picture from the body (headers/footers untouched).

    Donor templates embed scanned signatures, stamps and photographs of the
    donor's premises; none of them belong in another company's document.
    Text boxes and plain shapes are kept — a VML text box is also a ``w:pict``
    and in the legacy templates it holds the cover title. Returns the number
    of pictures removed.
    """
    removed = 0
    for tag in ("w:drawing", "w:pict"):
        for el in list(doc.element.body.iter(qn(tag))):
            has_textbox = el.find(".//" + qn("w:txbxContent")) is not None
            has_picture = (
                el.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip") is not None
                or el.find(".//{urn:schemas-microsoft-com:vml}imagedata") is not None
            )
            if has_textbox or not has_picture:
                continue
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                removed += 1
    return removed


def reset_table_rows(doc: Document, header_prefix: str, rows: list[list[str]], *, occurrence: int = 1) -> bool:
    """Keep a table's header row, drop the rest, and write ``rows`` in its
    place — used to replace a donor's revision history with this emission."""
    from docx.table import Table
    idx = find_table_index(doc, header_prefix, occurrence=occurrence)
    if idx is None:
        return False
    table = Table(_body_children(doc)[idx], doc)
    body_rows = list(table.rows)[1:]
    # Write in place rather than deleting rows: donor tables merge cells
    # vertically, and deleting a row can orphan a vMerge continuation, which
    # Word tolerates but LibreOffice chokes on. Surplus rows are blanked, so
    # the template's form keeps its shape (six revision lines, five empty).
    for r_index, row in enumerate(body_rows):
        values = rows[r_index] if r_index < len(rows) else []
        cells = row.cells
        seen = []
        for c_index, cell in enumerate(cells):
            if any(cell._tc is s for s in seen):
                continue
            seen.append(cell._tc)
            value = values[c_index] if c_index < len(values) else ""
            _set_cell_text_keep_format(cell, "" if value is None else str(value))
    for values in rows[len(body_rows):]:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = "" if value is None else str(value)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(TYPE_SCALE["table"])
    return True


def _set_cell_text_keep_format(cell, text: str) -> None:
    """Write ``text`` into a cell, reusing the first run's formatting."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return
    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def fill_label_table(doc: Document, values: dict) -> int:
    """Fill empty value cells next to (or under) matching labels in body tables.

    A label matches when the cell text, stripped of a trailing colon, equals a
    key (case, whitespace and apostrophe variants ignored). Two donor layouts
    are handled:

    - ``label | value`` rows: the cell to the right is written if empty. When
      the label cell is vertically merged over several rows (the "Sede Legale"
      cells that hold street and comune on two lines), each row takes the next
      item of a list value, so pass ``["Via Roma 1", "20100 Milano (MI)"]``.
    - stacked rows: a one-cell label row followed by a one-cell empty row.

    Each label is filled in the first table where it appears and skipped in
    later ones: a DUVRI carries the same "Ragione Sociale / Sede Legale" form
    for the committente and again for the appaltatore, and only the first is
    the client. Values may be strings or lists of strings. Returns the number
    of cells filled.
    """
    wanted: dict[str, list[str]] = {}
    for k, v in values.items():
        items = [x for x in (v if isinstance(v, (list, tuple)) else [v]) if x]
        if items:
            wanted[_norm_text(k).rstrip(":").strip()] = [str(x) for x in items]
    if not wanted:
        return 0

    def norm(cell) -> str:
        return _norm_text(cell.text).rstrip(":").strip()

    filled = 0
    done: set[str] = set()

    class _Seen:
        # Vertically merged label cells return the same <w:tc> for every row
        # they span. Hold the element references so identity is stable —
        # python-docx/lxml recycle proxy objects, so id() alone lies.
        def __init__(self):
            self.entries: list[tuple[object, int]] = []

        def bump(self, el) -> int:
            for k, (seen_el, n) in enumerate(self.entries):
                if seen_el is el:
                    self.entries[k] = (seen_el, n + 1)
                    return n
            self.entries.append((el, 1))
            return 0

    for table in doc.tables:
        rows = list(table.rows)
        used = _Seen()
        touched: set[str] = set()
        for r_index, row in enumerate(rows):
            cells = row.cells
            # stacked layout: single-cell label row, single-cell empty row below
            if len(set(id(c._tc) for c in cells)) == 1:
                label = norm(cells[0])
                if label in wanted and label not in done and r_index + 1 < len(rows):
                    below = rows[r_index + 1].cells
                    if len(set(id(c._tc) for c in below)) == 1 and not (below[0].text or "").strip():
                        _set_cell_text_keep_format(below[0], wanted[label][0])
                        filled += 1
                        touched.add(label)
                continue
            for i in range(len(cells) - 1):
                label = norm(cells[i])
                if label not in wanted or label in done:
                    continue
                target = cells[i + 1]
                if target._tc is cells[i]._tc:
                    continue
                # Position of this row within the merged label: the first
                # row takes the street, the second the comune, even when a
                # donor scrub already wrote the first line.
                n = used.bump(cells[i]._tc)
                if (target.text or "").strip():
                    touched.add(label)
                    break
                items = wanted[label]
                if n >= len(items):
                    break
                _set_cell_text_keep_format(target, items[n])
                filled += 1
                touched.add(label)
                break
        done |= touched
    return filled


def recolor_white_text(doc: Document, to: RGBColor = BRAND_NAVY) -> int:
    """Turn white runs navy. Donor covers set white type over a background
    picture that is no longer there, so the title printed white on white."""
    changed = 0
    for p in doc.paragraphs:
        for r in p.runs:
            try:
                rgb = r.font.color.rgb
            except Exception:
                rgb = None
            if rgb is not None and str(rgb).upper() == "FFFFFF":
                r.font.color.rgb = to
                changed += 1
    return changed
