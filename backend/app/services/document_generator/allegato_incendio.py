"""Allegato Rischio Incendio - D.M. 03/09/2021."""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_incendio
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    page_break,
    replace_placeholders,
    slugify,
)

TEMPLATE = TEMPLATES_DIR / "ALLEGATO RISCHIO INCENDIO.docx"
TIPO_DOC = "allegato_incendio"

# Six prevention areas per REFERENCE_DATA.md §4.5. Each area has tiered
# measures keyed by the area's livello_rischio so an inspector sees a
# concrete plan, not generic boilerplate.
PREVENTION_CATEGORIES: list[tuple[str, str]] = [
    ("1. Ridurre la probabilita di insorgenza di un incendio",
     "Controllo periodico degli impianti elettrici, separazione materiali infiammabili, divieti di fumo, manutenzione apparecchiature."),
    ("2. Garantire l'esodo delle persone in sicurezza",
     "Vie di fuga sgombre e segnalate, porte tagliafuoco verificate, illuminazione di emergenza, punto di raccolta esterno."),
    ("3. Sistemi di allarme e segnalazione rapida",
     "Rilevatori di fumo/calore, pulsanti manuali di allarme, sirena acustica/luminosa, procedura di chiamata 115."),
    ("4. Estinzione dell'incendio",
     "Estintori (verifica semestrale), idranti UNI 45/70 dove richiesti, attrezzature di primo intervento, addetti formati."),
    ("5. Efficienza dei sistemi di protezione antincendio",
     "Manutenzione periodica registro antincendio, verifiche annuali estintori/idranti/rivelatori, controllo porte REI."),
    ("6. Informazione e formazione dei lavoratori",
     "Corso antincendio (basso/medio/alto livello) D.M. 02/09/2021, esercitazione annuale, aggiornamento triennale."),
]

# Per-livello specific recommendations to layer on top of the generic
# categories. Higher risk -> more demanding measures.
LIVELLO_SPECIFIC_MEASURES: dict[str, list[str]] = {
    "BASSO": [
        "Esercitazione antincendio almeno annuale.",
        "Verifica estintori semestrale a cura di tecnico abilitato.",
        "Addetti antincendio formazione 4 ore (livello 1-FOR).",
    ],
    "MEDIO": [
        "Esercitazione antincendio almeno annuale con scenari multipli (esodo, primo intervento).",
        "Verifica estintori semestrale; impianto rivelazione + allarme periodicamente testato.",
        "Addetti antincendio formazione 8 ore (livello 2-FOR) + idoneita tecnica per gli ambienti soggetti a CPI.",
    ],
    "ALTO": [
        "Esercitazione antincendio semestrale; coinvolgimento VV.F. nei piani di emergenza complessi.",
        "Sistema di rivelazione automatica e spegnimento (sprinkler / aerosol) dove tecnicamente fattibile.",
        "Addetti antincendio formazione 16 ore (livello 3-FOR) + idoneita tecnica obbligatoria; CPI in corso di validita.",
        "Coordinamento con il responsabile CPI per aggiornamenti periodici.",
    ],
}


class AllegatoIncendioGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        incendi = await load_incendio(self.db, self.azienda_id)
        ambienti_map = {a.id: a for a in data["ambienti"]}

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"VALUTAZIONE SPECIFICA - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
            ("Riferimento normativo", "D.M. 03/09/2021 (Gestione della sicurezza antincendio nei luoghi di lavoro)"),
        ])

        add_heading(doc, "Metodologia di valutazione", level=2)
        add_paragraph(doc, "Il rischio incendio e valutato combinando tre indicatori, ciascuno in scala 1-3:")
        add_data_table(doc, ["Codice", "Indicatore", "Scala 1-3"], [
            ["INF", "Infiammabilita e carico d'incendio", "1=basso; 2=medio; 3=alto"],
            ["SI",  "Sorgenti di ignizione presenti", "1=assenti/rare; 2=discrete; 3=numerose"],
            ["PI",  "Presenza di persone e loro esodo", "1=semplice; 2=medio; 3=complesso"],
        ])
        add_paragraph(doc, "Classificazione del rischio = INF + SI + PI: 3-4 = BASSO, 5-7 = MEDIO, 8-9 = ALTO.")

        add_heading(doc, "Valutazione per ambiente", level=2)
        if not incendi:
            add_paragraph(doc, "Nessuna valutazione del rischio incendio disponibile.", italic=True)
        else:
            headers = ["Ambiente", "INF", "SI", "PI", "Totale", "Livello", "Uscite", "Estintori"]
            rows = []
            for v in incendi:
                amb_name = (
                ambienti_map[v.ambiente_id].nome
                if v.ambiente_id in ambienti_map
                else (v.nome_area or "—")
            )
                rows.append([
                    amb_name, str(v.inf), str(v.si), str(v.pi),
                    str(v.punteggio_totale or (v.inf + v.si + v.pi)),
                    v.livello_rischio or "",
                    str(v.uscite_emergenza), str(v.estintori_presenti),
                ])
            add_data_table(doc, headers, rows)

        # Company-wide aggregate (worst-case across areas + counts).
        if incendi:
            livelli = [v.livello_rischio or "" for v in incendi]
            counts = {
                "BASSO": livelli.count("BASSO"),
                "MEDIO": livelli.count("MEDIO"),
                "ALTO": livelli.count("ALTO"),
            }
            order = {"BASSO": 0, "MEDIO": 1, "ALTO": 2}
            worst = max((l for l in livelli if l in order), key=order.get, default="—")

            add_heading(doc, "Esito complessivo aziendale", level=2)
            add_kv_table(doc, [
                ("Ambienti valutati", str(len(incendi))),
                ("Aree a rischio BASSO", str(counts["BASSO"])),
                ("Aree a rischio MEDIO", str(counts["MEDIO"])),
                ("Aree a rischio ALTO", str(counts["ALTO"])),
                ("Livello aziendale (worst-case)", worst),
            ])

        add_heading(doc, "Misure di prevenzione e protezione per area", level=2)
        add_paragraph(
            doc,
            "Per ciascuna area di lavoro le misure sono organizzate nelle 6 categorie "
            "di prevenzione previste dal D.M. 03/09/2021 (REFERENCE_DATA §4.5). "
            "Alle misure standard si aggiungono prescrizioni specifiche calibrate "
            "sul livello di rischio assegnato.",
            italic=True,
            size=9,
        )
        for v in incendi:
            amb_name = (
                ambienti_map[v.ambiente_id].nome
                if v.ambiente_id in ambienti_map
                else (v.nome_area or "—")
            )
            livello = (v.livello_rischio or "BASSO").upper()
            add_heading(doc, f"{amb_name} — livello {livello}", level=3)

            # 6-category prevention table
            cat_rows = [[cat, contenuto] for cat, contenuto in PREVENTION_CATEGORIES]
            add_data_table(doc, ["Categoria di prevenzione", "Misure"], cat_rows)

            # Livello-specific addenda
            add_heading(doc, f"Prescrizioni aggiuntive — livello {livello}", level=4)
            for m in LIVELLO_SPECIFIC_MEASURES.get(livello, LIVELLO_SPECIFIC_MEASURES["BASSO"]):
                add_paragraph(doc, f"• {m}")

            # Operator-supplied measures, if any
            if v.misure_prevenzione:
                add_heading(doc, "Misure aggiuntive registrate", level=4)
                add_paragraph(doc, v.misure_prevenzione)

        add_heading(doc, "Gestione dell'emergenza", level=2)
        add_paragraph(doc, "Il personale e addestrato all'uso degli estintori. Il piano di emergenza (allegato PEE) descrive procedure dettagliate per ogni scenario d'incendio. E prevista esercitazione antincendio almeno annuale (D.M. 02/09/2021 art. 6).")

        add_heading(doc, "Sottoscrizione", level=2)
        add_data_table(doc, ["Ruolo", "Nominativo", "Firma"], [
            ["Datore di Lavoro", azienda.ragione_sociale or "", "________________________"],
            ["RSPP", "________________________", "________________________"],
            ["Addetto antincendio coordinatore", "________________________", "________________________"],
            ["Data", generated_at.strftime("%d/%m/%Y"), ""],
        ])

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento.in_([TIPO_DOC, "ALLEGATO_INCENDIO"]))
        )
        r = await self.db.execute(stmt)
        return (r.scalar() or 0) + 1
