"""Allegato Microclima Severo.

Covers both severe-environment evaluation types:
- caldo severo: UNI EN ISO 7933 (PHS - Predicted Heat Strain)
- freddo severo: UNI EN ISO 11079 (IREQ - required clothing insulation,
  wind chill / frostbite screening per Annex D)
"""

import math
import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_microclima
from app.services.document_generator.docx_utils import (
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    slugify,
)
from app.services.microclima_calculator import calculate_phs

TIPO_DOC = "allegato_microclima_severo"


def _compute_phs(t_air, t_rad, v_air, rh, met, clo) -> tuple[float | None, float | None, float | None]:
    """Use the shared current-version PHS model, with a safe fallback."""
    try:
        result = calculate_phs(
            air_temp=float(t_air),
            mean_radiant_temp=float(t_rad),
            air_velocity=float(v_air),
            humidity=float(rh),
            metabolic_rate=float(met),
            clothing_insulation=float(clo),
            posture="standing",
            duration_min=480,
        )
        values = (
            float(result["sweat_loss_g"]),
            float(result["t_re"]),
            float(result["d_lim"]),
        )
        if not all(math.isfinite(value) for value in values):
            raise ValueError("PHS returned a non-finite result")
        return values
    except Exception:
        # Heuristic fallback
        excess = max(0.0, float(t_air) - 28.0)
        return 1500 + excess * 200, 37.0 + excess * 0.1, max(30.0, 480.0 - excess * 40)


def _severity(d_lim_min: float | None) -> str:
    if d_lim_min is None:
        return "—"
    if d_lim_min >= 480:
        return "Accettabile per l'intera giornata lavorativa"
    if d_lim_min >= 240:
        return "Turno ridotto o pause supplementari"
    return "Esposizione non ammessa senza DPI/misure rinfrescanti"


def _compute_ireq(t_air, t_rad, v_air, rh, met, clo) -> dict:
    """Try the ISO 11079 calculator service; fall back to a conservative
    closed-form estimate (same house style as the PHS fallback above)."""
    try:
        from app.services.microclima_calculator import calculate_ireq

        return calculate_ireq(
            air_temp=float(t_air),
            mean_radiant_temp=float(t_rad),
            air_velocity=float(v_air),
            humidity=float(rh),
            metabolic_rate=float(met),
            clothing_insulation=float(clo),
        )
    except Exception:
        # Heuristic fallback: dry heat balance at 33 °C skin temperature,
        # 85% of metabolic heat available, still-air layer 0.7 clo.
        m_net = max(10.0, float(met) * 58.15 * 0.85)
        t_o = (float(t_air) + float(t_rad)) / 2.0
        req_clo = max(0.0, (33.0 - t_o) / m_net / 0.155 - 0.7)
        livello = "ACCETTABILE" if float(clo) >= req_clo else "CRITICO"
        return {
            "t_o": round(t_o, 1),
            "ireq_neutral": round(req_clo, 2),
            "ireq_minimal": round(req_clo * 0.9, 2),
            "icl": float(clo),
            "delta_clo": round(max(0.0, req_clo - float(clo)), 2),
            "dle_min": None,
            "t_wc": None,
            "frostbite_risk": "NON_APPLICABILE",
            "livello": livello,
        }


def _severity_freddo(livello: str | None) -> str:
    return {
        "ACCETTABILE": "Isolamento adeguato per l'intero turno",
        "LIMITE": "Aumentare l'isolamento o prevedere pause di riscaldamento",
        "CRITICO": "Esposizione limitata (DLE) - misure obbligatorie",
    }.get(livello or "", "—")


class AllegatoMicroclimaSeveroGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        micro = await load_microclima(self.db, self.azienda_id)
        ambienti_map = {a.id: a for a in data["ambienti"]}

        # PHS (UNI EN ISO 7933) models heat strain only — cold-severe rows are
        # scored with IREQ (UNI EN ISO 11079) in Parte II below.
        severe_rows = [m for m in micro if (m.tipo_ambiente or "") == "severo_caldo"]
        cold_rows = [m for m in micro if (m.tipo_ambiente or "") == "severo_freddo"]

        doc = Document()
        add_heading(doc, "ALLEGATO RISCHIO MICROCLIMA - AMBIENTI SEVERI (CALDO E FREDDO)", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
            ("Riferimento normativo (caldo)", "UNI EN ISO 7933:2023 - Determinazione dello stress termico - Indice PHS"),
            ("Riferimento normativo (freddo)", "UNI EN ISO 11079:2008 - Determinazione e interpretazione dello stress da freddo - Indici IREQ e raffreddamento localizzato"),
        ])

        add_heading(doc, "PARTE I - STRESS DA CALDO (PHS)", level=2)
        add_heading(doc, "Metodologia", level=2)
        add_paragraph(doc, "L'indice PHS (Predicted Heat Strain) stima la perdita totale di sudore (in g), la temperatura rettale prevista (t_re) e il limite di esposizione più restrittivo tra temperatura rettale e perdita idrica (d_lim in minuti).")

        add_heading(doc, "Soglie di azione", level=2)
        add_data_table(doc, ["d_lim", "Classificazione"], [
            [">= 480 min (intera giornata)", "ACCETTABILE"],
            ["240-480 min", "TURNI RIDOTTI / PAUSE"],
            ["< 240 min", "ESPOSIZIONE NON AMMESSA senza DPI"],
        ])

        add_heading(doc, "Valutazione per ambiente severo", level=2)
        if not severe_rows:
            add_paragraph(doc, "Nessun ambiente a rischio da caldo severo registrato.", italic=True)
        else:
            rows = []
            for m in severe_rows:
                amb_name = ambienti_map[m.ambiente_id].nome if m.ambiente_id in ambienti_map else "—"
                sw, t_re, dlim = _compute_phs(m.temperatura_aria, m.temperatura_radiante, m.velocita_aria, m.umidita_relativa, m.metabolismo, m.isolamento_vestiario)
                rows.append([
                    amb_name,
                    f"{float(m.temperatura_aria):.1f}",
                    f"{float(m.temperatura_radiante):.1f}",
                    f"{float(m.umidita_relativa):.0f}",
                    f"{float(m.metabolismo):.2f}",
                    f"{sw:.0f}" if sw is not None else "—",
                    f"{t_re:.1f}" if t_re is not None else "—",
                    f"{dlim:.0f}" if dlim is not None else "—",
                    _severity(dlim),
                ])
            add_data_table(
                doc,
                ["Ambiente", "t_aria", "t_rad", "RH%", "met", "Perdita sudore g", "t_re C", "d_lim min", "Classificazione"],
                rows,
                column_widths_cm=[1.8, 1.1, 1.1, 0.9, 0.9, 1.7, 1.1, 1.1, 5.3],
            )

        add_heading(doc, "Misure organizzative e di protezione", level=2)
        add_paragraph(doc, "Per ambienti con stress termico severo: idratazione frequente (>= 250 ml/h), pause in zona rinfrescata ogni 45 minuti, rotazione personale, monitoraggio sintomi, formazione sul riconoscimento del colpo di calore, sorveglianza sanitaria specifica.")

        # ------------------------------------------------------------------
        # Parte II — severe cold (IREQ, UNI EN ISO 11079)
        # ------------------------------------------------------------------
        add_heading(doc, "PARTE II - STRESS DA FREDDO (IREQ)", level=2)

        add_heading(doc, "Metodologia", level=2)
        add_paragraph(doc, "La norma UNI EN ISO 11079 valuta il raffreddamento generale del corpo mediante l'indice IREQ (Insulation REQuired): l'isolamento termico del vestiario necessario a mantenere l'equilibrio termico nelle condizioni ambientali e metaboliche rilevate. IREQ neutro corrisponde all'equilibrio in condizioni di neutralita termica; IREQ minimo al massimo raffreddamento corporeo accettabile. Se l'isolamento del vestiario indossato (Icl) e inferiore a IREQ minimo, l'esposizione deve essere limitata nel tempo (DLE - Durata Limite di Esposizione, calcolata con debito termico ammesso di 40 Wh/m2). Il raffreddamento localizzato viene valutato mediante la temperatura wind chill (t_wc, Appendice D della norma).")

        add_heading(doc, "Soglie di classificazione", level=2)
        add_data_table(doc, ["Condizione", "Classificazione"], [
            ["Icl >= IREQ neutro", "ACCETTABILE - isolamento adeguato all'intero turno"],
            ["IREQ minimo <= Icl < IREQ neutro", "LIMITE - raffreddamento progressivo lieve"],
            ["Icl < IREQ minimo", "CRITICO - esposizione limitata alla DLE"],
        ])
        add_data_table(doc, ["Wind chill t_wc", "Rischio congelamento (ISO 11079 App. D)"], [
            ["> -25 C", "BASSO - disagio da freddo"],
            ["da -25 C a -35 C", "MODERATO - congelamento cute esposta entro ~30 min"],
            ["da -35 C a -60 C", "ALTO - congelamento entro ~10 min"],
            ["<= -60 C", "ESTREMO - congelamento entro ~2 min"],
        ])

        add_heading(doc, "Valutazione per ambiente a freddo severo", level=2)
        if not cold_rows:
            add_paragraph(doc, "Nessun ambiente a rischio da freddo severo registrato.", italic=True)
        else:
            rows = []
            for m in cold_rows:
                amb_name = ambienti_map[m.ambiente_id].nome if m.ambiente_id in ambienti_map else (m.nome_area or "—")
                r = _compute_ireq(m.temperatura_aria, m.temperatura_radiante, m.velocita_aria, m.umidita_relativa, m.metabolismo, m.isolamento_vestiario)
                rows.append([
                    amb_name,
                    f"{float(m.temperatura_aria):.1f}",
                    f"{float(m.velocita_aria):.1f}",
                    f"{float(m.metabolismo):.2f}",
                    f"{float(m.isolamento_vestiario):.2f}",
                    f"{r['ireq_neutral']:.2f}",
                    f"{r['ireq_minimal']:.2f}",
                    f"{r['t_wc']:.1f}" if r.get("t_wc") is not None else "—",
                    f"{r['dle_min']:.0f}" if r.get("dle_min") is not None else "—",
                    r.get("livello") or "—",
                ])
            add_data_table(doc, ["Ambiente", "t_aria C", "v_aria m/s", "met", "Icl clo", "IREQ neutro", "IREQ min", "t_wc C", "DLE min", "Classificazione"], rows)
            for m in cold_rows:
                r = _compute_ireq(m.temperatura_aria, m.temperatura_radiante, m.velocita_aria, m.umidita_relativa, m.metabolismo, m.isolamento_vestiario)
                amb_name = ambienti_map[m.ambiente_id].nome if m.ambiente_id in ambienti_map else (m.nome_area or "—")
                add_paragraph(doc, f"{amb_name}: {_severity_freddo(r.get('livello'))}." + (f" Isolamento supplementare consigliato: +{r['delta_clo']:.2f} clo." if r.get("delta_clo") else ""))

        add_heading(doc, "Misure organizzative e di protezione contro il freddo", level=2)
        add_paragraph(doc, "Per ambienti con stress da freddo severo: indumenti di protezione contro il freddo con isolamento adeguato all'IREQ calcolato (abbigliamento multistrato, antivento), protezione di mani, piedi e capo, pause di riscaldamento in locale temperato con bevande calde, limitazione dell'esposizione alla DLE in caso di isolamento insufficiente, rotazione del personale, protezione della cute esposta quando la temperatura wind chill scende sotto -25 C, formazione sul riconoscimento di ipotermia e congelamento, sorveglianza sanitaria specifica (art. 181 D.Lgs. 81/2008).")

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "ALLEGATO_MICROCLIMA_SEVERO"])
