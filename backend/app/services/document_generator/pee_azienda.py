"""PEE - Piano di Emergenza ed Evacuazione (variante aziendale)."""

import logging
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy import func, select

from app.data.pee_procedures import (
    DEFAULT_TIPOLOGIA_ALLARME,
    NUE_LABEL,
    merge_with_overrides,
    normalize_emergency_number,
)
from app.models.ambiente import Ambiente
from app.models.ambiente_foto import AmbienteFoto
from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_pee
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    page_break,
    replace_placeholders,
    scrub_body,
    slugify,
)

logger = logging.getLogger(__name__)

TEMPLATE = TEMPLATES_DIR / "PIANO GESTIONE EMERGENZE - AZIENDA.docx"
TIPO_DOC = "pee_azienda"


def _remove_donor_collection_point_images(doc: Document) -> None:
    """Remove the legacy customer's collection-point photo and its overlay."""
    in_collection_point_block = False
    cleaned_paragraphs = []
    block_paragraphs = []
    donor_relationship_ids = set()
    donor_point = "parcheggio del polo commerciale"

    for paragraph in doc.paragraphs:
        text = paragraph.text.casefold()
        if "raggiungere il punto di raccolta esterno" in text and donor_point in text:
            in_collection_point_block = True
            continue
        if (
            in_collection_point_block
            and "il punto di raccolta del personale evacuato sarà" in text
            and donor_point in text
        ):
            break
        if not in_collection_point_block:
            continue

        block_paragraphs.append(paragraph)
        drawings = paragraph._p.xpath(".//w:drawing")
        for drawing in drawings:
            donor_relationship_ids.update(
                drawing.xpath(".//a:blip/@r:embed")
            )
            drawing.getparent().remove(drawing)
        if drawings:
            cleaned_paragraphs.append(paragraph)

    if cleaned_paragraphs:
        placeholder_paragraph = cleaned_paragraphs[0]
        run = placeholder_paragraph.add_run(
            "Immagine del punto di raccolta da configurare."
        )
        run.italic = True
        for paragraph in block_paragraphs:
            if paragraph is placeholder_paragraph or paragraph.text.strip():
                continue
            paragraph._p.getparent().remove(paragraph._p)

    for relationship_id in donor_relationship_ids:
        if not doc.element.body.xpath(
            f'.//a:blip[@r:embed="{relationship_id}"]'
        ):
            doc.part.drop_rel(relationship_id)


async def _find_planimetria_path(db, azienda_id) -> str | None:
    """Return the on-disk path of a planimetria photo for this azienda, if any.

    Heuristic: any ``ambienti_foto`` row whose filename (or path) contains
    the substring "planimetria" (case-insensitive) is treated as the floor
    plan. If multiple rows match we pick the most recent one. When no
    match exists we return ``None`` so the caller can render the placeholder.

    In fixture/test mode ``db`` is ``None`` (see scripts/verify_all_generators.py);
    we short-circuit so the test runner produces the placeholder output rather
    than crashing on the lookup.
    """
    if db is None:
        return None
    stmt = (
        select(AmbienteFoto)
        .join(Ambiente, Ambiente.id == AmbienteFoto.ambiente_id)
        .where(Ambiente.azienda_id == azienda_id)
        .where(
            func.lower(AmbienteFoto.filename).like("%planimetria%")
        )
        .order_by(AmbienteFoto.created_at.desc())
        .limit(1)
    )
    result = await db.execute(stmt)
    row = result.scalar_one_or_none()
    if row is None:
        return None
    # Only embed if the file still exists on disk.
    if not row.file_path or not os.path.exists(row.file_path):
        return None
    return row.file_path


class PeeAziendaGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        pee = await load_pee(self.db, self.azienda_id, tipo="azienda")

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {
                "RAGIONE SOCIALE": azienda.ragione_sociale or "",
                "[AZIENDA]": azienda.ragione_sociale or "",
            })
            _remove_donor_collection_point_images(doc)
            scrub_body(doc, {
                "Parcheggio del polo commerciale": (
                    (pee.punto_raccolta if pee else None) or "—"
                ),
                " (come illustrato sopra)": "",
            })
        else:
            doc = Document()

        drill_frequency = (pee.frequenza_prove if pee else None) or "non configurata"
        # Alarm type configured on the plan; the default keeps the document
        # coherent (and reviewable) before the operator picks one.
        tipologia_allarme = (
            (pee.tipologia_allarme if pee else None) or DEFAULT_TIPOLOGIA_ALLARME
        )
        page_break(doc)
        add_heading(doc, f"PIANO DI EMERGENZA - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Sede", format_sede(azienda, "legale")),
            ("Data emissione", generated_at.strftime("%d/%m/%Y")),
            ("Coordinatore emergenza", (pee.coordinatore_emergenza if pee else "—") or "—"),
            ("Tipologia di allarme", tipologia_allarme),
            ("Punto di raccolta", (pee.punto_raccolta if pee else "—") or "—"),
            ("Frequenza prove", drill_frequency),
            ("Orario di lavoro dichiarato", azienda.orario_lavoro or "—"),
            (
                "Lavoratori dichiarati dall'azienda",
                str(azienda.numero_dipendenti_dichiarati)
                if azienda.numero_dipendenti_dichiarati is not None
                else "—",
            ),
            ("Persone registrate nel DVR", str(len(data.get("persone") or []))),
            ("Tempo evacuazione stimato (min)", str(pee.tempo_evacuazione_stimato_min) if pee and pee.tempo_evacuazione_stimato_min else "—"),
            ("Riferimento normativo", "D.M. 02/09/2021 (Criteri gestione emergenza luoghi di lavoro)"),
        ])

        if pee:
            add_heading(doc, "Numeri telefonici di emergenza", level=2)
            # NUE reform: legacy national emergency numbers (113/115/118/...)
            # render as 112. Company-internal numbers pass through untouched.
            rows = [
                [k, normalize_emergency_number(v)]
                for k, v in (pee.telefoni_emergenza or {}).items()
            ]
            add_data_table(doc, ["Ente/Ruolo", "Numero"], rows or [[NUE_LABEL, "112"]])
            add_paragraph(
                doc,
                "Tutte le chiamate di soccorso confluiscono nel Numero Unico di "
                "Emergenza (NUE) 112.",
                italic=True,
            )

            add_heading(doc, "Squadra di emergenza", level=2)
            members = pee.squadra_emergenza or []
            if members:
                add_data_table(doc, ["Nominativo", "Ruolo"], [[m.get("nome", ""), m.get("ruolo", "")] for m in members])
            else:
                add_paragraph(doc, "Squadra non configurata.", italic=True)

            add_heading(doc, "Vie di fuga e punto di raccolta", level=2)
            add_paragraph(doc, pee.vie_fuga or "Vie di fuga indicate dalla segnaletica di sicurezza UNI EN ISO 7010.")
            add_paragraph(doc, f"Punto di raccolta: {pee.punto_raccolta or '—'}")

        # Planimetria (US-4.1 AC3): embed the uploaded floor plan if one exists
        # among ambienti_foto (filename containing "planimetria"); otherwise
        # render the placeholder text so the operator knows to attach one.
        add_heading(doc, "Planimetria di emergenza", level=2)
        planimetria_path = await _find_planimetria_path(self.db, self.azienda_id)
        if planimetria_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(planimetria_path, width=Inches(6.0))
            except Exception:
                # Any load-time error (corrupt image, unsupported format)
                # falls back to the placeholder so generation never breaks.
                logger.exception(
                    "Failed to embed planimetria for azienda %s", self.azienda_id
                )
                add_paragraph(
                    doc,
                    "Inserire planimetria (immagine allegata non leggibile).",
                    italic=True,
                )
            add_paragraph(
                doc,
                "Planimetria indicativa con percorsi di esodo, uscite di sicurezza e punto di raccolta.",
                italic=True,
            )
        else:
            add_paragraph(doc, "Inserire planimetria", italic=True)

        # Structured A-E procedures per event type (US-4.2). Standard procedures
        # from app.data.pee_procedures are merged with per-client overrides
        # persisted in pee.scenari. We always render the full 5×5 grid so the
        # operator gets consistent coverage even when no overrides exist.
        add_heading(doc, "Procedure di emergenza per scenario", level=2)
        merged_events = merge_with_overrides(
            pee.scenari if pee else None,
            tipologia_allarme=pee.tipologia_allarme if pee else None,
        )
        for event in merged_events:
            add_heading(doc, event["titolo"], level=3)
            for proc in event["procedure"]:
                suffix = " (personalizzata)" if proc.get("personalizzata") else ""
                add_paragraph(
                    doc,
                    f"{proc['lettera']}. {proc['titolo']}{suffix}",
                    bold=True,
                )
                add_paragraph(doc, proc["testo"])

        add_heading(doc, "Formazione e prove di evacuazione", level=2)
        add_paragraph(
            doc,
            "La squadra di emergenza riceve formazione specifica (primo soccorso D.M. 388/2003 "
            "e antincendio D.M. 02/09/2021).",
        )
        add_paragraph(
            doc,
            f"Le prove di evacuazione seguono la frequenza configurata: {drill_frequency}. "
            "Ogni prova viene registrata con il relativo esito.",
        )

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "PEE_AZIENDA"])
