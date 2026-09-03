"""HACCP — 16 schede di autocontrollo (SA-01 .. SA-16) bundled into one zip.

Output: a single .docx per form, then all assembled into a .zip that is
returned as the "file" for this doc type. If zip creation is not possible,
returns the main index .docx.

US-4.4: each form carries the consultancy letterhead (the organization's
logo, else the bundled default) and the client's ragione sociale, and the
operator can pick a subset of forms via ``options.selected_codes``.
"""

import io
import logging
import os
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.branding import Branding, resolve_logo_source
from app.services.document_generator.data_loader import load_haccp
from app.services.document_generator.design import (
    add_cover,
    finish_document,
    setup_document,
)
from app.services.document_generator.docx_utils import (
    BRAND_DEEP,
    BRAND_NAVY,
    BRAND_SLATE,
    TEMPLATES_DIR,
    TYPE_SCALE,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    replace_placeholders,
    slugify,
)

logger = logging.getLogger(__name__)

TIPO_DOC = "haccp_forms"
DOC_TITLE = "Schede di Autocontrollo HACCP"
HACCP_TEMPLATES_DIR = TEMPLATES_DIR / "haccp"

_MESI = [
    "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
    "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre",
]


def _normalize_code(code: str) -> str:
    """Compare codes ignoring case, hyphens and whitespace."""
    return (code or "").upper().replace("-", "").replace(" ", "").strip()


def _add_form_letterhead(doc, azienda, branding: Branding | None = None) -> None:
    """Compact letterhead at the top of a record form: logo, consultancy and
    client on three tight lines — a form is printed and filled by hand, so it
    gets no cover page (audit 2026-09-03: the old header was followed by a
    page break that left page one empty but for the logo).
    """
    branding = branding or Branding.default()
    logo_src = resolve_logo_source(branding)
    if logo_src is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        try:
            p.add_run().add_picture(logo_src, width=Cm(3.4))
        except Exception:
            # Corrupt or unreadable image — drop the picture and let the
            # firm-name line below act as the brand mark.
            logger.exception("HACCP form logo embed failed")
    firm = (branding.firm_name or "").strip()
    if firm:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{firm.upper()} · Scheda di autocontrollo HACCP")
        run.font.size = Pt(TYPE_SCALE["small"])
        run.font.color.rgb = BRAND_SLATE
    name = (azienda.ragione_sociale or "").strip()
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(TYPE_SCALE["h3"])
        run.font.color.rgb = BRAND_DEEP


class HaccpFormsGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        config, forms = await load_haccp(self.db, self.azienda_id)

        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        version = await self._next_version()

        # US-4.4 AC2: filter forms by the dialog-supplied selected_codes when
        # provided. Codes are normalised so the dialog can pass either
        # "SA-01" or "sa01"; missing/empty option means "all forms" so the
        # legacy "Genera Tutti" behaviour is preserved.
        selected_codes_raw = self.options.get("selected_codes")
        if selected_codes_raw:
            wanted = {_normalize_code(c) for c in selected_codes_raw}
            selected_forms = [f for f in forms if _normalize_code(f.form_code) in wanted]
        else:
            selected_forms = list(forms)

        responsabile = ((config.responsabile_haccp if config else None) or "").strip()
        periodo = f"{_MESI[generated_at.month - 1]} {generated_at.year}"

        # Build individual forms
        form_paths: list[str] = []
        for form in selected_forms:
            form_path = self._build_single_form(
                output_dir, slug, version, form, azienda, generated_at,
                responsabile=responsabile, periodo=periodo,
            )
            form_paths.append(form_path)

        # Build index document — the packet's cover sheet, furnished like every
        # other document so the zip opens on something a client recognises.
        index_doc = Document()
        setup_document(index_doc)
        add_cover(
            index_doc,
            title=DOC_TITLE,
            subtitle="Registrazioni del piano di autocontrollo",
            legal_basis="ai sensi del Reg. (CE) 852/2004 e del D.Lgs. 193/2007",
            eyebrow="Manuale di autocontrollo per l'igiene degli alimenti",
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )
        add_heading(index_doc, "Dati generali", level=1)
        add_kv_table(index_doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Data emissione", generated_at.strftime("%d/%m/%Y")),
            ("Responsabile HACCP", responsabile or "—"),
            ("Numero schede allegate", str(len(form_paths))),
        ])
        add_heading(index_doc, "Elenco schede", level=1)
        add_data_table(
            index_doc,
            ["Codice", "Titolo"],
            [[f.form_code, f.form_title] for f in selected_forms] or [["—", "—"]],
            column_widths_cm=[3.0, 13.5],
        )
        finish_document(
            index_doc,
            title=DOC_TITLE,
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )

        # Package all into a zip file
        zip_path = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            # Write index doc to zip
            idx_buf = io.BytesIO()
            index_doc.save(idx_buf)
            idx_buf.seek(0)
            zf.writestr("INDICE_schede_HACCP.docx", idx_buf.read())
            for fp in form_paths:
                zf.write(fp, arcname=os.path.basename(fp))
        return zip_path

    def _build_single_form(
        self, output_dir: str, slug: str, version: int, form, azienda, generated_at,
        *, responsabile: str = "", periodo: str = "",
    ) -> str:
        """Clone template form if available, or build from scratch; save .docx."""
        code = form.form_code
        template_candidate: Path | None = None
        if HACCP_TEMPLATES_DIR.exists():
            for p in HACCP_TEMPLATES_DIR.iterdir():
                if p.suffix.lower() == ".docx" and code.replace("-", "") in p.stem.replace("-", "").replace(" ", "").upper():
                    template_candidate = p
                    break
                if p.suffix.lower() == ".docx" and code.replace("-", "_") in p.stem.upper().replace("-", "_"):
                    template_candidate = p
                    break

        if template_candidate and template_candidate.exists():
            doc = Document(str(template_candidate))
            replace_placeholders(doc, {
                "RAGIONE SOCIALE": azienda.ragione_sociale or "",
                "[AZIENDA]": azienda.ragione_sociale or "",
            })
        else:
            doc = Document()
        setup_document(doc)

        # US-4.4 AC1: consultancy letterhead + client ragione sociale on
        # every form, regardless of whether the source template provided
        # its own placeholders.
        _add_form_letterhead(doc, azienda, self.branding)
        add_heading(doc, f"{code} — {form.form_title}", level=1)
        # Prefill what the operator would otherwise retype; both cells stay
        # editable ("solo una questione di revisione, non di inserimento").
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Responsabile compilazione", responsabile),
            ("Periodo di riferimento", periodo),
        ])
        add_heading(doc, "Registrazioni", level=2)
        righe = (form.data or {}).get("righe", [])
        if righe:
            headers = list(righe[0].keys()) if isinstance(righe[0], dict) else ["Riga"]
            rows = [[str(r.get(h, "")) for h in headers] if isinstance(r, dict) else [str(r)] for r in righe]
            add_data_table(doc, headers, rows)
        else:
            # Empty table with rows to compile by hand: a generous row height
            # so a pen fits, and columns sized for a date, a name and a value.
            table = add_data_table(
                doc,
                ["Data", "Responsabile", "Valore", "Note"],
                [["", "", "", ""] for _ in range(15)],
                column_widths_cm=[2.6, 4.4, 3.5, 6.0],
            )
            for row in table.rows[1:]:
                row.height = Cm(0.85)

        add_heading(doc, "Firma responsabile", level=2)
        add_paragraph(doc, "________________________")

        finish_document(
            doc,
            title=f"{code} · {form.form_title}",
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
            cover_is_clean=False,
        )

        filename = f"{TIPO_DOC}_{code}_{slug}_v{version}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "HACCP_FORMS"])
