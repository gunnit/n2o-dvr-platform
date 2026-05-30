"""PEE - Piano Gestione Emergenze variante Comune/Edificio."""

import logging
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy import func, select

from app.models.ambiente import Ambiente
from app.models.ambiente_foto import AmbienteFoto
from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_pee
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    page_break,
    replace_placeholders,
    slugify,
)

logger = logging.getLogger(__name__)


async def _find_planimetria_path(db, azienda_id) -> str | None:
    """Locate a planimetria photo among ambienti_foto (US-4.1 AC3).

    Returns ``None`` in fixture/test mode where ``db`` is None so the test
    runner can exercise the placeholder branch without a live session.
    """
    if db is None:
        return None
    stmt = (
        select(AmbienteFoto)
        .join(Ambiente, Ambiente.id == AmbienteFoto.ambiente_id)
        .where(Ambiente.azienda_id == azienda_id)
        .where(func.lower(AmbienteFoto.filename).like("%planimetria%"))
        .order_by(AmbienteFoto.created_at.desc())
        .limit(1)
    )
    result = await db.execute(stmt)
    row = result.scalar_one_or_none()
    if row is None or not row.file_path or not os.path.exists(row.file_path):
        return None
    return row.file_path

TEMPLATE = TEMPLATES_DIR / "PIANO GESTIONE EMERGENZE - COMUNE.docx"
TIPO_DOC = "pee_comune"


class PeeComuneGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        pee = await load_pee(self.db, self.azienda_id, tipo="comune") or await load_pee(self.db, self.azienda_id, tipo="azienda")

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"PIANO DI EMERGENZA EDIFICIO - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda riferimento", azienda.ragione_sociale or ""),
            ("Sede", format_sede(azienda, "legale")),
            ("Data emissione", generated_at.strftime("%d/%m/%Y")),
            ("Tipo", "Edificio multi-tenant"),
        ])

        add_heading(doc, "Obiettivo", level=2)
        add_paragraph(doc, "Il presente piano descrive la gestione coordinata delle emergenze in edificio condiviso, con attribuzione di ruoli e procedure tra le diverse aziende occupanti e l'amministratore condominiale.")

        if pee:
            add_heading(doc, "Numeri di emergenza", level=2)
            tel = pee.telefoni_emergenza or {"Numero Unico Europeo": "112"}
            add_data_table(doc, ["Ente/Ruolo", "Numero"], [[k, v] for k, v in tel.items()])
            add_heading(doc, "Coordinamento emergenze", level=2)
            add_kv_table(doc, [
                ("Coordinatore emergenza", pee.coordinatore_emergenza or "—"),
                ("Punto di raccolta", pee.punto_raccolta or "—"),
                ("Vie di fuga", pee.vie_fuga or "—"),
            ])

        add_heading(doc, "Procedure comuni multi-tenant", level=2)
        add_paragraph(doc, "In caso di attivazione dell'allarme generale dell'edificio, tutte le aziende interrompono le attivita, attivano il proprio coordinatore locale e procedono all'evacuazione verso il punto di raccolta condominiale.")

        # Planimetria (US-4.1 AC3): embed the uploaded floor plan if one exists,
        # otherwise a placeholder. Condominial buildings tend to share a single
        # planimetry across occupants so we look on the same azienda photo set.
        add_heading(doc, "Planimetria di emergenza edificio", level=2)
        planimetria_path = await _find_planimetria_path(self.db, self.azienda_id)
        if planimetria_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(planimetria_path, width=Inches(6.0))
            except Exception:
                logger.exception(
                    "Failed to embed planimetria for azienda %s", self.azienda_id
                )
                add_paragraph(
                    doc,
                    "Inserire planimetria (immagine allegata non leggibile).",
                    italic=True,
                )
            add_paragraph(
                doc,
                "Planimetria generale dell'edificio con percorsi di esodo e punto di raccolta.",
                italic=True,
            )
        else:
            add_paragraph(doc, "Inserire planimetria", italic=True)

        add_heading(doc, "Manutenzione dei presidi comuni", level=2)
        add_paragraph(doc, "Gli impianti antincendio comuni (rivelazione, idranti, porte REI) sono manutenuti dall'amministratore condominiale con cadenza almeno semestrale e documentazione conservata a disposizione.")

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento.in_([TIPO_DOC, "PEE_COMUNE"]))
        )
        r = await self.db.execute(stmt)
        return (r.scalar() or 0) + 1
