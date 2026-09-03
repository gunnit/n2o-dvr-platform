"""Allegato VDT - Videoterminali (D.Lgs. 81/2008 Titolo VII, artt. 172-179).

Generates the VDT attachment from scratch (no template loading) to avoid
leaking the donor template's pre-populated client data. Mirrors the
template structure end-to-end:

  1. Cover (logo, title, azienda, generation date + revision)
  2. Revision history
  3. TOC
  4. Introduzione (VDT theory)
  5. Anagrafica Aziendale
  6. Organigramma Dipendenti (same table the DVR Master renders — client
     feedback 2026-08: "inserire sempre la tabella organigramma dipendenti
     presente sul rischio master")
  7. Organizzazione Aziendale della Sicurezza
  8. Principali fattori di rischio (vista, postura, affaticamento)
  9. La postazione di lavoro (videoterminale, ambiente, posizionamento)
 10. Elenco postazioni VDT (Postazione | ATTIVITÀ)
 11. Tavole di Valutazione — grouped PER DIPENDENTE: one block per worker
     with all their postazioni, hours summed across devices, exposure
     classified on the TOTAL (client feedback 2026-08)
 12. Quadro sinottico di esposizione — one row per worker with total hours
 13. Misure di prevenzione
 14. Programma di attuazione (sorveglianza sanitaria: periodicità computed
     from the worker's age per art. 176 c.3; visit dates no longer printed)
 15. Dichiarazione del Datore di Lavoro
 16. Signature block (DdL / RSPP / MC / RLS)
"""

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.branding import resolve_logo_source
from app.services.document_generator.data_loader import load_vdt
from app.services.document_generator.design import (
    add_cover,
    add_revision_table,
    add_toc,
    finalize_toc,
    finish_document,
    setup_document,
)
from app.services.document_generator.docx_utils import (
    HEADER_BG,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_comune,
    format_sede,
    page_break,
    shade_cell,
    slugify,
    style_header_row,
)
from app.services.vdt_calculator import (
    VDT_EXPOSURE_THRESHOLD_HOURS,
    classify_total_exposure,
    total_weekly_hours,
)
from app.services.vdt_surveillance import age_on, surveillance_periodicita

TIPO_DOC = "allegato_vdt"
DOC_TITLE = "Allegato Rischio VDT"

_EXPOSURE_COLORS = {
    "Esposto": "F4CCCC",      # rose — surveillance triggered
    "Non Esposto": "D9EAD3",  # green
}

def group_rows_by_person(vdt_rows: list) -> list[tuple[object, list]]:
    """Group VDT rows by persona_id, preserving first-seen order.

    One person may use several devices/postazioni: their rows collapse into
    a single group so hours can be summed and exposure judged on the total
    (client feedback 2026-08). Generic rows (persona_id None) each form
    their own single-row group — anonymous workstations cannot be assumed
    to belong to the same worker.

    Returns a list of ``(persona_id | None, rows)`` tuples. Module-level and
    pure so the per-employee grouping is unit-testable without a DB.
    """
    by_person: dict = {}
    ordered: list[tuple[object, list]] = []
    for r in vdt_rows:
        pid = getattr(r, "persona_id", None)
        if pid is None:
            ordered.append((None, [r]))
            continue
        if pid not in by_person:
            bucket: list = []
            by_person[pid] = bucket
            ordered.append((pid, bucket))
        by_person[pid].append(r)
    return ordered


_CHECKLIST_ITEMS: list[tuple[str, str]] = [
    ("schermo_conforme", "Schermo conforme (leggibilità, stabilità, regolazioni)"),
    ("tastiera_separata", "Tastiera separata e inclinabile"),
    ("sedile_regolabile", "Sedile a 5 razze, altezza/schienale regolabili"),
    ("poggiapiedi_disponibile", "Poggiapiedi disponibile su richiesta"),
    ("illuminazione_adeguata", "Illuminazione adeguata (300-500 lux)"),
    ("riflessi_assenti", "Assenza di riflessi e abbagliamenti"),
    ("spazio_adeguato", "Spazio di lavoro sufficiente"),
    ("pause_previste", "Pause previste (15 min ogni 2 ore di applicazione continuativa)"),
]


class AllegatoVdtGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        persone = data["persone"]
        ambienti = data["ambienti"]
        generated_at: datetime = data["generated_at"]
        vdt_rows = await load_vdt(self.db, self.azienda_id)
        version = await self._next_version()

        ambiente_by_id = {a.id: a for a in ambienti}
        persona_by_id = {p.id: p for p in persone}

        doc = Document()
        setup_document(doc)

        # Shared cover / storico / indice (audit 2026-09-03): same furniture as
        # every other attachment, org logo instead of the invisible white mark,
        # "Revisione 00" for a first issue like the DVR.
        add_cover(
            doc,
            title=DOC_TITLE,
            subtitle="Valutazione del Rischio da Videoterminali",
            legal_basis="ai sensi del Titolo VII D.Lgs. 81/2008 e s.m.i. (D.Lgs. 106/09)",
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )
        add_revision_table(doc, version, generated_at)
        toc_anchors = add_toc(doc)
        self._add_introduzione(doc)
        self._add_anagrafica(doc, azienda)
        self._add_organigramma(doc, persone)
        self._add_organizzazione(doc, persone)
        self._add_fattori_rischio(doc)
        self._add_postazione_lavoro(doc)
        self._add_elenco_postazioni(doc, vdt_rows, ambiente_by_id)
        self._add_per_worker_assessments(
            doc, vdt_rows, persona_by_id, ambiente_by_id, generated_at
        )
        self._add_quadro_sinottico(doc, vdt_rows, persona_by_id)
        self._add_misure_prevenzione(doc)
        self._add_programma_attuazione(doc, vdt_rows, persona_by_id, generated_at)
        self._add_dichiarazione_ddl(doc, azienda, persone)
        self._add_signature_block(doc, persone)

        finalize_toc(doc, *toc_anchors)
        finish_document(
            doc,
            title=DOC_TITLE,
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )

        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    # ------------------------------------------------------------------
    # Versioning
    # ------------------------------------------------------------------

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "ALLEGATO_VDT"])

    # ------------------------------------------------------------------
    # Introduzione
    # ------------------------------------------------------------------

    def _add_introduzione(self, doc) -> None:
        add_heading(doc, "Introduzione", level=1)
        for txt in [
            "L'utilizzo del videoterminale, soprattutto se prolungato, può' provocare "
            "disturbi all'apparato muscolo-scheletrico, all'apparato visivo e fenomeni "
            "di affaticamento fisico o mentale. La rilevanza di tali disturbi e' "
            "strettamente correlata alla durata dell'esposizione e alle caratteristiche "
            "ergonomiche della postazione.",
            "Il Titolo VII del D.Lgs. 81/2008 (artt. 172-179) impone al datore di "
            "lavoro la valutazione del rischio per i lavoratori che utilizzano "
            "abitualmente videoterminali e definisce, all'art. 173, lavoratore esposto "
            "chi vi opera in modo sistematico o abituale per almeno "
            f"{int(VDT_EXPOSURE_THRESHOLD_HOURS)} ore settimanali, dedotte le pause.",
            "La presente valutazione classifica ciascuna postazione e ciascun "
            "lavoratore in funzione del tempo di utilizzo settimanale; per i "
            "lavoratori esposti viene predisposto il programma di sorveglianza "
            "sanitaria oculistica ai sensi dell'art. 176 (periodicità' "
            "quinquennale, biennale per gli over 50 o con prescrizioni).",
        ]:
            add_paragraph(doc, txt)
        page_break(doc)

    # ------------------------------------------------------------------
    # Anagrafica Aziendale
    # ------------------------------------------------------------------

    def _add_anagrafica(self, doc, azienda) -> None:
        add_heading(doc, "Anagrafica Aziendale", level=1)
        rows: list[tuple[str, str]] = [
            ("Azienda", azienda.ragione_sociale or "—"),
            ("Attività / Codice ATECO", getattr(azienda, "codice_ateco", "") or "—"),
            ("Partita IVA", getattr(azienda, "partita_iva", "") or "—"),
            ("Codice Fiscale", getattr(azienda, "codice_fiscale", "") or "—"),
            ("Sede Legale - Via", azienda.sede_legale_via or "—"),
            ("Sede Legale - Città", format_comune(
                getattr(azienda, "cap_legale", None),
                azienda.sede_legale_citta,
                getattr(azienda, "provincia_legale", None))),
            ("Sede Operativa - Via", getattr(azienda, "sede_operativa_via", "") or "—"),
            ("Sede Operativa - Città", format_comune(
                getattr(azienda, "cap_operativa", None),
                getattr(azienda, "sede_operativa_citta", None),
                getattr(azienda, "provincia_operativa", None))),
            ("Telefono", getattr(azienda, "telefono", "") or "—"),
            ("Email PEC", getattr(azienda, "email_pec", "") or "—"),
        ]
        add_kv_table(doc, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Organigramma Dipendenti — same table the DVR Master renders
    # ------------------------------------------------------------------

    def _add_organigramma(self, doc, persone: list) -> None:
        """Employee organigramma, mirroring the DVR Master's Table 5
        (``DvrGenerator._add_dati_occupazionali_table``): Nominativo |
        Mansione | Ambiente di Lavoro | Codice Fiscale | Tipologia
        contrattuale, uppercased, with the "Tutta l'azienda (ruolo
        trasversale)" label for cross-site roles.

        Client feedback 2026-08: "inserire sempre la tabella organigramma
        dipendenti presente sul rischio master" — always rendered, with an
        em-dash placeholder row when no persone exist (like the master).
        Logic is replicated rather than imported from dvr_master to keep
        the two generators decoupled.
        """
        add_heading(doc, "Organigramma Dipendenti", level=1)

        headers = [
            "Nominativo",
            "Mansione",
            "Ambiente di Lavoro",
            "Codice Fiscale",
            "Tipologia contrattuale",
        ]
        if not persone:
            add_data_table(doc, headers, [["—", "—", "—", "—", "—"]])
            page_break(doc)
            return

        rows = []
        for p in persone:
            ambienti_names = ", ".join(
                (a.nome or "")
                for a in (getattr(p, "ambienti", None) or [])
                if getattr(a, "nome", None)
            )
            if not ambienti_names:
                if (
                    p.ruolo_datore_lavoro
                    or p.ruolo_rspp
                    or p.ruolo_rls
                    or getattr(p, "ruolo_medico_competente", False)
                ):
                    ambienti_names = "Tutta l'azienda (ruolo trasversale)"
                else:
                    ambienti_names = "—"
            cf = getattr(p, "codice_fiscale", None) or "—"
            rows.append([
                (p.nominativo or "—").upper(),
                (p.mansione or "—").upper(),
                ambienti_names.upper(),
                cf,
                (p.tipologia_contrattuale or "—").upper(),
            ])
        add_data_table(doc, headers, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Organizzazione Sicurezza
    # ------------------------------------------------------------------

    def _add_organizzazione(self, doc, persone: list) -> None:
        add_heading(doc, "Organizzazione Aziendale della Sicurezza", level=1)

        def _names(predicate) -> str:
            matched = [p.nominativo for p in persone if predicate(p) and p.nominativo]
            return ", ".join(matched) if matched else "—"

        rows = [
            ("Datore di Lavoro", _names(lambda p: bool(p.ruolo_datore_lavoro))),
            ("RSPP", _names(lambda p: bool(p.ruolo_rspp))),
            ("RLS", _names(lambda p: bool(p.ruolo_rls))),
            ("Medico Competente", _names(lambda p: bool(p.ruolo_medico_competente))),
            ("Addetti Primo Soccorso", _names(lambda p: bool(p.ruolo_primo_soccorso))),
            ("Addetti Antincendio", _names(lambda p: bool(p.ruolo_antincendio))),
            ("Preposti", _names(lambda p: bool(p.ruolo_preposto))),
        ]
        add_kv_table(doc, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Principali fattori di rischio (static narrative)
    # ------------------------------------------------------------------

    def _add_fattori_rischio(self, doc) -> None:
        add_heading(doc, "Principali fattori di rischio", level=1)
        add_paragraph(
            doc,
            "I disturbi che i lavoratori addetti ai videoterminali possono accusare "
            "sono riconducibili a tre famiglie di rischio: sollecitazione degli "
            "organi della vista, posizione del corpo, affaticamento fisico e mentale.",
        )

        add_heading(doc, "Sollecitazione degli organi della vista", level=2)
        add_paragraph(
            doc,
            "Bruciore, lacrimazione, secchezza, fastidio alla luce, pesantezza, "
            "visione annebbiata o sdoppiata, stanchezza alla lettura. Sono dovuti a "
            "elevata sollecitazione e rapido affaticamento degli organi della vista "
            "causati da: errate condizioni di illuminazione; ubicazione sbagliata "
            "del videoterminale rispetto alle finestre; condizioni ambientali "
            "sfavorevoli (aria secca, correnti, temperatura); caratteristiche "
            "inadeguate di software o monitor; postazione non corretta; impegno "
            "visivo ravvicinato e protratto; difetti visivi non corretti.",
        )

        add_heading(doc, "Posizione del corpo", level=2)
        add_paragraph(
            doc,
            "Disturbi alla colonna vertebrale dovuti a errata postura e sedentarieta', "
            "disturbi muscolari da affaticamento e indolenzimento, disturbi a mano "
            "e avambraccio (dolore, formicolii, impaccio ai movimenti) per "
            "infiammazione di nervi e tendini sovraccaricati.",
        )

        add_heading(doc, "Affaticamento fisico o mentale", level=2)
        add_paragraph(
            doc,
            "Determinato da: cattiva organizzazione del lavoro (operazioni monotone "
            "ripetitive); cattive condizioni ambientali (temperatura, umidità', "
            "velocità' dell'aria); rumore ambientale che disturba l'attenzione; "
            "software non adeguato.",
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # La postazione di lavoro (static narrative)
    # ------------------------------------------------------------------

    def _add_postazione_lavoro(self, doc) -> None:
        add_heading(doc, "La postazione di lavoro", level=1)

        add_heading(doc, "Videoterminale, tastiera e mouse", level=2)
        add_paragraph(
            doc,
            "La postazione VDT deve essere allestita con attrezzature moderne e "
            "ergonomiche: monitor orientabile e inclinabile, con luminosità' e "
            "contrasto regolabili e privi di sfarfallii; tastiera indipendente, "
            "spostabile, di basso spessore e inclinabile, con tasti dotati di "
            "superficie opaca; mouse posizionato accanto alla tastiera con spazio "
            "sufficiente all'appoggio del polso. Software di facile uso, adeguato "
            "alla mansione, con velocità' di risposta congrua.",
        )

        add_heading(doc, "Condizioni ambientali", level=2)
        add_paragraph(
            doc,
            "Temperatura raccomandata 18-22 °C in inverno e 24-26 °C in estate; "
            "umidità' relativa 40-60%; assenza di correnti d'aria fastidiose. "
            "Illuminazione 300-500 lux, priva di abbagliamenti diretti o riflessi "
            "sullo schermo. Rumore ambientale tale da non disturbare l'attenzione "
            "ne' la comunicazione verbale.",
        )

        add_heading(doc, "Corretto posizionamento del videoterminale", level=2)
        add_paragraph(
            doc,
            "Il monitor va posizionato perpendicolarmente alle finestre per evitare "
            "abbagliamenti e riflessi; la direzione principale dello sguardo deve "
            "essere parallela al piano delle finestre. Distanza occhio-schermo "
            "50-70 cm. Il bordo superiore dello schermo deve trovarsi all'altezza "
            "degli occhi o leggermente sotto.",
        )

        add_heading(doc, "Piano di lavoro, sedia, poggiapiedi", level=2)
        add_paragraph(
            doc,
            "Piano di lavoro stabile, di profondità' adeguata e altezza regolabile "
            "(70-80 cm). Sedia a 5 razze con sedile e schienale regolabili in "
            "altezza e inclinazione. Poggiapiedi disponibile su richiesta del "
            "lavoratore quando l'altezza del piano di lavoro non consente l'appoggio "
            "completo dei piedi a terra.",
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Elenco postazioni VDT
    # ------------------------------------------------------------------

    @staticmethod
    def _attivita_label(r, ambiente_by_id: dict) -> str:
        """The ATTIVITÀ cell for a row: the operator-entered attività, or the
        ambiente name as a prefill for legacy rows (review, never re-enter)."""
        att = (getattr(r, "attivita", None) or "").strip()
        if att:
            return att
        ambiente = ambiente_by_id.get(r.ambiente_id) if r.ambiente_id else None
        return (ambiente.nome if ambiente else None) or "—"

    def _add_elenco_postazioni(self, doc, vdt_rows: list, ambiente_by_id: dict) -> None:
        add_heading(doc, "Elenco postazioni VDT", level=1)
        if not vdt_rows:
            add_paragraph(
                doc,
                "Nessuna postazione VDT e' stata valutata per questa azienda.",
                italic=True,
            )
            page_break(doc)
            return

        # Group by postazione name to deduplicate (multiple workers may share one).
        # Client feedback 2026-08: the "Ambienti di lavoro" column becomes
        # "ATTIVITÀ" (operator-entered; ambiente name as legacy fallback).
        postazioni: dict[str, str] = {}
        for r in vdt_rows:
            name = (r.postazione or "—").strip()
            if name in postazioni:
                continue
            postazioni[name] = self._attivita_label(r, ambiente_by_id)

        rows = [[name, attivita] for name, attivita in postazioni.items()]
        add_data_table(doc, ["Postazione VDT", "ATTIVITÀ"], rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Per-employee assessment grid — one block per worker, all their
    # postazioni inside, exposure classified on the SUMMED weekly hours
    # (client feedback 2026-08: "le tabelle devono essere divise per
    # dipendente e non per postazione" + "somma delle ore").
    # ------------------------------------------------------------------

    @staticmethod
    def _group_surveillance_inputs(rows: list) -> tuple:
        """Collapse per-row surveillance inputs to person level.

        ``data_nascita`` / ``idoneita_visiva`` come from the first row that
        carries a value; "con prescrizioni" anywhere wins (it forces the
        biennale cadence per art. 176 c.3); the legacy eta_50_plus flag is
        true if any row flags it.
        """
        data_nascita = next(
            (r.data_nascita for r in rows if getattr(r, "data_nascita", None)), None
        )
        eta_flag = any(bool(getattr(r, "eta_50_plus", False)) for r in rows)
        idoneita = next(
            (r.idoneita_visiva for r in rows if getattr(r, "idoneita_visiva", None)),
            None,
        )
        if any(
            getattr(r, "idoneita_visiva", None) == "con prescrizioni" for r in rows
        ):
            idoneita = "con prescrizioni"
        return data_nascita, eta_flag, idoneita

    def _add_per_worker_assessments(
        self,
        doc,
        vdt_rows: list,
        persona_by_id: dict,
        ambiente_by_id: dict,
        generated_at: datetime,
    ) -> None:
        add_heading(doc, "Tavole di Valutazione del Rischio VDT", level=1)
        if not vdt_rows:
            add_paragraph(
                doc,
                "Nessuna postazione VDT e' stata valutata per questa azienda.",
                italic=True,
            )
            page_break(doc)
            return

        add_paragraph(
            doc,
            "Le tavole sono suddivise per dipendente: per ciascun lavoratore "
            "sono elencate tutte le postazioni/dispositivi utilizzati e "
            "l'esposizione e' classificata sulla somma delle ore settimanali "
            "complessive (art. 173 D.Lgs. 81/2008).",
            italic=True,
            size=9,
        )

        for i, (pid, rows) in enumerate(group_rows_by_person(vdt_rows), 1):
            persona = persona_by_id.get(pid) if pid else None
            self._render_employee_assessment(
                doc, i, persona, rows, ambiente_by_id, generated_at
            )

    def _render_employee_assessment(
        self, doc, idx: int, persona, rows: list, ambiente_by_id: dict,
        generated_at: datetime,
    ) -> None:
        if persona is not None:
            nominativo = persona.nominativo or "—"
            mansione = persona.mansione or "—"
            heading = f"{idx}. {nominativo}"
        else:
            # Anonymous workstation — no worker to aggregate on.
            nominativo = "Postazione generica"
            mansione = "—"
            heading = f"{idx}. {rows[0].postazione or 'Postazione'} (generica)"

        total = total_weekly_hours(float(r.ore_settimanali or 0) for r in rows)
        esposto = classify_total_exposure(
            float(r.ore_settimanali or 0) for r in rows
        ) == "ESPOSTO"
        esp_label = "Esposto" if esposto else "Non Esposto"

        add_heading(doc, heading, level=2)

        # Header table: Nominativo / Mansione, then one row per postazione,
        # then the summed total + threshold + classification.
        table = doc.add_table(rows=0, cols=4)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_row = table.add_row().cells
        header_row[0].text = "Nominativo"
        header_row[1].text = nominativo
        header_row[2].text = "Mansione"
        header_row[3].text = mansione
        for cell in header_row:
            shade_cell(cell, "1A237E")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)

        # One row per postazione/device used by this worker.
        for r in rows:
            ore = float(r.ore_settimanali or 0)
            sub_row = table.add_row().cells
            sub_row[0].text = "Postazione"
            sub_row[1].text = r.postazione or "—"
            sub_row[2].text = f"Attività: {self._attivita_label(r, ambiente_by_id)}"
            sub_row[3].text = f"{ore:.1f} h/sett."
            for cell in sub_row:
                shade_cell(cell, "EEEEEE")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        body = [
            (
                "Tempo di utilizzo totale",
                "ore/settimana (somma di tutte le postazioni)",
                f"{total:.1f}",
                "",
            ),
            (
                "Soglia esposizione",
                "art. 173 D.Lgs. 81/2008",
                f">= {int(VDT_EXPOSURE_THRESHOLD_HOURS)} h/sett.",
                "",
            ),
            ("Rischio VDT", "Classificazione (sul totale ore)", esp_label, ""),
        ]
        for code, descr, val, _ in body:
            row = table.add_row().cells
            row[0].text = code
            row[1].text = descr
            row[2].text = val
            row[3].text = ""
            shade_cell(row[0], "F5F5F5")
            for p in row[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for cell in (row[1], row[2], row[3]):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # Tint the classification row by exposure
        cls_row = table.rows[-1]
        tint = _EXPOSURE_COLORS.get(esp_label, "F5F5F5")
        for cell in cls_row.cells:
            shade_cell(cell, tint)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

        # Ergonomic checklist — one per postazione (the checklist describes
        # the workstation, not the worker).
        for r in rows:
            add_paragraph(doc, "")
            add_paragraph(
                doc,
                f"Check-list ergonomica — {r.postazione or 'Postazione'}",
                bold=True,
                size=10,
            )
            check_rows = []
            for attr, label in _CHECKLIST_ITEMS:
                ok = bool(getattr(r, attr, False))
                check_rows.append([label, "SI" if ok else "NO"])
            add_data_table(doc, ["Requisito", "Conformità"], check_rows)

        # Surveillance summary (only meaningful when esposto). Client
        # feedback 2026-08: no Ultima/Prossima visita — the periodicità is
        # computed from the worker's age (art. 176 c.3).
        if esposto:
            data_nascita, eta_flag, idoneita = self._group_surveillance_inputs(rows)
            on = generated_at.date()
            periodicita = surveillance_periodicita(
                data_nascita=data_nascita,
                eta_50_plus=eta_flag,
                idoneita_visiva=idoneita,
                on=on,
            )
            if data_nascita is not None:
                eta_label = f"{age_on(data_nascita, on)} anni"
            elif eta_flag:
                eta_label = ">= 50 anni"
            else:
                eta_label = "—"
            add_paragraph(doc, "")
            add_paragraph(doc, "Sorveglianza sanitaria oculistica", bold=True, size=10)
            surv_rows = [
                ("Idoneità visiva", idoneita or "—"),
                ("Età del lavoratore", eta_label),
                ("Periodicità (in base all'età, art. 176 c.3)", periodicita),
            ]
            add_kv_table(doc, surv_rows)

        notes = [r.note for r in rows if getattr(r, "note", None)]
        for note in notes:
            add_paragraph(doc, f"Note: {note}", italic=True, size=9)
        page_break(doc)

    # ------------------------------------------------------------------
    # Quadro sinottico
    # ------------------------------------------------------------------

    def _add_quadro_sinottico(self, doc, vdt_rows: list, persona_by_id: dict) -> None:
        """One row PER DIPENDENTE with the hours summed across every
        postazione the worker uses, and the exposure classified on that
        total (client feedback 2026-08: "deve esserci il calcolo delle ore
        ... includendo tutte le postazioni che utilizza. deve essere
        visualizzato il totale su una riga")."""
        add_heading(doc, "Quadro sinottico di esposizione", level=1)
        if not vdt_rows:
            add_paragraph(doc, "Nessuna valutazione presente.", italic=True)
            page_break(doc)
            return

        table = doc.add_table(rows=1, cols=5)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        for i, h in enumerate([
            "Nominativo",
            "Mansione",
            "Postazioni utilizzate",
            "Ore totali (h/sett)",
            "Rischio VDT",
        ]):
            hdr.cells[i].text = h
        style_header_row(hdr)

        groups = group_rows_by_person(vdt_rows)
        esposti = 0
        for pid, rows in groups:
            persona = persona_by_id.get(pid) if pid else None
            nominativo = (persona.nominativo if persona else None) or (
                rows[0].postazione or "—"
            )
            mansione = (persona.mansione if persona else None) or "—"
            postazioni = ", ".join((r.postazione or "—") for r in rows)
            total = total_weekly_hours(float(r.ore_settimanali or 0) for r in rows)
            esposto = classify_total_exposure(
                float(r.ore_settimanali or 0) for r in rows
            ) == "ESPOSTO"
            esp_label = "Esposto" if esposto else "Non Esposto"
            if esposto:
                esposti += 1

            row = table.add_row()
            row.cells[0].text = nominativo
            row.cells[1].text = mansione
            row.cells[2].text = postazioni
            row.cells[3].text = f"{total:.1f}" if total > 0 else "—"
            row.cells[4].text = esp_label

            tint = _EXPOSURE_COLORS.get(esp_label, "FFFFFF")
            for cell in (row.cells[3], row.cells[4]):
                shade_cell(cell, tint)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            for cell in (row.cells[0], row.cells[1], row.cells[2]):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        # Counters — per worker, not per row.
        total_workers = len(groups)
        add_paragraph(doc, "")
        add_paragraph(
            doc,
            f"Totale lavoratori valutati: {total_workers} · Esposti: {esposti} "
            f"· Non esposti: {total_workers - esposti}",
            italic=True,
            size=10,
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Misure di prevenzione (static)
    # ------------------------------------------------------------------

    def _add_misure_prevenzione(self, doc) -> None:
        add_heading(doc, "Misure di prevenzione", level=1)

        add_heading(doc, "Pause", level=2)
        add_paragraph(
            doc,
            "Ai sensi dell'art. 175 del D.Lgs. 81/2008, il lavoratore esposto ha "
            "diritto a una pausa di 15 minuti ogni 120 minuti di applicazione "
            "continuativa al videoterminale. La pausa e' considerata a tutti gli "
            "effetti tempo di lavoro e non può' essere accumulata a inizio o fine "
            "turno.",
        )

        add_heading(doc, "Muoversi di più'", level=2)
        add_paragraph(
            doc,
            "Alternare la posizione seduta con brevi pause attive: alzarsi, fare "
            "qualche passo, sgranchire spalle e collo. La sedentarieta' prolungata "
            "amplifica i disturbi muscolo-scheletrici.",
        )

        add_heading(doc, "Training per gli occhi", level=2)
        add_paragraph(
            doc,
            "Ogni 20 minuti distogliere lo sguardo dallo schermo e fissare un "
            "punto distante (>= 6 metri) per circa 20 secondi (regola del 20-20-20). "
            "Sbattere consapevolmente le palpebre per ridurre la secchezza oculare.",
        )

        add_heading(doc, "Esercizi di stretching e rilassamento", level=2)
        add_paragraph(
            doc,
            "Eseguire periodicamente esercizi di mobilizzazione di collo, spalle, "
            "polsi e schiena. Mantenere la schiena appoggiata allo schienale, "
            "spalle rilassate, avambracci paralleli al pavimento durante la "
            "digitazione.",
        )

        add_heading(doc, "Lavoratrici gestanti", level=2)
        add_paragraph(
            doc,
            "Per le lavoratrici gestanti l'utilizzo del videoterminale non comporta "
            "rischi specifici da radiazioni; restano comunque applicabili le "
            "ordinarie tutele previste dal D.Lgs. 151/2001 (riduzione/adattamento "
            "delle mansioni in caso di affaticamento o disturbi posturali).",
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Programma di Attuazione (sorveglianza per esposto)
    # ------------------------------------------------------------------

    def _add_programma_attuazione(
        self, doc, vdt_rows: list, persona_by_id: dict, generated_at: datetime
    ) -> None:
        add_heading(doc, "Programma di Attuazione delle Misure di Prevenzione", level=1)
        add_paragraph(
            doc,
            "Tutti i lavoratori esposti al rischio da utilizzo di attrezzature "
            f"munite di videoterminali per almeno {int(VDT_EXPOSURE_THRESHOLD_HOURS)} "
            "ore settimanali sono sottoposti a sorveglianza sanitaria oculistica "
            "ai sensi dell'art. 176 del D.Lgs. 81/2008, integrando i protocolli "
            "predisposti dal Medico Competente. La periodicità' standard e' "
            "quinquennale; e' biennale per i lavoratori di età' pari o superiore "
            "a 50 anni e per quelli con prescrizioni o idoneità' parziale.",
        )
        add_paragraph(
            doc,
            "Tutti i dipendenti vengono sottoposti a formazione e informazione "
            "specifica sul rischio VDT (postura, ergonomia, gestione delle pause).",
        )

        # Exposure is judged per worker on the SUMMED hours across all their
        # postazioni. Client feedback 2026-08: the Ultima/Prossima visita
        # columns are gone and the periodicità is computed from the worker's
        # age (art. 176 c.3: biennale >= 50 anni o con prescrizioni,
        # quinquennale altrimenti).
        on = generated_at.date()
        esposti_groups = [
            (pid, rows)
            for pid, rows in group_rows_by_person(vdt_rows)
            if classify_total_exposure(
                float(r.ore_settimanali or 0) for r in rows
            )
            == "ESPOSTO"
        ]
        if not esposti_groups:
            add_paragraph(
                doc,
                "Nessun lavoratore risulta esposto: programma di sorveglianza "
                "sanitaria non attivato.",
                italic=True,
            )
            page_break(doc)
            return

        add_heading(doc, "Sorveglianza sanitaria - lavoratori esposti", level=2)
        rows_out = []
        for pid, rows in esposti_groups:
            persona = persona_by_id.get(pid) if pid else None
            nome = (persona.nominativo if persona else None) or (
                rows[0].postazione or "—"
            )
            data_nascita, eta_flag, idoneita = self._group_surveillance_inputs(rows)
            periodicita = surveillance_periodicita(
                data_nascita=data_nascita,
                eta_50_plus=eta_flag,
                idoneita_visiva=idoneita,
                on=on,
            )
            if data_nascita is not None:
                eta_label = f"{age_on(data_nascita, on)} anni"
            elif eta_flag:
                eta_label = ">= 50 anni"
            else:
                eta_label = "—"
            rows_out.append([nome, eta_label, periodicita])
        add_data_table(
            doc,
            ["Lavoratore", "Età", "Periodicità sorveglianza"],
            rows_out,
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Dichiarazione del Datore di Lavoro
    # ------------------------------------------------------------------

    def _add_dichiarazione_ddl(self, doc, azienda, persone) -> None:
        add_heading(doc, "Dichiarazione del Datore di Lavoro", level=1)
        ddl_names = [
            p.nominativo for p in persone if p.ruolo_datore_lavoro and p.nominativo
        ]
        ddl = ddl_names[0] if ddl_names else "il Datore di Lavoro"
        ragione = azienda.ragione_sociale or "l'Azienda"

        sede = format_sede(azienda, "legale")

        add_paragraph(
            doc,
            f"Il/la sottoscritto/a {ddl}, in qualita' di Datore di Lavoro di "
            f"{ragione}, con sede legale in {sede},",
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DICHIARA")
        run.bold = True
        run.font.size = Pt(13)

        add_paragraph(
            doc,
            "che il procedimento sulla valutazione dei rischi da uso di attrezzature "
            "munite di videoterminali ex Titolo VII del D.Lgs. n. 81/2008 e s.m.i. "
            "(D.Lgs. 106/09) e' stato attuato in collaborazione con il Responsabile "
            "del Servizio di Prevenzione e Protezione, con il Medico Competente ove "
            "nominato e previa consultazione del Rappresentante dei Lavoratori per "
            "la Sicurezza.",
        )
        add_paragraph(
            doc,
            "Le misure di prevenzione e protezione individuate nel Programma di "
            "Attuazione saranno adottate secondo il cronoprogramma concordato e "
            "verificate periodicamente. La presente valutazione sarà' aggiornata "
            "in occasione di modifiche significative del processo lavorativo o "
            "dell'organizzazione del lavoro.",
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Signature block
    # ------------------------------------------------------------------

    def _add_signature_block(self, doc, persone) -> None:
        add_heading(doc, "Firme", level=1)

        def _first_or_dash(predicate) -> str:
            for p in persone:
                if predicate(p) and p.nominativo:
                    return p.nominativo
            return "—"

        ddl = _first_or_dash(lambda p: bool(p.ruolo_datore_lavoro))
        rspp = _first_or_dash(lambda p: bool(p.ruolo_rspp))
        mc = _first_or_dash(lambda p: bool(p.ruolo_medico_competente))
        rls = _first_or_dash(lambda p: bool(p.ruolo_rls))

        table = doc.add_table(rows=2, cols=2)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        cells = [
            (0, 0, "Il Datore di Lavoro", ddl),
            (0, 1, "Il Responsabile del S.P.P.", rspp),
            (1, 0, "Il Medico Competente", mc),
            (1, 1, "Il Rappresentante dei Lavoratori (per consultazione)", rls),
        ]
        for r, c, label, name in cells:
            cell = table.rows[r].cells[c]
            cell.text = ""
            p1 = cell.paragraphs[0]
            run = p1.add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            p2 = cell.add_paragraph(f"({name})")
            for run in p2.runs:
                run.font.size = Pt(10)
                run.italic = True
            cell.add_paragraph("")
            cell.add_paragraph("__________________________")
