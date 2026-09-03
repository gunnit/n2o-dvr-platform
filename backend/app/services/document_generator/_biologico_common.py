"""Shared scaffolding for Biologico generators (alimentare/asilo/dentisti)."""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_biologico
from app.services.document_generator.design import (
    add_cover,
    add_revision_table,
    finish_document,
    setup_document,
)
from app.services.document_generator.docx_utils import (
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    slugify,
)
from app.services.document_generator.reference_data_biologico import (
    classify_biologico,
    get_checklist,
)

DOC_TITLE = "Allegato Rischio Biologico"


async def build_biologico_document(
    gen: BaseDocumentGenerator,
    *,
    settore_key: str,
    titolo: str,
    agenti_default: list,
    misure_default: list,
    dpi_default: list,
    protocollo_default: str,
    tipo_doc: str,
    tipo_aliases: list[str],
    settore_label: str | None = None,
) -> str:
    data = await gen.load_data()
    azienda = data["azienda"]
    generated_at = data["generated_at"]
    rows = await load_biologico(gen.db, gen.azienda_id, settore_key)
    row = rows[0] if rows else None
    version = await _next_version(gen, tipo_doc, tipo_aliases)
    settore_label = settore_label or settore_key.capitalize()

    doc = Document()
    setup_document(doc)
    # Audit 2026-09-03: the three biologico annexes shipped as 1-2 US-Letter
    # pages with no cover, header, footer or logo. Shared furniture now.
    add_cover(
        doc,
        title=DOC_TITLE,
        subtitle=f"Settore {settore_label}",
        legal_basis="ai sensi del D.Lgs. 81/2008 Titolo X (artt. 266-286) e Allegati XLIV-XLVI",
        azienda=azienda,
        branding=gen.branding,
        version=version,
        generated_at=generated_at,
    )
    add_revision_table(doc, version, generated_at)

    add_heading(doc, "Dati generali", level=1)
    add_kv_table(doc, [
        ("Azienda", azienda.ragione_sociale or ""),
        ("Sede", format_sede(azienda, "legale")),
        ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
        ("Settore di riferimento", settore_label),
        ("Riferimento normativo", "D.Lgs. 81/2008 Titolo X (artt. 266-286) e Allegati XLIV-XLVI - Esposizione ad agenti biologici"),
    ])

    add_heading(doc, "Inquadramento", level=1)
    add_paragraph(
        doc,
        "La valutazione segue il Titolo X (artt. 266-286) del D.Lgs. 81/2008. "
        "Il Titolo X classifica gli agenti biologici in quattro gruppi in base a patogenicità, "
        "trasmissibilità e disponibilità di misure profilattiche o terapeutiche; "
        "l'Allegato XLVI elenca gli agenti dei gruppi 2, 3 e 4. Le misure di contenimento "
        "e la segnaletica di rischio biologico sono definite negli Allegati XLIV e XLV.",
    )

    add_heading(doc, "Agenti biologici identificati", level=1)
    agenti = (row.agenti_identificati if row and row.agenti_identificati else None) or agenti_default
    add_data_table(
        doc,
        ["Agente", "Gruppo", "Via di esposizione", "Patologia"],
        [[a.get("nome", ""), a.get("gruppo", ""), a.get("via", ""), a.get("patologia", "")] for a in agenti],
        column_widths_cm=[5.0, 1.8, 4.2, 5.5],
    )

    add_heading(doc, "Misure di prevenzione e protezione collettive", level=1)
    misure = (row.misure_protettive if row and row.misure_protettive else None) or [{"descrizione": m} for m in misure_default]
    for m in misure:
        add_paragraph(doc, f"• {m.get('descrizione', '')}")

    add_heading(doc, "Dispositivi di protezione individuale (DPI)", level=1)
    dpi = (row.dpi_richiesti if row and row.dpi_richiesti else None) or [{"descrizione": d} for d in dpi_default]
    for d in dpi:
        add_paragraph(doc, f"• {d.get('descrizione', '')}")

    add_heading(doc, "Sorveglianza sanitaria e formazione", level=1)
    add_paragraph(doc, (row.protocollo_sanitario if row and row.protocollo_sanitario else protocollo_default))
    if row and row.formazione_specifica:
        add_heading(doc, "Formazione specifica", level=2)
        add_paragraph(doc, row.formazione_specifica)
    classification = None
    if row and row.risposte_checklist:
        # Audit 2026-09-03: the checklist printed bare ids ("AL.01 | NO") and
        # the overall level was a stored string or a hardcoded "MEDIO" while
        # classify_biologico() sat unused. Join the catalogue text and derive
        # the level from the answers, so the table and the verdict agree.
        catalogue = {item["id"]: item for item in get_checklist(settore_key)}
        classification = classify_biologico(settore_key, row.risposte_checklist)
        add_heading(doc, "Esiti checklist rischio biologico", level=1)
        add_data_table(
            doc,
            ["Codice", "Controllo", "Criticità", "Risposta"],
            [
                [
                    str(item.get("id", "")),
                    catalogue.get(item.get("id", ""), {}).get("descrizione", ""),
                    catalogue.get(item.get("id", ""), {}).get("criticita", "").capitalize(),
                    str(item.get("risposta", "")),
                ]
                for item in row.risposte_checklist
            ],
            column_widths_cm=[1.8, 10.2, 2.2, 2.3],
        )
        if classification["unanswered"]:
            add_paragraph(
                doc,
                "Controlli senza risposta: " + ", ".join(classification["unanswered"]) + ".",
                italic=True,
                size=9,
            )
    if row and row.note:
        add_heading(doc, "Note valutazione biologica", level=1)
        add_paragraph(doc, row.note)

    add_heading(doc, "Esito valutazione", level=1)
    livello = (
        (row.livello_rischio if row else None)
        or (classification["livello"] if classification else None)
        or "MEDIO"
    )
    esito = [("Livello di rischio complessivo", livello)]
    if classification:
        esito.append((
            "Indice di criticità",
            f"{classification['no_weight']} / {classification['max_weight']} "
            f"(rapporto {classification['ratio']:.2f}: ≥ 0,40 ALTO · ≥ 0,15 MEDIO · altrimenti BASSO)".replace(".", ","),
        ))
    esito.append(("Periodicità revisione", "Annuale o in caso di modifiche organizzative rilevanti"))
    add_kv_table(doc, esito)

    finish_document(
        doc,
        title=f"{DOC_TITLE} — {settore_label}",
        azienda=azienda,
        branding=gen.branding,
        version=version,
        generated_at=generated_at,
    )

    output_dir = gen._get_output_dir()
    slug = slugify(azienda.ragione_sociale or "azienda")
    filepath = os.path.join(output_dir, f"{tipo_doc}_{slug}_v{version}.docx")
    doc.save(filepath)
    return filepath


async def _next_version(gen: BaseDocumentGenerator, tipo_doc: str, aliases: list[str]) -> int:
    return await gen.resolve_version([tipo_doc] + aliases)
