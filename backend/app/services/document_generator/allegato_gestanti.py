"""Allegato Gestanti - D.Lgs. 151/2001."""

import os
from datetime import timedelta

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import (
    load_gestanti,
    load_gestanti_mansioni,
)
from app.services.document_generator.design import finish_document, insert_logo_at_top
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    page_break,
    reset_table_rows,
    scrub_n2o_legacy_donor,
    strip_body_images,
    slugify,
)


def _role_nominativo(persone: list, attr: str) -> str | None:
    """Return the first persona's nominativo whose role flag ``attr`` is True."""
    for p in persone:
        if getattr(p, attr, False):
            nome = (getattr(p, "nominativo", None) or "").strip()
            if nome:
                return nome
    return None

TEMPLATE = TEMPLATES_DIR / "ALLEGATO GESTANTI.docx"
TIPO_DOC = "allegato_gestanti"

ESITO_LABELS = {
    "compatibile": "Compatibile",
    "compatibile_con_limitazioni": "Compatibile con limitazioni",
    "non_compatibile": "Non compatibile",
}


def _mansione_rischi_summary(rischi: list | None) -> str:
    """Flatten the JSONB risk rows of a mansione into one printable cell.

    Reuses the same row shape as ``GestantiValutazione.rischi_vietati``:
    {risk_key, allegato, descrizione, misura?}.
    """
    parts: list[str] = []
    for r in rischi or []:
        if not isinstance(r, dict):
            continue
        desc = r.get("descrizione") or r.get("rischio") or r.get("risk_key") or ""
        if not desc:
            continue
        allegato = r.get("allegato") or ""
        parts.append(f"{desc} (All. {allegato})" if allegato else desc)
    return "; ".join(parts) or "Nessun rischio rilevato"


def render_mansioni_section(doc, mansioni_vals: list) -> None:
    """Render the preventive per-mansione valutazione (art. 11 D.Lgs. 151/2001).

    This section exists regardless of whether any lavoratrice is currently
    pregnant: the decree requires the objective assessment for every
    mansione in advance. ``mansioni_vals`` are objects with attributes
    mansione / esito / rischi / misure / note.
    """
    add_heading(
        doc,
        "Valutazione preventiva dei rischi per mansione (art. 11 D.Lgs. 151/2001)",
        level=2,
    )
    add_paragraph(
        doc,
        "La valutazione seguente e' condotta in via preventiva per ciascuna "
        "mansione presente in azienda, indipendentemente dalla presenza di "
        "lavoratrici in stato di gravidanza, puerperio o allattamento. In caso "
        "di successiva comunicazione dello stato di gravidanza, le misure qui "
        "individuate trovano immediata applicazione (artt. 11 e 12 D.Lgs. "
        "151/2001).",
    )
    if not mansioni_vals:
        add_paragraph(
            doc,
            "Nessuna valutazione preventiva per mansione risulta registrata.",
            italic=True,
        )
        return

    rows: list[list[str]] = []
    for m in mansioni_vals:
        esito = getattr(m, "esito", "") or ""
        misure = (getattr(m, "misure", None) or "").strip()
        note = (getattr(m, "note", None) or "").strip()
        # No "\n" in cells: python-docx keeps it as a literal char that Word
        # renders as whitespace, not as a line break.
        misure_cell = misure or "—"
        if note:
            misure_cell = f"{misure} — Note: {note}" if misure else f"Note: {note}"
        rows.append([
            getattr(m, "mansione", "") or "—",
            ESITO_LABELS.get(esito, esito or "—"),
            _mansione_rischi_summary(getattr(m, "rischi", None)),
            misure_cell,
        ])
    add_data_table(
        doc,
        ["Mansione", "Esito valutazione", "Rischi (Allegati A/B/C)", "Misure / limitazioni"],
        rows,
    )


class AllegatoGestantiGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        persone = data["persone"]
        generated_at = data["generated_at"]
        gestanti = await load_gestanti(self.db, self.azienda_id)

        # Preventive per-mansione valutazione (art. 11): rendered even with
        # zero pregnant workers.
        mansioni_vals = await load_gestanti_mansioni(self.db, self.azienda_id)

        # Resolve org-sicurezza nominatives once so each scheda's firma
        # table pre-fills the names; the operator only needs to add the
        # ink signature at the audit.
        nome_ddl = _role_nominativo(persone, "ruolo_datore_lavoro") or (azienda.ragione_sociale or "")
        nome_rspp = _role_nominativo(persone, "ruolo_rspp") or "Da nominare"
        nome_mc = _role_nominativo(persone, "ruolo_medico_competente") or "Da nominare"

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            # The template carries N2O's own anagrafica in the body as the
            # assessed company (the placeholder tokens it was meant to use were
            # never present). Scrub the donor identity to the client; the
            # header/footer letterhead is preserved.
            scrub_n2o_legacy_donor(doc, azienda)
            # Audit 2026-09-03: the body repeats the legacy ENNE.DUE.O logo
            # (cover and chapter heads) and ends with a donor-era
            # acknowledgment roster. Strip the pictures, put the
            # organization's mark above the title, and leave the roster
            # as ten blank rows for the client's own workers to sign.
            strip_body_images(doc)
            insert_logo_at_top(doc, self.branding)
            reset_table_rows(doc, "Nome e cognome | Mansione", [["", "", "", ""] for _ in range(10)])
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"VALUTAZIONE SPECIFICA - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
            ("Mansioni valutate in via preventiva", str(len(mansioni_vals))),
            ("Lavoratrici in stato di gravidanza/allattamento notificate", str(len(gestanti))),
            ("Riferimento normativo", "D.Lgs. 151/2001 (artt. 7, 11, 12; Allegati A-B-C), mod. D.Lgs. 105/2022 - Testo Unico maternità e paternità"),
        ])

        add_heading(doc, "Inquadramento normativo", level=2)
        add_paragraph(doc, "Il D.Lgs. 151/2001 (come modificato dal D.Lgs. 105/2022) tutela la salute delle lavoratrici in stato di gravidanza, puerperio (fino a 7 mesi dal parto) e durante l'allattamento. Gli Allegati A, B e C individuano rispettivamente i lavori vietati, quelli vietati salvo deroga e gli agenti nocivi cui non possono essere esposte.")
        add_paragraph(doc, "La presente valutazione e condotta ai sensi degli artt. 7 (lavori vietati), 11 (valutazione dei rischi) e 12 (conseguenze della valutazione) del D.Lgs. 151/2001, in connessione con gli artt. 28 e 17 del D.Lgs. 81/2008.")

        # Section 1 — objective, preventive, per mansione (no worker needed).
        render_mansioni_section(doc, mansioni_vals)

        # Section 2 — worker-specific schede, only for notified pregnancies.
        add_heading(doc, "Valutazioni specifiche delle lavoratrici notificate", level=2)
        if not gestanti:
            add_paragraph(doc, "Non sono presenti lavoratrici in stato di gravidanza o allattamento al momento della valutazione. Resta valida la valutazione preventiva per mansione riportata sopra (art. 11 D.Lgs. 151/2001).", italic=True)
        for idx, g in enumerate(gestanti, 1):
            page_break(doc)
            nome = g.persona.nominativo if getattr(g, "persona", None) else "—"
            add_heading(doc, f"{idx}. Scheda lavoratrice", level=2)

            # Astensione obbligatoria: 2 mesi pre-parto + 3 mesi post-parto
            # (art. 16 D.Lgs. 151/2001). Compute the window so the operator
            # has a concrete date range, not just a yes/no flag.
            astensione_window = "—"
            if g.data_presunto_parto:
                inizio = g.data_presunto_parto - timedelta(days=60)
                fine = g.data_presunto_parto + timedelta(days=90)
                astensione_window = f"{inizio.strftime('%d/%m/%Y')} → {fine.strftime('%d/%m/%Y')}"

            add_kv_table(doc, [
                ("Lavoratrice", nome),
                ("Mansione", getattr(g.persona, "mansione", None) if getattr(g, "persona", None) else "—"),
                ("Stato", (g.stato or "").capitalize()),
                ("Data notifica", g.data_notifica.strftime("%d/%m/%Y") if g.data_notifica else "—"),
                ("Data presunto parto", g.data_presunto_parto.strftime("%d/%m/%Y") if g.data_presunto_parto else "—"),
                ("Astensione obbligatoria (art. 16)", astensione_window),
                ("Mansione alternativa", g.mansione_alternativa or "—"),
                ("Astensione anticipata richiesta (art. 17)", "SI - richiesta ispettorato del lavoro" if g.richiesta_astensione_anticipata else "NO"),
            ])
            add_heading(doc, "Rischi identificati e misure di adeguamento", level=3)
            rischi = g.rischi_vietati or []
            if rischi:
                # Feedback #32: the gestanti API persists rows with keys
                # {risk_key, allegato, descrizione, action, justification,
                # misura_alternativa}. Old code read "rischio"/"misura" which
                # never existed, so the table came out empty. Map onto the
                # actual schema; the "Misura adottata" column shows the
                # alternative measure (action=reject) or the acceptance
                # justification (action=accept).
                rows = []
                for r in rischi:
                    descrizione = (
                        r.get("descrizione")
                        or r.get("rischio")
                        or r.get("risk_key")
                        or ""
                    )
                    allegato = r.get("allegato", "")
                    misura = (
                        r.get("misura_alternativa")
                        or r.get("justification")
                        or r.get("misura")
                        or ""
                    )
                    rows.append([descrizione, allegato, misura])
                add_data_table(doc, ["Rischio", "Allegato D.Lgs. 151/2001", "Misura adottata"], rows)
            else:
                add_paragraph(doc, "Nessun rischio vietato identificato.", italic=True)

            if g.misure_adeguamento:
                add_paragraph(doc, g.misure_adeguamento)

            add_heading(doc, "Firme", level=3)
            add_data_table(doc, ["Ruolo", "Nominativo", "Firma"], [
                ["Lavoratrice", g.firma_lavoratrice or nome, "________________________"],
                ["Datore di lavoro", g.firma_datore_lavoro or nome_ddl, "________________________"],
                ["RSPP", g.firma_rspp or nome_rspp, "________________________"],
                ["Medico competente", g.firma_medico_competente or nome_mc, "________________________"],
                ["Data", generated_at.strftime("%d/%m/%Y"), ""],
            ])

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        # Audit 2026-09-03: the donor template's header/footer carried the
        # legacy N2O letterhead and literal placeholders; rewrite them from
        # the organization's branding and set honest file properties.
        finish_document(
            doc,
            title='Valutazione del Rischio Lavoratrici Gestanti',
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
            fill_cover=True,
        )
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "ALLEGATO_GESTANTI"])
