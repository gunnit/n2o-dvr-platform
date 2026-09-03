"""Shared page furniture for every generated document.

Document audit 2026-09-03 found three visual families shipping under one
product: code-built documents on US Letter with no header, footer or page
numbers; donor-template documents carrying a 2000s-era N2O letterhead in
their header/footer XML; and one-page stubs with no cover at all. This module
is the single place the cover, running header/footer, revision table, table
of contents and file properties are drawn, so every document looks like the
same consultancy produced it — and so an organization that uploads its own
logo gets it on every page of every document.

Palette and type scale live in :mod:`docx_utils` (``BRAND_*``, ``TYPE_SCALE``)
because the table helpers there need them too. The DVR Master keeps its own
cover (the VERA mark is a client requirement, see
docs/superpowers/plans/2026-08-03-dvr-master-luca-improvements.md) but adopts
the running header/footer and page setup from here.
"""

from __future__ import annotations

import re
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.document_generator.branding import Branding, resolve_logo_source
from app.services.document_generator.docx_utils import (
    BRAND_DEEP,
    BRAND_NAVY,
    BRAND_NAVY_HEX,
    BRAND_RULE_HEX,
    BRAND_SLATE,
    FONT_FAMILY,
    TYPE_SCALE,
    add_data_table,
    fill_label_table,
    format_sede,
    insert_in_order,
    reset_table_rows,
)

# A4 with a binding-side margin; text width is 16.5 cm.
PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_TOP_CM = 2.0
MARGIN_BOTTOM_CM = 2.0
MARGIN_LEFT_CM = 2.5
MARGIN_RIGHT_CM = 2.0
TEXT_WIDTH_CM = PAGE_W_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM

_BACKSLASH = chr(92)

# CT_PPr child sequence (ECMA-376 17.3.1.26); used to insert in schema order.
_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl", "numPr",
    "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap",
    "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc",
    "textDirection", "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
    "sectPr", "pPrChange",
]


# ---------------------------------------------------------------------------
# Revision numbering — one source for all documents
# ---------------------------------------------------------------------------

def revision_label(version: int | None) -> str:
    """``DocumentoGenerato.versione`` (1-based) -> the printed revision ("00").

    The database counts emissions from 1; Italian safety documents number them
    from 0 because revision zero *is* the first issue. Before the audit the DVR
    used this rule and MMC/VDT printed ``Revisione 01`` for the same first
    issue; every generator now goes through here.
    """
    try:
        n = int(version or 0)
    except (TypeError, ValueError):
        n = 0
    return f"{max(n - 1, 0):02d}"


def revision_motivation(version: int | None) -> str:
    """Storico-revisioni wording for the current emission."""
    try:
        n = int(version or 0)
    except (TypeError, ValueError):
        n = 0
    return "Emissione" if n <= 1 else "Aggiornamento"


# ---------------------------------------------------------------------------
# Page setup and base styles
# ---------------------------------------------------------------------------

def _set_style_font(style, name: str) -> None:
    """Set the font on a style including the East-Asian/complex-script slots,
    otherwise Word can fall back to Cambria/Times for headings."""
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), name)


def _keep_with_next(style) -> None:
    ppr = style.element.get_or_add_pPr()
    for tag in ("w:keepNext", "w:keepLines"):
        for existing in ppr.findall(qn(tag)):
            ppr.remove(existing)
    after_keep_next = _PPR_ORDER[_PPR_ORDER.index("keepLines"):]
    insert_in_order(ppr, OxmlElement("w:keepNext"), tuple(after_keep_next))
    insert_in_order(ppr, OxmlElement("w:keepLines"), tuple(_PPR_ORDER[_PPR_ORDER.index("pageBreakBefore"):]))


def setup_document(doc: Document, *, landscape: bool = False) -> None:
    """A4, binding margins, Calibri body, navy headings that never orphan.

    Idempotent: safe to call on a document opened from a template too — it
    only rewrites page geometry and the ``Normal``/``Heading 1-3`` styles.
    """
    for section in doc.sections:
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(PAGE_H_CM)
            section.page_height = Cm(PAGE_W_CM)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(PAGE_W_CM)
            section.page_height = Cm(PAGE_H_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    _set_style_font(normal, FONT_FAMILY)
    normal.font.size = Pt(TYPE_SCALE["body"])
    normal.font.color.rgb = BRAND_DEEP
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    specs = {
        1: (TYPE_SCALE["h1"], BRAND_NAVY, 18, 6),
        2: (TYPE_SCALE["h2"], BRAND_NAVY, 14, 4),
        3: (TYPE_SCALE["h3"], BRAND_DEEP, 10, 2),
    }
    for level, (size, colour, before, after) in specs.items():
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        _set_style_font(style, FONT_FAMILY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        _keep_with_next(style)


# ---------------------------------------------------------------------------
# Small drawing primitives
# ---------------------------------------------------------------------------

def _paragraph_border(paragraph, edge: str, color_hex: str = BRAND_RULE_HEX, size: int = 6, space: int = 4) -> None:
    """Hairline on one edge of a paragraph (``top`` or ``bottom``)."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        insert_in_order(ppr, pbdr, tuple(_PPR_ORDER[_PPR_ORDER.index("shd"):]))
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color_hex)
    pbdr.append(el)


def _add_field(run, instruction: str, cached: str = "1") -> None:
    """Emit a Word field (PAGE, NUMPAGES, ...) with a cached result so the
    unrefreshed view is still sensible."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    text = OxmlElement("w:t")
    text.text = cached
    run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _spacer(doc: Document, lines: int = 1) -> None:
    for _ in range(lines):
        doc.add_paragraph("")


def _centered(
    doc: Document,
    text: str,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    colour: RGBColor | None = None,
    caps: bool = False,
    spacing: float | None = None,
    space_after: float | None = None,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper() if caps else text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if colour is not None:
        run.font.color.rgb = colour
    if spacing is not None:
        rpr = run._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))
        insert_in_order(rpr, sp, ("w", "kern", "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign", "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath"))
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def insert_logo_at_top(doc: Document, branding: Branding, *, width_cm: float = 4.5) -> bool:
    """Put the consultancy mark as the very first body element.

    For donor templates whose cover is a text box we cannot restyle: the
    donor's own logo is stripped with the other body pictures and the
    organization's mark takes its place above the title.
    """
    logo_src = resolve_logo_source(branding)
    if logo_src is None:
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    try:
        p.add_run().add_picture(logo_src, width=Cm(width_cm))
    except Exception:
        p._p.getparent().remove(p._p)
        return False
    body = doc.element.body
    body.remove(p._p)
    body.insert(0, p._p)
    return True


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover(
    doc: Document,
    *,
    title: str,
    azienda,
    branding: Branding,
    version: int | None,
    generated_at: datetime,
    subtitle: str | None = None,
    legal_basis: str | None = None,
    eyebrow: str | None = "Allegato al Documento di Valutazione dei Rischi",
    logo_width_cm: float = 5.0,
    show_consultancy: bool = True,
) -> None:
    """The cover every attachment shares.

    Top to bottom: consultancy mark (the organization's uploaded logo, else the
    bundled navy N2O mark), eyebrow, title, subtitle, legal basis, a hairline,
    the assessed company's identity block, then revision + date and — when
    ``show_consultancy`` — the "elaborato da" letterhead. Ends with a page
    break. Every field degrades to nothing when the data is missing; the cover
    never prints a placeholder.
    """
    _spacer(doc, 2)

    logo_src = resolve_logo_source(branding)
    if logo_src is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(logo_src, width=Cm(logo_width_cm))
        except Exception:
            # An unreadable upload must not break generation; the firm name
            # in the letterhead below still identifies the consultancy.
            p.add_run("")
    _spacer(doc, 1)

    if eyebrow:
        _centered(doc, eyebrow, size=TYPE_SCALE["small"], colour=BRAND_SLATE, caps=True, spacing=1.2, space_after=8)
    _centered(doc, title, size=TYPE_SCALE["cover_title"], bold=True, colour=BRAND_NAVY, space_after=4)
    if subtitle:
        _centered(doc, subtitle, size=TYPE_SCALE["cover_subtitle"], colour=BRAND_SLATE, space_after=2)
    if legal_basis:
        _centered(doc, legal_basis, size=TYPE_SCALE["body"], italic=True, colour=BRAND_SLATE, space_after=10)

    rule = doc.add_paragraph()
    rule.paragraph_format.left_indent = Cm(5.0)
    rule.paragraph_format.right_indent = Cm(5.0)
    rule.paragraph_format.space_after = Pt(0)
    _paragraph_border(rule, "bottom", BRAND_NAVY_HEX, size=8, space=1)
    _spacer(doc, 2)

    ragione = (getattr(azienda, "ragione_sociale", None) or "").strip()
    _centered(
        doc,
        ragione.upper() if ragione else "—",
        size=TYPE_SCALE["cover_client"],
        bold=True,
        colour=BRAND_DEEP,
        space_after=4,
    )
    sede = format_sede(azienda, "legale")
    if sede and sede != "—":
        _centered(doc, sede, size=TYPE_SCALE["body"], colour=BRAND_SLATE, space_after=2)
    bits = []
    piva = getattr(azienda, "partita_iva", None)
    ateco = getattr(azienda, "codice_ateco", None)
    if piva:
        bits.append(f"P.IVA {piva}")
    if ateco:
        bits.append(f"ATECO {ateco}")
    if bits:
        _centered(doc, " · ".join(bits), size=TYPE_SCALE["body"], colour=BRAND_SLATE)

    _spacer(doc, 4)
    _centered(
        doc,
        f"Revisione {revision_label(version)} · {generated_at.strftime('%d/%m/%Y')}",
        size=TYPE_SCALE["body"],
        bold=True,
        colour=BRAND_DEEP,
        space_after=14,
    )

    if show_consultancy:
        _centered(
            doc,
            "Documento elaborato da",
            size=TYPE_SCALE["small"],
            colour=BRAND_SLATE,
            caps=True,
            spacing=1.0,
            space_after=2,
        )
        _centered(
            doc,
            (branding.firm_name or "").upper(),
            size=TYPE_SCALE["body"],
            bold=True,
            colour=BRAND_NAVY,
            space_after=1,
        )
        for line in _letterhead_lines(branding):
            _centered(doc, line, size=TYPE_SCALE["small"], colour=BRAND_SLATE, space_after=0)

    doc.add_page_break()


def _letterhead_lines(branding: Branding) -> list[str]:
    lines: list[str] = []
    addr = branding.address_line()
    if addr:
        lines.append(addr)
    tax = []
    if branding.partita_iva:
        tax.append(f"P.IVA {branding.partita_iva}")
    if branding.codice_fiscale and branding.codice_fiscale != branding.partita_iva:
        tax.append(f"C.F. {branding.codice_fiscale}")
    if tax:
        lines.append(" · ".join(tax))
    contact = branding.contact_line()
    if contact:
        lines.append(contact)
    if branding.rspp_nome:
        lines.append(f"RSPP: {branding.rspp_nome}")
    return lines


# ---------------------------------------------------------------------------
# Running header / footer
# ---------------------------------------------------------------------------

def _clear(container) -> None:
    """Empty a header/footer container so re-application is idempotent."""
    for p in list(container.paragraphs):
        p._p.getparent().remove(p._p)
    for t in list(container.tables):
        t._tbl.getparent().remove(t._tbl)


def _two_sided(container, left: list[tuple[str, dict]], right: list[tuple[str, dict]], *, width_cm: float):
    """One paragraph with left text and a right-aligned tab: the classic
    running-head layout without a table."""
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(width_cm), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_after = Pt(0)

    def emit(parts):
        for text, fmt in parts:
            run = p.add_run()
            run.font.size = Pt(fmt.get("size", TYPE_SCALE["small"]))
            run.bold = fmt.get("bold", False)
            run.font.color.rgb = fmt.get("colour", BRAND_SLATE)
            field = fmt.get("field")
            if field:
                _add_field(run, field, fmt.get("cached", "1"))
            else:
                run.text = text

    emit(left)
    tab = p.add_run()
    tab.text = chr(9)
    emit(right)
    return p


def _write_running(header, footer, *, title: str, ragione: str, firm: str, rev: str, width_cm: float) -> None:
    hp = _two_sided(
        header,
        [(title, {"bold": True, "colour": BRAND_NAVY})],
        [(ragione, {"colour": BRAND_SLATE})],
        width_cm=width_cm,
    )
    _paragraph_border(hp, "bottom", BRAND_RULE_HEX, size=6, space=3)

    left = [(firm, {"bold": True, "colour": BRAND_NAVY})]
    if firm:
        left.append((" · ", {}))
    left.append((rev, {}))
    fp = _two_sided(
        footer,
        left,
        [("Pagina ", {}), ("", {"field": "PAGE"}), (" di ", {}), ("", {"field": "NUMPAGES"})],
        width_cm=width_cm,
    )
    _paragraph_border(fp, "top", BRAND_RULE_HEX, size=6, space=3)


def add_running_header_footer(
    doc: Document,
    *,
    title: str,
    azienda,
    branding: Branding,
    version: int | None,
    generated_at: datetime,
    cover_is_clean: bool = True,
) -> None:
    """Header: document title | client. Footer: consultancy · revision | Pagina X di Y.

    Applied to every section; with ``cover_is_clean`` the first page of the
    first section (the cover) gets no header/footer. Content is written fresh
    into the default, first-page and even-page header/footer parts, replacing
    whatever a donor template carried there — this is how the legacy N2O
    letterhead leaves the template family.
    """
    ragione = (getattr(azienda, "ragione_sociale", None) or "").strip()
    firm = (branding.firm_name or "").strip()
    rev = f"Rev. {revision_label(version)} del {generated_at.strftime('%d/%m/%Y')}"
    for index, section in enumerate(doc.sections):
        width_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
        section.different_first_page_header_footer = bool(cover_is_clean and index == 0)
        # Donor templates set the footer distance to 0 (DUVRI) so their
        # letterhead sat on the page edge; the running header/footer needs
        # room, and a 0 distance is also where renderers start to struggle.
        for attr in ("header_distance", "footer_distance"):
            current = getattr(section, attr, None)
            if current is None or current < Cm(0.8):
                setattr(section, attr, Cm(1.0))
        parts = (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        )
        for container in parts:
            container.is_linked_to_previous = False
            _clear(container)

        _write_running(section.header, section.footer, title=title, ragione=ragione, firm=firm, rev=rev, width_cm=width_cm)
        # Even pages only matter when the document asks for odd/even headers;
        # writing them keeps a donor's evenAndOddHeaders setting harmless.
        _write_running(section.even_page_header, section.even_page_footer, title=title, ragione=ragione, firm=firm, rev=rev, width_cm=width_cm)

        if section.different_first_page_header_footer:
            # Leave the cover clean: an empty paragraph is required so Word does
            # not fall back to the default header on page one.
            section.first_page_header.add_paragraph("")
            section.first_page_footer.add_paragraph("")
        else:
            _write_running(section.first_page_header, section.first_page_footer, title=title, ragione=ragione, firm=firm, rev=rev, width_cm=width_cm)


# ---------------------------------------------------------------------------
# Revision history, table of contents, file properties
# ---------------------------------------------------------------------------

def add_revision_table(
    doc: Document,
    version: int | None,
    generated_at: datetime,
    *,
    heading: str | None = "Storico delle revisioni",
    level: int = 2,
):
    """The single-row Storico for this emission, numbered like the cover."""
    if heading:
        doc.add_heading(heading, level=level)
    table = add_data_table(
        doc,
        ["Rev.", "Motivazione", "Data"],
        [[revision_label(version), revision_motivation(version), generated_at.strftime("%d/%m/%Y")]],
        column_widths_cm=[2.0, 10.5, 4.0],
    )
    doc.add_paragraph("")
    return table


def add_toc(doc: Document, *, title: str = "Indice"):
    """A real TOC field with a cached outline the reader sees before any F9.

    Returns ``(field_start_p, end_p)`` for :func:`finalize_toc`, which fills
    the cached body with the headings actually emitted. Word is told to
    refresh fields on open, which restores page numbers.
    """
    doc.add_heading(title, level=1)
    field_start_p = doc.add_paragraph()
    run = field_start_p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " TOC " + _BACKSLASH + 'o "1-3" ' + _BACKSLASH + "h " + _BACKSLASH + "z " + _BACKSLASH + "u "
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    placeholder = doc.add_paragraph()
    r = placeholder.add_run("Indice in fase di aggiornamento.")
    r.font.size = Pt(TYPE_SCALE["body"])
    r.italic = True
    r.font.color.rgb = BRAND_SLATE

    end_p = doc.add_paragraph()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_p.add_run()._r.append(end)

    try:
        settings_root = doc.settings.element
        if settings_root.find(qn("w:updateFields")) is None:
            uf = OxmlElement("w:updateFields")
            uf.set(qn("w:val"), "true")
            settings_root.append(uf)
    except Exception:
        pass

    doc.add_page_break()
    return field_start_p, end_p


def finalize_toc(doc: Document, field_start_p, end_p) -> None:
    """Rewrite the TOC's cached body from the Heading 1-3 paragraphs that
    follow it, so the document is navigable even before Word refreshes."""
    body = doc.element.body
    children = list(body)
    try:
        start_idx = children.index(field_start_p._p)
        end_idx = children.index(end_p._p)
    except ValueError:
        return

    entries: list[tuple[int, str]] = []
    for para in doc.paragraphs:
        name = (para.style.name if para.style else "") or ""
        if not name.startswith("Heading "):
            continue
        try:
            level = int(name.split(" ")[1])
        except (IndexError, ValueError):
            continue
        if level not in (1, 2, 3):
            continue
        try:
            idx = children.index(para._p)
        except ValueError:
            continue
        if idx <= end_idx:
            continue
        text = (para.text or "").strip()
        if text:
            entries.append((level, text))
    if not entries:
        return

    for el in children[start_idx + 1 : end_idx]:
        body.remove(el)
    for level, text in entries:
        new_p = doc.add_paragraph()
        indent = (level - 1) * 0.5
        if indent:
            new_p.paragraph_format.left_indent = Cm(indent)
        new_p.paragraph_format.space_after = Pt(2)
        run = new_p.add_run(text)
        run.font.size = Pt(TYPE_SCALE["body"] if level == 1 else TYPE_SCALE["table"])
        run.bold = level == 1
        run.font.color.rgb = BRAND_NAVY if level == 1 else BRAND_DEEP
        body.remove(new_p._p)
        body.insert(list(body).index(end_p._p), new_p._p)


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def set_core_properties(
    doc: Document,
    *,
    title: str,
    azienda,
    branding: Branding,
    generated_at: datetime,
    version: int | None,
) -> None:
    """Make File > Properties tell the truth.

    Donor templates arrived with their authors' names, a hardware vendor as
    Company and, in one case, "Ethan Frome" as the title. Everything a client
    could read in the properties pane is rewritten from this generation.
    """
    ragione = (getattr(azienda, "ragione_sociale", None) or "").strip()
    firm = (branding.firm_name or "").strip()
    cp = doc.core_properties
    cp.title = title
    cp.subject = ragione
    cp.author = firm
    cp.last_modified_by = firm
    cp.category = "Sicurezza sul lavoro"
    cp.keywords = ""
    cp.comments = ""
    cp.identifier = ""
    cp.content_status = ""
    cp.created = generated_at
    cp.modified = generated_at
    try:
        cp.revision = max(int(version or 1), 1)
    except (TypeError, ValueError):
        cp.revision = 1
    # app.xml (Company, Application, TotalTime...) has no python-docx API.
    try:
        for part in doc.part.package.iter_parts():
            if str(part.partname) != "/docProps/app.xml":
                continue
            blob = part.blob.decode("utf8", "replace")
            blob = re.sub(r"<Company>.*?</Company>", f"<Company>{_xml_escape(firm)}</Company>", blob, flags=re.S)
            blob = re.sub(r"<Manager>.*?</Manager>", "<Manager></Manager>", blob, flags=re.S)
            blob = re.sub(r"<Application>.*?</Application>", "<Application>Microsoft Office Word</Application>", blob, flags=re.S)
            blob = re.sub(r"<TotalTime>.*?</TotalTime>", "<TotalTime>0</TotalTime>", blob, flags=re.S)
            part._blob = blob.encode("utf8")
    except Exception:
        pass


def fill_cover_tables(
    doc: Document,
    *,
    azienda,
    version: int | None,
    generated_at: datetime,
    extra: dict[str, str] | None = None,
) -> None:
    """Fill a donor template's cover forms from the client's data.

    The template covers are label | value tables whose value cells are empty
    — there is no token to substitute, which is why ``replace_placeholders``
    never did anything on them (audit 2026-09-03). Only empty value cells are
    written, so a template that already carries a value is left alone. The
    revision history is reset to this emission.
    """
    ragione = (getattr(azienda, "ragione_sociale", None) or "").strip()

    def seat_lines(which: str) -> list[str]:
        # Donor covers keep street and comune on two merged rows; hand them
        # over as two items so each row gets its own line.
        from app.services.document_generator.docx_utils import format_comune
        via = (getattr(azienda, f"sede_{which}_via", None) or "").strip()
        comune = format_comune(
            getattr(azienda, f"cap_{which}", None),
            getattr(azienda, f"sede_{which}_citta", None),
            getattr(azienda, f"provincia_{which}", None),
        )
        return [x for x in (via, comune if comune != "—" else "") if x]

    legale_lines = seat_lines("legale")
    operativa_lines = seat_lines("operativa") or legale_lines
    values: dict = {
        "Azienda": ragione,
        "Ragione sociale": ragione,
        "Sede Legale": legale_lines,
        "Sede Operativa": operativa_lines,
        "Sede": legale_lines,
        "Data": generated_at.strftime("%d/%m/%Y"),
    }
    piva = getattr(azienda, "partita_iva", None)
    if piva:
        values["P.IVA"] = piva
        values["Partita IVA"] = piva
    if extra:
        values.update({k: v for k, v in extra.items() if v})
    fill_label_table(doc, values)
    reset_table_rows(
        doc,
        "Rev. | Motivazione",
        [[revision_label(version), revision_motivation(version), generated_at.strftime("%d/%m/%Y")]],
    )


def _referenced_ids(element) -> set[str]:
    ids: set[str] = set()
    for el in element.iter():
        for attr in (qn("r:embed"), qn("r:id"), qn("r:link")):
            value = el.get(attr)
            if value:
                ids.add(value)
    return ids


def _remove_external_links(part, root) -> int:
    """Delete pictures linked to the internet and unwrap web hyperlinks.

    The DUVRI donor template pulls its status icons from www.secofor.it and
    links a Google search: offline Word shows red boxes, LibreOffice hangs
    on the fetch, and a third party's site has no place in a client's file.
    """
    removed = 0
    for rId, rel in list(part.rels.items()):
        if not rel.is_external:
            continue
        kind = rel.reltype.rsplit("/", 1)[-1]
        if kind == "image":
            for el in list(root.iter()):
                if el.get(qn("r:link")) == rId or el.get(qn("r:id")) == rId:
                    holder = el
                    while holder is not None and holder.tag not in (qn("w:drawing"), qn("w:pict"), qn("w:r")):
                        holder = holder.getparent()
                    target = holder if holder is not None else el
                    parent = target.getparent()
                    if parent is not None:
                        parent.remove(target)
                        removed += 1
            part.rels.pop(rId, None)
        elif kind == "hyperlink":
            for link in list(root.iter(qn("w:hyperlink"))):
                if link.get(qn("r:id")) != rId:
                    continue
                parent = link.getparent()
                index = parent.index(link)
                for child in list(link):
                    parent.insert(index, child)
                    index += 1
                parent.remove(link)
                removed += 1
            part.rels.pop(rId, None)
    return removed


def prune_orphan_parts(doc: Document) -> int:
    """Drop header/footer/image relationships nothing points to any more,
    and everything that points outside the file.

    Removing a donor picture from the body, or rewriting a section's
    headers, leaves the old part in the package: not rendered, but still in
    the file a client receives (a stripped Street View photo travelled that
    way). Returns the number of relationships dropped.
    """
    dropped = _remove_external_links(doc.part, doc.element.body)
    used = _referenced_ids(doc.element.body)
    for rId, rel in list(doc.part.rels.items()):
        if rel.is_external:
            continue
        kind = rel.reltype.rsplit("/", 1)[-1]
        if kind in ("header", "footer", "image") and rId not in used:
            doc.part.drop_rel(rId)
            dropped += 1
    # Images inside the header/footer parts we kept (their paragraphs were
    # cleared, their picture relationships were not).
    for rel in list(doc.part.rels.values()):
        kind = rel.reltype.rsplit("/", 1)[-1]
        if rel.is_external or kind not in ("header", "footer"):
            continue
        part = rel.target_part
        element = getattr(part, "element", None)
        if element is None:
            continue
        inner_used = _referenced_ids(element)
        for rId, sub in list(part.rels.items()):
            kind = sub.reltype.rsplit("/", 1)[-1]
            if rId in inner_used or kind not in ("image", "hyperlink"):
                continue
            # A cleared donor footer keeps its mailto: relationship even
            # though no run points at it any more; drop it with the images.
            if sub.is_external:
                part.rels.pop(rId, None)
            else:
                part.drop_rel(rId)
            dropped += 1
    return dropped


def finish_document(
    doc: Document,
    *,
    title: str,
    azienda,
    branding: Branding,
    version: int | None,
    generated_at: datetime,
    cover_is_clean: bool = True,
    fill_cover: bool = False,
    cover_values: dict[str, str] | None = None,
) -> None:
    """Everything a generator must do before ``doc.save``: running header and
    footer on every section, honest file properties, no orphaned donor parts
    and — for documents opened from a donor template — the cover forms
    filled in."""
    if fill_cover:
        fill_cover_tables(doc, azienda=azienda, version=version, generated_at=generated_at, extra=cover_values)
    add_running_header_footer(
        doc,
        title=title,
        azienda=azienda,
        branding=branding,
        version=version,
        generated_at=generated_at,
        cover_is_clean=cover_is_clean,
    )
    set_core_properties(
        doc,
        title=title,
        azienda=azienda,
        branding=branding,
        generated_at=generated_at,
        version=version,
    )
    try:
        prune_orphan_parts(doc)
    except Exception:
        # Pruning is hygiene, never a reason to fail a generation.
        pass
