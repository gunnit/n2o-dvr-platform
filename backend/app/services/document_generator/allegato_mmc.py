"""Allegato MMC - Movimentazione Manuale dei Carichi (NIOSH method).

Generates the MMC attachment from scratch (no template loading) to avoid
leaking the donor template's pre-populated client data. Mirrors the
template structure end-to-end:

  1. Cover page (logo, title, azienda, generation date + revision)
  2. Revision history
  3. TOC field
  4. Introduzione (NIOSH overview, static)
  5. Anagrafica Aziendale (current azienda)
  6. Dati Occupazionali (current persone, no codice fiscale leaked)
  7. Organizzazione Aziendale della Sicurezza (DdL/RSPP/RLS/MC)
  8. Metodologia NIOSH (CP, A, B, C, D, E, F factor reference tables)
  9. Per-worker assessment grid (one 13x5 table per persona+task)
 10. Quadro sinottico (Nominativo / Mansione / IR / Area, color-coded)
 11. Programma di Attuazione delle Misure (suggested measures per zone)
 12. Dichiarazione del Datore di Lavoro
 13. Signature block (DdL / RSPP / MC / RLS)
"""

from __future__ import annotations

import os
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from sqlalchemy import func, select

from app.data.niosh_factors import (
    FACTOR_A,
    FACTOR_B,
    FACTOR_C,
    FACTOR_D,
    FACTOR_E,
    FACTOR_F_TABLE,
    classify_ir,
    compute_plr,
)
from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_mmc
from app.services.document_generator.docx_utils import (
    HEADER_BG,
    LOGO_PATH,
    RISK_COLORS,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_comune,
    format_sede,
    page_break,
    shade_cell,
    slugify,
    style_header_row,
)


TIPO_DOC = "allegato_mmc"

_AREA_COLORS = {
    "Verde": "D9EAD3",
    "Gialla": "FFF2CC",
    "Rossa": "F4CCCC",
}

_ZONE_HEADINGS = {
    "Verde": "Area Verde - Rischio trascurabile (IR <= 0,75)",
    "Gialla": "Area Gialla - Sorveglianza sanitaria (0,75 < IR <= 1,00)",
    "Rossa": "Area Rossa - Riprogettazione e intervento (IR > 1,00)",
}

_DEFAULT_MEASURES_BY_ZONE = {
    "Verde": (
        "Mantenere le condizioni operative attuali. Sorveglianza sanitaria "
        "ordinaria; rivalutazione periodica in occasione di modifiche del "
        "ciclo produttivo o introduzione di nuove attrezzature."
    ),
    "Gialla": (
        "Avviare sorveglianza sanitaria mirata (visita medica con focus "
        "rachide). Pianificare azioni di riduzione del rischio: ausili meccanici "
        "(carrelli, transpallet, sollevatori), riprogettazione delle altezze di "
        "presa/deposito, riduzione della frequenza, formazione specifica art. 169 "
        "D.Lgs. 81/08."
    ),
    "Rossa": (
        "Intervento immediato richiesto. Adottare ausili meccanici per la "
        "movimentazione, ridisegnare la postazione (altezza, distanza, angoli), "
        "ridurre frequenza e durata della movimentazione, rotazione delle mansioni. "
        "Sorveglianza sanitaria con periodicita ravvicinata. Verifica efficacia "
        "delle misure entro 6 mesi."
    ),
}


class AllegatoMmcGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        persone = data["persone"]
        ambienti = data["ambienti"]
        generated_at: datetime = data["generated_at"]
        mmc_rows = await load_mmc(self.db, self.azienda_id)
        version = await self._next_version()

        ambiente_by_id = {a.id: a for a in ambienti}
        persona_by_id = {p.id: p for p in persone}

        doc = Document()
        self._setup_styles(doc)

        self._add_cover(doc, azienda, generated_at, version)
        self._add_revision_table(doc, generated_at, version)
        self._add_toc(doc)
        self._add_introduzione(doc)
        self._add_anagrafica(doc, azienda)
        self._add_dati_occupazionali(doc, persone)
        self._add_organizzazione(doc, persone)
        self._add_metodologia(doc)
        self._add_per_worker_assessments(doc, mmc_rows, persona_by_id, ambiente_by_id)
        self._add_quadro_sinottico(doc, mmc_rows, persona_by_id)
        self._add_programma_attuazione(doc, mmc_rows, persona_by_id)
        self._add_dichiarazione_ddl(doc, azienda, persone)
        self._add_signature_block(doc, persone)

        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    # ------------------------------------------------------------------
    # Versioning
    # ------------------------------------------------------------------

    async def _next_version(self) -> int:
        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento.in_([TIPO_DOC, "ALLEGATO_MMC"]))
        )
        r = await self.db.execute(stmt)
        return (r.scalar() or 0) + 1

    # ------------------------------------------------------------------
    # Style setup
    # ------------------------------------------------------------------

    def _setup_styles(self, doc: Document) -> None:
        # Page margins
        for s in doc.sections:
            s.top_margin = Cm(2.0)
            s.bottom_margin = Cm(2.0)
            s.left_margin = Cm(2.5)
            s.right_margin = Cm(2.0)

        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)
        except Exception:
            pass

    # ------------------------------------------------------------------
    # Cover
    # ------------------------------------------------------------------

    def _add_cover(self, doc, azienda, generated_at: datetime, version: int) -> None:
        for _ in range(2):
            doc.add_paragraph("")

        if LOGO_PATH.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.0))
            except Exception:
                p.add_run("[LOGO AZIENDALE]").italic = True

        doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ALLEGATO RISCHIO MMC")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = HEADER_BG

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Movimentazione Manuale dei Carichi - Azioni di Sollevamento")
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "ai sensi del Titolo VI D.Lgs. 81/2008 e UNI EN ISO 11228-1 (NIOSH)"
        )
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for _ in range(3):
            doc.add_paragraph("")

        ragione = (azienda.ragione_sociale or "—").upper()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ragione)
        run.bold = True
        run.font.size = Pt(18)

        sede_legale = format_sede(azienda, "legale")
        if sede_legale and sede_legale != "—":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(sede_legale)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        ident_bits = []
        if getattr(azienda, "partita_iva", None):
            ident_bits.append(f"P.IVA {azienda.partita_iva}")
        if getattr(azienda, "codice_ateco", None):
            ident_bits.append(f"ATECO {azienda.codice_ateco}")
        if ident_bits:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" · ".join(ident_bits))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for _ in range(3):
            doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Revisione {version:02d} - {generated_at.strftime('%d/%m/%Y')}"
        )
        run.bold = True
        run.font.size = Pt(12)

        page_break(doc)

    # ------------------------------------------------------------------
    # Revision history
    # ------------------------------------------------------------------

    def _add_revision_table(self, doc, generated_at: datetime, version: int) -> None:
        add_heading(doc, "Storico delle Revisioni", level=2)
        headers = ["Rev.", "Motivazione", "Data"]
        rows = [[f"{version:02d}", "Emissione" if version == 1 else "Aggiornamento",
                 generated_at.strftime("%d/%m/%Y")]]
        add_data_table(doc, headers, rows)
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # TOC
    # ------------------------------------------------------------------

    def _add_toc(self, doc) -> None:
        add_heading(doc, "Indice", level=1)

        p = doc.add_paragraph()
        run = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

        add_paragraph(
            doc,
            "(Premere Ctrl+A poi F9 per aggiornare l'indice)",
            italic=True,
            size=9,
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Introduzione (static narrative on NIOSH method)
    # ------------------------------------------------------------------

    def _add_introduzione(self, doc) -> None:
        add_heading(doc, "Introduzione", level=1)
        for txt in [
            "Le affezioni cronico-degenerative della colonna vertebrale rappresentano "
            "una delle principali cause di assenza per malattia nei lavoratori adulti. "
            "Il National Institute of Occupational Safety and Health (NIOSH - USA) "
            "pone tali patologie al secondo posto nella lista delle malattie "
            "professionali correlate al lavoro.",
            "Il Titolo VI del D.Lgs. 81/2008 (artt. 167-171 e Allegato XXXIII) impone "
            "al datore di lavoro la valutazione del rischio da movimentazione manuale "
            "dei carichi. Lo standard tecnico di riferimento per le azioni di "
            "sollevamento e' la norma UNI EN ISO 11228-1, che recepisce il modello "
            "NIOSH (revisione 1993).",
            "La presente valutazione applica la formula NIOSH per stimare, per ciascun "
            "compito di sollevamento svolto in azienda, il Peso Limite Raccomandato "
            "(PLR) e l'Indice di Rischio (IR), classificando ogni mansione nelle aree "
            "Verde (rischio trascurabile, IR <= 0,75), Gialla (sorveglianza, "
            "0,75 < IR <= 1,00) o Rossa (riprogettazione, IR > 1,00).",
        ]:
            add_paragraph(doc, txt)
        page_break(doc)

    # ------------------------------------------------------------------
    # Anagrafica Aziendale
    # ------------------------------------------------------------------

    def _add_anagrafica(self, doc, azienda) -> None:
        add_heading(doc, "Anagrafica Aziendale", level=1)
        rows: list[tuple[str, str]] = [
            ("Azienda", azienda.ragione_sociale or "—"),
            ("Attivita / Codice ATECO", getattr(azienda, "codice_ateco", "") or "—"),
            ("Partita IVA", getattr(azienda, "partita_iva", "") or "—"),
            ("Codice Fiscale", getattr(azienda, "codice_fiscale", "") or "—"),
            ("Sede Legale - Via", azienda.sede_legale_via or "—"),
            ("Sede Legale - Citta", format_comune(
                getattr(azienda, "cap_legale", None),
                azienda.sede_legale_citta,
                getattr(azienda, "provincia_legale", None))),
            ("Sede Operativa - Via", getattr(azienda, "sede_operativa_via", "") or "—"),
            ("Sede Operativa - Citta", format_comune(
                getattr(azienda, "cap_operativa", None),
                getattr(azienda, "sede_operativa_citta", None),
                getattr(azienda, "provincia_operativa", None))),
            ("Telefono", getattr(azienda, "telefono", "") or "—"),
            ("Email PEC", getattr(azienda, "email_pec", "") or "—"),
        ]
        add_kv_table(doc, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Dati Occupazionali (no codice fiscale - GDPR)
    # ------------------------------------------------------------------

    def _add_dati_occupazionali(self, doc, persone: list) -> None:
        add_heading(doc, "Dati Occupazionali", level=1)
        if not persone:
            add_paragraph(doc, "Nessun lavoratore registrato per questa azienda.", italic=True)
            page_break(doc)
            return

        headers = ["Nominativo", "Mansione", "Sesso", "Fascia eta", "Tipologia contrattuale"]
        rows = []
        for p in persone:
            rows.append([
                p.nominativo or "—",
                p.mansione or "—",
                p.sesso or "—",
                p.fascia_eta or "—",
                p.tipologia_contrattuale or "—",
            ])
        add_data_table(doc, headers, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Organizzazione Aziendale Sicurezza
    # ------------------------------------------------------------------

    def _add_organizzazione(self, doc, persone: list) -> None:
        add_heading(doc, "Organizzazione Aziendale della Sicurezza", level=1)

        def _names(predicate) -> str:
            matched = [p.nominativo for p in persone if predicate(p) and p.nominativo]
            return ", ".join(matched) if matched else "—"

        rows = [
            ("Datore di Lavoro", _names(lambda p: bool(p.ruolo_datore_lavoro))),
            ("RSPP", _names(lambda p: bool(p.ruolo_rspp))),
            ("RLS", _names(lambda p: bool(p.ruolo_rls))),
            ("Medico Competente", _names(lambda p: bool(p.ruolo_medico_competente))),
            ("Addetti Primo Soccorso", _names(lambda p: bool(p.ruolo_primo_soccorso))),
            ("Addetti Antincendio", _names(lambda p: bool(p.ruolo_antincendio))),
            ("Preposti", _names(lambda p: bool(p.ruolo_preposto))),
        ]
        add_kv_table(doc, rows)
        page_break(doc)

    # ------------------------------------------------------------------
    # Metodologia NIOSH (CP, A, B, C, D, E, F reference tables)
    # ------------------------------------------------------------------

    def _add_metodologia(self, doc) -> None:
        add_heading(doc, "Metodologia - Il metodo NIOSH per azioni di sollevamento", level=1)
        add_paragraph(
            doc,
            "Per ogni compito di sollevamento si calcola il Peso Limite Raccomandato (PLR) "
            "moltiplicando la costante di peso CP (kg) per i sei fattori demoltiplicativi "
            "A, B, C, D, E, F. L'Indice di Rischio IR si ottiene come rapporto tra il peso "
            "effettivamente sollevato P e il PLR.",
        )
        add_paragraph(doc, "PLR = CP x A x B x C x D x E x F", bold=True)
        add_paragraph(doc, "IR = P / PLR", bold=True)
        add_paragraph(
            doc,
            "Soglie di accettabilita: IR <= 0,75 = VERDE (rischio trascurabile); "
            "0,75 < IR <= 1,00 = GIALLO (sorveglianza sanitaria); "
            "IR > 1,00 = ROSSO (riprogettazione del compito).",
            italic=True,
        )

        # CP table
        add_heading(doc, "CP - Costante di Peso (kg)", level=2)
        add_data_table(
            doc,
            ["Eta", "Maschi", "Femmine"],
            [
                ["> 18 anni", "25", "20"],
                ["15-18 anni", "20", "15"],
            ],
        )

        # Factor A
        add_heading(doc, "Fattore A - Altezza da terra delle mani all'inizio (cm)", level=2)
        add_paragraph(
            doc,
            "Formula: A = 1 - 0,003 * |V - 75|. Valore ottimale V = 75 cm (altezza nocche).",
            italic=True,
            size=9,
        )
        add_data_table(
            doc,
            ["Altezza V (cm)", "Fattore A"],
            [[str(int(x)), f"{y:.2f}"] for x, y in FACTOR_A],
        )

        # Factor B
        add_heading(doc, "Fattore B - Dislocazione verticale (cm)", level=2)
        add_data_table(
            doc,
            ["Dislocazione X (cm)", "Fattore B"],
            [[str(int(x)), f"{y:.2f}"] for x, y in FACTOR_B],
        )

        # Factor C
        add_heading(doc, "Fattore C - Distanza orizzontale (cm)", level=2)
        add_data_table(
            doc,
            ["Distanza H (cm)", "Fattore C"],
            [[str(int(x)), f"{y:.2f}"] for x, y in FACTOR_C],
        )

        # Factor D
        add_heading(doc, "Fattore D - Angolo di asimmetria (gradi)", level=2)
        add_data_table(
            doc,
            ["Angolo (gradi)", "Fattore D"],
            [[str(int(x)), f"{y:.2f}"] for x, y in FACTOR_D],
        )

        # Factor E
        add_heading(doc, "Fattore E - Giudizio sulla presa", level=2)
        add_data_table(
            doc,
            ["Giudizio", "Fattore E"],
            [[k, f"{v:.2f}"] for k, v in FACTOR_E.items()],
        )

        # Factor F
        add_heading(doc, "Fattore F - Frequenza dei gesti per durata del lavoro", level=2)
        f_rows = []
        for atti, breve, media, lunga in FACTOR_F_TABLE:
            label = f">{int(FACTOR_F_TABLE[-2][0])}" if atti >= 16 else (
                f"{atti:g}" if atti < 1 else str(int(atti))
            )
            f_rows.append([label, f"{breve:.2f}", f"{media:.2f}", f"{lunga:.2f}"])
        add_data_table(
            doc,
            [
                "Frequenza (atti/min)",
                "Breve durata (<1 ora)",
                "Media durata (1-2 ore)",
                "Lunga durata (>2 ore)",
            ],
            f_rows,
        )

        page_break(doc)

    # ------------------------------------------------------------------
    # Per-worker assessment grid (one 13x5 table per task, like template T14)
    # ------------------------------------------------------------------

    def _add_per_worker_assessments(
        self,
        doc,
        mmc_rows: list,
        persona_by_id: dict,
        ambiente_by_id: dict,
    ) -> None:
        add_heading(doc, "Tavole di Valutazione del Rischio MMC", level=1)
        if not mmc_rows:
            add_paragraph(
                doc,
                "Nessuna attivita di movimentazione manuale dei carichi e' stata "
                "valutata per questa azienda.",
                italic=True,
            )
            page_break(doc)
            return

        for i, r in enumerate(mmc_rows, 1):
            persona = persona_by_id.get(r.persona_id) if r.persona_id else None
            ambiente = ambiente_by_id.get(r.ambiente_id) if r.ambiente_id else None
            self._render_assessment_table(doc, i, r, persona, ambiente)
            page_break(doc)

    def _render_assessment_table(self, doc, idx: int, r, persona, ambiente) -> None:
        nominativo = (persona.nominativo if persona else "—") or "—"
        mansione = (persona.mansione if persona else "—") or "—"

        add_heading(doc, f"{idx}. {r.compito or 'Compito'}", level=2)
        if ambiente:
            add_paragraph(doc, f"Ambiente: {ambiente.nome or '—'}", italic=True, size=10)

        # Server-side recompute guard: if all 7 inputs are present, derive
        # multipliers + PLR + IR from the canonical lookup tables. This
        # protects against stored derivatives drifting out of sync with the
        # inputs (e.g. an old fixture that was never resaved through the API).
        # When inputs are absent we fall back to whatever the row has stored,
        # which is the legacy contract.
        recomputed = self._recompute_if_inputs_present(r)

        cp = float(recomputed.get("cp") if recomputed else r.cp or 25.0)
        peso = float(r.peso_kg or 0)
        plr = float(recomputed.get("plr") if recomputed else (r.plr or 0))
        ir = float(recomputed.get("ir") if recomputed else (r.indice_ir or 0))
        livello = (
            recomputed.get("livello") if recomputed
            else (r.livello_rischio or classify_ir(ir)).upper()
        )

        # Multipliers — recomputed values win over stored when both exist.
        mult_a = recomputed.get("fattore_a") if recomputed else r.fattore_a
        mult_b = recomputed.get("fattore_b") if recomputed else r.fattore_b
        mult_c = recomputed.get("fattore_c") if recomputed else r.fattore_c
        mult_d = recomputed.get("fattore_d") if recomputed else r.fattore_d
        mult_e = recomputed.get("fattore_e") if recomputed else r.fattore_e
        mult_f = recomputed.get("fattore_f") if recomputed else r.fattore_f

        # 13-row x 4-col table mirroring the template T14 layout:
        #   col0 = Fattore (CP/A/B/C/D/E/F/PLR/P/IR)
        #   col1 = Descrizione
        #   col2 = Input (cm/gradi/giudizio/atti per minuto/durata)
        #   col3 = Valore (multiplier or final number)
        table = doc.add_table(rows=0, cols=4)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header band: nominativo + mansione (merged-look via header style)
        header_row = table.add_row().cells
        header_row[0].text = "Nominativo"
        header_row[1].text = nominativo
        header_row[2].text = "Mansione"
        header_row[3].text = mansione
        for cell in header_row:
            shade_cell(cell, "1A237E")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)

        sex_label = (
            "Maschi > 18 anni" if (r.sesso or "M") == "M" and (r.fascia_eta or "").startswith(">")
            else "Femmine > 18 anni" if (r.sesso or "M") == "F" and (r.fascia_eta or "").startswith(">")
            else "Maschi 15-18" if (r.sesso or "M") == "M"
            else "Femmine 15-18"
        )

        def _cell_str(val) -> str:
            if val is None:
                return "—"
            return str(val)

        def _factor_str(val) -> str:
            if val is None:
                return "—"
            try:
                return f"{float(val):.2f}"
            except (TypeError, ValueError):
                return "—"

        body_rows = [
            ("CP", "Costante di peso (kg)", sex_label, f"{cp:.0f}"),
            ("A", "Altezza da terra delle mani all'inizio del sollevamento (cm)",
             _cell_str(r.altezza_cm), _factor_str(mult_a)),
            ("B", "Dislocazione verticale del peso fra inizio e fine (cm)",
             _cell_str(r.dislocazione_cm), _factor_str(mult_b)),
            ("C", "Distanza orizzontale tra mani e mezzo delle caviglie (cm)",
             _cell_str(r.distanza_cm), _factor_str(mult_c)),
            ("D", "Angolo di asimmetria del peso (gradi)",
             _cell_str(r.angolo_gradi), _factor_str(mult_d)),
            ("E", "Giudizio sulla presa del carico",
             _cell_str(r.giudizio_presa), _factor_str(mult_e)),
            ("F", "Frequenza dei gesti (atti/min)",
             _cell_str(r.frequenza_atti_min), _factor_str(mult_f)),
            ("", "Durata del lavoro (minuti)",
             _cell_str(r.durata_min), ""),
            ("PLR", "Peso Limite Raccomandato",
             "", f"{plr:.2f} kg" if plr > 0 else "—"),
            ("P", "Peso effettivamente sollevato",
             "", f"{peso:.1f} kg"),
            ("IR = P/PLR", "Indice di rischio",
             "", f"{ir:.2f}" if plr > 0 else "—"),
        ]

        for code, descr, input_val, multiplier in body_rows:
            row = table.add_row().cells
            row[0].text = code
            row[1].text = descr
            row[2].text = str(input_val)
            row[3].text = str(multiplier)
            shade_cell(row[0], "F5F5F5")
            for p in row[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for cell in (row[1], row[2], row[3]):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # Final result row, color-shaded by zone
        result_row = table.add_row().cells
        result_row[0].text = "Risultato"
        result_row[1].text = "Classificazione"
        result_row[2].text = livello
        result_row[3].text = (
            "Verde" if livello == "VERDE"
            else "Gialla" if livello == "GIALLO"
            else "Rossa" if livello == "ROSSO"
            else "—"
        )
        area = result_row[3].text
        shade_hex = _AREA_COLORS.get(area, "F5F5F5")
        for cell in result_row:
            shade_cell(cell, shade_hex)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        if r.note:
            add_paragraph(doc, f"Note: {r.note}", italic=True, size=9)

    # ------------------------------------------------------------------
    # Quadro sinottico
    # ------------------------------------------------------------------

    def _add_quadro_sinottico(self, doc, mmc_rows: list, persona_by_id: dict) -> None:
        add_heading(doc, "Quadro sinottico di esposizione", level=1)
        if not mmc_rows:
            add_paragraph(doc, "Nessuna valutazione presente.", italic=True)
            page_break(doc)
            return

        # Worst-case per persona — using recomputed IR when inputs allow it,
        # so the synopsis matches what the per-worker grid actually shows.
        worst_by_persona: dict = {}
        ir_by_row: dict = {}
        area_by_row: dict = {}
        for r in mmc_rows:
            recomp = self._recompute_if_inputs_present(r)
            ir = float(recomp["ir"]) if recomp else float(r.indice_ir or 0)
            area = (
                recomp["area"] if recomp
                else (r.area_classificazione or self._area_from_livello(r.livello_rischio))
            )
            ir_by_row[id(r)] = ir
            area_by_row[id(r)] = area
            key = r.persona_id or f"_anon_{id(r)}"
            cur = worst_by_persona.get(key)
            if cur is None or ir > ir_by_row[id(cur)]:
                worst_by_persona[key] = r

        table = doc.add_table(rows=1, cols=4)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        for i, h in enumerate(["Nominativo", "Mansione", "IR", "Area"]):
            hdr.cells[i].text = h
        style_header_row(hdr)

        for key, r in worst_by_persona.items():
            persona = persona_by_id.get(r.persona_id) if r.persona_id else None
            nominativo = (persona.nominativo if persona else "—") or "—"
            mansione = (persona.mansione if persona else "—") or "—"
            ir = ir_by_row[id(r)]
            area = area_by_row[id(r)]

            row = table.add_row()
            row.cells[0].text = nominativo
            row.cells[1].text = mansione
            row.cells[2].text = f"{ir:.2f}"
            row.cells[3].text = area or "—"

            # Color the IR + Area cells by zone
            shade_hex = _AREA_COLORS.get(area or "", "FFFFFF")
            for cell in (row.cells[2], row.cells[3]):
                shade_cell(cell, shade_hex)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            for cell in (row.cells[0], row.cells[1]):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        page_break(doc)

    @staticmethod
    def _recompute_if_inputs_present(r) -> dict | None:
        """Return derived multipliers + PLR + IR + livello when all 7 NIOSH
        inputs are present on the row; otherwise None.

        Used as a generation-time guard so the document never displays stored
        PLR/IR values that have drifted out of sync with the inputs (e.g. when
        a fixture was hand-edited or the row predates a factor-table fix).
        """
        inputs = (
            r.altezza_cm,
            r.dislocazione_cm,
            r.distanza_cm,
            r.angolo_gradi,
            r.giudizio_presa,
            r.frequenza_atti_min,
            r.durata_min,
        )
        if any(v is None for v in inputs):
            return None
        cp = float(r.cp or 25.0)
        out = compute_plr(
            cp=cp,
            altezza_cm=float(r.altezza_cm),
            dislocazione_cm=float(r.dislocazione_cm),
            distanza_cm=float(r.distanza_cm),
            angolo_gradi=float(r.angolo_gradi),
            giudizio_presa=str(r.giudizio_presa),
            frequenza_atti_min=float(r.frequenza_atti_min),
            durata_min=float(r.durata_min),
        )
        peso = float(r.peso_kg or 0)
        plr = float(out["plr"])
        ir = peso / plr if plr > 0 else 0.0
        livello = classify_ir(ir)
        return {
            "cp": cp,
            "fattore_a": out["fattore_a"],
            "fattore_b": out["fattore_b"],
            "fattore_c": out["fattore_c"],
            "fattore_d": out["fattore_d"],
            "fattore_e": out["fattore_e"],
            "fattore_f": out["fattore_f"],
            "plr": plr,
            "ir": round(ir, 4),
            "livello": livello,
            "area": {"VERDE": "Verde", "GIALLO": "Gialla", "ROSSO": "Rossa"}[livello],
        }

    @staticmethod
    def _area_from_livello(livello: str | None) -> str:
        if not livello:
            return "—"
        u = livello.upper()
        if u == "VERDE":
            return "Verde"
        if u == "GIALLO":
            return "Gialla"
        if u == "ROSSO":
            return "Rossa"
        return "—"

    # ------------------------------------------------------------------
    # Programma di Attuazione delle Misure
    # ------------------------------------------------------------------

    def _add_programma_attuazione(
        self, doc, mmc_rows: list, persona_by_id: dict
    ) -> None:
        add_heading(doc, "Programma di Attuazione delle Misure di Prevenzione", level=1)
        if not mmc_rows:
            add_paragraph(doc, "Nessuna valutazione presente.", italic=True)
            page_break(doc)
            return

        zones: dict[str, list] = {"Verde": [], "Gialla": [], "Rossa": []}
        for r in mmc_rows:
            recomp = self._recompute_if_inputs_present(r)
            area = (
                recomp["area"] if recomp
                else (r.area_classificazione or self._area_from_livello(r.livello_rischio))
            )
            if area in zones:
                zones[area].append(r)

        for zone, rows in zones.items():
            add_heading(doc, _ZONE_HEADINGS.get(zone, zone), level=2)
            if not rows:
                add_paragraph(
                    doc,
                    f"Nessun lavoratore classificato in {zone.lower()}.",
                    italic=True,
                    size=10,
                )
                continue

            # Table: lavoratore | compito | misure
            data_rows = []
            for r in rows:
                persona = persona_by_id.get(r.persona_id) if r.persona_id else None
                nome = (persona.nominativo if persona else "—") or "—"
                misure = r.misure_proposte or _DEFAULT_MEASURES_BY_ZONE.get(zone, "")
                data_rows.append([nome, r.compito or "—", misure])

            tbl = add_data_table(
                doc,
                ["Lavoratore", "Compito", "Misure di prevenzione e protezione"],
                data_rows,
            )
            # Tint every body cell so the zone is visually obvious even when
            # the table spans a page break (bullet from audit P3 #15).
            tint = _AREA_COLORS.get(zone, "F5F5F5")
            for row in tbl.rows[1:]:
                for cell in row.cells:
                    shade_cell(cell, tint)

        page_break(doc)

    # ------------------------------------------------------------------
    # Dichiarazione del Datore di Lavoro
    # ------------------------------------------------------------------

    def _add_dichiarazione_ddl(self, doc, azienda, persone) -> None:
        add_heading(doc, "Dichiarazione del Datore di Lavoro", level=1)
        ddl_names = [p.nominativo for p in persone if p.ruolo_datore_lavoro and p.nominativo]
        ddl = ddl_names[0] if ddl_names else "il Datore di Lavoro"
        ragione = azienda.ragione_sociale or "l'Azienda"
        add_paragraph(
            doc,
            f"Il sottoscritto {ddl}, in qualita di Datore di Lavoro di {ragione}, "
            "dichiara di aver effettuato la valutazione del rischio da movimentazione "
            "manuale dei carichi ai sensi del Titolo VI del D.Lgs. 81/2008 e della "
            "norma UNI EN ISO 11228-1 (metodo NIOSH), in collaborazione con il "
            "Responsabile del Servizio di Prevenzione e Protezione, il Medico "
            "Competente ove nominato e previa consultazione del Rappresentante dei "
            "Lavoratori per la Sicurezza.",
        )
        add_paragraph(
            doc,
            "Le misure di prevenzione e protezione individuate nel Programma di "
            "Attuazione saranno adottate secondo il cronoprogramma concordato e "
            "verificate periodicamente. La presente valutazione sara aggiornata in "
            "occasione di modifiche significative del processo lavorativo, "
            "dell'organizzazione del lavoro o di insorgenza di patologie correlate.",
        )
        page_break(doc)

    # ------------------------------------------------------------------
    # Signature block
    # ------------------------------------------------------------------

    def _add_signature_block(self, doc, persone) -> None:
        add_heading(doc, "Firme", level=1)

        def _first_or_dash(predicate) -> str:
            for p in persone:
                if predicate(p) and p.nominativo:
                    return p.nominativo
            return "—"

        ddl = _first_or_dash(lambda p: bool(p.ruolo_datore_lavoro))
        rspp = _first_or_dash(lambda p: bool(p.ruolo_rspp))
        mc = _first_or_dash(lambda p: bool(p.ruolo_medico_competente))
        rls = _first_or_dash(lambda p: bool(p.ruolo_rls))

        # 2x2 signature grid
        table = doc.add_table(rows=2, cols=2)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        cells = [
            (0, 0, "Il Datore di Lavoro", ddl),
            (0, 1, "Il Responsabile del S.P.P.", rspp),
            (1, 0, "Il Medico Competente", mc),
            (1, 1, "Il Rappresentante dei Lavoratori (per consultazione)", rls),
        ]
        for r, c, label, name in cells:
            cell = table.rows[r].cells[c]
            cell.text = ""
            p1 = cell.paragraphs[0]
            run = p1.add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            p2 = cell.add_paragraph(f"({name})")
            for run in p2.runs:
                run.font.size = Pt(10)
                run.italic = True
            cell.add_paragraph("")
            cell.add_paragraph("__________________________")
