"""Allegato Stress Lavoro-Correlato — metodologia INAIL."""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_stress
from app.services.document_generator.docx_utils import (
    TEMPLATES_DIR,
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    page_break,
    replace_placeholders,
    slugify,
)
from app.services.stress_calculator import (
    FINAL_THRESHOLDS,
    TOTALE_B_THRESHOLDS,
    TOTALE_C_THRESHOLDS,
    _area_a_converted,
    _azione_per_livello,
    _band,
    get_default_measures,
)

TEMPLATE = TEMPLATES_DIR / "ALLEGATO STRESS DA LAVORO CORRELATO.docx"
TIPO_DOC = "allegato_stress"


class AllegatoStressGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        stress = await load_stress(self.db, self.azienda_id)

        if TEMPLATE.exists():
            doc = Document(str(TEMPLATE))
            replace_placeholders(doc, {"RAGIONE SOCIALE": azienda.ragione_sociale or "", "[AZIENDA]": azienda.ragione_sociale or ""})
        else:
            doc = Document()

        page_break(doc)
        add_heading(doc, f"VALUTAZIONE SPECIFICA - {azienda.ragione_sociale}", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Gruppo omogeneo", stress.gruppo_omogeneo if stress else "N/D"),
            ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
            ("Metodologia", "INAIL 2011 - 3 aree: A Indicatori Aziendali, B Contesto del lavoro, C Contenuto del lavoro"),
        ])

        add_heading(doc, "Inquadramento normativo", level=2)
        add_paragraph(doc, "Art. 28 comma 1-bis del D.Lgs. 81/2008 prevede la valutazione del rischio stress lavoro-correlato secondo le indicazioni della Commissione Consultiva Permanente (metodologia INAIL 2011).")

        if stress:
            # INAIL bands per REFERENCE_DATA §3.4:
            #   Area A: raw 0-40 → converted 0/2/5 with band BASSO/MEDIO/ALTO
            #   Area B (Contesto, max 26): BASSO 0-8 / MEDIO 9-17 / ALTO 18-26
            #   Area C (Contenuto, max 36): BASSO 0-13 / MEDIO 14-25 / ALTO 26-36
            #   Total: BASSO 0-17 / MEDIO 18-34 / ALTO 35-67
            a_raw = stress.punteggio_a or 0
            b_raw = stress.punteggio_b or 0
            c_raw = stress.punteggio_c or 0
            a_conv, a_livello = _area_a_converted(a_raw)
            b_livello = _band(max(b_raw, 0), TOTALE_B_THRESHOLDS)
            c_livello = _band(max(c_raw, 0), TOTALE_C_THRESHOLDS)
            totale = stress.punteggio_totale if stress.punteggio_totale is not None else (a_conv + b_raw + c_raw)
            livello_final = stress.livello_rischio or _band(max(totale, 0), FINAL_THRESHOLDS)

            add_heading(doc, "Punteggi per area", level=2)
            add_data_table(doc, ["Area", "Punteggio grezzo", "Convertito", "Livello parziale"], [
                ["A - Indicatori aziendali (infortuni, assenze, turnover)", f"{a_raw} / 40", str(a_conv), a_livello],
                ["B - Contesto del lavoro (organizzazione, ruoli, rapporti)", f"{b_raw} / 26", str(max(b_raw, 0)), b_livello],
                ["C - Contenuto del lavoro (ambiente, compiti, ritmo, orario)", f"{c_raw} / 36", str(c_raw), c_livello],
            ])
            add_paragraph(
                doc,
                "Nota: il punteggio dell'Area A è convertito sulla scala 0/2/5 prima di essere sommato a B e C (max teorico totale = 67). Soglie: 0-17 BASSO, 18-34 MEDIO, 35-67 ALTO.",
                italic=True,
                size=9,
            )

            add_heading(doc, "Esito complessivo", level=2)
            add_kv_table(doc, [
                ("Punteggio totale (A conv + B + C)", f"{totale} / 67"),
                ("Livello di rischio", livello_final),
                ("Azione conseguente", _azione_per_livello(livello_final)),
            ])

            add_heading(doc, "Misure correttive", level=2)
            if stress.misure_correttive:
                add_paragraph(doc, stress.misure_correttive)
            else:
                for m in get_default_measures(livello_final):
                    add_paragraph(doc, f"• {m}")

            add_heading(doc, "Dettaglio indicatori per area", level=2)
            # NB: model field names are swapped vs INAIL labels (area_b_contenuto_lavoro
            # actually holds Contesto answers and vice versa). Show under the
            # correct INAIL heading regardless of column name.
            _area_detail(doc, "Area A - Indicatori Aziendali", stress.area_a_eventi_sentinella or {})
            _area_detail(doc, "Area B - Contesto del lavoro", stress.area_b_contenuto_lavoro or {})
            _area_detail(doc, "Area C - Contenuto del lavoro", stress.area_c_contesto_lavoro or {})
        else:
            add_paragraph(doc, "Nessuna valutazione stress lavoro-correlato disponibile per questa azienda.", italic=True)

        add_heading(doc, "Sottoscrizione", level=2)
        add_data_table(doc, ["Ruolo", "Nominativo", "Firma"], [
            ["Datore di Lavoro", azienda.ragione_sociale or "", "________________________"],
            ["RSPP", "________________________", "________________________"],
            ["Medico Competente", "________________________", "________________________"],
            ["RLS", "________________________", "________________________"],
            ["Data", generated_at.strftime("%d/%m/%Y"), ""],
        ])

        version = await self._next_version()
        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        stmt = (
            select(func.coalesce(func.max(DocumentoGenerato.versione), 0))
            .where(DocumentoGenerato.azienda_id == self.azienda_id)
            .where(DocumentoGenerato.tipo_documento.in_([TIPO_DOC, "ALLEGATO_STRESS"]))
        )
        r = await self.db.execute(stmt)
        return (r.scalar() or 0) + 1


def _area_detail(doc, title: str, payload: dict) -> None:
    add_heading(doc, title, level=3)
    if not payload:
        add_paragraph(doc, "Nessun dato registrato.", italic=True, size=9)
        return
    rows = [[str(k), str(v)] for k, v in payload.items()]
    add_data_table(doc, ["Indicatore", "Risposta"], rows)
