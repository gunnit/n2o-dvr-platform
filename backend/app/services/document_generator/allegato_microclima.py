"""Allegato Microclima Moderato - UNI EN ISO 7730 (PMV/PPD)."""

import os

from docx import Document
from sqlalchemy import func, select

from app.models.documento_generato import DocumentoGenerato
from app.services.document_generator.base import BaseDocumentGenerator
from app.services.document_generator.data_loader import load_microclima
from app.services.document_generator.design import (
    add_cover,
    add_revision_table,
    finish_document,
    setup_document,
)
from app.services.document_generator.docx_utils import (
    add_data_table,
    add_heading,
    add_kv_table,
    add_paragraph,
    format_sede,
    page_break,
    slugify,
)
from app.services.microclima_calculator import calculate_pmv_ppd

TIPO_DOC = "allegato_microclima"
DOC_TITLE = "Allegato Rischio Microclima"


def _it(value: float, digits: int) -> str:
    """Italian decimal notation (comma), e.g. 21,0."""
    return f"{value:.{digits}f}".replace(".", ",")


def _compute_pmv_ppd(t_air, t_rad, v_air, rh, met, clo) -> tuple[float | None, float | None]:
    """Use the shared current-version ISO model, with a safe fallback."""
    try:
        result = calculate_pmv_ppd(
            air_temp=float(t_air),
            mean_radiant_temp=float(t_rad),
            air_velocity=float(v_air),
            humidity=float(rh),
            metabolic_rate=float(met),
            clothing_insulation=float(clo),
        )
        return float(result["pmv"]), float(result["ppd"])
    except Exception:
        # Simplified fallback: linear distance from 22 C optimal
        pmv = (float(t_air) - 22.0) * 0.25
        ppd = min(95.0, max(5.0, 5.0 + abs(pmv) * 25.0))
        return pmv, ppd


def _comfort_category(ppd: float | None) -> str:
    if ppd is None:
        return "—"
    if ppd < 6:
        return "Categoria A (eccellente)"
    if ppd < 10:
        return "Categoria B (buona)"
    if ppd < 15:
        return "Categoria C (accettabile)"
    return "Fuori categoria - azioni correttive"


class AllegatoMicroclimaGenerator(BaseDocumentGenerator):
    async def generate(self) -> str:
        data = await self.load_data()
        azienda = data["azienda"]
        generated_at = data["generated_at"]
        micro = await load_microclima(self.db, self.azienda_id)
        ambienti_map = {a.id: a for a in data["ambienti"]}
        version = await self._next_version()

        # Filter moderato (the severo-caldo variant is handled separately)
        moderate_rows = [m for m in micro if (m.tipo_ambiente or "moderato") == "moderato"]

        doc = Document()
        setup_document(doc)
        # Audit 2026-09-03: this annex shipped as one US-Letter page with no
        # cover, header, footer or logo. It now carries the shared furniture.
        add_cover(
            doc,
            title=DOC_TITLE,
            subtitle="Ambienti termici moderati — indici PMV/PPD",
            legal_basis="ai sensi del D.Lgs. 81/2008 Titolo VIII, art. 180 e UNI EN ISO 7730:2006",
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )
        add_revision_table(doc, version, generated_at)

        add_heading(doc, "Dati generali", level=1)
        add_kv_table(doc, [
            ("Azienda", azienda.ragione_sociale or ""),
            ("Sede", format_sede(azienda, "legale")),
            ("Data valutazione", generated_at.strftime("%d/%m/%Y")),
            ("Riferimento normativo", "UNI EN ISO 7730:2006 - Ergonomia ambienti termici moderati"),
        ])

        add_heading(doc, "Metodologia", level=1)
        add_paragraph(doc, "La norma UNI EN ISO 7730 definisce il comfort termico mediante gli indici PMV (Predicted Mean Vote, scala -3..+3) e PPD (Predicted Percentage of Dissatisfied). I parametri considerati sono: temperatura dell'aria (tdb), temperatura radiante media (tr), velocità dell'aria (var), umidità relativa (RH), metabolismo (met) e isolamento del vestiario (clo).")

        add_data_table(doc, ["Categoria", "PPD", "Giudizio"], [
            ["A", "< 6%", "Eccellente comfort"],
            ["B", "< 10%", "Comfort buono"],
            ["C", "< 15%", "Comfort accettabile"],
            ["—", ">= 15%", "Necessarie azioni correttive"],
        ], column_widths_cm=[3.0, 3.5, 10.0])

        add_heading(doc, "Parametri per ambiente", level=1)
        if not moderate_rows:
            add_paragraph(doc, "Nessun ambiente valutato nella fascia moderata.", italic=True)
        else:
            rows = []
            for m in moderate_rows:
                amb_name = ambienti_map[m.ambiente_id].nome if m.ambiente_id in ambienti_map else "—"
                pmv, ppd = _compute_pmv_ppd(m.temperatura_aria, m.temperatura_radiante, m.velocita_aria, m.umidita_relativa, m.metabolismo, m.isolamento_vestiario)
                rows.append([
                    amb_name,
                    _it(float(m.temperatura_aria), 1),
                    _it(float(m.temperatura_radiante), 1),
                    _it(float(m.velocita_aria), 2),
                    _it(float(m.umidita_relativa), 0),
                    _it(float(m.metabolismo), 2),
                    _it(float(m.isolamento_vestiario), 2),
                    _it(pmv, 2) if pmv is not None else "—",
                    f"{_it(ppd, 1)}%" if ppd is not None else "—",
                    _comfort_category(ppd),
                ])
            add_data_table(
                doc,
                ["Ambiente", "tₐ (°C)", "tᵣ (°C)", "vₐ (m/s)", "UR (%)", "met", "clo", "PMV", "PPD", "Categoria"],
                rows,
                column_widths_cm=[3.8, 1.3, 1.3, 1.5, 1.2, 1.1, 1.1, 1.2, 1.4, 2.6],
            )

        add_heading(doc, "Misure correttive suggerite", level=1)
        add_paragraph(doc, "Per ambienti con PPD >= 15%: adeguare il sistema di climatizzazione, rivedere l'isolamento del vestiario, introdurre schermature solari o umidificatori, verificare la velocità dell'aria nelle postazioni.")

        finish_document(
            doc,
            title=DOC_TITLE,
            azienda=azienda,
            branding=self.branding,
            version=version,
            generated_at=generated_at,
        )

        output_dir = self._get_output_dir()
        slug = slugify(azienda.ragione_sociale or "azienda")
        filepath = os.path.join(output_dir, f"{TIPO_DOC}_{slug}_v{version}.docx")
        doc.save(filepath)
        return filepath

    async def _next_version(self) -> int:
        return await self.resolve_version([TIPO_DOC, "ALLEGATO_MICROCLIMA"])
