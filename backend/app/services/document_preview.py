"""In-browser preview + inline editing support for generated .docx files.

Parses a document body into a JSON-serialisable block model (consumed by the
frontend preview) and applies per-paragraph plain-text overrides back onto
the original bytes (consumed by the download and save-edited-version
endpoints). The .docx stays the source of truth — overrides only ever
replace run text, never styles, tables, images or headers/footers.

Address scheme (shared contract with the frontend):
- Top-level block: the index over the body's direct CT_P / CT_Tbl children
  in document order, skipping everything else (sectPr, bookmarks) — "12".
- Cell paragraph: "{table}:{row}:{cell}:{para}" where row enumerates the
  table's direct ``w:tr`` elements, cell enumerates the row's actual
  ``w:tc`` XML children (NOT python-docx ``row.cells``, which repeats
  merged cells), and para enumerates the cell's direct ``w:p`` children.

Parse and apply share the same private enumeration helpers below so the
addresses always agree between the two directions.
"""

from __future__ import annotations

import io
import logging
import mimetypes
import posixpath
import re
import zipfile
from copy import deepcopy
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)

# EMU per inch / CSS reference pixels per inch — used to convert wp:extent
# (EMU) into the px dimensions the browser preview renders images at.
_EMU_PER_INCH = 914400
_PX_PER_INCH = 96

# Heading level from the paragraph style name. Generators use python-docx's
# built-in "Heading N" styles (docx_utils.add_heading), but templates saved
# from Italian-localised Word carry "Titolo N" as the style name instead —
# accept both, levels 1-4 per the contract.
_HEADING_STYLE_RE = re.compile(r"^(?:heading|titolo)\s+([1-4])$", re.IGNORECASE)

_ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


# ---------------------------------------------------------------------------
# Shared enumeration helpers — the single source of truth for addresses.
# Both parse_docx_to_blocks and apply_overrides_to_docx MUST go through
# these so an address handed out by parse always resolves the same element
# on apply.
# ---------------------------------------------------------------------------


def _iter_body_blocks(body):
    """Yield ``(index, element)`` for the body's direct CT_P/CT_Tbl children.

    The index over THIS filtered enumeration is the top-level block address.
    Everything else (w:sectPr, bookmarks, comments) is skipped and does not
    consume an index.
    """
    idx = 0
    for child in body.iterchildren():
        if isinstance(child, (CT_P, CT_Tbl)):
            yield idx, child
            idx += 1


def _row_elements(tbl_el) -> list:
    """Direct ``w:tr`` children of a table (nested tables never leak in)."""
    return tbl_el.findall(qn("w:tr"))


def _cell_elements(tr_el) -> list:
    """Direct ``w:tc`` children of a row — the ACTUAL XML cells.

    Deliberately not python-docx ``row.cells``: that API repeats merged
    cells to normalise the grid, which would make addresses ambiguous.
    """
    return tr_el.findall(qn("w:tc"))


def _cell_paragraph_elements(tc_el) -> list:
    """Direct ``w:p`` children of a cell (nested-table paragraphs excluded)."""
    return tc_el.findall(qn("w:p"))


def _paragraph_is_locked(p_el) -> bool:
    """True when the paragraph must not be text-replaced.

    Drawings/pictures live inside runs we'd delete, and fldChar/instrText
    are field machinery (TOC, page refs) that Word re-evaluates — replacing
    their runs would corrupt the field. w:hyperlink and w:fldSimple wrap
    runs the removal loop would destroy (link target and field code lost
    for good), and a run-level page/column break would silently vanish —
    so those paragraphs stay visible but read-only too. Same rule on parse
    (editable=False) and on apply (skip silently).
    """
    return bool(
        p_el.xpath(
            ".//w:drawing | .//w:pict | .//w:fldChar | .//w:instrText"
            " | .//w:hyperlink | .//w:fldSimple"
            ' | .//w:br[@w:type="page" or @w:type="column"]'
        )
    )


# ---------------------------------------------------------------------------
# Parse: .docx bytes -> block list
# ---------------------------------------------------------------------------


def _paragraph_run_elements(p_el) -> list:
    """All ``w:r`` elements of a paragraph in visual order.

    ``Paragraph.runs`` only yields direct ``w:r`` children, which makes text
    inside ``w:hyperlink`` / ``w:fldSimple`` wrappers (e.g. "www.inail.it")
    invisible in the preview. Walk the paragraph's direct children in
    document order and descend into those wrapper elements so the full text
    is shown — such paragraphs are locked (see _paragraph_is_locked), so
    visibility never turns into editability.
    """
    wrapper_tags = (qn("w:hyperlink"), qn("w:fldSimple"))
    run_tag = qn("w:r")
    runs: list = []
    for child in p_el.iterchildren():
        if child.tag == run_tag:
            runs.append(child)
        elif child.tag in wrapper_tags:
            runs.extend(child.iter(run_tag))
    return runs


def _run_payload(run) -> dict:
    """Serialise one run. color/size come from run.font — None when inherited."""
    font = run.font
    color = None
    try:
        if font.color is not None and font.color.rgb is not None:
            color = str(font.color.rgb)
    except (AttributeError, ValueError):  # theme colors without an RGB value
        color = None
    return {
        "text": run.text,
        "bold": bool(font.bold),
        "italic": bool(font.italic),
        "underline": bool(font.underline),
        "color": color,
        "size": float(font.size.pt) if font.size is not None else None,
    }


def _image_payloads(p_el) -> list[dict]:
    """Inline images in the paragraph's runs, as rId + px dimensions."""
    images: list[dict] = []
    for drawing in p_el.xpath(".//w:drawing"):
        rids = drawing.xpath(".//a:blip/@r:embed")
        if not rids:
            continue
        width_px = height_px = None
        extents = drawing.xpath(".//wp:extent")
        if extents:
            try:
                width_px = round(int(extents[0].get("cx")) / _EMU_PER_INCH * _PX_PER_INCH)
                height_px = round(int(extents[0].get("cy")) / _EMU_PER_INCH * _PX_PER_INCH)
            except (TypeError, ValueError):
                width_px = height_px = None
        images.append({"image_id": rids[0], "width_px": width_px, "height_px": height_px})
    return images


def _paragraph_payload(p_el, addr: str, parent, *, in_nested_table: bool = False) -> dict:
    """Common paragraph fields shared by top-level blocks and cell paragraphs."""
    para = Paragraph(p_el, parent)
    style_name = None
    try:
        style_name = para.style.name if para.style is not None else None
    except Exception:  # defensive: malformed styles part must not kill preview
        style_name = None
    heading_level = None
    if style_name:
        match = _HEADING_STYLE_RE.match(style_name)
        if match:
            heading_level = int(match.group(1))
    return {
        "addr": addr,
        "style": style_name,
        "heading_level": heading_level,
        "alignment": _ALIGNMENT_MAP.get(para.alignment),
        "editable": not in_nested_table and not _paragraph_is_locked(p_el),
        "runs": [
            _run_payload(Run(r_el, para))
            for r_el in _paragraph_run_elements(p_el)
        ],
        "images": _image_payloads(p_el),
    }


def _paragraph_has_page_break(p_el, para: Paragraph) -> bool:
    if para.paragraph_format.page_break_before:
        return True
    return bool(p_el.xpath('.//w:br[@w:type="page"]'))


def _cell_payload(tc_el, cell_addr: str, parent) -> dict:
    paragraphs: list[dict] = []
    para_idx = 0
    for p_el in _cell_paragraph_elements(tc_el):
        paragraphs.append(
            _paragraph_payload(p_el, f"{cell_addr}:{para_idx}", parent)
        )
        para_idx += 1

    # Graceful degradation for nested tables: no recursion — flatten their
    # text into extra read-only paragraphs appended to the cell. Addresses
    # continue the paraIdx sequence, so they can never collide with the
    # direct w:p addresses the apply side resolves, and editable=False
    # keeps them locked regardless.
    for nested_tbl in tc_el.findall(qn("w:tbl")):
        for nested_p in nested_tbl.iter(qn("w:p")):
            text = "".join(t.text or "" for t in nested_p.iter(qn("w:t")))
            paragraphs.append(
                {
                    "addr": f"{cell_addr}:{para_idx}",
                    "style": None,
                    "heading_level": None,
                    "alignment": None,
                    "editable": False,
                    "runs": (
                        [
                            {
                                "text": text,
                                "bold": False,
                                "italic": False,
                                "underline": False,
                                "color": None,
                                "size": None,
                            }
                        ]
                        if text
                        else []
                    ),
                    "images": [],
                }
            )
            para_idx += 1

    shading = None
    col_span = 1
    v_merge = None
    tc_pr = tc_el.find(qn("w:tcPr"))
    if tc_pr is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            if fill and fill.lower() != "auto":
                shading = fill.upper()
        grid_span = tc_pr.find(qn("w:gridSpan"))
        if grid_span is not None:
            try:
                col_span = int(grid_span.get(qn("w:val")) or 1)
            except (TypeError, ValueError):
                col_span = 1
        v_merge_el = tc_pr.find(qn("w:vMerge"))
        if v_merge_el is not None:
            # A bare <w:vMerge/> (no val) means "continue" per OOXML.
            v_merge = v_merge_el.get(qn("w:val")) or "continue"

    return {
        "addr": cell_addr,
        "paragraphs": paragraphs,
        "shading": shading,
        "col_span": col_span,
        "v_merge": v_merge,
    }


def _table_payload(tbl_el, addr: str, parent) -> dict:
    rows: list[list[dict]] = []
    for row_idx, tr_el in enumerate(_row_elements(tbl_el)):
        rows.append(
            [
                _cell_payload(tc_el, f"{addr}:{row_idx}:{cell_idx}", parent)
                for cell_idx, tc_el in enumerate(_cell_elements(tr_el))
            ]
        )
    return {"kind": "table", "addr": addr, "rows": rows}


def parse_docx_to_blocks(docx_bytes: bytes) -> list[dict]:
    """Parse a .docx body into the addressed block model.

    Empty paragraphs are kept — they carry vertical spacing the preview
    must reproduce. Returns plain dicts (JSON-ready) validated by the
    DocumentPreviewResponse schema at the endpoint boundary.
    """
    document = Document(io.BytesIO(docx_bytes))
    blocks: list[dict] = []
    for idx, element in _iter_body_blocks(document.element.body):
        addr = str(idx)
        if isinstance(element, CT_P):
            payload = _paragraph_payload(element, addr, document)
            payload["kind"] = "paragraph"
            payload["page_break_before"] = _paragraph_has_page_break(
                element, Paragraph(element, document)
            )
            blocks.append(payload)
        else:  # CT_Tbl
            blocks.append(_table_payload(element, addr, document))
    return blocks


# ---------------------------------------------------------------------------
# Apply: overrides -> new .docx bytes
# ---------------------------------------------------------------------------


def _resolve_paragraph_element(blocks: dict[int, object], addr: str):
    """Resolve an override address to a CT_P element, or None.

    Only 1-part (top-level paragraph) and 4-part (cell paragraph) addresses
    can name a paragraph; anything else — including addresses pointing past
    the end of the document or at a table — resolves to None.
    """
    try:
        parts = [int(piece) for piece in addr.split(":")]
    except ValueError:
        return None
    if any(part < 0 for part in parts):
        return None

    if len(parts) == 1:
        element = blocks.get(parts[0])
        return element if isinstance(element, CT_P) else None

    if len(parts) != 4:
        return None
    tbl_el = blocks.get(parts[0])
    if not isinstance(tbl_el, CT_Tbl):
        return None
    rows = _row_elements(tbl_el)
    if parts[1] >= len(rows):
        return None
    cells = _cell_elements(rows[parts[1]])
    if parts[2] >= len(cells):
        return None
    paragraphs = _cell_paragraph_elements(cells[parts[2]])
    if parts[3] >= len(paragraphs):
        return None
    return paragraphs[parts[3]]


def _replace_paragraph_text(p_el, text: str, document) -> None:
    """Replace the paragraph's runs with ``text``, preserving formatting.

    The first run's rPr is captured as the formatting template and carried
    onto every new run. "\\n" in the override becomes a <w:br/> line break
    WITHIN the paragraph — never a new paragraph, so block addresses stay
    stable across edits.
    """
    # Normalize Windows/mac line endings so "\r\n" yields ONE <w:br/>,
    # not a doubled break. The PATCH endpoint normalizes stored values
    # too — this is the safety net for values persisted before that.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    rpr_template = None
    first_run = p_el.find(qn("w:r"))
    if first_run is not None:
        rpr = first_run.find(qn("w:rPr"))
        if rpr is not None:
            rpr_template = deepcopy(rpr)

    # Remove the existing runs. Hyperlink wrappers carry runs too — drop
    # them as well so the override fully replaces the visible text.
    # (Hyperlink paragraphs are locked upstream, so this branch is
    # defensive only — a lock bypass must not leave duplicated text.)
    for child in list(p_el):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p_el.remove(child)

    para = Paragraph(p_el, document)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line)
        if rpr_template is not None:
            run._r.insert(0, deepcopy(rpr_template))
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def apply_overrides_to_docx(docx_bytes: bytes, overrides: dict[str, str]) -> bytes:
    """Apply per-paragraph text overrides and return new .docx bytes.

    Unresolvable addresses and locked (non-editable) paragraphs are skipped
    silently (debug log only) — a stale override must never break the
    download. Styles, tables, images and headers/footers are untouched.
    """
    document = Document(io.BytesIO(docx_bytes))
    blocks = dict(_iter_body_blocks(document.element.body))

    for addr, text in overrides.items():
        if not isinstance(text, str):
            logger.debug("Override %s has a non-string value — skipped", addr)
            continue
        p_el = _resolve_paragraph_element(blocks, addr)
        if p_el is None:
            logger.debug("Override address %s does not resolve — skipped", addr)
            continue
        if _paragraph_is_locked(p_el):
            logger.debug("Override address %s targets a locked paragraph — skipped", addr)
            continue
        _replace_paragraph_text(p_el, text, document)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Inline image extraction (preview <img> endpoint)
# ---------------------------------------------------------------------------


def extract_image(docx_bytes: bytes, image_id: str) -> tuple[bytes, str] | None:
    """Return ``(blob, content_type)`` for a relationship id, or None.

    ``image_id`` is the rId on the main document part, exactly as emitted
    in the parse output's ``images[].image_id``. Non-image parts (an rId
    pointing at styles, numbering, …) return None as well — the endpoint
    only ever serves actual pictures.
    """
    document = Document(io.BytesIO(docx_bytes))
    try:
        part = document.part.related_parts[image_id]
    except KeyError:
        return None
    content_type = getattr(part, "content_type", None)
    if not content_type or not content_type.startswith("image/"):
        return None
    return part.blob, content_type


# OPC relationships namespace (the .rels part, NOT the document r: namespace).
_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

# Extensions mimetypes doesn't reliably know but Word embeds routinely.
_IMAGE_CT_FALLBACK = {
    ".emf": "image/x-emf",
    ".wmf": "image/x-wmf",
    ".svg": "image/svg+xml",
}


def extract_image_fast(docx_bytes: bytes, image_id: str) -> tuple[bytes, str] | None:
    """Zip-level equivalent of :func:`extract_image` — no python-docx parse.

    ``Document(bytes)`` parses the entire main part (~seconds on the 4.8MB
    DVR) just to serve one picture; here we only read the main part's .rels
    to map ``image_id`` -> Target and pull that entry out of the archive.
    Returns None for unknown ids, non-image relationships, external targets
    and missing parts — exactly like the slow path, which is kept for
    compatibility.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
            except KeyError:
                return None
            try:
                root = ElementTree.fromstring(rels_xml)
            except ElementTree.ParseError:
                return None
            for rel in root.iter(f"{{{_RELS_NS}}}Relationship"):
                if rel.get("Id") != image_id:
                    continue
                if (rel.get("TargetMode") or "").lower() == "external":
                    return None
                if not (rel.get("Type") or "").endswith("/image"):
                    return None
                target = rel.get("Target") or ""
                if target.startswith("/"):
                    # Absolute part name ("/word/media/x.png").
                    part_name = target.lstrip("/")
                else:
                    # Relative to the document part's folder; normpath
                    # resolves "../" segments ("../media/x.png" -> "media/x.png").
                    part_name = posixpath.normpath(posixpath.join("word", target))
                try:
                    blob = zf.read(part_name)
                except KeyError:
                    return None
                content_type, _ = mimetypes.guess_type(part_name)
                if not content_type or not content_type.startswith("image/"):
                    ext = posixpath.splitext(part_name)[1].lower()
                    content_type = _IMAGE_CT_FALLBACK.get(ext, "image/png")
                return blob, content_type
    except zipfile.BadZipFile:
        return None
    return None
