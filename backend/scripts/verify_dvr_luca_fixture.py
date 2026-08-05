"""Build and audit a synthetic DVR Master integration fixture.

The harness deliberately avoids database and customer inputs.  It exercises the
complete existing Acme fixture and a richer, deterministic fixture containing
only fictional identities and generated image pixels.
"""

import argparse
import asyncio
import importlib
import json
import shutil
import sys
import uuid
from collections import Counter
from contextlib import contextmanager
from datetime import datetime
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

BACKEND_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BACKEND_ROOT))

import pillow_heif
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image

from app.services.ambiente_photo import normalize_document_image
from app.services.document_generator.dvr_master import (
    DVRMasterGenerator,
    _PART_IV_PROCEDURAL_SECTIONS,
    _saved_order_key,
)
from scripts.verify_all_generators import build_fixture, mk, patch_generators, run_one

_H1_HEADINGS = [
    "INDICE",
    "Premessa",
    "PARTE I — DATI GENERALI DELL'AZIENDA",
    "PARTE II — DESCRIZIONE DELL'ATTIVITÀ E METODOLOGIA DI VALUTAZIONE",
    "PARTE III — VALUTAZIONE DEI RISCHI PER AMBIENTE DI LAVORO",
    "PARTE IV — PROGRAMMA DI MIGLIORAMENTO",
]
_PART_CONTEXT_LABELS = frozenset(
    heading for heading in _H1_HEADINGS if heading.startswith("PARTE ")
)
_VERA_COVER_SHA256 = (
    "d69929b5cb8981836db8e0d0610e3ff62e91702581392567ffe0247aa616fe15"
)
_VERA_COVER_DIMENSIONS = (640, 358)
_REPORT_KEYS = (
    "acme_regression",
    "vera_cover",
    "saved_people_order",
    "saved_environment_order",
    "external_roles",
    "grouped_equipment",
    "all_ten_photos",
    "effective_risks",
    "person_specific_risks",
    "dpi_dash_alignment",
    "topic_separators",
    "complete_improvements",
    "declaration_signatures",
)


def _image_bytes(image: Image.Image, image_format: str) -> bytes:
    output = BytesIO()
    image.save(output, format=image_format)
    return output.getvalue()


def _next_paragraph_has_page_break(paragraph) -> bool:
    following = paragraph._p.getnext()
    return bool(
        following is not None
        and following.xpath('.//w:br[@w:type="page"]')
    )


def _paragraph_for_element(doc: Document, element):
    return next(
        (paragraph for paragraph in doc.paragraphs if paragraph._p is element),
        None,
    )


def _is_part_heading_or_context(paragraph) -> bool:
    if paragraph.text not in _PART_CONTEXT_LABELS:
        return False
    if paragraph.style.name == "Heading 1":
        return True
    populated_runs = [run for run in paragraph.runs if run.text]
    return (
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        and bool(populated_runs)
        and all(run.bold for run in populated_runs)
    )


def _element_starts_new_page(element) -> bool:
    if element is None or not element.tag.endswith("}p"):
        return False
    if element.xpath("./w:pPr/w:sectPr"):
        return True

    meaningful_tags = {
        qn("w:drawing"),
        qn("w:instrText"),
        qn("w:object"),
        qn("w:pict"),
        qn("w:sym"),
        qn("w:tab"),
    }

    def is_meaningful(node) -> bool:
        if node.tag == qn("w:t"):
            return bool((node.text or "").strip())
        return node.tag in meaningful_tags

    page_break_before = element.xpath("./w:pPr/w:pageBreakBefore")
    if page_break_before:
        value = page_break_before[0].get(qn("w:val"))
        enabled = value is None or value.lower() not in {"0", "false", "off", "no"}
        if enabled and not any(is_meaningful(node) for node in element.iter()):
            return True

    page_breaks = element.xpath('.//w:br[@w:type="page"]')
    if not page_breaks:
        return False
    last_page_break = page_breaks[-1]
    nodes_after_break = iter(element.iter())
    for node in nodes_after_break:
        if node is last_page_break:
            break
    return not any(is_meaningful(node) for node in nodes_after_break)


def _topic_has_start_boundary(doc: Document, heading) -> bool:
    previous = heading._p.getprevious()
    while previous is not None:
        paragraph = _paragraph_for_element(doc, previous)
        if paragraph is None or not _is_part_heading_or_context(paragraph):
            break
        previous = previous.getprevious()
    return _element_starts_new_page(previous)


def _assign_order_metadata(items: list, *, first_uuid: int) -> None:
    for index, item in enumerate(items):
        item.id = uuid.UUID(int=first_uuid + index)
        item.ordine = index
        item.created_at = datetime(2026, 1, index + 1)


def build_acme_fixture() -> dict:
    """Return the existing full Acme fixture with deterministic saved order."""
    fixture = build_fixture()
    fixture["azienda"].id = uuid.UUID(int=100)
    fixture["azienda"].organization_id = uuid.UUID(int=101)
    fixture["generated_at"] = datetime(2026, 8, 3, 9, 0, 0)
    _assign_order_metadata(fixture["ambienti"], first_uuid=200)
    _assign_order_metadata(fixture["persone"], first_uuid=300)
    _assign_order_metadata(fixture["attrezzature"], first_uuid=400)
    _assign_order_metadata(fixture["sostanze_chimiche"], first_uuid=500)
    for environment_index, environment in enumerate(fixture["ambienti"]):
        _assign_order_metadata(
            environment.valutazioni_rischio,
            first_uuid=600 + environment_index * 10,
        )

    fixture["expected_people_rows"] = [
        [
            "MARIO ROSSI",
            "DATORE DI LAVORO",
            "TUTTA L'AZIENDA (RUOLO TRASVERSALE)",
            "—",
            "—",
        ],
        [
            "LUCA BIANCHI",
            "RSPP",
            "TUTTA L'AZIENDA (RUOLO TRASVERSALE)",
            "—",
            "—",
        ],
        [
            "GIULIA VERDI",
            "RLS",
            "TUTTA L'AZIENDA (RUOLO TRASVERSALE)",
            "—",
            "—",
        ],
        ["ANTONIO MARRONE", "OPERAIO TORNITORE", "—", "—", "—"],
        ["VALENTINA RINALDI", "IMPIEGATA (GESTANTE)", "—", "—", "—"],
    ]
    fixture["expected_environments"] = [
        "UFFICI AMMINISTRATIVI E TECNICI",
        "OFFICINA MECCANICA",
        "MAGAZZINO",
        "MENSA AZIENDALE CON CUCINA",
        "DEPOSITO CHIMICI",
        "AREA ESTERNA",
    ]
    fixture["expected_environment_rows"] = [
        ["UFFICI AMMINISTRATIVI E TECNICI", "Ufficio", "220 mq", "0"],
        ["OFFICINA MECCANICA", "Officina", "850 mq", "0"],
        ["MAGAZZINO", "Magazzino", "620 mq", "0"],
        ["MENSA AZIENDALE CON CUCINA", "Cucina", "180 mq", "0"],
        ["DEPOSITO CHIMICI", "Magazzino", "90 mq", "0"],
        ["AREA ESTERNA", "Esterno", "440 mq", "0"],
    ]
    fixture["expected_global_equipment_rows"] = [
        ["TORNIO PARALLELO CNC", "—", "SI", "SI"],
        ["FRESATRICE CNC", "—", "SI", "SI"],
        ["CARRELLO ELEVATORE", "—", "SI", "SI"],
        ["POSTAZIONE VDT", "—", "SI", "NO"],
    ]
    fixture["expected_table_inventory"] = {
        "total": 71,
        "dimensions": Counter(
            {
                (2, 1): 5,
                (2, 2): 6,
                (2, 3): 13,
                (2, 5): 12,
                (2, 7): 1,
                (3, 2): 7,
                (3, 6): 1,
                (5, 2): 6,
                (5, 3): 2,
                (5, 4): 2,
                (6, 5): 1,
                (7, 4): 3,
                (9, 2): 1,
                (10, 5): 1,
                (15, 2): 6,
                (16, 2): 1,
                (17, 2): 1,
                (28, 2): 1,
                (33, 2): 1,
            }
        ),
    }
    return fixture


def build_luca_fixture() -> dict:
    """Return a deterministic, wholly synthetic staging-equivalent fixture."""
    fixture = build_acme_fixture()
    fixture["azienda"].ragione_sociale = "LUCA FIXTURE INDUSTRIA SRL"
    fixture["generated_at"] = datetime(2026, 8, 3, 9, 0, 0)
    zulu_env, alpha_env = fixture["ambienti"][:2]
    zulu_env.nome = "Zulu Reparto"
    zulu_env.ordine = 1
    zulu_env.created_at = datetime(2026, 1, 1)
    alpha_env.nome = "Alpha Reparto"
    alpha_env.ordine = 2
    alpha_env.created_at = datetime(2026, 1, 2)
    fixture["ambienti"] = [alpha_env, zulu_env]

    ddl = mk(
        id=uuid.UUID(int=700),
        nominativo="Datore Fixture",
        mansione="Datore",
        ordine=0,
        created_at=datetime(2026, 1, 1),
        is_esterno=False,
        ruolo_datore_lavoro=True,
        ruolo_rspp=False,
        ruolo_rls=False,
        ruolo_medico_competente=False,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ambienti=[zulu_env, alpha_env],
        dpi_codes=[],
        rischi_specifici_codes=[],
        attrezzature_speciali=[],
    )
    zulu_worker = mk(
        id=uuid.UUID(int=701),
        nominativo="Zulu Worker",
        mansione="Operaio",
        ordine=1,
        created_at=datetime(2026, 1, 2),
        is_esterno=False,
        ruolo_datore_lavoro=False,
        ruolo_rspp=False,
        ruolo_rls=False,
        ruolo_medico_competente=False,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ambienti=[zulu_env],
        dpi_codes=[],
        rischi_specifici_codes=["af_rumore"],
        attrezzature_speciali=[],
    )
    alpha_worker = mk(
        id=uuid.UUID(int=702),
        nominativo="Alpha Worker",
        mansione="Operaio",
        ordine=2,
        created_at=datetime(2026, 1, 3),
        is_esterno=False,
        ruolo_datore_lavoro=False,
        ruolo_rspp=False,
        ruolo_rls=True,
        ruolo_medico_competente=False,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ambienti=[alpha_env],
        dpi_codes=[],
        rischi_specifici_codes=["mmc"],
        attrezzature_speciali=[],
    )
    external_rspp = mk(
        id=uuid.UUID(int=703),
        nominativo="Consulente RSPP",
        mansione="RSPP",
        ordine=3,
        created_at=datetime(2026, 1, 4),
        is_esterno=True,
        ruolo_datore_lavoro=False,
        ruolo_rspp=True,
        ruolo_rls=False,
        ruolo_medico_competente=False,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ambienti=[],
        dpi_codes=[],
        rischi_specifici_codes=[],
        attrezzature_speciali=[],
    )
    external_medico = mk(
        id=uuid.UUID(int=704),
        nominativo="Consulente Medico",
        mansione="Medico",
        ordine=4,
        created_at=datetime(2026, 1, 5),
        is_esterno=True,
        ruolo_datore_lavoro=False,
        ruolo_rspp=False,
        ruolo_rls=False,
        ruolo_medico_competente=True,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ambienti=[],
        dpi_codes=[],
        rischi_specifici_codes=[],
        attrezzature_speciali=[],
    )
    fixture["persone"] = [
        external_medico,
        alpha_worker,
        external_rspp,
        zulu_worker,
        ddl,
    ]
    zulu_env.persone = [ddl, zulu_worker]
    alpha_env.persone = [ddl, alpha_worker]

    fixture["attrezzature"] = [
        mk(
            id=uuid.UUID(int=800),
            descrizione=" Trapano   a colonna ",
            ambiente_id=alpha_env.id,
            marcatura_ce=True,
            verifiche_periodiche=False,
            ordine=1,
            created_at=datetime(2026, 1, 1),
        ),
        mk(
            id=uuid.UUID(int=801),
            descrizione="TRAPANO A COLONNA",
            ambiente_id=zulu_env.id,
            marcatura_ce=False,
            verifiche_periodiche=False,
            ordine=0,
            created_at=datetime(2026, 1, 2),
        ),
    ]
    applicable = mk(
        id=uuid.UUID(int=900),
        pericolo="RISK APPLICABLE SENTINEL",
        applicabile=True,
        ordine=1,
        created_at=datetime(2026, 1, 1),
        condizioni_esposizione="Condizione",
        rischio="Rischio",
        misure_prevenzione="Misura",
        probabilita_p=2,
        danno_d=3,
        livello_rischio="GRAVE",
    )
    disabled = mk(
        id=uuid.UUID(int=901),
        pericolo="RISK DISABLED SENTINEL",
        applicabile=False,
        ordine=2,
        created_at=datetime(2026, 1, 2),
    )
    zulu_env.valutazioni_rischio = [
        mk(
            id=uuid.UUID(int=902),
            categoria_rischio="Macchine",
            applicabile=False,
            ordine=0,
            created_at=datetime(2026, 1, 1),
            pericoli=[applicable, disabled],
        )
    ]
    alpha_env.valutazioni_rischio = []

    photos = []
    for index in range(10):
        image = Image.new(
            "RGB", (60 + index, 40 + index), (20 * index, 80, 160)
        )
        if index % 3 == 0:
            filename = f"foto-{index}.jpg"
            source = _image_bytes(image, "JPEG")
        elif index % 3 == 1:
            filename = f"foto-{index}.png"
            source = _image_bytes(image, "PNG")
        else:
            heif = pillow_heif.from_pillow(image)
            output = BytesIO()
            heif.save(output)
            filename = f"foto-{index}.heic"
            source = output.getvalue()
        normalized = normalize_document_image(source)
        photos.append(
            mk(
                id=uuid.UUID(int=1000 + index),
                filename=filename,
                file_path="/not-shared-on-worker/" + filename,
                document_image_bytes=normalized.content,
                document_image_content_type=normalized.content_type,
                created_at=datetime(2026, 1, 1, 9, index),
            )
        )

    measures = [
        mk(
            id=uuid.UUID(int=2),
            ordine=2,
            created_at=datetime(2026, 1, 2),
            priorita="MODESTO",
            misura="R2",
            misura_miglioramento="M2",
            procedura="P2",
            risorse="S2",
            responsabile="A2",
            scadenza="D2",
        ),
        mk(
            id=uuid.UUID(int=1),
            ordine=1,
            created_at=datetime(2026, 1, 1),
            priorita="GRAVE",
            misura="R1",
            misura_miglioramento="M1",
            procedura="P1",
            risorse="S1",
            responsabile="A1",
            scadenza="D1",
        ),
    ]
    fixture["dvr_extras"] = {
        "foto_by_ambiente": {zulu_env.id: photos},
        "vdt_esposti_persona_ids": set(),
        "allegati_presenti": [],
        "misure_miglioramento": measures,
    }
    fixture["expected_people"] = [
        "DATORE FIXTURE",
        "ZULU WORKER",
        "ALPHA WORKER",
    ]
    fixture["expected_environments"] = ["ZULU REPARTO", "ALPHA REPARTO"]
    fixture["expected_photo_captions"] = [
        f"Fig. {index + 1} — {photo.filename}"
        for index, photo in enumerate(photos)
    ]
    fixture["expected_signature_cells"] = [
        [
            "Il Datore di Lavoro\n(DATORE FIXTURE)\n"
            "___________________________",
            "\n\n___________________________",
            "Il Responsabile del S.P.P.\n(CONSULENTE RSPP)\n"
            "___________________________",
        ],
        [
            "Il Medico Competente\n(CONSULENTE MEDICO)\n"
            "___________________________",
            "\n\n___________________________",
            "Per consultazione\nIl Rappresentante dei Lavoratori\n"
            "(ALPHA WORKER)\n___________________________",
        ],
    ]
    return fixture


def _table_with_headers(doc: Document, headers: list[str]):
    return next(
        table
        for table in doc.tables
        if table.rows
        and [cell.text.strip() for cell in table.rows[0].cells] == headers
    )


def _text(doc: Document) -> str:
    return "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def _heading_texts(doc: Document, level: int) -> list[str]:
    style = f"Heading {level}"
    return [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style.name == style
    ]


def _expected_topic_headings(fixture: dict) -> list[str]:
    environment_names = fixture.get("expected_environments")
    if environment_names is None:
        environment_names = [
            (environment.nome or "—").upper()
            for environment in sorted(fixture["ambienti"], key=_saved_order_key)
        ]
    return [
        "1. Presentazione dell'Azienda",
        "2. Anagrafica Aziendale",
        "3. Dati Occupazionali",
        "4. Organizzazione Aziendale della Sicurezza",
        "5. Ambienti di Lavoro",
        "6. Servizi Igienico-Assistenziali",
        "7. Macchine, Attrezzature ed Impianti",
        "8. Sostanze, Prodotti e Preparati Chimici",
        "9. Elenco Fattori di Pericolo (Riferimento)",
        "2.1 Descrizione dell'Attività",
        "2.2 Definizioni",
        "2.3 Metodologia di Valutazione dei Rischi",
        "2.4 Scala di Probabilità (P)",
        "2.5 Scala del Danno (D)",
        *[
            "Identificazione dell'Ambiente di Lavoro e degli Addetti — "
            + name
            for name in environment_names
        ],
        "Mansioni che espongono i lavoratori a rischi specifici",
        "DPI in dotazione per Mansione",
        "Segnaletica di Sicurezza",
        "Programma di Informazione, Formazione e Addestramento",
        "4.1 Programma e Procedure di attuazione delle Misure di Miglioramento",
        *[spec.heading for spec in _PART_IV_PROCEDURAL_SECTIONS],
        "Documenti correlati al presente DVR",
        "4.13 Dichiarazione del Datore di Lavoro",
    ]


def _audit_outline(doc: Document, fixture: dict) -> bool:
    headings = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Heading 2"
    ]
    flags = [
        bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))
        for paragraph in doc.paragraphs
    ]
    expected = _expected_topic_headings(fixture)
    return (
        _heading_texts(doc, 1) == _H1_HEADINGS
        and [paragraph.text for paragraph in headings] == expected
        and Counter(paragraph.text for paragraph in headings)
        == Counter(expected)
        and all(_topic_has_start_boundary(doc, heading) for heading in headings)
        and all(_next_paragraph_has_page_break(heading) for heading in headings)
        and not any(left and right for left, right in zip(flags, flags[1:]))
    )


def _audit_vera_cover(path: Path, doc: Document) -> bool:
    with ZipFile(path) as archive:
        media = [
            archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/media/")
        ]
    legacy = (BACKEND_ROOT / "assets" / "logo.png").read_bytes()
    matching = [
        content
        for content in media
        if sha256(content).hexdigest() == _VERA_COVER_SHA256
    ]
    dimensions_match = False
    if len(matching) == 1:
        try:
            with Image.open(BytesIO(matching[0])) as image:
                dimensions_match = image.size == _VERA_COVER_DIMENSIONS
        except OSError:
            dimensions_match = False
    return (
        len(matching) == 1
        and dimensions_match
        and legacy not in media
        and "Documento elaborato da" not in _text(doc)
    )


def _audit_saved_people_order(doc: Document, fixture: dict) -> bool:
    table = _table_with_headers(
        doc,
        [
            "Nominativo",
            "Mansione",
            "Ambiente di Lavoro",
            "Codice Fiscale",
            "Tipologia contrattuale",
        ],
    )
    return [
        row.cells[0].text.strip() for row in table.rows[1:]
    ] == fixture["expected_people"]


def _audit_saved_environment_order(doc: Document, fixture: dict) -> bool:
    prefix = "Identificazione dell'Ambiente di Lavoro e degli Addetti — "
    actual = [
        text.removeprefix(prefix)
        for text in _heading_texts(doc, 2)
        if text.startswith(prefix)
    ]
    return actual == fixture["expected_environments"]


def _single_column_table_values(doc: Document, title: str) -> list[str]:
    table = next(
        table
        for table in doc.tables
        if table.rows
        and len(table.rows[0].cells) == 1
        and table.rows[0].cells[0].text.strip() == title
    )
    return [row.cells[0].text.strip() for row in table.rows[1:]]


def _audit_external_roles(doc: Document, fixture: dict) -> bool:
    occupational = _table_with_headers(
        doc,
        [
            "Nominativo",
            "Mansione",
            "Ambiente di Lavoro",
            "Codice Fiscale",
            "Tipologia contrattuale",
        ],
    )
    occupational_names = [
        row.cells[0].text.strip() for row in occupational.rows[1:]
    ]
    return (
        occupational_names == fixture["expected_people"]
        and _single_column_table_values(
            doc, "Responsabile del Servizio di Prevenzione e Protezione"
        )
        == ["CONSULENTE RSPP (ESTERNO)"]
        and _single_column_table_values(doc, "Medico Competente")
        == ["CONSULENTE MEDICO (ESTERNO)"]
    )


def _audit_grouped_equipment(doc: Document, fixture: dict) -> bool:
    table = _table_with_headers(
        doc,
        [
            "Macchine, Attrezzature ed Impianti",
            "Ambiente",
            "Marcata CE",
            "Verifiche Periodiche",
        ],
    )
    return [
        [cell.text for cell in row.cells] for row in table.rows[1:]
    ] == [["TRAPANO A COLONNA", "ZULU REPARTO, ALPHA REPARTO", "MISTO", "NO"]]


def _audit_all_ten_photos(doc: Document, fixture: dict) -> bool:
    paragraphs = doc.paragraphs
    caption_positions = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.startswith("Fig. ")
    ]
    captions = [paragraphs[index].text for index in caption_positions]
    paired_images = [
        paragraphs[index - 1]
        for index in caption_positions
        if index > 0
        and len(paragraphs[index - 1]._p.xpath(".//w:drawing")) == 1
        and not paragraphs[index - 1].text.strip()
    ]
    return (
        captions == fixture["expected_photo_captions"]
        and len(paired_images) == 10
        and len({id(paragraph._p) for paragraph in paired_images}) == 10
        and "[Foto non disponibile:" not in _text(doc)
    )


def _audit_effective_risks(doc: Document, fixture: dict) -> bool:
    full = _text(doc)
    checklist_header = ["Categoria di Rischio", "Applicabile"]
    checklists = [
        table
        for table in doc.tables
        if table.rows
        and [cell.text.strip() for cell in table.rows[0].cells]
        == checklist_header
    ]
    machine_rows = [
        [cell.text.strip() for cell in row.cells]
        for table in checklists
        for row in table.rows[1:]
        if row.cells[0].text.strip() == "Macchine"
    ]
    risk_header = [
        "PERICOLO",
        "CONDIZIONI DI IMPIEGO O DI ESPOSIZIONE",
        "RISCHIO",
        "MISURE DI PREVENZIONE E PROTEZIONE ATTUATE E DPI",
        "I = P + 2*D",
    ]
    risk_tables = [
        table
        for table in doc.tables
        if table.rows
        and [cell.text.strip() for cell in table.rows[0].cells] == risk_header
    ]
    risk_rows = [
        [cell.text.strip() for cell in row.cells]
        for table in risk_tables
        for row in table.rows[1:]
    ]
    return (
        full.count("RISK APPLICABLE SENTINEL") == 1
        and "RISK DISABLED SENTINEL" not in full
        and machine_rows == [["Macchine", "SI"], ["Macchine", "NO"]]
        and len(risk_tables) == 1
        and risk_rows
        == [[
            "RISK APPLICABLE SENTINEL",
            "Condizione",
            "Rischio",
            "Misura",
            "P = 2; D = 3; I = 8; GRAVE",
        ]]
    )


def _audit_person_specific_risks(doc: Document, fixture: dict) -> bool:
    table = _table_with_headers(
        doc, ["Nominativo", "Mansione", "Rischio specifico"]
    )
    return [
        [cell.text for cell in row.cells] for row in table.rows[1:]
    ] == [
        ["ZULU WORKER", "OPERAIO", "Agenti fisici - Rumore"],
        [
            "ALPHA WORKER",
            "OPERAIO",
            "Movimentazione manuale dei carichi (MMC)",
        ],
    ]


def _audit_dpi_dash_alignment(doc: Document) -> bool:
    headers = ["Descrizione DPI", "Marca / Modello", "Note"]
    tables = [
        table
        for table in doc.tables
        if table.rows
        and [cell.text.strip() for cell in table.rows[0].cells] == headers
    ]
    cells = [row.cells[1] for table in tables for row in table.rows[1:]]
    return (
        len(tables) == 2
        and [cell.text.strip() for cell in cells] == ["—", "—"]
        and all(
            cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
            and all(
                paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs
            )
            for cell in cells
        )
    )


def _audit_topic_separators(doc: Document, fixture: dict) -> bool:
    return _audit_outline(doc, fixture)


def _audit_complete_improvements(doc: Document, fixture: dict) -> bool:
    headers = [
        "Priorità",
        "Rischio",
        "Misura di Miglioramento",
        "Attività / Procedura",
        "Risorse",
        "Responsabile",
        "Scadenza",
    ]
    table = _table_with_headers(doc, headers)
    rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
    return (
        rows
        == [
            ["GRAVE", "R1", "M1", "P1", "S1", "A1", "D1"],
            ["MODESTO", "R2", "M2", "P2", "S2", "A2", "D2"],
        ]
        and any(
            section.orientation == WD_ORIENT.LANDSCAPE
            for section in doc.sections
        )
        and doc.sections[-1].orientation == WD_ORIENT.PORTRAIT
    )


def _audit_declaration_signatures(doc: Document, fixture: dict) -> bool:
    declaration = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Heading 2"
        and paragraph.text == "4.13 Dichiarazione del Datore di Lavoro"
    )
    signature = next(
        table
        for table in doc.tables
        if len(table.rows) == 2
        and len(table.rows[0].cells) == 3
        and "Il Datore di Lavoro"
        in " ".join(cell.text for cell in table.rows[0].cells)
    )
    signature_cells = [
        [cell.text for cell in row.cells] for row in signature.rows
    ]
    final_clause = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.text.startswith("di impegnarsi a rielaborare")
        ),
        None,
    )
    place_date = next(
        (paragraph for paragraph in doc.paragraphs if ", li " in paragraph.text),
        None,
    )
    declaration_text = _text(doc)
    return (
        _next_paragraph_has_page_break(declaration)
        and "DATORE FIXTURE, in qualita di Datore di Lavoro della "
        "LUCA FIXTURE INDUSTRIA SRL" in declaration_text
        and signature_cells == fixture["expected_signature_cells"]
        and final_clause is not None
        and place_date is not None
        and bool(final_clause._p.xpath("./w:pPr/w:keepNext"))
        and bool(place_date._p.xpath("./w:pPr/w:keepNext"))
        and final_clause._p.getnext() is place_date._p
        and place_date._p.getnext() is signature._tbl
        and all(
            row.height >= Cm(3)
            and row.height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST
            and row._tr.xpath("./w:trPr/w:cantSplit")
            for row in signature.rows
        )
    )


def _generator_patch_objects() -> list[object]:
    from app.services.document_generator import (
        _biologico_common,
        base,
        data_loader,
    )

    module_classes = [
        ("allegato_mmc", "AllegatoMmcGenerator"),
        ("allegato_vdt", "AllegatoVdtGenerator"),
        ("allegato_stress", "AllegatoStressGenerator"),
        ("allegato_gestanti", "AllegatoGestantiGenerator"),
        ("allegato_incendio", "AllegatoIncendioGenerator"),
        ("allegato_microclima", "AllegatoMicroclimaGenerator"),
        ("allegato_microclima_severo", "AllegatoMicroclimaSeveroGenerator"),
        ("pee_azienda", "PeeAziendaGenerator"),
        ("pee_comune", "PeeComuneGenerator"),
        ("duvri", "DuvriGenerator"),
        ("pos", "PosGenerator"),
        ("haccp_manuale", "HaccpManualeGenerator"),
        ("haccp_forms", "HaccpFormsGenerator"),
    ]
    modules = [
        importlib.import_module(
            f"app.services.document_generator.{module_name}"
        )
        for module_name, _class_name in module_classes
    ]
    version_classes = [
        getattr(module, class_name)
        for module, (_module_name, class_name) in zip(modules, module_classes)
    ]
    return [
        base.BaseDocumentGenerator,
        DVRMasterGenerator,
        *version_classes,
        data_loader,
        _biologico_common,
        *modules,
    ]


_MISSING_ATTRIBUTE = object()


def _changed_attributes(snapshots: list[tuple[object, dict]]) -> list[tuple]:
    changes = []
    for owner, before in snapshots:
        current = vars(owner)
        for name in before.keys() | current.keys():
            original = before.get(name, _MISSING_ATTRIBUTE)
            replacement = current.get(name, _MISSING_ATTRIBUTE)
            if replacement is not original:
                changes.append((owner, name, original))
    return changes


def _restore_attributes(changes: list[tuple]) -> None:
    for owner, name, original in reversed(changes):
        if original is _MISSING_ATTRIBUTE:
            if name in vars(owner):
                delattr(owner, name)
        else:
            setattr(owner, name, original)


@contextmanager
def _patched_generator_fixture(fixture: dict, output_dir: Path):
    snapshots = [
        (owner, dict(vars(owner))) for owner in _generator_patch_objects()
    ]
    try:
        patch_generators(fixture, str(output_dir))
    except BaseException:
        _restore_attributes(_changed_attributes(snapshots))
        raise
    changes = _changed_attributes(snapshots)
    try:
        yield
    finally:
        _restore_attributes(changes)


def _audit_acme_regression(path: Path, fixture: dict) -> bool:
    doc = Document(path)
    occupational = _table_with_headers(
        doc,
        [
            "Nominativo",
            "Mansione",
            "Ambiente di Lavoro",
            "Codice Fiscale",
            "Tipologia contrattuale",
        ],
    )
    environment_summary = _table_with_headers(
        doc, ["Ambiente", "Tipo", "Metratura", "N. Lavoratori"]
    )
    global_equipment = _table_with_headers(
        doc,
        [
            "Macchine, Attrezzature ed Impianti",
            "Ambiente",
            "Marcata CE",
            "Verifiche Periodiche",
        ],
    )
    actual_inventory = {
        "total": len(doc.tables),
        "dimensions": Counter(
            (len(table.rows), len(table.rows[0].cells))
            for table in doc.tables
        ),
    }
    return (
        fixture["azienda"].ragione_sociale in _text(doc)
        and _audit_outline(doc, fixture)
        and [
            [cell.text.strip() for cell in row.cells]
            for row in occupational.rows[1:]
        ]
        == fixture["expected_people_rows"]
        and [
            [cell.text.strip() for cell in row.cells]
            for row in environment_summary.rows[1:]
        ]
        == fixture["expected_environment_rows"]
        and [
            [cell.text.strip() for cell in row.cells]
            for row in global_equipment.rows[1:]
        ]
        == fixture["expected_global_equipment_rows"]
        and actual_inventory == fixture["expected_table_inventory"]
    )


def build_and_audit(output_dir: Path) -> dict[str, bool]:
    """Generate the Acme and Luca DOCX files, then return all audit gates."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    acme = build_acme_fixture()
    with _patched_generator_fixture(acme, output_dir):
        acme_ok, acme_path, acme_message = asyncio.run(
            run_one("DVR_MASTER", acme["azienda"].id)
        )
        if not acme_ok:
            raise AssertionError(acme_message)
        stable_acme = output_dir / "DVR_Acme_Regression.docx"
        shutil.copy2(acme_path, stable_acme)

        fixture = build_luca_fixture()
        patch_generators(fixture, str(output_dir))

        async def rich_dvr_extras(self, data):
            return fixture["dvr_extras"]

        DVRMasterGenerator._load_dvr_extras = rich_dvr_extras
        ok, generated_path, message = asyncio.run(
            run_one("DVR_MASTER", fixture["azienda"].id)
        )
        if not ok:
            raise AssertionError(message)
        stable_path = output_dir / "DVR_Luca_Fixture.docx"
        shutil.copy2(generated_path, stable_path)

        report = audit_luca_docx(stable_path, fixture)
        report["acme_regression"] = _audit_acme_regression(stable_acme, acme)
    return {key: report[key] for key in _REPORT_KEYS}


def audit_luca_docx(path: Path, fixture: dict) -> dict[str, bool]:
    """Audit exact structural acceptance behavior in a generated Luca DOCX."""
    doc = Document(path)
    return {
        "vera_cover": _audit_vera_cover(path, doc),
        "saved_people_order": _audit_saved_people_order(doc, fixture),
        "saved_environment_order": _audit_saved_environment_order(doc, fixture),
        "external_roles": _audit_external_roles(doc, fixture),
        "grouped_equipment": _audit_grouped_equipment(doc, fixture),
        "all_ten_photos": _audit_all_ten_photos(doc, fixture),
        "effective_risks": _audit_effective_risks(doc, fixture),
        "person_specific_risks": _audit_person_specific_risks(doc, fixture),
        "dpi_dash_alignment": _audit_dpi_dash_alignment(doc),
        "topic_separators": _audit_topic_separators(doc, fixture),
        "complete_improvements": _audit_complete_improvements(doc, fixture),
        "declaration_signatures": _audit_declaration_signatures(doc, fixture),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    report = build_and_audit(args.output_dir)
    print(json.dumps(report, indent=2, sort_keys=True))
    return 0 if all(report.values()) else 1


if __name__ == "__main__":
    sys.exit(main())
