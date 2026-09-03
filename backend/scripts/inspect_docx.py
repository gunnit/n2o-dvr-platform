"""Dump the visual + structural anatomy of a generated .docx for audit.

Usage:
    python -m scripts.inspect_docx <file.docx> [--body N]

Prints: embedded images (name/size/where), section headers & footers with
their runs, fonts/sizes/colours actually used, table inventory with header
shading, style usage, page setup, and the first N body paragraphs with their
style + formatting. This is the shared evidence format for the document audit.
"""
from __future__ import annotations

import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import RGBColor  # noqa: F401


def _fmt_color(run):
    try:
        c = run.font.color
        if c is not None and c.rgb is not None:
            return f"#{c.rgb}"
    except Exception:
        pass
    return None


def _runinfo(p, limit=6):
    bits = []
    for r in p.runs[:limit]:
        d = []
        if r.bold:
            d.append("b")
        if r.italic:
            d.append("i")
        if r.underline:
            d.append("u")
        if r.font.size:
            d.append(f"{r.font.size.pt:g}pt")
        if r.font.name:
            d.append(r.font.name)
        col = _fmt_color(r)
        if col:
            d.append(col)
        txt = (r.text or "")[:60].replace("\n", "\n")
        bits.append(f"[{','.join(d) or '-'}]{txt!r}")
    if len(p.runs) > limit:
        bits.append(f"...(+{len(p.runs)-limit} runs)")
    return " ".join(bits)


def _cell_fill(cell):
    from docx.oxml.ns import qn
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _textbox_texts(doc):
    """Text inside VML/DrawingML text boxes (w:txbxContent), which
    ``paragraph.text`` never sees. Cover titles in the legacy templates live
    here (e.g. Gestanti's 28pt Castellar title), so a body dump that shows an
    'empty' paragraph may actually be the document's title."""
    from docx.oxml.ns import qn
    out = []
    for tx in doc.element.body.iter(qn("w:txbxContent")):
        txt = " / ".join(
            "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
            for p in tx.iter(qn("w:p"))
        ).strip(" /")
        if txt:
            out.append(txt)
    return out


def _fields(part_el):
    """Field codes (PAGE, NUMPAGES, DATE, TOC, ...) in an XML part, with their
    CACHED result. The cached <w:t> between fldChar separate/end is what the
    template last computed -- Word/LibreOffice recompute it on open, so a
    cached 'Pag. 67' is NOT a frozen literal. A date or revision printed as a
    plain <w:t> with no field around it IS a literal and will not update."""
    from docx.oxml.ns import qn
    found = []
    for instr in part_el.iter(qn("w:instrText")):
        code = (instr.text or "").strip()
        found.append(code.split(" ")[0] if code else "?")
    return found


def main():
    path = Path(sys.argv[1])
    body_n = 40
    if "--body" in sys.argv:
        body_n = int(sys.argv[sys.argv.index("--body") + 1])

    print("=" * 78)
    print(f"FILE: {path.name}  ({path.stat().st_size:,} bytes)")
    print("=" * 78)

    # --- embedded media straight from the zip ---
    with zipfile.ZipFile(path) as z:
        media = [i for i in z.infolist() if "/media/" in i.filename]
        print(f"\n--- EMBEDDED IMAGES ({len(media)}) ---")
        for i in media:
            print(f"  {i.filename}  {i.file_size:,} bytes")
        rels = [n for n in z.namelist() if n.endswith(".rels") and ("header" in n or "footer" in n)]
        for n in rels:
            data = z.read(n).decode("utf8", "replace")
            if "image" in data:
                print(f"  ! image referenced from {n}")
        hdrs = [n for n in z.namelist() if "header" in n and n.endswith(".xml")]
        ftrs = [n for n in z.namelist() if "footer" in n and n.endswith(".xml")]
        print(f"  header parts: {hdrs}")
        print(f"  footer parts: {ftrs}")

    doc = Document(str(path))

    # --- page setup ---
    print(f"\n--- SECTIONS ({len(doc.sections)}) ---")
    for si, s in enumerate(doc.sections):
        try:
            print(
                f"  [{si}] page {s.page_width.cm:.1f}x{s.page_height.cm:.1f}cm "
                f"margins L{s.left_margin.cm:.1f} R{s.right_margin.cm:.1f} "
                f"T{s.top_margin.cm:.1f} B{s.bottom_margin.cm:.1f} "
                f"orient={s.orientation}"
            )
        except Exception as e:
            print(f"  [{si}] page setup unreadable: {e}")
        for label, part in (("HEADER", s.header), ("FOOTER", s.footer)):
            txts = [p for p in part.paragraphs if (p.text or "").strip() or p.runs]
            print(f"    {label}: {len(part.paragraphs)} paragraphs, {len(part.tables)} tables")
            for p in txts[:8]:
                print(f"      · {p.style.name if p.style else '?'} | {_runinfo(p)}")
            for ti, t in enumerate(part.tables[:3]):
                print(f"      · TABLE {ti}: {len(t.rows)}x{len(t.columns)}")
                for r in t.rows[:3]:
                    print(f"          | " + " | ".join((c.text or "").strip()[:40] for c in r.cells))

    # --- text boxes (cover titles hide here) ---
    tb = _textbox_texts(doc)
    print(f"\n--- TEXT BOXES ({len(tb)}) -- content python-docx paragraph.text does NOT show ---")
    for t in tb[:10]:
        print(f"  · {t[:160]!r}")

    # --- fields: cached values re-render; plain literals do not ---
    print("\n--- FIELD CODES (a cached PAGE/NUMPAGES/DATE value is NOT a frozen literal) ---")
    print(f"  body: {_fields(doc.element.body)}")
    for si, s in enumerate(doc.sections):
        hf = _fields(s.header._element) + _fields(s.footer._element)
        if hf:
            print(f"  section[{si}] header/footer: {hf}")
    print("  NOTE: a date/revision that appears in header/footer text but NOT as a field above is a hardcoded literal.")

    # --- fonts / sizes / colours actually used in the body ---
    fonts, sizes, colors, styles = Counter(), Counter(), Counter(), Counter()
    for p in doc.paragraphs:
        styles[p.style.name if p.style else "?"] += 1
        for r in p.runs:
            fonts[r.font.name] += 1
            sizes[r.font.size.pt if r.font.size else None] += 1
            colors[_fmt_color(r)] += 1
    print("\n--- TYPOGRAPHY (body runs) ---")
    print(f"  fonts:  {dict(fonts.most_common(8))}")
    print(f"  sizes:  {dict(sizes.most_common(10))}")
    print(f"  colors: {dict(colors.most_common(10))}")
    print(f"  para styles: {dict(styles.most_common(12))}")

    # --- tables ---
    print(f"\n--- TABLES ({len(doc.tables)}) ---")
    for ti, t in enumerate(doc.tables):
        try:
            stl = t.style.name if t.style else None
        except Exception:
            stl = "?"
        hdr_fill = _cell_fill(t.rows[0].cells[0]) if t.rows else None
        first = " | ".join((c.text or "").strip()[:28] for c in t.rows[0].cells) if t.rows else ""
        print(f"  [{ti:>2}] {len(t.rows)}x{len(t.columns)} style={stl} hdrfill={hdr_fill} :: {first[:110]}")

    # --- body ---
    print(f"\n--- BODY (first {body_n} of {len(doc.paragraphs)} paragraphs) ---")
    shown = 0
    for p in doc.paragraphs:
        if shown >= body_n:
            break
        txt = (p.text or "").strip()
        if not txt and not p.runs:
            continue
        shown += 1
        style = p.style.name if p.style else "?"
        align = p.alignment
        print(f"  [{style}|{align}] {_runinfo(p)}")

    # --- suspicious content scan ---
    print("\n--- LEAKAGE / PLACEHOLDER SCAN ---")
    needles = [
        "N2O", "CIARAMITARO", "GORGONZOLA", "GESSATE", "VIA DEI CHIOSI",
        "VIA MONZA", "{{", "}}", "XXX", "TODO", "LOREM", "NON DISPONIBILE",
        "None", "N/D", "[", "…",
    ]
    alltext = []
    for p in doc.paragraphs:
        alltext.append(p.text or "")
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                alltext.append(c.text or "")
    for s in doc.sections:
        for part in (s.header, s.footer):
            for p in part.paragraphs:
                alltext.append(p.text or "")
    blob = "\n".join(alltext)
    for n in needles:
        cnt = blob.count(n)
        if cnt:
            # show a sample line
            sample = next((l.strip()[:100] for l in blob.splitlines() if n in l), "")
            print(f"  {n!r}: {cnt}x   e.g. {sample!r}")

    empties = sum(1 for l in alltext if not l.strip())
    print(f"  empty text nodes: {empties}/{len(alltext)}")


if __name__ == "__main__":
    main()
