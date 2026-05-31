"""Data-fidelity audit for generated documents — DEPLOYED PLATFORM ONLY.

Unlike ``verify_all_generators.py`` (which checks *structural* validity against
an in-memory fixture with the DB bypassed), this harness answers the question
the client actually cares about:

    "Does the data stored in the database render correctly into the documents?"

It talks **only to the deployed API** (no local DB, no local generation):

  1. Logs in, then reads back the full live state of an azienda — anagrafica
     plus every related entity (persone, ambienti, attrezzature, sostanze,
     rischi, pericoli, and the assessment tables). Every *renderable* leaf
     value becomes an "expected" item, tagged critical / normal / info.
  2. For each requested document type: (re)generates on the platform, polls to
     completion, and downloads the produced .docx / .zip bytes.
  3. Parses the document into a text corpus (paragraphs + every table cell,
     recursing into nested tables and into .zip members).
  4. Cross-checks: is each expected DB value actually present in the rendered
     text? Flags missing *critical* values, and scans for placeholder / leak
     tokens ({{...}}, raw UUIDs, "[inserire ...]", bare None/null, etc.).
  5. Writes a per-document PASS / PARTIAL / FAIL report (Markdown + JSON).

Usage (from anywhere; uses only stdlib + python-docx):

    python backend/scripts/audit_documents.py \
        --azienda 1a5e6164-3ce7-4636-a6ba-ef65a9cd376d \
        --types dvr_master \
        [--regenerate] [--report .audit/data_fidelity]

Credentials default to the long-lived QA account (see credentials/test-account.json);
override with --email / --password or AUDIT_EMAIL / AUDIT_PASSWORD env vars.
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
import zipfile
from dataclasses import dataclass, field

DEFAULT_API = "https://n2o-dvr-api.onrender.com/api/v1"
DEFAULT_EMAIL = "claude-verify-20260525@niuexa.ai"
DEFAULT_PASSWORD = "TstyLthG8NdyhBwyIgclPrm!9"

# ---------------------------------------------------------------------------
# Patterns
# ---------------------------------------------------------------------------

UUID_RE = re.compile(
    r"\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b"
)
ISO_DT_RE = re.compile(r"^\d{4}-\d{2}-\d{2}([T ]\d{2}:\d{2}.*)?$")

# Keys whose values are internal plumbing, never meant to be read by a human in
# the document. We skip these when harvesting "expected" values.
_SKIP_KEY_RE = re.compile(
    r"(^id$|_id$|_at$|_url$|^ordine$|status$|^source$|password|token|hash|"
    r"organization|_key$|^versione$|^applicabile$|_codes?$|pittogrammi|"
    r"_error$|confidence|sds_file|ai_extracted|human_reviewed)",
    re.IGNORECASE,
)

# Hard leaks: if any of these appear in a rendered document it's almost
# certainly a bug — an unfilled template token, a leaked object repr, etc.
_HARD_LEAK_PATTERNS: list[tuple[str, re.Pattern]] = [
    ("jinja/handlebars token", re.compile(r"\{\{.*?\}\}")),
    ("square-bracket placeholder", re.compile(r"\[\s*(inseri|nome|indirizz|indica|compila|dato|valore|xxx)", re.IGNORECASE)),
    ("python None/object repr", re.compile(r"(?<![A-Za-z])None(?![A-Za-z])|<[a-z_]+\.[A-Za-z0-9_.]+ object at 0x")),
    ("null literal", re.compile(r"(?<![A-Za-z])null(?![A-Za-z])")),
    ("NaN", re.compile(r"(?<![A-Za-z])nan(?![A-Za-z])")),
    ("lorem ipsum", re.compile(r"lorem ipsum", re.IGNORECASE)),
    ("raw UUID in body", UUID_RE),
    ("XXXX placeholder", re.compile(r"\bX{4,}\b")),
]

# Soft markers: legitimate when the underlying data really is absent, but worth
# surfacing so a reviewer can decide. NOT a failure by themselves.
_SOFT_MARKER_RE = re.compile(
    r"(non comunicat|non disponibil|non specificat|da definir|da designar|"
    r"da compilar|non indicat|n\.d\.|n/d)",
    re.IGNORECASE,
)

# Only the company NAME must appear in *every* document. Italian safety
# annexes are bound to the DVR dossier and are not required to repeat the full
# registry block — they only need to identify the company. The fuller block is
# upgraded to critical per-document (see _is_critical_for) for the master and
# the annexes that actually carry a full anagrafica table.
_CRITICAL_AZIENDA_KEYS = {"ragione_sociale"}

# Registry fields expected on documents that carry a full company-identity table.
_REGISTRY_KEYS = {
    "partita_iva", "codice_fiscale", "codice_ateco",
    "sede_legale_via", "sede_legale_citta", "provincia_legale",
    "sede_operativa_via", "sede_operativa_citta",
}
# Documents that render a full anagrafica table (registry block critical for
# them). Every other annex only needs the company name + its own assessment.
_FULL_ANAGRAFICA_DOCS = {"dvr_master", "allegato_mmc", "allegato_vdt"}


def _is_critical_for(e: "Expected", tipo: str) -> bool:
    """Per-document criticality.

    ``ragione_sociale`` (weight ``critical``) is critical everywhere. The DVR
    master additionally must carry the full registry block + the persone roster
    + the ambienti; the MMC/VDT annexes must carry the registry block. All other
    fields are non-failing ``normal`` misses.
    """
    if e.weight == "critical":
        return True
    key = e.path.split(".")[-1]
    if tipo == "dvr_master":
        if e.kind == "anagrafica" and key in _REGISTRY_KEYS:
            return True
        if e.kind == "persona" and key == "nominativo":
            return True
        if e.kind == "ambiente" and key == "nome":
            return True
    elif tipo in _FULL_ANAGRAFICA_DOCS:
        if e.kind == "anagrafica" and key in _REGISTRY_KEYS:
            return True
    return False


# ---------------------------------------------------------------------------
# API client
# ---------------------------------------------------------------------------

class APIClient:
    def __init__(self, api: str, email: str, password: str):
        self.api = api.rstrip("/")
        self.email = email
        self.password = password
        self.token: str | None = None

    def login(self) -> None:
        status, body = self._raw("POST", "/auth/login",
                                 body={"email": self.email, "password": self.password})
        if status != 200:
            raise SystemExit(f"Login failed [{status}]: {body}")
        self.token = body["access_token"]

    def _raw(self, method, path, body=None, raw_bytes=False):
        url = path if path.startswith("http") else f"{self.api}{path}"
        data = json.dumps(body).encode() if body is not None else None
        req = urllib.request.Request(url, data=data, method=method)
        if data is not None:
            req.add_header("Content-Type", "application/json")
        if self.token:
            req.add_header("Authorization", f"Bearer {self.token}")
        try:
            with urllib.request.urlopen(req, timeout=120) as r:
                payload = r.read()
                if raw_bytes:
                    return r.status, payload
                return r.status, (json.loads(payload.decode()) if payload else None)
        except urllib.error.HTTPError as e:
            detail = e.read()
            if raw_bytes:
                return e.code, detail
            try:
                return e.code, json.loads(detail.decode())
            except Exception:
                return e.code, detail.decode()[:500]
        except urllib.error.URLError as e:
            return 0, str(e)

    def request(self, method, path, body=None, raw_bytes=False, _retry=True):
        status, payload = self._raw(method, path, body, raw_bytes)
        if status == 401 and _retry:  # token expired — re-login once
            self.login()
            return self.request(method, path, body, raw_bytes, _retry=False)
        return status, payload

    def get(self, path, raw_bytes=False):
        return self.request("GET", path, raw_bytes=raw_bytes)

    def post(self, path, body=None):
        return self.request("POST", path, body=body)


# ---------------------------------------------------------------------------
# Expected-value harvesting
# ---------------------------------------------------------------------------

@dataclass
class Expected:
    path: str           # where it came from, e.g. "persone[0].nominativo"
    value: object       # original value
    weight: str         # "critical" | "normal" | "info"
    kind: str           # entity bucket: anagrafica | persona | ambiente | ...


def _is_renderable_str(v: str) -> bool:
    s = v.strip()
    if len(s) < 2:
        return False
    if UUID_RE.fullmatch(s):
        return False
    if ISO_DT_RE.match(s):
        return False
    return True


def harvest(obj, kind: str, path: str, out: list[Expected],
            critical_keys: set[str] | None = None) -> None:
    """Recursively collect human-renderable leaf values from an entity."""
    critical_keys = critical_keys or set()
    if isinstance(obj, dict):
        for k, v in obj.items():
            if _SKIP_KEY_RE.search(k):
                continue
            child_path = f"{path}.{k}" if path else k
            weight = "critical" if k in critical_keys else None
            harvest(v, kind, child_path, out, critical_keys)
            # tag the just-added leaves from this key as critical if needed
            if weight == "critical":
                for e in out:
                    if e.path == child_path and e.weight == "normal":
                        e.weight = "critical"
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            harvest(item, kind, f"{path}[{i}]", out, critical_keys)
    elif isinstance(obj, bool):
        return  # booleans are rendered as SI/NO etc., not matchable verbatim
    elif isinstance(obj, str):
        if _is_renderable_str(obj):
            out.append(Expected(path, obj.strip(), "normal", kind))
    elif isinstance(obj, (int, float)):
        # Skip the small ints that are almost always enum/scale noise rather
        # than distinctive data (0,1) — keep 2+ which include P/D/I scores,
        # counts, surfaces. They are matched as "info" weight (never fail).
        if isinstance(obj, bool):
            return
        out.append(Expected(path, obj, "info", kind))


@dataclass
class Inventory:
    azienda: dict
    expected: list[Expected] = field(default_factory=list)
    counts: dict = field(default_factory=dict)

    def by_kind(self, *kinds: str) -> list[Expected]:
        return [e for e in self.expected if e.kind in kinds]


def build_inventory(client: APIClient, azienda_id: str) -> Inventory:
    status, az = client.get(f"/aziende/{azienda_id}")
    if status != 200:
        raise SystemExit(f"Cannot read azienda {azienda_id} [{status}]: {az}")

    inv = Inventory(azienda=az)
    harvest(az, "anagrafica", "azienda", inv.expected, _CRITICAL_AZIENDA_KEYS)

    # Persone — nominativo is critical for the DVR organigramma (upgraded
    # per-doc in _is_critical_for); a normal-weight expectation elsewhere.
    _, persone = client.get(f"/aziende/{azienda_id}/persone")
    persone = persone if isinstance(persone, list) else []
    harvest(persone, "persona", "persone", inv.expected)

    # Ambienti — nome is critical for the DVR (upgraded per-doc).
    _, ambienti = client.get(f"/aziende/{azienda_id}/ambienti")
    ambienti = ambienti if isinstance(ambienti, list) else []
    harvest(ambienti, "ambiente", "ambienti", inv.expected)

    _, attrezzature = client.get(f"/aziende/{azienda_id}/attrezzature")
    attrezzature = attrezzature if isinstance(attrezzature, list) else []
    harvest(attrezzature, "attrezzatura", "attrezzature", inv.expected)

    _, sostanze = client.get(f"/aziende/{azienda_id}/sostanze-chimiche")
    sostanze = sostanze if isinstance(sostanze, list) else []
    harvest(sostanze, "sostanza", "sostanze", inv.expected)

    # Rischi (flat, carry ambiente_id) + their 1:N pericoli children
    _, rischi = client.get(f"/aziende/{azienda_id}/rischi")
    rischi = rischi if isinstance(rischi, list) else []
    harvest(rischi, "rischio", "rischi", inv.expected)
    n_pericoli = 0
    for r in rischi:
        amb = r.get("ambiente_id")
        rid = r.get("id")
        if not amb or not rid:
            continue
        s, peric = client.get(
            f"/aziende/{azienda_id}/ambienti/{amb}/rischi/{rid}/pericoli"
        )
        if isinstance(peric, list):
            n_pericoli += len(peric)
            harvest(peric, "pericolo", f"pericoli[{rid[:8]}]", inv.expected)

    # Assessment tables (best-effort; absent ones return [] / 404)
    for kind, path in [
        ("mmc", f"/aziende/{azienda_id}/mmc"),
        ("vdt", f"/aziende/{azienda_id}/vdt"),
        ("duvri", f"/aziende/{azienda_id}/duvri"),
        ("pos", f"/aziende/{azienda_id}/pos"),
    ]:
        s, rows = client.get(path)
        if isinstance(rows, list) and rows:
            harvest(rows, kind, kind, inv.expected)
            inv.counts[kind] = len(rows)

    inv.counts.update({
        "persone": len(persone), "ambienti": len(ambienti),
        "attrezzature": len(attrezzature), "sostanze": len(sostanze),
        "rischi": len(rischi), "pericoli": n_pericoli,
    })
    return inv


# ---------------------------------------------------------------------------
# Document parsing -> text corpus
# ---------------------------------------------------------------------------

def _iter_table_text(table) -> list[str]:
    out = []
    for row in table.rows:
        for cell in row.cells:
            out.append(cell.text)
            for nested in cell.tables:  # recurse into nested tables
                out.extend(_iter_table_text(nested))
    return out


def docx_to_corpus(data: bytes) -> tuple[str, dict]:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    n_tables = len(doc.tables)
    for t in doc.tables:
        parts.extend(_iter_table_text(t))
    # headers/footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            parts.extend(p.text for p in hf.paragraphs)
    corpus = "\n".join(x for x in parts if x and x.strip())
    meta = {"paragraphs": len(doc.paragraphs), "tables": n_tables,
            "inline_shapes": len(doc.inline_shapes)}
    return corpus, meta


def bytes_to_corpus(data: bytes, filename: str) -> tuple[str, dict]:
    if filename.endswith(".zip") or (data[:2] == b"PK" and not _looks_like_docx(data)):
        if zipfile.is_zipfile(io.BytesIO(data)):
            texts, members = [], []
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                for name in z.namelist():
                    members.append(name)
                    if name.endswith(".docx"):
                        try:
                            c, _ = docx_to_corpus(z.read(name))
                            texts.append(c)
                        except Exception:
                            pass
            return "\n".join(texts), {"zip_members": len(members),
                                      "docx_members": sum(1 for m in members if m.endswith(".docx"))}
    return docx_to_corpus(data)


def _looks_like_docx(data: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            return "word/document.xml" in z.namelist()
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Normalisation + matching
# ---------------------------------------------------------------------------

def normalize(s: str) -> str:
    s = s.casefold()
    s = (s.replace("’", "'").replace("‘", "'")
           .replace("“", '"').replace("”", '"')
           .replace("–", "-").replace("—", "-").replace("−", "-")
           .replace(" ", " "))
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _number_variants(v) -> list[str]:
    out = []
    if isinstance(v, float) and v.is_integer():
        v = int(v)
    if isinstance(v, int):
        out = [str(v)]
    elif isinstance(v, float):
        out = [repr(v), f"{v:.1f}", f"{v:.2f}",
               repr(v).replace(".", ","), f"{v:.1f}".replace(".", ","),
               f"{v:.2f}".replace(".", ",")]
        if v.is_integer():
            out.append(str(int(v)))
    return list(dict.fromkeys(out))


def matches(expected: Expected, corpus_norm: str) -> bool:
    v = expected.value
    if isinstance(v, bool):
        return True
    if isinstance(v, (int, float)):
        for variant in _number_variants(v):
            if re.search(rf"(?<!\d){re.escape(variant)}(?!\d)", corpus_norm):
                return True
        return False
    token = normalize(str(v))
    if not token:
        return True
    if len(token) <= 3:  # short tokens (province codes) need word boundaries
        return re.search(rf"(?<![a-z0-9]){re.escape(token)}(?![a-z0-9])", corpus_norm) is not None
    return token in corpus_norm


# ---------------------------------------------------------------------------
# Per-document audit
# ---------------------------------------------------------------------------

# Which entity kinds each document type is expected to render. anagrafica is
# implicitly checked for every type (company identity on the cover).
DOC_RELEVANCE: dict[str, set[str]] = {
    "dvr_master": {"anagrafica", "persona", "ambiente", "attrezzatura",
                   "sostanza", "rischio", "pericolo"},
    "allegato_mmc": {"anagrafica", "mmc", "persona"},
    "allegato_vdt": {"anagrafica", "vdt", "persona"},
    "allegato_stress": {"anagrafica"},
    "allegato_gestanti": {"anagrafica", "persona"},
    "allegato_incendio": {"anagrafica", "ambiente"},
    "allegato_microclima": {"anagrafica", "ambiente"},
    "allegato_microclima_severo": {"anagrafica", "ambiente"},
    "allegato_biologico_alimentare": {"anagrafica"},
    "allegato_biologico_asilo": {"anagrafica"},
    "allegato_biologico_dentisti": {"anagrafica"},
    "pee_azienda": {"anagrafica", "persona", "ambiente"},
    "pee_comune": {"anagrafica", "persona"},
    "haccp": {"anagrafica"},
    "haccp_forms": {"anagrafica"},
    "duvri": {"anagrafica", "duvri"},
    "pos": {"anagrafica", "pos"},
}


@dataclass
class DocResult:
    tipo: str
    verdict: str = "PASS"
    doc_id: str | None = None
    version: int | None = None
    meta: dict = field(default_factory=dict)
    checked: int = 0
    matched: int = 0
    critical_misses: list[Expected] = field(default_factory=list)
    normal_misses: list[Expected] = field(default_factory=list)
    hard_leaks: list[tuple[str, str]] = field(default_factory=list)
    soft_markers: list[str] = field(default_factory=list)
    error: str | None = None


def find_or_generate(client, azienda_id, tipo, regenerate) -> tuple[str | None, int | None, str | None]:
    """Return (document_id, version, error) for a completed document."""
    s, docs = client.get(f"/aziende/{azienda_id}/documents")
    docs = docs if isinstance(docs, list) else []
    existing = [d for d in docs if d.get("tipo_documento") == tipo]
    completed = [d for d in existing if d.get("status") == "completed"]

    if completed and not regenerate:
        latest = max(completed, key=lambda d: d.get("versione", 0))
        return latest["id"], latest.get("versione"), None

    # Generate a fresh one on the platform
    s, created = client.post(f"/aziende/{azienda_id}/documents/generate",
                             body={"tipo_documento": tipo})
    if s not in (200, 202):
        return None, None, f"generate returned {s}: {created}"
    doc_id = created["id"]
    # Poll status
    for _ in range(120):  # up to ~6 min
        time.sleep(3)
        s, st = client.get(f"/aziende/{azienda_id}/documents/{doc_id}/status")
        if not isinstance(st, dict):
            continue
        if st.get("status") == "completed":
            return doc_id, st.get("versione"), None
        if st.get("status") == "failed":
            return None, None, f"generation failed: {st.get('error_message')}"
    return None, None, "timeout waiting for generation"


def audit_document(client, azienda_id, tipo, inv: Inventory,
                   regenerate: bool, save_dir: str | None) -> DocResult:
    res = DocResult(tipo=tipo)
    doc_id, version, err = find_or_generate(client, azienda_id, tipo, regenerate)
    if err:
        res.verdict, res.error = "ERROR", err
        return res
    res.doc_id, res.version = doc_id, version

    s, data = client.get(f"/documenti/{doc_id}/download", raw_bytes=True)
    if s != 200 or not isinstance(data, (bytes, bytearray)):
        res.verdict, res.error = "ERROR", f"download returned {s}"
        return res

    # Determine filename for zip-vs-docx routing
    s, meta_doc = client.get(f"/aziende/{azienda_id}/documents/{doc_id}/status")
    fname = (meta_doc or {}).get("file_path") or f"{tipo}.docx"
    if save_dir:
        os.makedirs(save_dir, exist_ok=True)
        ext = ".zip" if (fname.endswith(".zip")) else ".docx"
        out_path = os.path.join(save_dir, f"{tipo}_v{version}{ext}")
        with open(out_path, "wb") as f:
            f.write(data)

    try:
        corpus, res.meta = bytes_to_corpus(bytes(data),
                                           ".zip" if fname.endswith(".zip") else ".docx")
    except Exception as e:
        res.verdict, res.error = "ERROR", f"parse failed: {e}"
        return res
    corpus_norm = normalize(corpus)

    relevant = DOC_RELEVANCE.get(tipo, {"anagrafica"})
    to_check = [e for e in inv.expected if e.kind in relevant]
    res.checked = len(to_check)
    for e in to_check:
        if matches(e, corpus_norm):
            res.matched += 1
        else:
            if _is_critical_for(e, tipo):
                res.critical_misses.append(e)
            elif e.weight in ("critical", "normal"):
                res.normal_misses.append(e)
            # info-weight misses are tolerated (small ints/floats)

    # Leak scans on the RAW corpus (not normalised)
    for label, pat in _HARD_LEAK_PATTERNS:
        for m in pat.finditer(corpus):
            res.hard_leaks.append((label, m.group(0)[:80]))
            if len(res.hard_leaks) >= 25:
                break
    res.soft_markers = sorted({m.group(0) for m in _SOFT_MARKER_RE.finditer(corpus)})

    if res.error:
        res.verdict = "ERROR"
    elif res.critical_misses or res.hard_leaks:
        res.verdict = "FAIL"
    elif res.normal_misses:
        res.verdict = "PARTIAL"
    else:
        res.verdict = "PASS"
    return res


# ---------------------------------------------------------------------------
# Reporting
# ---------------------------------------------------------------------------

def _fmt_expected(e: Expected) -> str:
    return f"`{e.path}` = `{e.value!r}`"


def write_report(inv: Inventory, results: list[DocResult], report_base: str) -> str:
    az = inv.azienda
    lines = [
        f"# Document Data-Fidelity Audit — {az.get('ragione_sociale')}",
        "",
        f"- **Azienda ID**: `{az.get('id')}`",
        f"- **Survey status**: `{az.get('survey_status')}`",
        f"- **Platform**: deployed API (read-back + generation + download)",
        f"- **Entity counts**: " + ", ".join(f"{k}={v}" for k, v in sorted(inv.counts.items())),
        f"- **Expected values harvested**: {len(inv.expected)} "
        f"(critical={sum(1 for e in inv.expected if e.weight=='critical')}, "
        f"normal={sum(1 for e in inv.expected if e.weight=='normal')}, "
        f"info={sum(1 for e in inv.expected if e.weight=='info')})",
        "",
        "## Verdict summary",
        "",
        "| Document | Verdict | Matched/Checked | Critical misses | Hard leaks | Soft markers |",
        "|---|---|---|---|---|---|",
    ]
    for r in results:
        lines.append(
            f"| `{r.tipo}` (v{r.version}) | **{r.verdict}** | {r.matched}/{r.checked} | "
            f"{len(r.critical_misses)} | {len(r.hard_leaks)} | {len(r.soft_markers)} |"
        )
    lines.append("")

    for r in results:
        lines.append(f"## `{r.tipo}` — {r.verdict}")
        lines.append("")
        if r.error:
            lines.append(f"- **ERROR**: {r.error}")
            lines.append("")
            continue
        lines.append(f"- doc_id `{r.doc_id}`, version {r.version}")
        lines.append(f"- structure: {r.meta}")
        lines.append(f"- data coverage: **{r.matched}/{r.checked}** relevant values found in the rendered text")
        if r.critical_misses:
            lines.append(f"- **CRITICAL MISSES ({len(r.critical_misses)})** — DB data that should appear but doesn't:")
            for e in r.critical_misses:
                lines.append(f"    - {_fmt_expected(e)}")
        if r.hard_leaks:
            lines.append(f"- **HARD LEAKS ({len(r.hard_leaks)})** — placeholder/garbage rendered:")
            for label, snippet in r.hard_leaks[:25]:
                lines.append(f"    - [{label}] `{snippet}`")
        if r.normal_misses:
            lines.append(f"- normal misses ({len(r.normal_misses)}):")
            for e in r.normal_misses[:40]:
                lines.append(f"    - {_fmt_expected(e)}")
        if r.soft_markers:
            lines.append(f"- soft markers (legit if data truly absent): {', '.join(repr(m) for m in r.soft_markers)}")
        lines.append("")

    md = "\n".join(lines)
    md_path = report_base + ".md"
    json_path = report_base + ".json"
    os.makedirs(os.path.dirname(md_path) or ".", exist_ok=True)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    payload = {
        "azienda": {"id": az.get("id"), "ragione_sociale": az.get("ragione_sociale"),
                    "survey_status": az.get("survey_status")},
        "counts": inv.counts,
        "results": [
            {
                "tipo": r.tipo, "verdict": r.verdict, "doc_id": r.doc_id,
                "version": r.version, "meta": r.meta,
                "checked": r.checked, "matched": r.matched,
                "critical_misses": [{"path": e.path, "value": e.value} for e in r.critical_misses],
                "normal_misses": [{"path": e.path, "value": e.value} for e in r.normal_misses],
                "hard_leaks": r.hard_leaks,
                "soft_markers": r.soft_markers,
                "error": r.error,
            }
            for r in results
        ],
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)
    return md_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--azienda", required=True, help="azienda UUID")
    ap.add_argument("--types", default="dvr_master",
                    help="comma-separated tipo_documento, or 'all'")
    ap.add_argument("--regenerate", action="store_true",
                    help="force fresh generation even if a completed doc exists")
    ap.add_argument("--report", default=".audit/data_fidelity",
                    help="report path base (writes .md and .json)")
    ap.add_argument("--save-docs", default=None, help="dir to save downloaded docs")
    ap.add_argument("--api", default=os.environ.get("AUDIT_API", DEFAULT_API))
    ap.add_argument("--email", default=os.environ.get("AUDIT_EMAIL", DEFAULT_EMAIL))
    ap.add_argument("--password", default=os.environ.get("AUDIT_PASSWORD", DEFAULT_PASSWORD))
    args = ap.parse_args()

    all_types = list(DOC_RELEVANCE.keys())
    types = all_types if args.types == "all" else [t.strip() for t in args.types.split(",") if t.strip()]

    client = APIClient(args.api, args.email, args.password)
    client.login()

    print(f"Reading live state of azienda {args.azienda} ...")
    inv = build_inventory(client, args.azienda)
    print(f"  harvested {len(inv.expected)} expected values; counts={inv.counts}")

    results = []
    for tipo in types:
        print(f"Auditing {tipo} ...", flush=True)
        r = audit_document(client, args.azienda, tipo, inv, args.regenerate, args.save_docs)
        print(f"  -> {r.verdict}  ({r.matched}/{r.checked} matched, "
              f"{len(r.critical_misses)} critical misses, {len(r.hard_leaks)} leaks)"
              + (f"  ERROR: {r.error}" if r.error else ""))
        results.append(r)

    path = write_report(inv, results, args.report)
    print(f"\nReport written to {path}")
    n_fail = sum(1 for r in results if r.verdict in ("FAIL", "ERROR"))
    print(f"Verdicts: " + ", ".join(f"{r.tipo}={r.verdict}" for r in results))
    sys.exit(1 if n_fail else 0)


if __name__ == "__main__":
    main()
