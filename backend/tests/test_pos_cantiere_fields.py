"""Unit tests for the POS cantiere fields (client request 2026-08-13).

Covers the new surface end-to-end at the unit level, mirroring the pure
(no-DB) style of ``test_pos_phases.py``:

1. Schema round-trip — subappalti / figure di sicurezza / dipendenti in
   cantiere / sostanze flag survive PosCreate → model_dump and
   PosUpdate → exclude_unset, and PosResponse parses an ORM-shaped row.
2. Generator rendering — the new docx sections print the operator's data
   (and the default diciture from the N2O original when the flags say no).
"""

from __future__ import annotations

import uuid
from datetime import date, datetime
from types import SimpleNamespace

import pytest
from docx import Document
from pydantic import ValidationError

from app.schemas.pos import (
    FIGURE_SICUREZZA_RUOLI,
    FiguraSicurezza,
    PosCreate,
    PosResponse,
    PosUpdate,
    Subappaltatore,
)


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------


def _persona(**overrides) -> SimpleNamespace:
    base = dict(
        id=str(uuid.uuid4()),
        nominativo="Mario Rossi",
        mansione="Elettricista",
        ruolo_rspp=False,
        ruolo_rls=False,
        ruolo_primo_soccorso=False,
        ruolo_antincendio=False,
        ruolo_preposto=False,
        ruolo_datore_lavoro=False,
        ruolo_medico_competente=False,
    )
    base.update(overrides)
    return SimpleNamespace(**base)


def _all_text(doc) -> str:
    """Paragraph + table-cell text, flattened for asserts."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Schema round-trip
# ---------------------------------------------------------------------------


def test_pos_create_roundtrips_new_fields():
    body = PosCreate(
        cantiere_indirizzo="Via Roma 12, Milano",
        data_inizio=date(2026, 9, 1),
        data_fine=date(2026, 12, 31),
        subappalti_presenti=True,
        subappaltatori=[{"ragione_sociale": "Scavi SRL", "lavori": "Opere di scavo"}],
        dipendenti_cantiere=["11111111-1111-1111-1111-111111111111"],
        figure_sicurezza=[
            {"ruolo": "rspp", "persona_id": None, "nominativo": "Ing. Bianchi"}
        ],
        sostanze_pericolose_presenti=True,
        sostanze_pericolose=[{"nome": "Malta cementizia", "uso": "Allettamento"}],
    )
    payload = body.model_dump(exclude_none=True)
    # Nested models dump to plain dicts (JSONB-safe payload for Pos(**payload)).
    assert payload["subappaltatori"] == [
        {"ragione_sociale": "Scavi SRL", "lavori": "Opere di scavo"}
    ]
    assert payload["figure_sicurezza"][0]["ruolo"] == "rspp"
    assert payload["dipendenti_cantiere"] == [
        "11111111-1111-1111-1111-111111111111"
    ]
    assert payload["subappalti_presenti"] is True
    assert payload["sostanze_pericolose_presenti"] is True


def test_pos_update_exclude_unset_leaves_other_surfaces_alone():
    """The cantiere card PUTs only its own fields — the dump must not leak
    defaults for untouched columns (that would clobber the DPI matrix)."""
    body = PosUpdate(subappalti_presenti=False, subappaltatori=[])
    dumped = body.model_dump(exclude_unset=True)
    assert set(dumped) == {"subappalti_presenti", "subappaltatori"}


def test_subappaltatore_requires_ragione_sociale():
    with pytest.raises(ValidationError):
        Subappaltatore(ragione_sociale="", lavori="x")


def test_figure_ruoli_catalog_is_stable():
    """The frontend dropdown mirrors these keys — renaming one silently
    orphans persisted rows."""
    assert FIGURE_SICUREZZA_RUOLI == (
        "datore_lavoro",
        "direttore_tecnico_cantiere",
        "capocantiere_preposto",
        "rspp",
        "rls",
        "medico_competente",
        "addetto_primo_soccorso",
        "addetto_antincendio",
    )
    # Free-text assignment: persona_id stays None.
    f = FiguraSicurezza(ruolo="direttore_tecnico_cantiere", nominativo="Geom. Verdi")
    assert f.persona_id is None


def test_pos_response_parses_orm_shaped_row():
    row = SimpleNamespace(
        id=uuid.uuid4(),
        azienda_id=uuid.uuid4(),
        cantiere_indirizzo="Via Roma 12",
        created_at=datetime(2026, 8, 13, 10, 0, 0),
        updated_at=datetime(2026, 8, 13, 10, 0, 0),
        subappalti_presenti=True,
        subappaltatori=[{"ragione_sociale": "Scavi SRL", "lavori": None}],
        dipendenti_cantiere=["abc"],
        figure_sicurezza=[
            {"ruolo": "rspp", "persona_id": "p1", "nominativo": "Ing. Bianchi"}
        ],
        sostanze_pericolose_presenti=False,
        sostanze_pericolose=[],
    )
    resp = PosResponse.model_validate(row, from_attributes=True)
    assert resp.subappaltatori[0].ragione_sociale == "Scavi SRL"
    assert resp.figure_sicurezza[0].nominativo == "Ing. Bianchi"
    assert resp.dipendenti_cantiere == ["abc"]


# ---------------------------------------------------------------------------
# Generator — subappalti
# ---------------------------------------------------------------------------


def test_generator_subappalti_default_dicitura_when_flag_off():
    from app.services.document_generator.pos import _render_subappalti

    doc = Document()
    pos = SimpleNamespace(subappalti_presenti=False, subappaltatori=[])
    _render_subappalti(doc, pos)
    text = _all_text(doc)
    assert "Non è previsto l'affidamento di lavorazioni in subappalto" in text


def test_generator_subappalti_table_when_flag_on():
    from app.services.document_generator.pos import _render_subappalti

    doc = Document()
    pos = SimpleNamespace(
        subappalti_presenti=True,
        subappaltatori=[
            {"ragione_sociale": "Scavi SRL", "lavori": "Opere di scavo"},
            {"ragione_sociale": "Ponteggi SNC", "lavori": None},
        ],
    )
    _render_subappalti(doc, pos)
    text = _all_text(doc)
    assert "Scavi SRL" in text
    assert "Opere di scavo" in text
    assert "Ponteggi SNC" in text
    assert "art. 97" in text  # vigilanza dicitura


# ---------------------------------------------------------------------------
# Generator — sostanze pericolose
# ---------------------------------------------------------------------------


def test_generator_sostanze_default_dicitura_when_flag_off():
    from app.services.document_generator.pos import _render_sostanze_pericolose

    doc = Document()
    pos = SimpleNamespace(sostanze_pericolose_presenti=False, sostanze_pericolose=[])
    _render_sostanze_pericolose(doc, pos, "ELECTRA SRL")
    text = _all_text(doc)
    assert "ELECTRA SRL non utilizzerà sostanze chimiche" in text


def test_generator_sostanze_table_when_flag_on():
    from app.services.document_generator.pos import _render_sostanze_pericolose

    doc = Document()
    pos = SimpleNamespace(
        sostanze_pericolose_presenti=True,
        sostanze_pericolose=[{"nome": "Malta cementizia", "uso": "Allettamento"}],
    )
    _render_sostanze_pericolose(doc, pos, "ELECTRA SRL")
    text = _all_text(doc)
    assert "Malta cementizia" in text
    assert "Allettamento" in text
    assert "schede di sicurezza" in text.lower()


def test_generator_sostanze_legacy_list_renders_without_flag():
    """Pre-2026-08 rows have a populated list but the flag defaulted to
    False — the list must still win over the default dicitura."""
    from app.services.document_generator.pos import _render_sostanze_pericolose

    doc = Document()
    pos = SimpleNamespace(
        sostanze_pericolose_presenti=False,
        sostanze_pericolose=[{"nome": "Primer bituminoso", "uso": "Guaine"}],
    )
    _render_sostanze_pericolose(doc, pos, "ELECTRA SRL")
    text = _all_text(doc)
    assert "Primer bituminoso" in text
    assert "non utilizzerà sostanze chimiche" not in text


# ---------------------------------------------------------------------------
# Generator — figure di sicurezza
# ---------------------------------------------------------------------------


def test_generator_figure_from_operator_selection():
    from app.services.document_generator.pos import _render_figure_sicurezza

    persona = _persona(nominativo="Luca Bianchi", ruolo_rspp=True)
    doc = Document()
    pos = SimpleNamespace(
        figure_sicurezza=[
            {"ruolo": "rspp", "persona_id": persona.id, "nominativo": None},
            {
                "ruolo": "direttore_tecnico_cantiere",
                "persona_id": None,
                "nominativo": "Geom. Verdi",
            },
        ]
    )
    _render_figure_sicurezza(doc, pos, [persona])
    text = _all_text(doc)
    # persona_id reference resolves to the Persona nominativo.
    assert "Luca Bianchi" in text
    # Free-text assignment prints verbatim.
    assert "Geom. Verdi" in text
    assert "Figure di sicurezza sul cantiere" in text
    # Mansionario boilerplate for the selected figures only.
    assert "Servizio di Prevenzione e Protezione provvede" in text
    assert "direttore tecnico di cantiere è una figura apicale" in text
    # Not auto-derived → no "precompilate" note.
    assert "precompilate automaticamente" not in text


def test_generator_figure_prefilled_from_organigramma_when_empty():
    from app.services.document_generator.pos import _render_figure_sicurezza

    ddl = _persona(nominativo="Anna Neri", ruolo_datore_lavoro=True)
    aps = _persona(nominativo="Paolo Blu", ruolo_primo_soccorso=True)
    doc = Document()
    pos = SimpleNamespace(figure_sicurezza=[])
    _render_figure_sicurezza(doc, pos, [ddl, aps])
    text = _all_text(doc)
    assert "Anna Neri" in text
    assert "Paolo Blu" in text
    assert "precompilate automaticamente" in text.replace("\n", " ")
    # Roles nobody holds still appear for review.
    assert "Medico Competente" in text


# ---------------------------------------------------------------------------
# Generator — dipendenti in cantiere
# ---------------------------------------------------------------------------


def test_generator_dipendenti_filtered_by_selection():
    from app.services.document_generator.pos import _render_dipendenti_table

    in_cantiere = _persona(nominativo="Mario In Cantiere")
    escluso = _persona(nominativo="Franco Escluso")
    doc = Document()
    _render_dipendenti_table(doc, [in_cantiere, escluso], [in_cantiere.id])
    text = _all_text(doc)
    assert "Dipendenti impegnati in cantiere" in text
    assert "Mario In Cantiere" in text
    assert "Franco Escluso" not in text


def test_generator_dipendenti_empty_selection_prints_everyone():
    from app.services.document_generator.pos import _render_dipendenti_table

    a = _persona(nominativo="Mario Rossi")
    b = _persona(nominativo="Franco Verdi")
    doc = Document()
    _render_dipendenti_table(doc, [a, b], [])
    text = _all_text(doc)
    assert "Dipendenti dell'azienda" in text
    assert "Mario Rossi" in text
    assert "Franco Verdi" in text


def test_generator_dipendenti_stale_selection_falls_back_to_everyone():
    """Persona deleted after selection → never print an empty table in a
    safety document."""
    from app.services.document_generator.pos import _render_dipendenti_table

    a = _persona(nominativo="Mario Rossi")
    doc = Document()
    _render_dipendenti_table(doc, [a], [str(uuid.uuid4())])
    text = _all_text(doc)
    assert "Mario Rossi" in text


# ---------------------------------------------------------------------------
# Generator — default diciture blocks
# ---------------------------------------------------------------------------


def test_generator_documentazione_and_misure_blocks():
    from app.services.document_generator.pos import (
        _render_documentazione_cantiere,
        _render_misure_prevenzione,
    )

    doc = Document()
    _render_documentazione_cantiere(doc)
    _render_misure_prevenzione(doc)
    text = _all_text(doc)
    assert "Documentazione da conservare in cantiere" in text
    assert "Piano operativo di sicurezza" in text
    assert "Documento Unico di Regolarità Contributiva" in text
    assert "I lavoratori devono:" in text
    assert "I lavoratori non devono:" in text
    assert "Dispositivi di protezione individuale" in text


def test_generator_dichiarazione_prefills_datore_di_lavoro():
    from app.services.document_generator.pos import _render_dichiarazione

    azienda = SimpleNamespace(
        ragione_sociale="ELECTRA SRL",
        sede_legale_via="Via Dei Chiosi 4",
        sede_legale_citta="Gorgonzola",
    )
    ddl = _persona(nominativo="Anna Neri", ruolo_datore_lavoro=True)
    doc = Document()
    _render_dichiarazione(doc, azienda, [ddl], datetime(2026, 8, 13))
    text = _all_text(doc)
    assert "Anna Neri" in text
    assert "ELECTRA SRL" in text
    assert "D I C H I A R A" in text
    assert "Rappresentante dei lavoratori per la sicurezza" in text
    assert "13/08/2026" in text
