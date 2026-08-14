"""Regression coverage for the live MMC document feedback
(2026-08-05 grouping fixes + 2026-08-13 dati occupazionali / AZIONE requests).
"""

from __future__ import annotations

import uuid
from datetime import datetime, timezone
from types import SimpleNamespace

from docx import Document
from docx.shared import Pt

from app.services.codice_fiscale import extract_age
from app.services.document_generator.allegato_mmc import (
    AllegatoMmcGenerator,
    group_mmc_rows,
)


def _person(
    number: int,
    name: str,
    *,
    order: int,
    sex: str | None = None,
    cf: str | None = None,
    age_band: str | None = ">18",
    external: bool = False,
):
    return SimpleNamespace(
        id=uuid.UUID(int=number),
        nominativo=name,
        mansione="Magazziniere",
        sesso=sex,
        codice_fiscale=cf,
        fascia_eta=age_band,
        tipologia_contrattuale="Tempo pieno",
        ordine=order,
        created_at=datetime(2026, 1, number, tzinfo=timezone.utc),
        ambienti=[],
        ruolo_datore_lavoro=False,
        ruolo_rspp=external,
        ruolo_rls=False,
        ruolo_medico_competente=False,
        is_esterno=external,
    )


def _task(
    number: int,
    person_id: uuid.UUID | None,
    name: str,
    *,
    created_day: int,
    ir: float,
):
    return SimpleNamespace(
        id=uuid.UUID(int=number),
        persona_id=person_id,
        ambiente_id=None,
        compito=name,
        created_at=datetime(2026, 2, created_day, tzinfo=timezone.utc),
        peso_kg=10,
        sesso="M",
        fascia_eta=">18",
        altezza_cm=None,
        dislocazione_cm=None,
        distanza_cm=None,
        angolo_gradi=None,
        giudizio_presa=None,
        frequenza_atti_min=None,
        durata_min=None,
        cp=25,
        fattore_a=1,
        fattore_b=1,
        fattore_c=1,
        fattore_d=1,
        fattore_e=1,
        fattore_f=1,
        plr=25,
        indice_ir=ir,
        livello_rischio="VERDE" if ir <= 0.75 else "ROSSO",
        area_classificazione="Verde" if ir <= 0.75 else "Rossa",
        note=None,
        misure_proposte=None,
    )


def _table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def test_mmc_groups_interleaved_tasks_and_renders_dvr_master_roster():
    first = _person(
        1,
        "Anna Rossi",
        order=1,
        sex=None,
        cf="BNCMRA90E62H501W",
    )
    second = _person(2, "Bruno Bianchi", order=2, sex="M", cf="malformed")
    external = _person(3, "Dott. Esterno", order=0, external=True)
    rows = [
        _task(22, second.id, "Scarico pallet", created_day=2, ir=1.20),
        _task(12, first.id, "Deposito cassa", created_day=2, ir=0.70),
        _task(11, first.id, "Prelievo cassa", created_day=1, ir=0.60),
    ]

    groups = group_mmc_rows(rows, [second, external, first])
    assert [group[0].nominativo for group in groups] == [
        "Anna Rossi",
        "Bruno Bianchi",
    ]
    assert [[task.compito for task in group[1]] for group in groups] == [
        ["Prelievo cassa", "Deposito cassa"],
        ["Scarico pallet"],
    ]

    doc = Document()
    generator = object.__new__(AllegatoMmcGenerator)
    generator._add_dati_occupazionali(
        doc, [second, external, first], [first, second]
    )

    # Client request 2026-08-13: two tables — the DVR-Master-identical dati
    # occupazionali grid (codice fiscale included) plus the reduced roster.
    assert len(doc.tables) == 2

    dvr_clone = _table_text(doc.tables[0])
    assert dvr_clone[0] == [
        "Nominativo",
        "Mansione",
        "Ambiente di Lavoro",
        "Codice Fiscale",
        "Tipologia contrattuale",
    ]
    rendered = "\n".join(cell for row in dvr_clone for cell in row)
    # Identical to the Risk Master means the CF column IS present now.
    assert "BNCMRA90E62H501W" in rendered
    # External safety consultants are excluded, exactly like the DVR Master.
    assert "DOTT. ESTERNO" not in rendered
    assert "ANNA ROSSI" in rendered
    assert "BRUNO BIANCHI" in rendered


def test_mmc_worker_roster_has_exactly_name_sex_age_mansione_columns():
    anna = _person(1, "Anna Rossi", order=1, sex=None, cf="BNCMRA90E62H501W")
    bruno = _person(2, "Bruno Bianchi", order=2, sex="M", cf="malformed")

    doc = Document()
    generator = object.__new__(AllegatoMmcGenerator)
    generator._add_dati_occupazionali(doc, [anna, bruno], [anna, bruno])

    roster = _table_text(doc.tables[1])
    assert roster[0] == ["Nome e Cognome", "Sesso", "Età", "Mansione"]
    # One row per assessed worker.
    assert len(roster) == 1 + 2
    # Anna: sesso derived from her (valid, female) CF; età computed from it.
    assert roster[1] == [
        "Anna Rossi",
        "F",
        str(extract_age("BNCMRA90E62H501W")),
        "Magazziniere",
    ]
    # Bruno: malformed CF → declared sesso kept, età falls back to the band.
    assert roster[2] == ["Bruno Bianchi", "M", ">18", "Magazziniere"]


def test_all_tasks_appear_in_grouped_detail_and_complete_synopsis():
    anna = _person(1, "Anna Rossi", order=1)
    rows = [
        _task(12, anna.id, "Deposito cassa", created_day=2, ir=1.20),
        _task(11, anna.id, "Prelievo cassa", created_day=1, ir=0.60),
        _task(31, None, "Valutazione anonima", created_day=1, ir=0.80),
        _task(32, uuid.UUID(int=99), "Persona eliminata", created_day=2, ir=0.90),
    ]
    groups = group_mmc_rows(rows, [anna])

    assert groups[-1][0] is None
    assert [task.compito for task in groups[-1][1]] == [
        "Valutazione anonima",
        "Persona eliminata",
    ]

    generator = object.__new__(AllegatoMmcGenerator)
    detail = Document()
    generator._add_per_worker_assessments(detail, groups, {})
    detail_text = [paragraph.text for paragraph in detail.paragraphs]
    assert detail_text.index("Anna Rossi — Magazziniere") < detail_text.index(
        "1. Prelievo cassa"
    ) < detail_text.index("2. Deposito cassa")
    assert detail_text.index("Valutazioni non associate") < detail_text.index(
        "3. Valutazione anonima"
    ) < detail_text.index("4. Persona eliminata")

    synopsis = Document()
    generator._add_quadro_sinottico(synopsis, groups)
    table = _table_text(synopsis.tables[0])
    assert table[0] == ["Nominativo", "Mansione", "Compito", "IR", "Area"]
    assert [row[2] for row in table[1:]] == [
        "Prelievo cassa",
        "Deposito cassa",
        "Valutazione anonima",
        "Persona eliminata",
    ]
    assert len(table) - 1 == len(rows)


def test_programma_attuazione_uses_azione_header_and_one_row_per_employee():
    anna = _person(1, "Anna Rossi", order=1)
    bruno = _person(2, "Bruno Bianchi", order=2, sex="M")
    green_1 = _task(11, anna.id, "Prelievo cassa", created_day=1, ir=0.60)
    green_2 = _task(12, anna.id, "Deposito cassa", created_day=2, ir=0.70)
    red = _task(21, bruno.id, "Scarico pallet", created_day=3, ir=1.20)
    red.misure_proposte = "Adottare transpallet elettrico."

    generator = object.__new__(AllegatoMmcGenerator)
    doc = Document()
    generator._add_programma_attuazione(
        doc,
        [green_1, green_2, red],
        {anna.id: anna, bruno.id: bruno},
    )

    green_table = _table_text(doc.tables[0])
    assert green_table[0] == [
        "Lavoratore",
        "AZIONE",
        "Misure di prevenzione e protezione",
    ]
    # Anna's two green lifting actions collapse into a single row, with the
    # (identical) default zone measure deduplicated to one copy.
    assert len(green_table) == 2
    assert green_table[1][0] == "Anna Rossi"
    assert green_table[1][1] == "Prelievo cassa\nDeposito cassa"
    assert green_table[1][2].count("Mantenere le condizioni operative") == 1

    red_table = _table_text(doc.tables[1])
    assert red_table[0] == [
        "Lavoratore",
        "AZIONE",
        "Misure di prevenzione e protezione",
    ]
    assert red_table[1] == [
        "Bruno Bianchi",
        "Scarico pallet",
        "Adottare transpallet elettrico.",
    ]


def test_programma_attuazione_keeps_unassigned_assessments_on_separate_rows():
    orphan_1 = _task(31, None, "Valutazione anonima", created_day=1, ir=0.60)
    orphan_2 = _task(32, uuid.UUID(int=99), "Persona eliminata", created_day=2, ir=0.70)

    generator = object.__new__(AllegatoMmcGenerator)
    doc = Document()
    generator._add_programma_attuazione(doc, [orphan_1, orphan_2], {})

    green_table = _table_text(doc.tables[0])
    # No associated worker → never merged into one anonymous mega-row.
    assert [row[0] for row in green_table[1:]] == ["—", "—"]
    assert [row[1] for row in green_table[1:]] == [
        "Valutazione anonima",
        "Persona eliminata",
    ]


def test_empty_prevention_zone_notice_has_explicit_spacing_after_long_heading():
    anna = _person(1, "Anna Rossi", order=1)
    green = _task(11, anna.id, "Prelievo cassa", created_day=1, ir=0.60)
    generator = object.__new__(AllegatoMmcGenerator)
    doc = Document()

    generator._add_programma_attuazione(doc, [green], {anna.id: anna})

    notice = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text == "Nessun lavoratore classificato in gialla."
    )
    assert notice.paragraph_format.space_before == Pt(6)
