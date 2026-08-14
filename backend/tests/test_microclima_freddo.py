"""Cold stress (freddo severo) — UNI EN ISO 11079 (IREQ).

Covers the third microclima evaluation type added next to comfort (PMV/PPD)
and severe heat (PHS):

* calculator classification bands (ACCETTABILE / LIMITE / CRITICO + DLE) and
  the Annex D wind-chill / frostbite screening, including its escalation of
  the overall level;
* persistence schemas accepting ``severo_freddo`` rows with cold-environment
  input ranges;
* the Allegato Microclima Severo rendering a full Parte II (freddo) section,
  generated off-DB via the same harness as ``test_generators.py``.
"""

import asyncio
import importlib.util
import sys
from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError

BACKEND_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BACKEND_ROOT))

from app.schemas.microclima import IreqCalcRequest, MicroclimaCreate  # noqa: E402
from app.services.microclima_calculator import calculate_ireq  # noqa: E402


# ---------------------------------------------------------------------------
# Calculator — classification bands (ISO 11079 decision scheme)
# ---------------------------------------------------------------------------


class TestIreqClassificationBands:
    def test_adequate_insulation_in_mild_cold_is_accettabile(self):
        """+5 °C, hard work, winter clothing: balance holds for the shift."""
        r = calculate_ireq(
            air_temp=5.0,
            mean_radiant_temp=5.0,
            air_velocity=0.2,
            humidity=70.0,
            metabolic_rate=2.5,
            clothing_insulation=1.5,
        )
        assert r["livello"] == "ACCETTABILE"
        assert r["ireq_neutral"] < 1.5  # clothing exceeds the requirement
        assert r["dle_min"] is None  # no exposure limit when balance holds
        assert r["delta_clo"] == 0.0

    def test_insulation_between_minimal_and_neutral_is_limite(self):
        """−5 °C, light work: Icl between IREQminimal and IREQneutral."""
        r = calculate_ireq(
            air_temp=-5.0,
            mean_radiant_temp=-5.0,
            air_velocity=0.5,
            humidity=70.0,
            metabolic_rate=1.6,
            clothing_insulation=2.75,
        )
        assert r["livello"] == "LIMITE"
        assert r["ireq_minimal"] <= 2.75 < r["ireq_neutral"]
        assert r["dle_min"] is None  # DLE binds only below IREQminimal
        assert r["delta_clo"] > 0  # extra insulation is recommended

    def test_freezer_room_with_thin_clothing_is_critico_with_dle(self):
        """−25 °C freezer with 2.0 clo: below IREQminimal → time-limited."""
        r = calculate_ireq(
            air_temp=-25.0,
            mean_radiant_temp=-25.0,
            air_velocity=0.4,
            humidity=85.0,
            metabolic_rate=2.0,
            clothing_insulation=2.0,
        )
        assert r["livello"] == "CRITICO"
        assert 2.0 < r["ireq_minimal"] < r["ireq_neutral"]
        assert r["dle_min"] is not None
        # Body heat debt of 40 Wh/m² is reached within roughly an hour here.
        assert 20 <= r["dle_min"] <= 90
        # Calm air at −25 °C: t_wc equals the air temperature → MODERATO band.
        assert r["t_wc"] == -25.0
        assert r["frostbite_risk"] == "MODERATO"

    def test_extreme_wind_chill_escalates_to_critico(self):
        """−40 °C with 10 m/s wind: insulation is adequate for the general
        balance, but exposed skin freezes within minutes (Annex D) — the
        local-cooling screening must govern."""
        r = calculate_ireq(
            air_temp=-40.0,
            mean_radiant_temp=-40.0,
            air_velocity=10.0,
            humidity=70.0,
            metabolic_rate=3.0,
            clothing_insulation=5.0,
        )
        assert r["ireq_neutral"] <= 5.0  # general cooling would be acceptable
        assert r["t_wc"] is not None and r["t_wc"] <= -60
        assert r["frostbite_risk"] == "ESTREMO"
        assert r["livello"] == "CRITICO"
        assert r["dle_min"] is None  # limit driven by frostbite, not heat debt

    def test_moderate_frostbite_band_escalates_accettabile_to_limite(self):
        r = calculate_ireq(
            air_temp=-30.0,
            mean_radiant_temp=-30.0,
            air_velocity=0.2,
            humidity=70.0,
            metabolic_rate=3.0,
            clothing_insulation=3.0,
        )
        assert r["frostbite_risk"] == "MODERATO"
        assert r["livello"] == "LIMITE"


class TestIreqWindChill:
    def test_calm_air_wind_chill_equals_air_temperature(self):
        """Below the 4.8 km/h validity floor of the JAG/TI formula the index
        is undefined; calm air gives no aggravation."""
        r = calculate_ireq(
            air_temp=5.0,
            mean_radiant_temp=5.0,
            air_velocity=0.2,
            humidity=50.0,
            metabolic_rate=2.0,
            clothing_insulation=1.5,
        )
        assert r["t_wc"] == 5.0
        assert r["frostbite_risk"] == "BASSO"

    def test_windy_minus_ten_gives_low_frostbite_risk(self):
        """−10 °C with 5 m/s wind → t_wc ≈ −19 °C, still above the −25 °C
        Annex D threshold."""
        r = calculate_ireq(
            air_temp=-10.0,
            mean_radiant_temp=-10.0,
            air_velocity=5.0,
            humidity=70.0,
            metabolic_rate=2.5,
            clothing_insulation=2.5,
        )
        assert r["t_wc"] is not None
        assert -21 < r["t_wc"] < -17
        assert r["frostbite_risk"] == "BASSO"


# ---------------------------------------------------------------------------
# Schemas — severo_freddo rows and the IREQ preview request
# ---------------------------------------------------------------------------


class TestFreddoSchemas:
    def test_create_schema_accepts_cold_environment_row(self):
        row = MicroclimaCreate(
            nome_area="Cella surgelati",
            tipo_ambiente="severo_freddo",
            temperatura_aria=-40.0,
            temperatura_radiante=-40.0,
            velocita_aria=12.0,
            umidita_relativa=85.0,
            metabolismo=2.0,
            isolamento_vestiario=3.5,
            ireq_neutral=3.7,
            ireq_minimal=3.5,
            t_wind_chill=-40.0,
            dle_freddo=120.0,
            livello_rischio="CRITICO",
        )
        assert row.tipo_ambiente == "severo_freddo"
        assert row.dle_freddo == 120.0

    def test_ireq_request_rejects_warm_environment(self):
        with pytest.raises(ValidationError):
            IreqCalcRequest(
                air_temp=20.0,  # above the 10 °C cold-environment bound
                mean_radiant_temp=10.0,
                air_velocity=0.5,
                humidity=50.0,
                metabolic_rate=2.0,
                clothing_insulation=1.0,
            )


# ---------------------------------------------------------------------------
# Document generation — Allegato Microclima Severo, Parte II (freddo)
# ---------------------------------------------------------------------------


def _load_verify():
    spec = importlib.util.spec_from_file_location(
        "verify_all_generators",
        str(BACKEND_ROOT / "scripts" / "verify_all_generators.py"),
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _document_text(path: str) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _generate_severo(module, fixture, output_dir) -> str:
    module.patch_generators(fixture, str(output_dir))
    ok, path, message = asyncio.run(
        module.run_one("ALLEGATO_MICROCLIMA_SEVERO", fixture["azienda"].id)
    )
    assert ok, message
    return path


def test_allegato_severo_renders_freddo_valutazione(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()
    fixture["microclima"] = fixture["microclima"] + [
        module.mk(
            ambiente_id=None,
            nome_area="Cella surgelati -25",
            tipo_ambiente="severo_freddo",
            temperatura_aria=-25.0,
            temperatura_radiante=-25.0,
            velocita_aria=0.4,
            umidita_relativa=85.0,
            metabolismo=2.0,
            isolamento_vestiario=2.0,
        ),
        module.mk(
            ambiente_id=None,
            nome_area="Cella frigo +2",
            tipo_ambiente="severo_freddo",
            temperatura_aria=2.0,
            temperatura_radiante=2.0,
            velocita_aria=0.3,
            umidita_relativa=85.0,
            metabolismo=2.0,
            isolamento_vestiario=2.0,
        ),
    ]

    path = _generate_severo(module, fixture, tmp_path)
    text = _document_text(path)

    # Parte II exists, cites the norm, and carries both areas with results.
    assert "STRESS DA FREDDO" in text
    assert "11079" in text
    assert "IREQ" in text
    assert "Cella surgelati -25" in text
    assert "Cella frigo +2" in text
    # The freezer row is under-insulated → CRITICO with a numeric DLE; the
    # +2 °C room with the same clothing is fine.
    assert "CRITICO" in text
    assert "ACCETTABILE" in text
    # The heat part is still rendered (fixture ships one severo_caldo row).
    assert "PHS" in text
    assert "Mensa aziendale con cucina" in text


def test_allegato_severo_without_freddo_rows_says_so(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()  # default fixture: no severo_freddo rows

    path = _generate_severo(module, fixture, tmp_path)
    text = _document_text(path)

    assert "STRESS DA FREDDO" in text  # the section is always present
    assert "Nessun ambiente a rischio da freddo severo registrato." in text
