"""Unit tests for the in-browser preview + inline editing service.

Pins the contract shared with the frontend document editor: the address
scheme emitted by ``parse_docx_to_blocks``, the run/cell/table payload
shapes, the editable=False locking rules, and the override application
semantics (``apply_overrides_to_docx`` preserves formatting, turns "\\n"
into line breaks — never new paragraphs — and skips stale addresses).

Tests stay unit-level — no live DB — matching test_document_gdoc_editing.py.
All fixture .docx files are built in-memory with python-docx.
"""

from __future__ import annotations

import base64
import io
import uuid
from datetime import datetime

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pydantic import ValidationError

from app.models.documento_generato import DocumentoGenerato
from app.schemas.document import (
    DocumentPreviewResponse,
    OverridesPatchRequest,
    OverridesResponse,
)
from app.services.document_preview import (
    apply_overrides_to_docx,
    extract_image,
    extract_image_fast,
    parse_docx_to_blocks,
)

# 1x1 red-pixel PNG — enough for python-docx's image header parser.
_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _para_text(paragraph_payload: dict) -> str:
    return "".join(run["text"] for run in paragraph_payload["runs"])


def _add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_v_merge(cell, val: str | None) -> None:
    """val="restart" starts a vertical merge; None appends a bare <w:vMerge/>."""
    tc_pr = cell._tc.get_or_add_tcPr()
    vm = OxmlElement("w:vMerge")
    if val is not None:
        vm.set(qn("w:val"), val)
    tc_pr.append(vm)


def _add_toc_field_char(paragraph) -> None:
    """Append a run carrying a fldChar — the TOC machinery marker."""
    run = paragraph.add_run("Sommario")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append a real external <w:hyperlink r:id=...> wrapping one run."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_fld_simple(paragraph, instr: str, text: str) -> None:
    """Append a <w:fldSimple w:instr=...> with its cached-result run."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    fld.append(run)
    paragraph._p.append(fld)


def _build_sample_docx() -> bytes:
    """Mixed-content document exercising every parse feature.

    Block layout (addresses):
      0 — Heading 1
      1 — paragraph with formatted runs (bold+size+color / italic+underline)
      2 — empty paragraph (must be kept — carries spacing)
      3 — centered paragraph
      4 — 2x3 table: shading on (0,0), horizontal merge (0,1)+(0,2)
      5 — paragraph with page-break-before (pPr)
      6 — 2x1 table with a vertical merge (restart / continue)
      7 — paragraph whose run carries a <w:br w:type="page"/>
    """
    doc = Document()
    doc.add_heading("Titolo principale", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Grassetto")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run2 = p.add_run(" corsivo")
    run2.italic = True
    run2.underline = True

    doc.add_paragraph("")

    centered = doc.add_paragraph("Centrato")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Intestazione"
    _add_shading(table.cell(0, 0), "1a237e")
    merged = table.cell(0, 1).merge(table.cell(0, 2))
    merged.text = "Unita"
    table.cell(1, 0).text = "Riga2"

    pb = doc.add_paragraph("Nuova pagina")
    pb.paragraph_format.page_break_before = True

    vt = doc.add_table(rows=2, cols=1)
    vt.cell(0, 0).text = "Verticale"
    _add_v_merge(vt.cell(0, 0), "restart")
    _add_v_merge(vt.cell(1, 0), None)

    br_p = doc.add_paragraph("Prima del salto")
    br_p.add_run().add_break(WD_BREAK.PAGE)

    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Parse: block layout, address stability, formatting fields
# ---------------------------------------------------------------------------


def test_parse_block_layout_and_addresses():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    # sectPr never consumes an index — exactly the 8 authored blocks.
    assert len(blocks) == 8
    assert [b["addr"] for b in blocks] == [str(i) for i in range(8)]
    kinds = [b["kind"] for b in blocks]
    assert kinds == [
        "paragraph", "paragraph", "paragraph", "paragraph",
        "table", "paragraph", "table", "paragraph",
    ]


def test_parse_heading_and_alignment():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    assert blocks[0]["style"] == "Heading 1"
    assert blocks[0]["heading_level"] == 1
    assert _para_text(blocks[0]) == "Titolo principale"
    assert blocks[1]["heading_level"] is None
    assert blocks[3]["alignment"] == "center"


def test_parse_empty_paragraph_kept():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    assert blocks[2]["kind"] == "paragraph"
    assert blocks[2]["runs"] == []
    assert blocks[2]["editable"] is True


def test_parse_run_formatting_fields():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    bold_run, italic_run = blocks[1]["runs"]
    assert bold_run["text"] == "Grassetto"
    assert bold_run["bold"] is True
    assert bold_run["size"] == 14.0
    assert bold_run["color"] == "FF0000"
    assert bold_run["italic"] is False
    assert italic_run["italic"] is True
    assert italic_run["underline"] is True
    # Inherited formatting stays None — the frontend falls back to style CSS.
    assert italic_run["color"] is None
    assert italic_run["size"] is None


def test_parse_table_shading_and_merges():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    table = blocks[4]
    # Horizontal merge removed one w:tc from row 0 — addresses follow XML.
    assert len(table["rows"][0]) == 2
    assert len(table["rows"][1]) == 3
    header = table["rows"][0][0]
    assert header["addr"] == "4:0:0"
    assert header["shading"] == "1A237E"
    assert _para_text(header["paragraphs"][0]) == "Intestazione"
    merged = table["rows"][0][1]
    assert merged["addr"] == "4:0:1"
    assert merged["col_span"] == 2
    assert merged["paragraphs"][0]["addr"] == "4:0:1:0"

    v_table = blocks[6]
    assert v_table["rows"][0][0]["v_merge"] == "restart"
    assert v_table["rows"][1][0]["v_merge"] == "continue"


def test_parse_page_breaks():
    blocks = parse_docx_to_blocks(_build_sample_docx())
    assert blocks[5]["page_break_before"] is True  # pPr pageBreakBefore
    assert blocks[7]["page_break_before"] is True  # run-level <w:br type="page"/>
    assert blocks[0]["page_break_before"] is False


# ---------------------------------------------------------------------------
# Parse: editable=False locking rules (TOC fields, images, nested tables)
# ---------------------------------------------------------------------------


def test_toc_field_paragraph_not_editable():
    doc = Document()
    doc.add_paragraph("Normale")
    _add_toc_field_char(doc.add_paragraph())
    blocks = parse_docx_to_blocks(_to_bytes(doc))
    assert blocks[0]["editable"] is True
    assert blocks[1]["editable"] is False


def test_image_paragraph_not_editable_and_lists_image():
    doc = Document()
    doc.add_paragraph("Prima")
    doc.add_picture(io.BytesIO(_PNG_1PX))
    blocks = parse_docx_to_blocks(_to_bytes(doc))
    img_block = blocks[1]
    assert img_block["editable"] is False
    assert len(img_block["images"]) == 1
    image = img_block["images"][0]
    assert image["image_id"].startswith("rId")
    # 1x1 px at PNG-default 72dpi -> 12700 EMU -> round(12700/914400*96) == 1
    assert image["width_px"] == 1
    assert image["height_px"] == 1


def test_nested_table_flattened_as_locked_paragraphs():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Diretto"
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Annidato"
    blocks = parse_docx_to_blocks(_to_bytes(doc))
    paragraphs = blocks[0]["rows"][0][0]["paragraphs"]
    # Direct w:p children come first (addressable), the nested table's text
    # is appended after as locked paragraphs continuing the index sequence.
    assert _para_text(paragraphs[0]) == "Diretto"
    assert paragraphs[0]["editable"] is True
    nested_paras = [p for p in paragraphs if _para_text(p) == "Annidato"]
    assert len(nested_paras) == 1
    assert nested_paras[0]["editable"] is False
    addrs = [p["addr"] for p in paragraphs]
    assert addrs == [f"0:0:0:{i}" for i in range(len(paragraphs))]


def test_page_break_run_paragraph_locked_and_break_survives():
    """A run-level <w:br w:type="page"/> locks the paragraph (B1).

    Editing would delete the break's run and the following chapter would
    silently stop starting on a new page — so apply must skip it.
    """
    docx_bytes = _build_sample_docx()
    blocks = parse_docx_to_blocks(docx_bytes)
    assert blocks[7]["editable"] is False
    # pPr page_break_before lives OUTSIDE the runs — replacement can't
    # destroy it, so that paragraph stays editable.
    assert blocks[5]["editable"] is True

    edited = apply_overrides_to_docx(docx_bytes, {"7": "Testo sostitutivo"})
    reparsed = parse_docx_to_blocks(edited)
    assert _para_text(reparsed[7]).startswith("Prima del salto")
    assert reparsed[7]["page_break_before"] is True
    document = Document(io.BytesIO(edited))
    target = next(
        p for p in document.paragraphs if p.text.startswith("Prima del salto")
    )
    assert len(target._p.xpath('.//w:br[@w:type="page"]')) == 1


def test_column_break_run_paragraph_locked():
    doc = Document()
    p = doc.add_paragraph("Colonna")
    p.add_run().add_break(WD_BREAK.COLUMN)
    blocks = parse_docx_to_blocks(_to_bytes(doc))
    assert blocks[0]["editable"] is False


def test_plain_line_break_does_not_lock():
    """Typeless <w:br/> (the apply path's own output) must stay editable."""
    doc = Document()
    p = doc.add_paragraph("Prima riga")
    p.add_run().add_break(WD_BREAK.LINE)
    p.add_run("Seconda riga")
    blocks = parse_docx_to_blocks(_to_bytes(doc))
    assert blocks[0]["editable"] is True


def test_hyperlink_paragraph_text_visible_and_locked():
    """Runs inside <w:hyperlink> show in the preview but stay read-only (B2)."""
    doc = Document()
    p = doc.add_paragraph("Visita ")
    _add_hyperlink(p, "https://www.inail.it", "www.inail.it")
    docx_bytes = _to_bytes(doc)

    blocks = parse_docx_to_blocks(docx_bytes)
    assert _para_text(blocks[0]) == "Visita www.inail.it"
    assert blocks[0]["editable"] is False

    # An override attempt is skipped — link element AND text survive.
    edited = apply_overrides_to_docx(docx_bytes, {"0": "distruggi il link"})
    reparsed = parse_docx_to_blocks(edited)
    assert _para_text(reparsed[0]) == "Visita www.inail.it"
    document = Document(io.BytesIO(edited))
    assert len(document.paragraphs[0]._p.xpath(".//w:hyperlink")) == 1


def test_fld_simple_paragraph_text_visible_and_locked():
    """<w:fldSimple> field codes lock the paragraph, text stays visible (B4)."""
    doc = Document()
    p = doc.add_paragraph()
    _add_fld_simple(p, " PAGE ", "42")
    docx_bytes = _to_bytes(doc)

    blocks = parse_docx_to_blocks(docx_bytes)
    assert _para_text(blocks[0]) == "42"
    assert blocks[0]["editable"] is False

    # Apply skips — no duplicated out-of-order text, field survives intact.
    edited = apply_overrides_to_docx(docx_bytes, {"0": "43"})
    assert _para_text(parse_docx_to_blocks(edited)[0]) == "42"
    document = Document(io.BytesIO(edited))
    assert len(document.paragraphs[0]._p.xpath(".//w:fldSimple")) == 1


# ---------------------------------------------------------------------------
# Image extraction (preview <img> endpoint)
# ---------------------------------------------------------------------------


def test_extract_image_round_trip():
    doc = Document()
    doc.add_picture(io.BytesIO(_PNG_1PX))
    docx_bytes = _to_bytes(doc)
    rid = parse_docx_to_blocks(docx_bytes)[0]["images"][0]["image_id"]

    extracted = extract_image(docx_bytes, rid)
    assert extracted is not None
    blob, content_type = extracted
    assert blob == _PNG_1PX
    assert content_type == "image/png"


def test_extract_image_unknown_or_non_image_rid():
    docx_bytes = _to_bytes(Document())
    assert extract_image(docx_bytes, "rId999") is None
    # Every default-template rId points at a non-image part (styles, ...).
    document = Document(io.BytesIO(docx_bytes))
    non_image_rid = next(iter(document.part.rels))
    assert extract_image(docx_bytes, non_image_rid) is None


def test_extract_image_fast_matches_document_parse_path():
    """The zip-level fast path (B6) agrees with the python-docx slow path."""
    doc = Document()
    doc.add_picture(io.BytesIO(_PNG_1PX))
    docx_bytes = _to_bytes(doc)
    rid = parse_docx_to_blocks(docx_bytes)[0]["images"][0]["image_id"]

    assert extract_image_fast(docx_bytes, rid) == extract_image(docx_bytes, rid)
    assert extract_image_fast(docx_bytes, rid) == (_PNG_1PX, "image/png")

    # Unknown / non-image relationship ids miss, exactly like the slow path.
    assert extract_image_fast(docx_bytes, "rId999") is None
    document = Document(io.BytesIO(docx_bytes))
    non_image_rid = next(
        r_id
        for r_id, rel in document.part.rels.items()
        if not rel.reltype.endswith("/image")
    )
    assert extract_image_fast(docx_bytes, non_image_rid) is None
    # Garbage bytes don't raise — they just miss.
    assert extract_image_fast(b"not a zip", "rId1") is None


# ---------------------------------------------------------------------------
# Apply: formatting preservation, line breaks, cell targeting, skip rules
# ---------------------------------------------------------------------------


def test_apply_single_line_preserves_bold_formatting():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Vecchio testo")
    run.bold = True
    run.font.size = Pt(14)

    edited = apply_overrides_to_docx(_to_bytes(doc), {"0": "Nuovo testo"})
    blocks = parse_docx_to_blocks(edited)
    assert _para_text(blocks[0]) == "Nuovo testo"
    assert blocks[0]["runs"][0]["bold"] is True
    assert blocks[0]["runs"][0]["size"] == 14.0


def test_apply_multiline_produces_line_breaks_not_paragraphs():
    doc = Document()
    doc.add_paragraph("Prima")
    doc.add_paragraph("Seconda")
    original = parse_docx_to_blocks(_to_bytes(doc))

    edited = apply_overrides_to_docx(
        _to_bytes(doc), {"0": "Linea 1\nLinea 2\nLinea 3"}
    )
    blocks = parse_docx_to_blocks(edited)
    # Same block count — line breaks never split into new paragraphs.
    assert len(blocks) == len(original)
    # python-docx renders <w:br/> as "\n" in run text, so the round-trip
    # reproduces the override verbatim.
    assert _para_text(blocks[0]) == "Linea 1\nLinea 2\nLinea 3"
    assert _para_text(blocks[1]) == "Seconda"

    document = Document(io.BytesIO(edited))
    brs = document.paragraphs[0]._p.xpath('.//w:br')
    assert len(brs) == 2


def test_apply_normalizes_crlf_to_single_line_breaks():
    """CRLF/CR overrides produce ONE <w:br/> per visual line break (B3)."""
    doc = Document()
    doc.add_paragraph("Vecchio")

    edited = apply_overrides_to_docx(
        _to_bytes(doc), {"0": "Linea 1\r\nLinea 2\rLinea 3"}
    )
    blocks = parse_docx_to_blocks(edited)
    assert _para_text(blocks[0]) == "Linea 1\nLinea 2\nLinea 3"
    assert "\r" not in _para_text(blocks[0])
    # The edited paragraph carries only typeless breaks — still editable.
    assert blocks[0]["editable"] is True

    document = Document(io.BytesIO(edited))
    brs = document.paragraphs[0]._p.xpath(".//w:br")
    assert len(brs) == 2  # not 4 — no doubled breaks from \r\n


def test_apply_cell_paragraph_override_lands_in_right_cell():
    doc = Document()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(1, 1).text = "Vecchio"

    edited = apply_overrides_to_docx(_to_bytes(doc), {"1:1:1:0": "Nuovo"})
    blocks = parse_docx_to_blocks(edited)
    rows = blocks[1]["rows"]
    assert _para_text(rows[1][1]["paragraphs"][0]) == "Nuovo"
    # Neighbouring cells untouched.
    assert _para_text(rows[0][0]["paragraphs"][0]) == "A"
    assert _para_text(rows[1][0]["paragraphs"][0]) == ""


def test_apply_skips_locked_and_bogus_addresses():
    doc = Document()
    _add_toc_field_char(doc.add_paragraph())  # addr 0 — locked (fldChar)
    doc.add_picture(io.BytesIO(_PNG_1PX))     # addr 1 — locked (drawing)
    doc.add_paragraph("Modificabile")         # addr 2
    docx_bytes = _to_bytes(doc)

    edited = apply_overrides_to_docx(
        docx_bytes,
        {
            "0": "NON deve comparire",     # locked: fldChar
            "1": "NON deve comparire",     # locked: drawing
            "99": "fuori documento",       # unresolvable index
            "0:0:0:0": "non una tabella",  # 4-part addr on a paragraph
            "abc": "non numerico",         # malformed
            "2": "Sostituito",
        },
    )
    blocks = parse_docx_to_blocks(edited)
    assert _para_text(blocks[0]) == "Sommario"
    assert len(blocks[1]["images"]) == 1
    assert _para_text(blocks[2]) == "Sostituito"


def test_round_trip_new_text_at_same_address():
    docx_bytes = _build_sample_docx()
    original = parse_docx_to_blocks(docx_bytes)
    target_addr = original[3]["addr"]  # the centered paragraph

    edited = apply_overrides_to_docx(docx_bytes, {target_addr: "Testo aggiornato"})
    reparsed = parse_docx_to_blocks(edited)
    assert [b["addr"] for b in reparsed] == [b["addr"] for b in original]
    assert _para_text(reparsed[3]) == "Testo aggiornato"
    # Alignment (paragraph-level formatting) survives the run replacement.
    assert reparsed[3]["alignment"] == "center"
    # Untouched blocks are byte-for-byte semantically identical.
    assert _para_text(reparsed[0]) == "Titolo principale"
    assert reparsed[4]["rows"][0][1]["col_span"] == 2


# ---------------------------------------------------------------------------
# Model column + schema contract
# ---------------------------------------------------------------------------


def test_documento_generato_has_content_overrides_column():
    cols = {c.name: c for c in DocumentoGenerato.__table__.columns}
    assert "content_overrides" in cols, "content_overrides column missing"
    assert cols["content_overrides"].nullable is True


def test_preview_response_validates_parse_output():
    """The parser's dicts must satisfy the response schema 1:1."""
    blocks = parse_docx_to_blocks(_build_sample_docx())
    resp = DocumentPreviewResponse(
        id=uuid.uuid4(),
        azienda_id=uuid.uuid4(),
        azienda_nome="ACME SRL",
        tipo_documento="dvr_master",
        versione=1,
        file_name="DVR_v1.docx",
        stale_snapshot=False,
        generated_at="2026-07-17T10:00:00",
        blocks=blocks,
        overrides={"3": "Testo salvato"},
    )
    assert resp.blocks[0].kind == "paragraph"
    assert resp.blocks[4].kind == "table"
    assert resp.blocks[4].rows[0][1].col_span == 2
    assert resp.overrides == {"3": "Testo salvato"}


def test_overrides_patch_request_accepts_null_deletes():
    body = OverridesPatchRequest(set={"12": "Nuovo", "13:0:2:0": None})
    assert body.set["12"] == "Nuovo"
    assert body.set["13:0:2:0"] is None
    resp = OverridesResponse(overrides={})
    assert resp.overrides == {}


def test_overrides_patch_request_size_limits():
    """B11 — schema-level caps: 20k chars per value, 500 entries per request."""
    # At-limit payloads pass.
    OverridesPatchRequest(set={"0": "x" * 20000})
    OverridesPatchRequest(set={str(i): "x" for i in range(500)})
    # Over-limit payloads are rejected before reaching the endpoint.
    with pytest.raises(ValidationError):
        OverridesPatchRequest(set={"0": "x" * 20001})
    with pytest.raises(ValidationError):
        OverridesPatchRequest(set={str(i): "x" for i in range(501)})


class _FakeDocRow:
    """Stand-in for DocumentoGenerato with the attrs `_doc_to_response` reads."""

    def __init__(self, options=None):
        self.id = uuid.uuid4()
        self.azienda_id = uuid.uuid4()
        self.tipo_documento = "dvr_master"
        self.versione = 2
        self.status = "completed"
        self.file_path = None
        self.gdrive_file_id = None
        self.gdoc_file_id = None
        self.options = options
        self.error_message = None
        self.created_at = datetime(2026, 7, 17, 10, 0, 0)
        self.stale_snapshot = False


def test_doc_to_response_derives_edited_inline():
    """B9 — the "Modificato" badge flag must ride on DocumentResponse."""
    from app.api.v1.documents import _doc_to_response

    edited = _doc_to_response(_FakeDocRow(options={"edited_inline": True}), None)
    assert edited.edited_inline is True
    assert edited.edited_in_gdocs is False

    assert _doc_to_response(_FakeDocRow(options=None), None).edited_inline is False
    unrelated = _doc_to_response(
        _FakeDocRow(options={"selected_codes": ["SA-01"]}), None
    )
    assert unrelated.edited_inline is False


def test_override_address_regex_grammar():
    from app.api.v1.documents import _OVERRIDE_ADDR_RE

    for good in ("0", "12", "13:0", "13:0:2", "13:0:2:0"):
        assert _OVERRIDE_ADDR_RE.fullmatch(good), good
    for bad in ("", "a", "1:a", "-1", "1:", ":1", "1:2:3:4:5", "1.5"):
        assert not _OVERRIDE_ADDR_RE.fullmatch(bad), bad


# ---------------------------------------------------------------------------
# Route registration — the frontend calls these exact paths
# ---------------------------------------------------------------------------


def test_router_registers_preview_and_override_endpoints():
    from tests.conftest import route_pairs

    from app.api.v1.router import api_router

    # Walks nested routers, so this survives FastAPI changing whether
    # include_router() flattens child routes onto the parent. See conftest.
    paths = route_pairs(api_router)
    assert ("GET", "/api/v1/documenti/{document_id}/preview") in paths
    assert (
        "GET",
        "/api/v1/documenti/{document_id}/preview/images/{image_id}",
    ) in paths
    assert ("PATCH", "/api/v1/documenti/{document_id}/overrides") in paths
    assert ("POST", "/api/v1/documenti/{document_id}/save-edited-version") in paths
    # The download endpoint stays where it was — overrides ride on it.
    assert ("GET", "/api/v1/documenti/{document_id}/download") in paths
