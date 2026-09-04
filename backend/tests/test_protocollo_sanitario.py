"""Protocollo sanitario per mansione (segnalazione 2026-08-25).

Covers, without a DB:
  * the occupational-disease reference table (shape, uniqueness, vocabulary);
  * the per-mansione aggregation and the overview merge;
  * the privacy contract of the AI prompt (no name / codice fiscale);
  * the suggester's server-side filtering of disease codes;
  * the endpoint shape (routes registered, writes gated, AI metered);
  * the DVR §4.3 renderer with and without a saved protocol.
"""

from __future__ import annotations

import ast
import asyncio
import re
from pathlib import Path
from types import SimpleNamespace
import uuid

import pytest
from docx import Document

from app.data.malattie_professionali import (
    MALATTIE_BY_CODICE,
    MALATTIE_PROFESSIONALI,
    REQUIRED_KEYS,
    malattie_per_rischi,
)
from app.services.ai import protocollo_sanitario_suggester as suggester
from app.services.ai.protocollo_sanitario_suggester import (
    AccertamentoSuggerito,
    MalattiaScelta,
    ProtocolloSuggerito,
    build_prompt,
    suggest_protocollo,
)
from app.services.document_generator.dvr_master import (
    DVRMasterGenerator,
    _format_protocol_periodicita_cell,
)
from app.services.protocollo_sanitario import aggregate_per_mansione, mansione_key
from app.services.reference_data import (
    RISCHI_SPECIFICI_CATALOG,
    RISK_CATEGORY_SHORT_NAMES,
)

BACKEND = Path(__file__).resolve().parents[1]
ROUTER_SRC = (BACKEND / "app" / "api" / "v1" / "protocollo_sanitario.py").read_text(
    encoding="utf-8"
)


# --- reference table -------------------------------------------------------


def test_reference_table_entries_have_required_keys_and_unique_codes():
    assert 25 <= len(MALATTIE_PROFESSIONALI) <= 40
    seen: set[str] = set()
    for entry in MALATTIE_PROFESSIONALI:
        assert REQUIRED_KEYS <= set(entry), f"{entry.get('codice')} misses keys"
        assert entry["codice"] not in seen, f"duplicate codice {entry['codice']}"
        seen.add(entry["codice"])
        assert entry["malattia"].strip() and entry["tabella"].strip()
        assert isinstance(entry["tabellata"], bool)
        assert entry["rischi_specifici_codes"], entry["codice"]
    assert set(MALATTIE_BY_CODICE) == seen


def test_reference_table_uses_the_apps_vocabulary():
    """Every rischi code is a RISCHI_SPECIFICI_CATALOG code and every
    categoria a canonical DVR short name — otherwise the prefill never
    matches anything the operator flagged."""
    for entry in MALATTIE_PROFESSIONALI:
        for code in entry["rischi_specifici_codes"]:
            assert code in RISCHI_SPECIFICI_CATALOG, f"{entry['codice']}: {code}"
        for cat in entry["categorie"]:
            assert cat in RISK_CATEGORY_SHORT_NAMES, f"{entry['codice']}: {cat}"


def test_vdt_diseases_are_flagged_non_tabellate_with_art_176():
    vdt = [e for e in MALATTIE_PROFESSIONALI if "vdt" in e["rischi_specifici_codes"]]
    assert vdt
    for e in vdt:
        assert e["tabellata"] is False
        assert "176" in e["tabella"]


def test_malattie_per_rischi_matches_codes_and_long_or_short_categories():
    assert malattie_per_rischi([]) == []
    assert malattie_per_rischi(None, None) == []
    rumore = malattie_per_rischi(["af_rumore"])
    assert [e["codice"] for e in rumore] == ["ipoacusia_rumore"]
    by_long = {e["codice"] for e in malattie_per_rischi([], ["Agenti Biologici"])}
    by_short = {e["codice"] for e in malattie_per_rischi([], ["Biologici"])}
    assert by_long == by_short and "epatite_b_c" in by_long
    # union, catalogue order, no duplicates
    both = malattie_per_rischi(["af_rumore", "amianto"], ["Fisici"])
    codes = [e["codice"] for e in both]
    assert len(codes) == len(set(codes))
    assert codes.index("ipoacusia_rumore") < codes.index("asbestosi")


# --- aggregation ------------------------------------------------------------


def _persona(**kw):
    base = dict(
        nominativo="Mario Rossi",
        codice_fiscale="RSSMRA80A01H501U",
        mansione=None,
        dpi_codes=[],
        rischi_specifici_codes=[],
        is_esterno=False,
        ruolo_rspp=False,
        ruolo_medico_competente=False,
    )
    base.update(kw)
    return SimpleNamespace(**base)


def test_aggregate_per_mansione_unions_codes_case_insensitively_and_skips_externals():
    persone = [
        _persona(mansione="Saldatore", dpi_codes=["caschi_industria"], rischi_specifici_codes=["af_rumore"]),
        _persona(mansione=" saldatore ", dpi_codes=["guanti_meccanici"], rischi_specifici_codes=["mmc"]),
        _persona(mansione="Impiegata", rischi_specifici_codes=["vdt"]),
        _persona(mansione=None),
        _persona(mansione="Medico Competente", is_esterno=True, ruolo_medico_competente=True,
                 rischi_specifici_codes=["agenti_biologici"]),
    ]
    agg = aggregate_per_mansione(persone)
    assert set(agg) == {"saldatore", "impiegata"}
    sald = agg["saldatore"]
    assert sald.mansione == "Saldatore" and sald.num_persone == 2
    assert sald.rischi_codes == {"af_rumore", "mmc"}
    assert sald.dpi_codes == {"caschi_industria", "guanti_meccanici"}
    assert [m["codice"] for m in sald.malattie_riferimento()] == [
        "ipoacusia_rumore", "ernia_discale_mmc", "tendinopatie_arto_superiore", "tunnel_carpale",
    ]
    assert [r["code"] for r in agg["impiegata"].rischi_items()] == ["vdt"]
    assert mansione_key("  Sal  datore ") == "sal datore"


def test_build_overview_merges_persone_with_saved_and_orphan_protocols():
    from app.api.v1.protocollo_sanitario import build_overview

    now = "2026-09-04T10:00:00"
    saved = [
        SimpleNamespace(
            id=uuid.uuid4(), azienda_id=uuid.uuid4(), mansione="SALDATORE",
            rischi_specifici=[{"code": "af_rumore", "etichetta": "Rumore"}],
            accertamenti=[{"esame": "Audiometria", "periodicita": "annuale"}],
            periodicita="annuale", malattie_correlate=[], note=None, fonte="ai",
            created_at=now, updated_at=now,
        ),
        SimpleNamespace(
            id=uuid.uuid4(), azienda_id=uuid.uuid4(), mansione="Magazziniere",
            rischi_specifici=[], accertamenti=[], periodicita=None,
            malattie_correlate=[], note=None, fonte="manuale",
            created_at=now, updated_at=now,
        ),
    ]
    persone = [
        _persona(mansione="Saldatore", rischi_specifici_codes=["af_rumore"]),
        _persona(mansione="Impiegata", rischi_specifici_codes=["vdt"]),
    ]
    items = build_overview(persone, saved)
    assert [it.mansione for it in items] == ["Impiegata", "Magazziniere", "Saldatore"]
    sald = items[2]
    assert sald.protocollo is not None and sald.protocollo.fonte == "ai"
    assert sald.num_persone == 1
    assert [m.codice for m in sald.malattie_riferimento] == ["ipoacusia_rumore"]
    orphan = items[1]
    assert orphan.num_persone == 0 and orphan.protocollo is not None
    assert items[0].protocollo is None


# --- privacy -----------------------------------------------------------------


def _fake_azienda():
    return SimpleNamespace(
        id=uuid.uuid4(),
        ragione_sociale="OFFICINA TEST SRL",
        attivita="Carpenteria metallica",
        codice_ateco="25.11",
        descrizione_attivita="Saldatura e assemblaggio di strutture in acciaio.",
    )


def test_prompt_never_carries_a_persons_name_or_codice_fiscale(monkeypatch):
    """The persona holding the mansione has a name and a CF; the prompt is
    built from the mansione aggregate only, so neither can appear."""
    nome = "Giuseppe Verdi"
    cf = "VRDGPP75C15F205X"
    persone = [
        _persona(nominativo=nome, codice_fiscale=cf, mansione="Saldatore",
                 dpi_codes=["caschi_industria"], rischi_specifici_codes=["af_rumore", "agenti_chimici"]),
    ]
    agg = aggregate_per_mansione(persone)["saldatore"]

    captured: dict = {}

    async def fake_generate_structured(prompt, *, schema, system=None, model=None, reasoning_effort="low"):
        captured["prompt"] = prompt
        captured["system"] = system or ""
        return ProtocolloSuggerito(
            accertamenti=[AccertamentoSuggerito(esame="Visita medica", periodicita="annuale", motivazione="art. 41")],
            periodicita="annuale",
            malattie_correlate=[MalattiaScelta(codice="ipoacusia_rumore", motivazione="rumore")],
            motivazione="ok",
        )

    monkeypatch.setattr(suggester, "generate_structured", fake_generate_structured)

    result = asyncio.run(
        suggest_protocollo(
            mansione=agg.mansione,
            rischi_codes=sorted(agg.rischi_codes),
            dpi_codes=sorted(agg.dpi_codes),
            azienda=_fake_azienda(),
            malattie_riferimento=agg.malattie_riferimento(),
        )
    )
    text = captured["prompt"] + captured["system"]
    for forbidden in (nome, "Giuseppe", "Verdi", cf):
        assert forbidden not in text
    assert "Saldatore" in captured["prompt"]
    assert "Agenti fisici - Rumore" in captured["prompt"]
    assert "codice=ipoacusia_rumore" in captured["prompt"]
    assert "25.11" in captured["prompt"]
    assert result.malattie_correlate[0].codice == "ipoacusia_rumore"


def test_build_prompt_reads_only_role_level_facts():
    prompt = build_prompt(
        mansione="Impiegata",
        rischi_codes=["vdt"],
        dpi_codes=[],
        azienda=_fake_azienda(),
        malattie_riferimento=malattie_per_rischi(["vdt"]),
    )
    assert "Lavori ai videoterminali (VDT)" in prompt
    assert "DPI assegnati alla mansione: nessuno." in prompt
    assert "NON tabellata" in prompt
    assert not re.search(r"\b[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]\b", prompt)


def test_suggester_drops_disease_codes_outside_the_reference_subset(monkeypatch):
    async def fake_generate_structured(prompt, *, schema, system=None, model=None, reasoning_effort="low"):
        return ProtocolloSuggerito(
            accertamenti=[AccertamentoSuggerito(esame="Visita medica", periodicita="annuale", motivazione="x")],
            periodicita="annuale",
            malattie_correlate=[
                MalattiaScelta(codice="ipoacusia_rumore", motivazione="ok"),
                MalattiaScelta(codice="ipoacusia_rumore", motivazione="dup"),
                MalattiaScelta(codice="asbestosi", motivazione="not offered"),
                MalattiaScelta(codice="inventata", motivazione="hallucinated"),
            ],
            motivazione="x",
        )

    monkeypatch.setattr(suggester, "generate_structured", fake_generate_structured)
    result = asyncio.run(
        suggest_protocollo(
            mansione="Saldatore",
            rischi_codes=["af_rumore"],
            dpi_codes=[],
            azienda=_fake_azienda(),
            malattie_riferimento=malattie_per_rischi(["af_rumore"]),
        )
    )
    assert [m.codice for m in result.malattie_correlate] == ["ipoacusia_rumore"]


def test_ai_module_never_touches_persona_fields():
    src = (BACKEND / "app" / "services" / "ai" / "protocollo_sanitario_suggester.py").read_text(
        encoding="utf-8"
    )
    for term in ("nominativo", "codice_fiscale", "sesso", "fascia_eta", "app.models.persona"):
        assert term not in src, f"{term} must not be referenced by the AI module"


# --- endpoint shape ----------------------------------------------------------


def test_routes_are_registered():
    from tests.conftest import route_pairs
    from app.api.v1.router import api_router

    paths = route_pairs(api_router)
    base = "/api/v1/aziende/{azienda_id}/protocollo-sanitario/mansioni"
    assert ("GET", base) in paths
    assert ("PUT", base) in paths
    assert ("DELETE", base + "/{protocollo_id}") in paths
    assert ("POST", base + "/suggerisci") in paths


def _endpoint(name: str) -> ast.AsyncFunctionDef:
    tree = ast.parse(ROUTER_SRC)
    for node in ast.walk(tree):
        if isinstance(node, ast.AsyncFunctionDef) and node.name == name:
            return node
    raise AssertionError(f"{name} not found")


def _guarded_capabilities(node: ast.AsyncFunctionDef) -> set[str]:
    """Capability constants named in `require_capability(...)` calls on the
    endpoint — decorators included, which `get_source_segment` leaves out."""
    found: set[str] = set()
    for sub in ast.walk(ast.Module(body=[node], type_ignores=[])):
        if (
            isinstance(sub, ast.Call)
            and isinstance(sub.func, ast.Name)
            and sub.func.id == "require_capability"
        ):
            found.update(a.id for a in sub.args if isinstance(a, ast.Name))
    return found


@pytest.mark.parametrize(
    "name",
    ["upsert_protocollo_mansione", "delete_protocollo_mansione", "suggerisci_protocollo_mansione"],
)
def test_write_and_ai_endpoints_are_capability_gated(name):
    assert "ASSESSMENTS_WRITE" in _guarded_capabilities(_endpoint(name))


def test_read_endpoint_is_not_capability_gated():
    node = _endpoint("list_protocolli_mansioni")
    assert _guarded_capabilities(node) == set()
    names = {n.id for n in ast.walk(ast.Module(body=[node], type_ignores=[])) if isinstance(n, ast.Name)}
    assert "get_entitlements" not in names


def test_ai_endpoint_is_metered_with_the_mansione_key():
    node = _endpoint("suggerisci_protocollo_mansione")
    args = {a.arg for a in node.args.args + node.args.kwonlyargs}
    assert "ent" in args
    src = ast.get_source_segment(ROUTER_SRC, node)
    assert 'metered(org_id, "reasoning", f"protocollo-sanitario:{azienda_id}:{key}", db, ent)' in src
    # The AI call is the only thing inside the metered block.
    inner = [n for n in ast.walk(node) if isinstance(n, ast.AsyncWith)]
    assert len(inner) == 1
    body_src = "".join(ast.get_source_segment(ROUTER_SRC, b) for b in inner[0].body)
    assert "suggest_protocollo(" in body_src and "db.add" not in body_src


def test_dvr_extras_loader_hydrates_protocolli():
    src = (BACKEND / "app" / "services" / "document_generator" / "dvr_master.py").read_text(
        encoding="utf-8"
    )
    tree = ast.parse(src)
    for node in ast.walk(tree):
        if isinstance(node, ast.AsyncFunctionDef) and node.name == "_load_dvr_extras":
            body = ast.get_source_segment(src, node)
            assert "ProtocolloSanitarioMansione" in body
            assert '"protocolli_sanitari": protocolli_sanitari' in body
            break
    else:
        raise AssertionError("_load_dvr_extras not found")


# --- DVR §4.3 renderer -------------------------------------------------------


def _generator() -> DVRMasterGenerator:
    return DVRMasterGenerator.__new__(DVRMasterGenerator)


def _cells(table) -> list[list[str]]:
    return [[c.text.strip() for c in row.cells] for row in table.rows]


def _protocollo(**kw):
    base = dict(
        mansione="Saldatore",
        rischi_specifici=[{"code": "af_rumore", "etichetta": "Agenti fisici - Rumore"}],
        accertamenti=[
            {"esame": "Visita medica", "periodicita": "annuale"},
            {"esame": "Audiometria", "periodicita": "annuale"},
        ],
        periodicita="annuale",
        malattie_correlate=[
            {"codice": "ipoacusia_rumore", "malattia": "Ipoacusia da rumore",
             "riferimento": "D.M. 9/4/2008 — Tab. Industria voce 75"},
        ],
        note=None,
        fonte="ai_modificato",
    )
    base.update(kw)
    return SimpleNamespace(**base)


def test_periodicita_cell_prefers_the_saved_protocol():
    assert _format_protocol_periodicita_cell(None) == "[DA COMPILARE — MC]"
    cell = _format_protocol_periodicita_cell(_protocollo())
    assert cell.startswith("Annuale")
    assert "Audiometria (annuale)" in cell
    assert _format_protocol_periodicita_cell(
        _protocollo(periodicita=None, accertamenti=[])
    ) == "[DA COMPILARE — MC]"


def test_section_43_uses_saved_protocol_and_reference_table_per_mansione():
    persone = [
        _persona(mansione="Saldatore", dpi_codes=["caschi_industria"], rischi_specifici_codes=["af_rumore"]),
        _persona(mansione="Impiegata", rischi_specifici_codes=["vdt"]),
        _persona(mansione="Custode", dpi_codes=["caschi_industria"]),
    ]
    doc = Document()
    _generator()._add_sorveglianza_protocol_table(doc, persone, [_protocollo()])

    assert len(doc.tables) == 2
    protocol, malattie = doc.tables
    rows = _cells(protocol)
    assert rows[0][0] == "Mansione"
    by_mansione = {r[0]: r for r in rows[1:]}
    assert set(by_mansione) == {"CUSTODE", "IMPIEGATA", "SALDATORE"}
    assert by_mansione["SALDATORE"][3].startswith("Annuale")
    assert "Audiometria (annuale)" in by_mansione["SALDATORE"][3]
    assert by_mansione["IMPIEGATA"][3] == "[DA COMPILARE — MC]"

    mrows = _cells(malattie)
    assert mrows[0] == ["Mansione", "Malattie correlate", "Riferimento tabellare"]
    m = {r[0]: r for r in mrows[1:]}
    # Saved protocol: its own list, no "da confermare" caveat.
    assert m["SALDATORE"][1] == "Ipoacusia da rumore"
    assert "voce 75" in m["SALDATORE"][2] and "da confermare" not in m["SALDATORE"][2]
    # No protocol: reference rows by rischi codes, flagged for the MC.
    assert "Astenopia" in m["IMPIEGATA"][1]
    assert "da confermare dal MC" in m["IMPIEGATA"][2] and "176" in m["IMPIEGATA"][2]
    # Flags but no matching disease: still a row, still the MC's call.
    assert "Nessuna malattia tabellata" in m["CUSTODE"][1]
    assert m["CUSTODE"][2] == "da confermare dal MC"

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Malattie professionali correlate per mansione" in text
    assert "D.M. 9 aprile 2008" in text


def test_section_43_lists_a_saved_protocol_whose_persone_carry_no_flags():
    persone = [_persona(mansione="Saldatore")]
    doc = Document()
    _generator()._add_sorveglianza_protocol_table(doc, persone, [_protocollo(mansione="saldatore")])
    rows = _cells(doc.tables[0])
    assert [r[0] for r in rows[1:]] == ["SALDATORE"]
    assert "Agenti fisici - Rumore" in rows[1][1]


def test_section_43_without_flags_or_protocols_keeps_the_no_data_paragraph():
    doc = Document()
    _generator()._add_sorveglianza_protocol_table(doc, [_persona(mansione="Custode")], [])
    assert not doc.tables
    assert "Nessuna mansione con sorveglianza sanitaria configurata" in "\n".join(
        p.text for p in doc.paragraphs
    )


def test_section_43_default_argument_keeps_the_legacy_call_shape():
    """Callers that pass only persone (the existing test-suite shape) still
    get both tables."""
    doc = Document()
    _generator()._add_sorveglianza_protocol_table(
        doc, [_persona(mansione="Saldatore", rischi_specifici_codes=["af_rumore"])]
    )
    assert len(doc.tables) == 2
    assert _cells(doc.tables[0])[1][3] == "[DA COMPILARE — MC]"
