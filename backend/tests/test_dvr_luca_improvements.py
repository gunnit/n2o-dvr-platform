"""Regression coverage for Luca's DVR cover corrections."""

import asyncio
import collections
import importlib.util
import logging
import subprocess
import sys
from datetime import datetime
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
import uuid
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt
from PIL import Image

from app.services.document_generator.branding import Branding
from app.services.document_generator import dvr_master
from app.services.document_generator.dvr_master import (
    DVRMasterGenerator,
    _PART_IV_PROCEDURAL_SECTIONS,
    _employee_persons,
    _global_equipment_rows,
    _saved_order_key,
)
from app.services.ambiente_photo import normalize_document_image
from scripts import verify_dvr_luca_fixture as luca_fixture_verifier
from scripts.verify_dvr_luca_fixture import build_and_audit


BACKEND_ROOT = Path(__file__).resolve().parents[1]


def _load_verify():
    path = BACKEND_ROOT / "scripts" / "verify_all_generators.py"
    spec = importlib.util.spec_from_file_location("verify_all_generators_luca", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


@pytest.fixture
def full_dvr_doc(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()
    fixture["persone"].append(module.mk(
        nominativo="Dott.ssa Test Medico", mansione="Medico Competente",
        ordine=99, created_at=datetime(2026, 1, 10), is_esterno=True,
        ruolo_medico_competente=True, ruolo_datore_lavoro=False,
        ruolo_rspp=False, ruolo_rls=False, ruolo_primo_soccorso=False,
        ruolo_antincendio=False, ruolo_preposto=False, ambienti=[],
        dpi_codes=[], rischi_specifici_codes=[], attrezzature_speciali=[],
    ))
    module.patch_generators(fixture, str(tmp_path))
    ok, path, message = asyncio.run(module.run_one("DVR_MASTER", fixture["azienda"].id))
    assert ok, message
    return Document(path), fixture


@pytest.fixture
def empty_environment_dvr_doc(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()
    fixture["ambienti"] = []
    module.patch_generators(fixture, str(tmp_path))
    ok, path, message = asyncio.run(module.run_one("DVR_MASTER", fixture["azienda"].id))
    assert ok, message
    return Document(path)


def _next_paragraph_has_page_break(paragraph) -> bool:
    following = paragraph._p.getnext()
    return bool(following is not None and following.xpath('.//w:br[@w:type="page"]'))


def _document_xml_has_adjacent_page_breaks(doc: Document) -> bool:
    flags = [bool(paragraph._p.xpath('.//w:br[@w:type="page"]')) for paragraph in doc.paragraphs]
    return any(left and right for left, right in zip(flags, flags[1:]))


def test_table_ending_topic_uses_page_break_before_boundary_to_avoid_a_blank_page():
    doc = Document()
    doc.add_table(rows=1, cols=1)
    generator = _new_generator()

    generator._ensure_page_boundary(doc)

    boundary = generator._last_content_element(doc)
    assert boundary.xpath("./w:pPr/w:pageBreakBefore")
    assert not boundary.xpath('.//w:br[@w:type="page"]')
    body_length = len(doc._element.body)

    generator._ensure_page_boundary(doc)

    assert len(doc._element.body) == body_length


def _new_generator() -> DVRMasterGenerator:
    return DVRMasterGenerator.__new__(DVRMasterGenerator)


def _cover_azienda():
    return SimpleNamespace(
        ragione_sociale="AZIENDA TEST SRL",
        sede_legale_via="Via Test 1",
        sede_legale_citta="Roma",
        cap_legale="00100",
        provincia_legale="RM",
        partita_iva="00000000000",
        codice_ateco="00.00.00",
    )


def _all_text(doc: Document) -> str:
    paragraph_text = [paragraph.text for paragraph in doc.paragraphs]
    table_text = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraph_text + table_text)


def _jpeg_bytes(color: str = "blue") -> bytes:
    output = BytesIO()
    Image.new("RGB", (24, 24), color).save(output, "JPEG")
    return output.getvalue()


def test_dvr_cover_embeds_vera_asset_and_omits_consultancy(tmp_path):
    """Configured organization branding must not leak into DVR covers."""
    gen = _new_generator()
    gen.branding = Branding(
        firm_name="CONSULTANCY SENTINEL",
        indirizzo="SENTINEL ADDRESS",
        logo_bytes=(BACKEND_ROOT / "assets" / "logo.png").read_bytes(),
        logo_content_type="image/png",
    )
    doc = Document()
    gen._add_cover_page(doc, _cover_azienda(), datetime(2026, 8, 3), 1)
    target = tmp_path / "cover.docx"
    doc.save(target)

    with ZipFile(target) as archive:
        media = [archive.read(name) for name in archive.namelist() if name.startswith("word/media/")]
    expected = (BACKEND_ROOT / "assets" / "n2o_vera_dvr.png").read_bytes()
    assert expected in media
    assert (BACKEND_ROOT / "assets" / "logo.png").read_bytes() not in media
    text = _all_text(doc)
    assert "DOCUMENTO DI VALUTAZIONE DEI RISCHI" in text
    assert "Documento elaborato da" not in text
    assert "CONSULTANCY SENTINEL" not in text


def test_saved_order_places_null_after_explicit_order_and_uses_stable_ties():
    rows = [
        SimpleNamespace(nome="Zulu", ordine=2, created_at=datetime(2026, 1, 3), id=uuid.UUID(int=3)),
        SimpleNamespace(nome="Alpha", ordine=1, created_at=datetime(2026, 1, 2), id=uuid.UUID(int=2)),
        SimpleNamespace(nome="Beta", ordine=None, created_at=datetime(2026, 1, 1), id=uuid.UUID(int=1)),
    ]
    assert [row.nome for row in sorted(rows, key=_saved_order_key)] == ["Alpha", "Zulu", "Beta"]


def test_external_rspp_and_medico_are_not_employees_but_remain_role_holders():
    internal_rspp = SimpleNamespace(nominativo="Interno", is_esterno=False, ruolo_rspp=True, ruolo_medico_competente=False)
    external_rspp = SimpleNamespace(nominativo="RSPP Esterno", is_esterno=True, ruolo_rspp=True, ruolo_medico_competente=False)
    external_medico = SimpleNamespace(nominativo="Medico Esterno", is_esterno=True, ruolo_rspp=False, ruolo_medico_competente=True)
    employees = _employee_persons([external_medico, internal_rspp, external_rspp])
    assert employees == [internal_rspp]


def test_global_equipment_groups_description_and_orders_environment_names():
    ambienti = [
        SimpleNamespace(id=uuid.UUID(int=1), nome="Reparto B"),
        SimpleNamespace(id=uuid.UUID(int=2), nome="Reparto A"),
    ]
    rows = _global_equipment_rows(
        [
            SimpleNamespace(descrizione=" Trapano   a colonna ", ambiente_id=ambienti[1].id, marcatura_ce=True, verifiche_periodiche=False),
            SimpleNamespace(descrizione="TRAPANO A COLONNA", ambiente_id=ambienti[0].id, marcatura_ce=False, verifiche_periodiche=False),
        ],
        ambienti,
    )
    assert rows == [["TRAPANO A COLONNA", "REPARTO B, REPARTO A", "MISTO", "NO"]]


def test_employee_tables_keep_saved_order_and_external_consultants_keep_roles():
    gen = DVRMasterGenerator.__new__(DVRMasterGenerator)
    internal = SimpleNamespace(
        nominativo="Zulu Interno", mansione="Operaio", is_esterno=False,
        ruolo_datore_lavoro=True, ruolo_rspp=False, ruolo_rls=False,
        ruolo_medico_competente=False, ruolo_primo_soccorso=False,
        ruolo_antincendio=False, ambienti=[], codice_fiscale=None,
        tipologia_contrattuale=None,
    )
    external = SimpleNamespace(
        nominativo="Alpha RSPP", mansione="RSPP", is_esterno=True,
        ruolo_datore_lavoro=False, ruolo_rspp=True, ruolo_rls=False,
        ruolo_medico_competente=False, ruolo_primo_soccorso=False,
        ruolo_antincendio=False, ambienti=[], codice_fiscale=None,
        tipologia_contrattuale=None,
    )
    employees = _employee_persons([internal, external])
    doc = Document()
    gen._add_dati_occupazionali_table(doc, employees)
    gen._add_single_role_title_table(doc, "Responsabile del Servizio di Prevenzione e Protezione", [external])
    assert [row.cells[0].text for row in doc.tables[0].rows[1:]] == ["ZULU INTERNO"]
    assert "ALPHA RSPP (ESTERNO)" in doc.tables[1].cell(1, 0).text


def test_photo_image_sources_put_database_bytes_before_local_path(tmp_path):
    local = tmp_path / "legacy.jpg"
    local.write_bytes(_jpeg_bytes("green"))
    photo = SimpleNamespace(
        document_image_bytes=_jpeg_bytes("blue"), file_path=str(local)
    )

    sources = dvr_master._photo_image_sources(photo)

    assert len(sources) == 2
    assert isinstance(sources[0], BytesIO)
    assert sources[0].read() == photo.document_image_bytes
    assert sources[1] == str(local)


def test_dvr_embeds_all_ten_database_photos():
    photos = [
        SimpleNamespace(
            id=uuid.uuid4(),
            filename=f"foto-{index}.jpg",
            document_image_bytes=_jpeg_bytes(),
            file_path=None,
        )
        for index in range(10)
    ]
    doc = Document()

    _new_generator()._add_env_foto_block(doc, "REPARTO", photos)

    assert len(doc.inline_shapes) == 10
    assert "Fig. 10 — foto-9.jpg" in _all_text(doc)


def test_dvr_embeds_normalized_heic_photo_derivative():
    import pillow_heif

    source = BytesIO()
    pillow_heif.from_pillow(Image.new("RGB", (24, 24), "orange")).save(source)
    derivative = normalize_document_image(source.getvalue())
    photo = SimpleNamespace(
        id=uuid.UUID(int=10),
        filename="sopralluogo.heic",
        document_image_bytes=derivative.content,
        file_path=None,
    )
    doc = Document()

    _new_generator()._add_env_foto_block(doc, "REPARTO", [photo])

    assert len(doc.inline_shapes) == 1
    assert "Fig. 1 — sopralluogo.heic" in _all_text(doc)


def test_dvr_renders_one_filename_specific_marker_per_unavailable_photo(caplog):
    photos = [
        SimpleNamespace(
            id=uuid.UUID(int=1),
            filename="folder/assente-a.heic",
            document_image_bytes=None,
            file_path="/private/missing-a.heic",
        ),
        SimpleNamespace(
            id=uuid.UUID(int=2),
            filename="assente-b.jpg",
            document_image_bytes=None,
            file_path="/private/missing-b.jpg",
        ),
    ]
    doc = Document()

    with caplog.at_level(logging.WARNING):
        _new_generator()._add_env_foto_block(doc, "REPARTO", photos)

    text = _all_text(doc)
    assert "[Foto non disponibile: assente-a.heic]" in text
    assert "[Foto non disponibile: assente-b.jpg]" in text
    warnings = [record for record in caplog.records if record.levelno >= logging.WARNING]
    assert {
        (record.photo_id, record.photo_filename) for record in warnings
    } == {
        (str(uuid.UUID(int=1)), "assente-a.heic"),
        (str(uuid.UUID(int=2)), "assente-b.jpg"),
    }
    assert all("/private/" not in record.getMessage() for record in warnings)
    assert all("folder/" not in record.getMessage() for record in warnings)


def test_corrupt_database_derivative_falls_back_to_valid_local_file(tmp_path, caplog):
    local = tmp_path / "legacy.jpg"
    local.write_bytes(_jpeg_bytes("green"))
    photo = SimpleNamespace(
        id=uuid.UUID(int=3),
        filename="folder/legacy.jpg",
        document_image_bytes=b"corrupt",
        file_path=str(local),
    )
    doc = Document()

    with caplog.at_level(logging.WARNING):
        _new_generator()._add_env_foto_block(doc, "REPARTO", [photo])

    assert len(doc.inline_shapes) == 1
    assert "[Foto non disponibile" not in _all_text(doc)
    warnings = [record for record in caplog.records if record.levelno >= logging.WARNING]
    assert [(record.photo_id, record.photo_filename) for record in warnings] == [
        (str(uuid.UUID(int=3)), "legacy.jpg")
    ]
    assert all(str(tmp_path) not in record.getMessage() for record in warnings)


def test_parent_false_applicable_children_are_effective_in_relationship_order():
    parent = SimpleNamespace(
        applicabile=False,
        pericoli=[
            SimpleNamespace(pericolo="A", applicabile=True),
            SimpleNamespace(pericolo="B", applicabile=False),
            SimpleNamespace(pericolo="C", applicabile=True),
        ],
    )
    assert [
        row.pericolo for row in dvr_master._effective_risk_sources(parent)
    ] == ["A", "C"]


def test_all_disabled_children_do_not_resurrect_parent():
    parent = SimpleNamespace(
        applicabile=True,
        pericoli=[SimpleNamespace(pericolo="A", applicabile=False)],
    )
    assert dvr_master._effective_risk_sources(parent) == []


def test_childless_applicable_parent_remains_effective():
    parent = SimpleNamespace(applicabile=True, pericoli=[])
    assert dvr_master._effective_risk_sources(parent) == [parent]


def test_parent_false_applicable_children_render_once_and_mark_checklist_yes():
    applicable = SimpleNamespace(
        pericolo="CHILD APPLICABLE SENTINEL",
        applicabile=True,
        condizioni_esposizione="COND",
        rischio="RISK",
        misure_prevenzione="MEASURE",
        probabilita_p=2,
        danno_d=3,
        livello_rischio="GRAVE",
    )
    disabled = SimpleNamespace(
        pericolo="CHILD DISABLED SENTINEL", applicabile=False
    )
    parent = SimpleNamespace(
        categoria_rischio="Macchine",
        applicabile=False,
        pericoli=[applicable, disabled],
    )
    ambiente = SimpleNamespace(valutazioni_rischio=[parent])
    doc = Document()
    gen = _new_generator()
    gen._add_env_risk_checklist(doc, ambiente)
    gen._add_env_risk_tables(doc, ambiente)
    text = _all_text(doc)
    assert text.count("CHILD APPLICABLE SENTINEL") == 1
    assert "CHILD DISABLED SENTINEL" not in text
    assert any(
        "SI" in cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )


def test_every_level_two_heading_is_a_single_separator_without_adjacent_page_breaks(full_dvr_doc):
    doc, fixture = full_dvr_doc
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    expected = {
        "1. Presentazione dell'Azienda", "2. Anagrafica Aziendale",
        "3. Dati Occupazionali", "4. Organizzazione Aziendale della Sicurezza",
        "5. Ambienti di Lavoro", "6. Servizi Igienico-Assistenziali",
        "7. Macchine, Attrezzature ed Impianti", "8. Sostanze, Prodotti e Preparati Chimici",
        "9. Elenco Fattori di Pericolo (Riferimento)", "2.1 Descrizione dell'Attività",
        "2.2 Definizioni", "2.3 Metodologia di Valutazione dei Rischi",
        "2.4 Scala di Probabilità (P)", "2.5 Scala del Danno (D)",
        "Mansioni che espongono i lavoratori a rischi specifici",
        "DPI in dotazione per Mansione", "Segnaletica di Sicurezza",
        "Programma di Informazione, Formazione e Addestramento",
        "4.1 Programma e Procedure di attuazione delle Misure di Miglioramento",
        "Documenti correlati al presente DVR", "4.13 Dichiarazione del Datore di Lavoro",
    }
    expected.update(spec.heading for spec in _PART_IV_PROCEDURAL_SECTIONS)
    expected.update(
        f"Identificazione dell'Ambiente di Lavoro e degli Addetti — {(ambiente.nome or '—').upper()}"
        for ambiente in fixture["ambienti"]
    )
    assert collections.Counter(paragraph.text for paragraph in headings) == collections.Counter(expected)
    for heading in headings:
        assert _next_paragraph_has_page_break(heading)
    assert not _document_xml_has_adjacent_page_breaks(doc)


def test_each_part_h1_occurs_once_and_shares_its_first_topic_separator(full_dvr_doc):
    doc, fixture = full_dvr_doc
    paragraphs = doc.paragraphs
    h1 = [p for p in paragraphs if p.style.name == "Heading 1"]
    h1_text = [p.text for p in h1]
    assert h1_text.count("PARTE I — DATI GENERALI DELL'AZIENDA") == 1
    assert h1_text.count("PARTE II — DESCRIZIONE DELL'ATTIVITÀ E METODOLOGIA DI VALUTAZIONE") == 1
    assert h1_text.count("PARTE III — VALUTAZIONE DEI RISCHI PER AMBIENTE DI LAVORO") == 1
    assert h1_text.count("PARTE IV — PROGRAMMA DI MIGLIORAMENTO") == 1
    assert all(paragraph.runs[0].font.size == Pt(11) for paragraph in h1 if paragraph.text.startswith("PARTE "))
    first_environment = min(fixture["ambienti"], key=_saved_order_key)
    first_topics = {
        "PARTE I — DATI GENERALI DELL'AZIENDA": "1. Presentazione dell'Azienda",
        "PARTE II — DESCRIZIONE DELL'ATTIVITÀ E METODOLOGIA DI VALUTAZIONE": "2.1 Descrizione dell'Attività",
        "PARTE III — VALUTAZIONE DEI RISCHI PER AMBIENTE DI LAVORO": f"Identificazione dell'Ambiente di Lavoro e degli Addetti — {(first_environment.nome or '—').upper()}",
        "PARTE IV — PROGRAMMA DI MIGLIORAMENTO": "4.1 Programma e Procedure di attuazione delle Misure di Miglioramento",
    }
    for part_heading, first_topic in first_topics.items():
        part = next(paragraph for paragraph in h1 if paragraph.text == part_heading)
        following = paragraphs[paragraphs.index(part) + 1]
        assert following.style.name == "Heading 2"
        assert following.text == first_topic


def test_empty_environment_part_keeps_its_h1_with_all_inline_tail_topics(empty_environment_dvr_doc):
    doc = empty_environment_dvr_doc
    paragraphs = doc.paragraphs
    part = next(
        paragraph for paragraph in paragraphs
        if paragraph.style.name == "Heading 1"
        and paragraph.text == "PARTE III — VALUTAZIONE DEI RISCHI PER AMBIENTE DI LAVORO"
    )
    part_index = paragraphs.index(part)
    assert paragraphs[part_index + 1].style.name == "Heading 2"
    assert paragraphs[part_index + 1].text == "Mansioni che espongono i lavoratori a rischi specifici"
    expected_tail = {
        "Mansioni che espongono i lavoratori a rischi specifici",
        "DPI in dotazione per Mansione",
        "Segnaletica di Sicurezza",
        "Programma di Informazione, Formazione e Addestramento",
    }
    tail = [
        paragraph for paragraph in paragraphs
        if paragraph.style.name == "Heading 2" and paragraph.text in expected_tail
    ]
    assert collections.Counter(paragraph.text for paragraph in tail) == collections.Counter(expected_tail)
    assert all(paragraph.style.name == "Heading 2" for paragraph in tail)
    assert all(_next_paragraph_has_page_break(paragraph) for paragraph in tail)
    assert any(paragraph.text == "Nessun ambiente di lavoro registrato." for paragraph in paragraphs)


def test_improvement_table_prints_all_saved_fields_in_order_and_restores_portrait():
    rows = [
        SimpleNamespace(id=uuid.UUID(int=2), created_at=datetime(2026, 1, 2), ordine=2, priorita="MODESTO", misura="R2", misura_miglioramento="M2", procedura="P2", risorse="S2", responsabile="A2", scadenza="D2"),
        SimpleNamespace(id=uuid.UUID(int=1), created_at=datetime(2026, 1, 1), ordine=1, priorita="GRAVE", misura="R1", misura_miglioramento="M1", procedura="P1", risorse="S1", responsabile="A1", scadenza="D1"),
    ]
    doc = Document()
    _new_generator()._add_improvement_program_table(doc, rows)
    table = doc.tables[-1]
    assert [c.text for c in table.rows[0].cells] == [
        "Priorità", "Rischio", "Misura di Miglioramento", "Attività / Procedura", "Risorse", "Responsabile", "Scadenza"
    ]
    assert [c.text for c in table.rows[1].cells] == ["GRAVE", "R1", "M1", "P1", "S1", "A1", "D1"]
    assert doc.sections[-2].orientation == WD_ORIENT.LANDSCAPE
    assert doc.sections[-1].orientation == WD_ORIENT.PORTRAIT


def test_declaration_has_fresh_content_page_and_signature_rows_are_signable(full_dvr_doc):
    doc, _fixture = full_dvr_doc
    declaration = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text == "4.13 Dichiarazione del Datore di Lavoro"
        and paragraph.style.name == "Heading 2"
    )
    assert _next_paragraph_has_page_break(declaration)
    signature = next(
        table for table in doc.tables
        if len(table.rows) == 2
        and len(table.rows[0].cells) == 3
        and "Il Datore di Lavoro" in " ".join(cell.text for cell in table.rows[0].cells)
    )
    signature_text = " ".join(cell.text for row in signature.rows for cell in row.cells)
    for expected in ("MARIO ROSSI", "LUCA BIANCHI", "DOTT.SSA TEST MEDICO", "GIULIA VERDI"):
        assert expected in signature_text
    for row in signature.rows:
        assert row.height >= Cm(3)
        assert row.height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST
        assert row._tr.xpath("./w:trPr/w:cantSplit")
    final_clause = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("di impegnarsi a rielaborare"))
    place_date = next(paragraph for paragraph in doc.paragraphs if ", li " in paragraph.text)
    assert final_clause._p.xpath("./w:pPr/w:keepNext")
    assert place_date._p.xpath("./w:pPr/w:keepNext")


def test_luca_fixture_auditor_reports_all_acceptance_checks(tmp_path):
    report = build_and_audit(tmp_path)
    assert report == {
        "acme_regression": True,
        "vera_cover": True,
        "saved_people_order": True,
        "saved_environment_order": True,
        "external_roles": True,
        "grouped_equipment": True,
        "all_ten_photos": True,
        "effective_risks": True,
        "person_specific_risks": True,
        "dpi_dash_alignment": True,
        "topic_separators": True,
        "complete_improvements": True,
        "declaration_signatures": True,
    }


def test_luca_fixture_cli_is_directly_runnable_outside_backend(tmp_path):
    result = subprocess.run(
        [
            sys.executable,
            str(BACKEND_ROOT / "scripts" / "verify_dvr_luca_fixture.py"),
            "--help",
        ],
        cwd=tmp_path,
        check=False,
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr


def _generator_patch_surface():
    from app.services.document_generator import (
        _biologico_common,
        allegato_gestanti,
        allegato_incendio,
        allegato_microclima,
        allegato_microclima_severo,
        allegato_mmc,
        allegato_stress,
        allegato_vdt,
        base,
        data_loader,
        duvri,
        haccp_forms,
        haccp_manuale,
        pee_azienda,
        pee_comune,
        pos,
    )

    version_classes = [
        allegato_mmc.AllegatoMmcGenerator,
        allegato_vdt.AllegatoVdtGenerator,
        allegato_stress.AllegatoStressGenerator,
        allegato_gestanti.AllegatoGestantiGenerator,
        allegato_incendio.AllegatoIncendioGenerator,
        allegato_microclima.AllegatoMicroclimaGenerator,
        allegato_microclima_severo.AllegatoMicroclimaSeveroGenerator,
        pee_azienda.PeeAziendaGenerator,
        pee_comune.PeeComuneGenerator,
        duvri.DuvriGenerator,
        pos.PosGenerator,
        haccp_manuale.HaccpManualeGenerator,
        haccp_forms.HaccpFormsGenerator,
    ]
    loader_modules = [
        (allegato_mmc, "load_mmc"),
        (allegato_vdt, "load_vdt"),
        (allegato_stress, "load_stress"),
        (allegato_incendio, "load_incendio"),
        (allegato_microclima, "load_microclima"),
        (allegato_microclima_severo, "load_microclima"),
        (allegato_gestanti, "load_gestanti"),
        (duvri, "load_duvri"),
        (pos, "load_pos"),
        (haccp_manuale, "load_haccp"),
        (haccp_forms, "load_haccp"),
        (pee_azienda, "load_pee"),
        (pee_comune, "load_pee"),
    ]
    data_loaders = [
        "load_mmc",
        "load_vdt",
        "load_stress",
        "load_incendio",
        "load_microclima",
        "load_gestanti",
        "load_biologico",
        "load_haccp",
        "load_pee",
        "load_duvri",
        "load_pos",
    ]
    return [
        (base.BaseDocumentGenerator, "load_data"),
        (base.BaseDocumentGenerator, "_get_output_dir"),
        (DVRMasterGenerator, "_load_dvr_extras"),
        (DVRMasterGenerator, "_get_next_version"),
        *[(generator, "_next_version") for generator in version_classes],
        (_biologico_common, "_next_version"),
        *[(data_loader, name) for name in data_loaders],
        *loader_modules,
        (_biologico_common, "load_biologico"),
    ]


def _snapshot_patch_surface():
    return [
        (owner, name, getattr(owner, name))
        for owner, name in _generator_patch_surface()
    ]


def _assert_patch_surface_restored(snapshot):
    assert [
        (owner, name)
        for owner, name, original in snapshot
        if getattr(owner, name) is not original
    ] == []


def _ordered_acme_doc_content(path: Path):
    doc = Document(path)
    roster = next(
        table
        for table in doc.tables
        if table.rows
        and [cell.text.strip() for cell in table.rows[0].cells]
        == [
            "Nominativo",
            "Mansione",
            "Ambiente di Lavoro",
            "Codice Fiscale",
            "Tipologia contrattuale",
        ]
    )
    prefix = "Identificazione dell'Ambiente di Lavoro e degli Addetti — "
    return (
        [row.cells[0].text.strip() for row in roster.rows[1:]],
        [
            paragraph.text.removeprefix(prefix)
            for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 2"
            and paragraph.text.startswith(prefix)
        ],
    )


def test_repeat_builds_keep_exact_acme_roster_and_environment_order(tmp_path):
    expected = (
        [
            "MARIO ROSSI",
            "LUCA BIANCHI",
            "GIULIA VERDI",
            "ANTONIO MARRONE",
            "VALENTINA RINALDI",
        ],
        [
            "UFFICI AMMINISTRATIVI E TECNICI",
            "OFFICINA MECCANICA",
            "MAGAZZINO",
            "MENSA AZIENDALE CON CUCINA",
            "DEPOSITO CHIMICI",
            "AREA ESTERNA",
        ],
    )
    first = tmp_path / "first"
    second = tmp_path / "second"

    build_and_audit(first)
    build_and_audit(second)

    assert _ordered_acme_doc_content(first / "DVR_Acme_Regression.docx") == expected
    assert _ordered_acme_doc_content(second / "DVR_Acme_Regression.docx") == expected


def test_build_and_audit_restores_every_generator_patch_on_success(tmp_path):
    before = _snapshot_patch_surface()

    build_and_audit(tmp_path)

    _assert_patch_surface_restored(before)


def test_build_and_audit_restores_every_generator_patch_on_exception(
    tmp_path, monkeypatch
):
    before = _snapshot_patch_surface()

    def fail_audit(path, fixture):
        raise RuntimeError("forced audit failure")

    monkeypatch.setattr(luca_fixture_verifier, "audit_luca_docx", fail_audit)
    with pytest.raises(RuntimeError, match="forced audit failure"):
        build_and_audit(tmp_path)

    _assert_patch_surface_restored(before)


def test_photo_audit_rejects_captions_without_adjacent_drawings():
    doc = Document()
    expected = [f"Fig. {index} — foto-{index - 1}.jpg" for index in range(1, 11)]
    for caption in expected:
        doc.add_paragraph(caption)

    assert not luca_fixture_verifier._audit_all_ten_photos(
        doc, {"expected_photo_captions": expected}
    )


def test_effective_risk_audit_rejects_sentinel_text_without_risk_tables():
    doc = Document()
    doc.add_paragraph("RISK APPLICABLE SENTINEL")

    assert not luca_fixture_verifier._audit_effective_risks(doc, {})


def test_acme_audit_rejects_wrong_roster_order(tmp_path):
    build_and_audit(tmp_path)
    fixture = luca_fixture_verifier.build_acme_fixture()
    path = tmp_path / "DVR_Acme_Regression.docx"
    doc = Document(path)
    roster = next(
        table
        for table in doc.tables
        if table.rows
        and table.rows[0].cells[0].text.strip() == "Nominativo"
    )
    first = roster.rows[1].cells[0].text
    second = roster.rows[2].cells[0].text
    roster.rows[1].cells[0].text = second
    roster.rows[2].cells[0].text = first
    doc.save(path)

    assert not luca_fixture_verifier._audit_acme_regression(path, fixture)


def test_declaration_audit_rejects_swapped_rspp_and_medico_cells():
    fixture = luca_fixture_verifier.build_luca_fixture()
    doc = Document()
    doc.add_heading("4.13 Dichiarazione del Datore di Lavoro", level=2)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph(
        "DATORE FIXTURE, in qualita di Datore di Lavoro della "
        "LUCA FIXTURE INDUSTRIA SRL"
    )
    _new_generator()._add_signature_table(doc, fixture["persone"])
    signature = doc.tables[-1]
    rspp = signature.cell(0, 2).text
    medico = signature.cell(1, 0).text
    signature.cell(0, 2).text = medico
    signature.cell(1, 0).text = rspp

    assert not luca_fixture_verifier._audit_declaration_signatures(
        doc, fixture
    )
