"""Integration test: every generator produces a valid .docx for Acme fixture.

This reuses scripts/verify_all_generators.py which monkey-patches the DB
loaders to run without a live Postgres. Under real deployment, generators
would run via Celery with the actual DB.
"""

import asyncio
import hashlib
import importlib
import os
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document


BACKEND_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BACKEND_ROOT))
sys.path.insert(0, str(BACKEND_ROOT / "scripts"))


def _load_verify():
    spec = importlib.util.spec_from_file_location(
        "verify_all_generators",
        str(BACKEND_ROOT / "scripts" / "verify_all_generators.py"),
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _document_text(path: str) -> str:
    """Return all visible DOCX text, including tables and headers/footers."""
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for part in (section.header, section.footer):
            chunks.extend(p.text for p in part.paragraphs)
            for table in part.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _find_table(doc: Document, expected_header: tuple[str, ...]):
    for table in doc.tables:
        if not table.rows:
            continue
        header = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if header == expected_header:
            return table
    return None


def _all_key_value_pairs(doc: Document) -> dict[str, str]:
    pairs: dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) == 2:
                key = row.cells[0].text.strip()
                if key:
                    pairs[key] = row.cells[1].text.strip()
    return pairs


@pytest.fixture(scope="module")
def generated_outputs(tmp_path_factory):
    out = tmp_path_factory.mktemp("gen_out")
    module = _load_verify()
    fixture = module.build_fixture()
    module.patch_generators(fixture, str(out))

    from app.services.document_generator.dispatcher import ALL_DOCUMENT_TYPES

    results = {}

    async def run_all():
        for tipo in ALL_DOCUMENT_TYPES:
            try:
                ok, path, msg = await module.run_one(tipo, fixture["azienda"].id)
                results[tipo] = (ok, path, msg)
            except Exception as e:
                results[tipo] = (False, "", str(e))

    asyncio.run(run_all())
    return results


def test_all_17_generators_pass(generated_outputs):
    failed = {k: v for k, v in generated_outputs.items() if not v[0]}
    assert not failed, f"Failed generators: {failed}"
    assert len(generated_outputs) == 17


def test_dvr_master_has_acme_name(generated_outputs):
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "ACME" in full_text.upper()


def test_duvri_renders_current_equipment_without_legacy_donor_equipment(
    generated_outputs,
):
    ok, path, _ = generated_outputs["DUVRI"]
    assert ok and path
    doc = Document(path)
    text = _document_text(path)

    company_equipment = _find_table(
        doc, ("Attrezzatura del committente", "Ambiente")
    )
    assert company_equipment is not None
    company_rows = [
        tuple(cell.text.strip() for cell in row.cells)
        for row in company_equipment.rows[1:]
    ]
    assert sorted(company_rows) == sorted([
        ("Tornio parallelo CNC", "Officina meccanica"),
        ("Fresatrice CNC", "Officina meccanica"),
        ("Carrello elevatore", "Magazzino"),
        ("Postazione VDT", "Uffici amministrativi e tecnici"),
    ])

    contractor_equipment = _find_table(doc, ("Tipo", "Descrizione"))
    assert contractor_equipment is not None
    assert [
        tuple(cell.text.strip() for cell in row.cells)
        for row in contractor_equipment.rows[1:]
    ] == [("Pulizia ordinaria", "Lavasciuga pavimenti")]

    for donor_text in (
        "RECOM",
        "Carrello elevatore (muletto)",
        "Transpallet elettrico",
        "Lavatrice ad uso operativo",
        "Sabbiatrice",
        "Forno per lavorazioni",
    ):
        assert donor_text not in text
    assert text.count("Carrello elevatore") == 1


def test_haccp_forms_produces_zip(generated_outputs):
    ok, path, _ = generated_outputs["HACCP_FORMS"]
    assert ok and path.endswith(".zip")
    assert zipfile.is_zipfile(path)
    with zipfile.ZipFile(path) as z:
        # Index + 16 forms = 17 entries
        assert len(z.namelist()) >= 16


def test_pos_has_significant_content(generated_outputs):
    """POS is the most complex document — should have many tables."""
    ok, path, _ = generated_outputs["POS"]
    assert ok and path
    doc = Document(path)
    assert len(doc.tables) >= 5


def test_biologico_all_three_variants_generate(generated_outputs):
    for key in ("ALLEGATO_BIOLOGICO_ALIMENTARE", "ALLEGATO_BIOLOGICO_ASILO", "ALLEGATO_BIOLOGICO_DENTISTI"):
        ok, path, _ = generated_outputs[key]
        assert ok and path.endswith(".docx"), f"{key} failed"


def test_fire_renders_current_pi_meaning_and_modeled_hydrants(generated_outputs):
    ok, path, _ = generated_outputs["ALLEGATO_INCENDIO"]
    assert ok and path
    doc = Document(path)
    text = _document_text(path)

    assert "Propagazione dell'incendio" in text
    assert "Presenza di persone e loro esodo" not in text

    header = (
        "Ambiente", "INF", "SI", "PI", "Totale", "Livello",
        "Uscite", "Estintori", "Idranti",
    )
    table = _find_table(doc, header)
    assert table is not None, "fire environment table must expose modeled hydrants"
    rows = [tuple(cell.text.strip() for cell in row.cells) for row in table.rows[1:]]
    assert (
        "Officina meccanica", "3", "3", "2", "8", "ALTO", "2", "4", "2"
    ) in rows


def test_fire_training_and_drill_text_uses_supported_cadence(generated_outputs):
    ok, path, _ = generated_outputs["ALLEGATO_INCENDIO"]
    assert ok and path
    text = _document_text(path)

    assert "aggiornamento almeno quinquennale" in text
    assert "aggiornamento triennale" not in text
    assert "D.M. 02/09/2021 art. 6" not in text
    assert "cadenza almeno annuale" in text
    assert "Allegato I, punto 1.3" in text


def test_fire_does_not_claim_a_drill_cadence_for_small_employer(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()
    fixture["azienda"].numero_dipendenti_dichiarati = 4
    fixture["persone"] = fixture["persone"][:4]
    module.patch_generators(fixture, str(tmp_path))

    ok, path, msg = asyncio.run(
        module.run_one("ALLEGATO_INCENDIO", fixture["azienda"].id)
    )
    assert ok, msg
    text = _document_text(path)

    assert "cadenza almeno annuale" not in text
    assert "Esercitazione antincendio almeno annuale" not in text
    assert "Esercitazione antincendio semestrale" not in text
    assert "verificata in base all'applicabilità" in text


def test_fire_does_not_count_registered_consultants_as_declared_workers(tmp_path):
    module = _load_verify()
    fixture = module.build_fixture()
    fixture["azienda"].numero_dipendenti_dichiarati = None
    fixture["persone"] = [
        module.mk(
            nominativo=f"Consulente esterno {index}",
            is_esterno=True,
            ruolo_rspp=index == 0,
            ruolo_medico_competente=index == 1,
        )
        for index in range(10)
    ]
    module.patch_generators(fixture, str(tmp_path))

    ok, path, msg = asyncio.run(
        module.run_one("ALLEGATO_INCENDIO", fixture["azienda"].id)
    )
    assert ok, msg
    text = _document_text(path)

    assert "cadenza almeno annuale" not in text
    assert "verificata in base all'applicabilità" in text


def test_pee_azienda_renders_existing_workforce_data_with_exact_labels(
    generated_outputs,
):
    ok, path, _ = generated_outputs["PEE_AZIENDA"]
    assert ok and path
    pairs = _all_key_value_pairs(Document(path))

    assert pairs["Orario di lavoro dichiarato"] == "Lun-Ven 08:00-17:00"
    assert pairs["Lavoratori dichiarati dall'azienda"] == "37"
    assert pairs["Persone registrate nel DVR"] == "5"
    assert "Affollamento massimo" not in pairs
    assert "Occupazione massima" not in pairs


def test_pee_azienda_uses_only_the_configured_drill_frequency(generated_outputs):
    ok, path, _ = generated_outputs["PEE_AZIENDA"]
    assert ok and path
    doc = Document(path)
    pairs = _all_key_value_pairs(doc)
    text = _document_text(path)

    assert pairs["Frequenza prove"] == "semestrale"
    assert "Le prove di evacuazione seguono la frequenza configurata: semestrale." in text
    assert "Prove di evacuazione con cadenza almeno annuale" not in text


def test_pee_azienda_scrubs_donor_collection_point(generated_outputs):
    ok, path, _ = generated_outputs["PEE_AZIENDA"]
    assert ok and path
    doc = Document(path)
    text = _document_text(path)

    assert "Parcheggio del polo commerciale" not in text
    assert "Piazzale ingresso" in text

    start = next(
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if "Raggiungere il punto di raccolta esterno (Piazzale ingresso)"
        in paragraph.text
    )
    end = next(
        index
        for index, paragraph in enumerate(doc.paragraphs[start + 1 :], start + 1)
        if "Il punto di raccolta del personale evacuato sarà il Piazzale ingresso"
        in paragraph.text
    )
    assert not any(
        paragraph._p.xpath(".//w:drawing")
        for paragraph in doc.paragraphs[start + 1 : end]
    ), "donor collection-point images must not remain in the configured output"
    assert all(
        paragraph.text.strip() for paragraph in doc.paragraphs[start + 1 : end]
    ), "removing donor media must not leave its image-spacer paragraphs behind"
    assert "come illustrato sopra" not in text

    template_doc = Document(
        BACKEND_ROOT.parent
        / "templates"
        / "PIANO GESTIONE EMERGENZE - AZIENDA.docx"
    )
    template_start = next(
        index
        for index, paragraph in enumerate(template_doc.paragraphs)
        if "Raggiungere il punto di raccolta esterno " in paragraph.text
        and "Parcheggio del polo commerciale" in paragraph.text
    )
    template_end = next(
        index
        for index, paragraph in enumerate(
            template_doc.paragraphs[template_start + 1 :], template_start + 1
        )
        if "Il punto di raccolta del personale evacuato sarà" in paragraph.text
        and "Parcheggio del polo commerciale" in paragraph.text
    )
    donor_relationship_ids = {
        relationship_id
        for paragraph in template_doc.paragraphs[template_start + 1 : template_end]
        for relationship_id in paragraph._p.xpath(".//a:blip/@r:embed")
    }
    assert len(donor_relationship_ids) == 2
    donor_media_hashes = {
        hashlib.sha256(
            template_doc.part.related_parts[relationship_id].blob
        ).hexdigest()
        for relationship_id in donor_relationship_ids
    }

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        generated_media_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
    assert "stampa discount" not in document_xml.casefold()
    assert donor_media_hashes.isdisjoint(generated_media_hashes)


def test_biologico_renders_all_already_stored_assessment_content(
    generated_outputs,
):
    ok, path, _ = generated_outputs["ALLEGATO_BIOLOGICO_ALIMENTARE"]
    assert ok and path
    text = _document_text(path)

    expected_values = (
        "Salmonella spp.",
        "ingestione",
        "Salmonellosi",
        "Catena del freddo ≤ 4°C",
        "Guanti monouso",
        "Sorveglianza annuale.",
        "Corso HACCP base.",
        "AL.01",
        "NO",
        "AL.02",
        "SI",
        "Verificare lavaggio mani e formazione del turno serale.",
    )
    missing = [value for value in expected_values if value not in text]
    assert not missing, f"biological annex omitted stored values: {missing}"


def test_biologico_describes_title_x_and_annex_xlvi_accurately(
    generated_outputs,
):
    ok, path, _ = generated_outputs["ALLEGATO_BIOLOGICO_ALIMENTARE"]
    assert ok and path
    text = _document_text(path)

    assert "Il Titolo X classifica gli agenti biologici in quattro gruppi" in text
    assert "l'Allegato XLVI elenca gli agenti dei gruppi 2, 3 e 4" in text
    assert "gruppi da 1 a 4 dell'Allegato XLVI" not in text


def test_biological_sections_remain_isolated_from_fire_and_pee(generated_outputs):
    bio_path = generated_outputs["ALLEGATO_BIOLOGICO_ALIMENTARE"][1]
    unique_headings = (
        "Esiti checklist rischio biologico",
        "Note valutazione biologica",
    )
    bio_text = _document_text(bio_path)
    for heading in unique_headings:
        assert heading in bio_text

    for key in ("ALLEGATO_INCENDIO", "PEE_AZIENDA", "PEE_COMUNE"):
        text = _document_text(generated_outputs[key][1])
        leaked = [heading for heading in unique_headings if heading in text]
        assert not leaked, f"{key} contains biological-only sections: {leaked}"


def test_vdt_emits_full_template_sections(generated_outputs):
    """Allegato VDT must emit every template section, not just a header +
    summary. The audit on 2026-04-29 caught the pre-rewrite generator only
    producing a few sections; this guards against regression to that state.
    """
    ok, path, _ = generated_outputs["ALLEGATO_VDT"]
    assert ok and path.endswith(".docx")
    doc = Document(path)

    headings = [
        p.text.strip() for p in doc.paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]
    required = [
        "Indice",
        "Introduzione",
        "Anagrafica Aziendale",
        "Dati Occupazionali",
        "Organizzazione Aziendale della Sicurezza",
        "Principali fattori di rischio",
        "La postazione di lavoro",
        "Elenco postazioni VDT",
        "Tavole di Valutazione del Rischio VDT",
        "Quadro sinottico di esposizione",
        "Misure di prevenzione",
        "Programma di Attuazione delle Misure di Prevenzione",
        "Dichiarazione del Datore di Lavoro",
        "Firme",
    ]
    missing = [r for r in required if r not in headings]
    assert not missing, f"Allegato VDT missing sections: {missing}"


def test_vdt_quadro_sinottico_emits_classification(generated_outputs):
    """The quadro sinottico must show every valutazione row with its
    Esposto/Non Esposto classification — that's the per-worker summary
    the medico competente reads first.
    """
    ok, path, _ = generated_outputs["ALLEGATO_VDT"]
    assert ok and path
    doc = Document(path)

    sinottico_header = ("Nominativo", "Mansione", "Tempo di utilizzo (h/sett)", "Rischio VDT")
    matched = False
    for t in doc.tables:
        if not t.rows:
            continue
        header = tuple(c.text.strip() for c in t.rows[0].cells)
        if header == sinottico_header:
            matched = True
            assert len(t.rows) >= 2, "quadro sinottico has no data rows"
            risk_col = {row.cells[3].text.strip() for row in t.rows[1:]}
            assert risk_col & {"Esposto", "Non Esposto"}, (
                f"quadro sinottico Rischio VDT col missing classification: {risk_col}"
            )
            break
    assert matched, "VDT quadro sinottico table not found"


def test_dvr_total_table_count_hits_template_parity(generated_outputs):
    """US-2.8 AC1: DVR .docx emits enough tables to match the master template
    (Pre-Parte I + Parte I + II + III + IV).

    For the 6-env Acme fixture: 3 pre + 15 Parte I + 5 Parte II +
    (1 azienda + 6 envs × (identity + addetti + checklist + 2 cat)) +
    3 Parte IV = 57 tables. Real clients with richer per-env risk data
    climb toward the template's 111 organically.
    """
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)
    count = len(doc.tables)
    assert count >= 50, (
        f"DVR Master emitted only {count} tables; expected ≥50 for the "
        f"Acme fixture. Regression in Parte I/II/III/IV parity."
    )


def test_dvr_parte_i_has_anagrafica_and_hazard_library(generated_outputs):
    """US-2.8 AC1: Parte I must emit the anagrafica block, single-role title
    tables, and the 3-group static hazard library (Tables 4, 6–9, 15–17)."""
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)

    headers_seen = []
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        headers_seen.append(tuple(header_cells))

    assert ("Datore di Lavoro",) in headers_seen, (
        "missing single-role Datore di Lavoro table (Template Table 6)"
    )
    assert ("Responsabile del Servizio di Prevenzione e Protezione",) in headers_seen, (
        "missing RSPP title table (Template Table 7)"
    )
    assert ("Rappresentante dei Lavoratori per la Sicurezza",) in headers_seen, (
        "missing RLS title table (Template Table 8)"
    )

    macro_headers = {"Rischi per la Sicurezza", "Rischi per la Salute", "Rischi Trasversali"}
    static_library_headers = [
        h for h in headers_seen
        if len(h) == 2 and h[1] in macro_headers and h[0] == "Categoria"
    ]
    assert len(static_library_headers) == 3, (
        f"expected 3 static hazard-library tables (Templates 15/16/17), "
        f"got {len(static_library_headers)}"
    )


def test_dvr_parte_ii_has_definizioni_and_criteria(generated_outputs):
    """US-2.8 AC1: Parte II must emit the Definizioni glossary and full
    P/D criteria tables (Templates 19, 21, 22)."""
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)

    found_definizioni = False
    found_prob = False
    found_danno = False
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if header_cells == ("Termine", "Definizione"):
            found_definizioni = len(table.rows) >= 10
        if header_cells == ("P", "Livello", "Criteri"):
            found_prob = True
        if header_cells == ("D", "Livello", "Criteri"):
            found_danno = True

    assert found_definizioni, "missing Definizioni glossary (Template Table 19) with ≥10 rows"
    assert found_prob, "missing Scala di Probabilita with criteri column (Template Table 21)"
    assert found_danno, "missing Scala del Danno with criteri column (Template Table 22)"


def test_dvr_parte_iv_has_signature_table(generated_outputs):
    """US-2.8 AC1: Parte IV emits the improvement program grid and the 2×3
    signature block as a real table (Templates 109, 110)."""
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)

    found_program = False
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if header_cells == [
            "Priorità",
            "Rischio",
            "Misura di Miglioramento",
            "Attività / Procedura",
            "Risorse",
            "Responsabile",
            "Scadenza",
        ]:
            found_program = True
            break
    assert found_program, (
        "missing improvement-program grid (Template Table 109)"
    )

    signature_match = False
    for table in doc.tables:
        if len(table.rows) != 2 or len(table.rows[0].cells) != 3:
            continue
        row0_text = " ".join(cell.text for cell in table.rows[0].cells)
        row1_text = " ".join(cell.text for cell in table.rows[1].cells)
        if "Datore di Lavoro" in row0_text and "Rappresentante dei Lavoratori" in row1_text:
            signature_match = True
            break
    assert signature_match, (
        "missing 2×3 signature table (Template Table 110)"
    )


def test_dvr_parte_iii_env_block_structure(generated_outputs):
    """US-2.8 AC1: each environment in Parte III emits the full template block.

    Expected per env: 1 identity table (Table 24), 1 addetti table (Table 25),
    1 SI/NO risk-category checklist (Table 26), plus 1 per-category 5-col risk
    table for every applicable macro-category. The Acme fixture has 6 envs
    each with 2 applicable risk categories → at least 6 × (3 + 2) = 30
    template-shaped tables in Parte III alone.
    """
    ok, path, _ = generated_outputs["DVR_MASTER"]
    assert ok and path
    doc = Document(path)

    headings = [
        p.text for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    env_identity_headers = [
        h for h in headings
        if h.startswith("Identificazione dell'Ambiente di Lavoro")
    ]
    assert len(env_identity_headers) == 6, (
        f"expected one env-identity heading per Acme env (6), "
        f"got {len(env_identity_headers)}: {env_identity_headers}"
    )

    si_no_tables = 0
    macro_label_set = {"Rischi per la Sicurezza", "Rischi per la Salute", "Rischi Trasversali"}
    for table in doc.tables:
        cell_texts = {cell.text.strip() for row in table.rows for cell in row.cells}
        if macro_label_set.issubset(cell_texts) and "Applicabile" in cell_texts:
            si_no_tables += 1
    assert si_no_tables == 6, (
        f"expected one SI/NO risk-category checklist per env (6), "
        f"got {si_no_tables}"
    )

    # Template Tables 27-33 column structure (DVR_TEMPLATE_MAPPING.md
    # "Risk assessment table column structure"). The index column label must
    # match the N2O template's own table header ordering — the template uses
    # "I = P + 2*D" in its risk tables and "I = 2*D + P" in its methodology
    # prose, and the generated DVR reproduces both. 7937654 collapsed the
    # header to the prose form; this pins it back to the template's.
    risk_table_header = [
        "PERICOLO",
        "CONDIZIONI DI IMPIEGO O DI ESPOSIZIONE",
        "RISCHIO",
        "MISURE DI PREVENZIONE E PROTEZIONE ATTUATE E DPI",
        "I = P + 2*D",
    ]
    category_headers_seen = 0
    for table in doc.tables:
        if not table.rows:
            continue
        header_texts = [cell.text.strip() for cell in table.rows[0].cells]
        if header_texts == risk_table_header:
            category_headers_seen += 1
    assert category_headers_seen >= 12, (
        f"expected ≥12 per-category risk tables (6 envs × 2 cats), "
        f"got {category_headers_seen}"
    )
