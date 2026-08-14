"""Preventive per-mansione GESTANTI assessment (art. 11 D.Lgs. 151/2001).

Client request (2026-08): "Nell'allegato gestanti devo avere la possibilità
di poter fare una valutazione oggettiva dei rischi legati alla mansione/i
senza che qualche dipendente sia già in fase di gestazione."

Coverage, following the suite convention of exercising pure helpers over
HTTP round-trips (see test_vdt_api.py / test_gestanti_cross_reference.py):

 1. Schema: the upsert payload has NO worker/persona field — the assessment
    is creatable with zero pregnant workers by construction — and enforces
    the misure-required-unless-compatibile rule.
 2. Overview builder: distinct mansioni prefilled from the organigramma with
    catalog risks, merged with saved rows, kept when staff turns over.
 3. Docx: render_mansioni_section really renders the per-mansione table
    (esito labels, risk summaries) into a python-docx Document.
 4. Route ordering: /gestanti/mansioni must stay registered before
    /gestanti/{valutazione_id} or the literal segment 422s.
"""

from __future__ import annotations

from types import SimpleNamespace
from uuid import uuid4

import pytest
from pydantic import ValidationError

from app.api.v1.gestanti import build_mansioni_overview, router
from app.schemas.gestanti import GestantiMansioneUpsert
from app.services.document_generator.allegato_gestanti import (
    ESITO_LABELS,
    _mansione_rischi_summary,
    render_mansioni_section,
)


# ---------------------------------------------------------------------------
# 1. Upsert schema — objective assessment needs no worker
# ---------------------------------------------------------------------------


def test_upsert_payload_has_no_worker_reference():
    """The valutazione is per-mansione: no persona_id / worker_id anywhere."""
    fields = set(GestantiMansioneUpsert.model_fields)
    assert fields == {"mansione", "esito", "rischi", "misure", "note"}
    # Creatable with just a mansione — zero pregnant workers required.
    body = GestantiMansioneUpsert(mansione="Impiegata amministrativa")
    assert body.esito == "compatibile"
    assert body.rischi is None  # None = "server, prefill from the catalog"


def test_mansione_is_normalized():
    body = GestantiMansioneUpsert(mansione="  Aiuto   cuoco ")
    assert body.mansione == "Aiuto cuoco"


def test_mansione_too_short_rejected():
    with pytest.raises(ValidationError):
        GestantiMansioneUpsert(mansione="  a  ")


def test_esito_enum_enforced():
    with pytest.raises(ValidationError):
        GestantiMansioneUpsert(mansione="Cuoco", esito="forse")  # type: ignore[arg-type]


def test_limitazioni_require_misure():
    with pytest.raises(ValidationError):
        GestantiMansioneUpsert(
            mansione="Cuoco", esito="compatibile_con_limitazioni"
        )
    with pytest.raises(ValidationError):
        GestantiMansioneUpsert(
            mansione="Cuoco", esito="non_compatibile", misure="corte"
        )
    ok = GestantiMansioneUpsert(
        mansione="Cuoco",
        esito="compatibile_con_limitazioni",
        misure="Esonero dalla movimentazione di pentole oltre 3 kg",
    )
    assert ok.misure.startswith("Esonero")


def test_compatibile_needs_no_misure():
    ok = GestantiMansioneUpsert(mansione="Segretaria", esito="compatibile")
    assert ok.misure is None


# ---------------------------------------------------------------------------
# 2. Overview builder — prefill from the organigramma, no pregnancy needed
# ---------------------------------------------------------------------------


def _persona(mansione, sesso="F"):
    return SimpleNamespace(mansione=mansione, sesso=sesso)


def _saved_row(mansione, esito="compatibile", rischi=None, misure=None, note=None):
    return SimpleNamespace(
        id=uuid4(),
        azienda_id=uuid4(),
        mansione=mansione,
        esito=esito,
        rischi=rischi or [],
        misure=misure,
        note=note,
        created_at=__import__("datetime").datetime(2026, 8, 1, 12, 0),
    )


def test_overview_lists_mansioni_with_zero_pregnant_workers():
    """The core of the client request: mansioni are assessable even though
    nobody notified a pregnancy (there is no gestanti row anywhere)."""
    persone = [_persona("Saldatrice"), _persona("Impiegata amministrativa")]
    items = build_mansioni_overview(persone, saved=[])
    assert [it.mansione for it in items] == [
        "Impiegata amministrativa",
        "Saldatrice",
    ]
    assert all(it.valutazione is None for it in items)


def test_overview_prefills_catalog_risks_per_mansione():
    items = build_mansioni_overview([_persona("Saldatrice")], saved=[])
    keys = {r.risk_key for r in items[0].suggested_risks}
    # saldat* hits at least CMR chemicals, noise, hand-arm vibrations.
    assert {"chemical_exposure_cmr", "noise_exposure", "hand_arm_vibrations"} <= keys
    # Office job: clean prefill.
    clean = build_mansioni_overview([_persona("Addetta segreteria")], saved=[])
    assert clean[0].suggested_risks == []


def test_overview_deduplicates_mansioni_case_and_whitespace():
    persone = [
        _persona("Cuoco"),
        _persona("cuoco ", sesso="M"),
        _persona("  CUOCO", sesso="F"),
    ]
    items = build_mansioni_overview(persone, saved=[])
    assert len(items) == 1
    assert items[0].num_persone == 3
    assert items[0].num_lavoratrici == 2


def test_overview_merges_saved_valutazione():
    saved = _saved_row(
        "Saldatrice",
        esito="compatibile_con_limitazioni",
        misure="Divieto di saldatura; adibizione a controllo qualita'",
    )
    items = build_mansioni_overview([_persona("saldatrice")], [saved])
    assert len(items) == 1
    assert items[0].valutazione is not None
    assert items[0].valutazione.esito == "compatibile_con_limitazioni"


def test_overview_keeps_saved_rows_after_staff_turnover():
    """A mansione assessed in the past stays listed even when nobody holds
    it any more — the preventive valutazione outlives the organigramma."""
    saved = _saved_row("Magazziniera", esito="non_compatibile", misure="x" * 12)
    items = build_mansioni_overview([_persona("Impiegata")], [saved])
    mansioni = {it.mansione for it in items}
    assert "Magazziniera" in mansioni
    orphan = next(it for it in items if it.mansione == "Magazziniera")
    assert orphan.num_persone == 0
    assert orphan.valutazione is not None


def test_overview_skips_blank_mansioni():
    items = build_mansioni_overview(
        [_persona(""), _persona(None), _persona("   ")], saved=[]
    )
    assert items == []


# ---------------------------------------------------------------------------
# 3. Docx rendering — the allegato carries the per-mansione section
# ---------------------------------------------------------------------------


def _doc_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_render_mansioni_section_renders_table():
    from docx import Document

    doc = Document()
    rows = [
        SimpleNamespace(
            mansione="Saldatrice",
            esito="compatibile_con_limitazioni",
            rischi=[
                {
                    "risk_key": "noise_exposure",
                    "allegato": "C",
                    "descrizione": "Esposizione a rumore",
                }
            ],
            misure="Adibizione temporanea a mansioni di ufficio",
            note=None,
        ),
        SimpleNamespace(
            mansione="Impiegata",
            esito="compatibile",
            rischi=[],
            misure=None,
            note="Nessuna limitazione",
        ),
    ]
    render_mansioni_section(doc, rows)
    text = _doc_text(doc)
    assert "Valutazione preventiva dei rischi per mansione" in text
    assert "Saldatrice" in text
    assert "Compatibile con limitazioni" in text
    assert "Esposizione a rumore (All. C)" in text
    assert "Adibizione temporanea a mansioni di ufficio" in text
    assert "Nessun rischio rilevato" in text
    assert "Note: Nessuna limitazione" in text


def test_render_mansioni_section_empty_state():
    from docx import Document

    doc = Document()
    render_mansioni_section(doc, [])
    text = _doc_text(doc)
    assert "Valutazione preventiva dei rischi per mansione" in text
    assert "Nessuna valutazione preventiva per mansione" in text


def test_esito_labels_cover_all_esiti():
    from app.schemas.gestanti import EsitoMansione
    from typing import get_args

    assert set(ESITO_LABELS) == set(get_args(EsitoMansione))


def test_rischi_summary_tolerates_legacy_and_junk_rows():
    assert _mansione_rischi_summary(None) == "Nessun rischio rilevato"
    assert _mansione_rischi_summary([{"rischio": "Lavoro notturno"}]) == "Lavoro notturno"
    assert (
        _mansione_rischi_summary(["garbage", {"allegato": "A"}])
        == "Nessun rischio rilevato"
    )


def test_generator_renders_section_in_generate():
    """The generate() path must call the section renderer and query the
    per-mansione table — pins the wiring, same style as
    test_allegato_gestanti_keys.test_generator_uses_same_logic."""
    from pathlib import Path

    src = (
        Path(__file__).resolve().parents[1]
        / "app"
        / "services"
        / "document_generator"
        / "allegato_gestanti.py"
    )
    text = src.read_text(encoding="utf-8")
    assert "render_mansioni_section(doc, mansioni_vals)" in text
    assert "load_gestanti_mansioni(self.db, self.azienda_id)" in text


# ---------------------------------------------------------------------------
# 4. Routing — literal /mansioni must beat the {valutazione_id} UUID param
# ---------------------------------------------------------------------------


def test_mansioni_routes_registered_before_uuid_route():
    paths = [getattr(r, "path", None) for r in router.routes]
    mansioni_get = paths.index("/aziende/{azienda_id}/gestanti/mansioni")
    uuid_get = paths.index("/aziende/{azienda_id}/gestanti/{valutazione_id}")
    assert mansioni_get < uuid_get, (
        "/gestanti/mansioni must be registered before /gestanti/{valutazione_id} "
        "or 'mansioni' is captured by the UUID path param and 422s"
    )


def test_mansioni_endpoints_exist_with_expected_methods():
    from tests.conftest import route_pairs

    pairs = route_pairs(router)
    assert ("GET", "/aziende/{azienda_id}/gestanti/mansioni") in pairs
    assert ("PUT", "/aziende/{azienda_id}/gestanti/mansioni") in pairs
    assert (
        "DELETE",
        "/aziende/{azienda_id}/gestanti/mansioni/{mansione_valutazione_id}",
    ) in pairs
