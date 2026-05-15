"""Tests for the DVR Master TOC cached-body finalizer.

The TOC field is emitted before the rest of the doc, then post-filled by
``_finalize_table_of_contents`` so the pre-F9 view in Word shows the
real outline instead of a generic 4-line placeholder
(user feedback 2026-05-12: "indice deve essere esploso con paginazione
e piu dettagliato").
"""

from docx import Document
from docx.oxml.ns import qn

from app.services.document_generator.dvr_master import DVRMasterGenerator


def _new_generator() -> DVRMasterGenerator:
    return DVRMasterGenerator.__new__(DVRMasterGenerator)


def _cached_body_paragraphs(doc: Document, field_start_p, end_p) -> list[str]:
    """Return the text of every paragraph between field_start_p and end_p
    in the body — i.e. the TOC field's cached body the user sees before
    pressing F9.
    """
    body = doc.element.body
    children = list(body)
    start = children.index(field_start_p._p)
    end = children.index(end_p._p)
    out: list[str] = []
    for el in children[start + 1 : end]:
        if el.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t")))
        out.append(text)
    return out


def test_toc_cached_body_replaced_with_real_headings() -> None:
    """After the rest of the doc emits real headings, finalize should
    swap the placeholder line out for the actual outline."""
    gen = _new_generator()
    doc = Document()
    field_start_p, end_p = gen._add_table_of_contents(doc)

    # Before finalize: just the single placeholder line.
    initial = _cached_body_paragraphs(doc, field_start_p, end_p)
    assert len(initial) == 1
    assert "Ctrl+A" in initial[0]

    # Simulate what the rest of the generator does — add a few headings
    # at multiple levels. INDICE is already a Heading 1 above the field
    # but it sits before end_p so the finalizer must skip it.
    doc.add_heading("Premessa", level=1)
    doc.add_heading("PARTE I — DATI GENERALI DELL'AZIENDA", level=1)
    doc.add_heading("1. Presentazione dell'Azienda", level=2)
    doc.add_heading("2. Anagrafica Aziendale", level=2)
    doc.add_heading("Storico Revisioni", level=3)
    doc.add_heading("PARTE II — DESCRIZIONE", level=1)
    doc.add_heading("2.1 Descrizione dell'Attivita", level=2)

    gen._finalize_table_of_contents(doc, field_start_p, end_p)

    cached = _cached_body_paragraphs(doc, field_start_p, end_p)
    # Placeholder is gone, real headings are in.
    assert "Ctrl+A" not in "\n".join(cached)
    assert cached == [
        "Premessa",
        "PARTE I — DATI GENERALI DELL'AZIENDA",
        "1. Presentazione dell'Azienda",
        "2. Anagrafica Aziendale",
        "Storico Revisioni",
        "PARTE II — DESCRIZIONE",
        "2.1 Descrizione dell'Attivita",
    ]


def test_toc_skips_indice_heading_itself() -> None:
    """The "INDICE" heading sits ABOVE the field's end paragraph in the
    body, so the finalizer must not include it in the cached outline —
    otherwise the TOC would list itself."""
    gen = _new_generator()
    doc = Document()
    field_start_p, end_p = gen._add_table_of_contents(doc)

    doc.add_heading("Premessa", level=1)

    gen._finalize_table_of_contents(doc, field_start_p, end_p)

    cached = _cached_body_paragraphs(doc, field_start_p, end_p)
    assert "INDICE" not in cached
    assert cached == ["Premessa"]


def test_toc_preserves_field_chars() -> None:
    """The fldChar begin / separate / end triplet must survive the
    cached-body rewrite — without it Word treats the TOC as plain text
    and F9 refresh becomes a no-op."""
    gen = _new_generator()
    doc = Document()
    field_start_p, end_p = gen._add_table_of_contents(doc)
    doc.add_heading("Premessa", level=1)
    gen._finalize_table_of_contents(doc, field_start_p, end_p)

    fld_chars = [
        el.get(qn("w:fldCharType"))
        for el in doc.element.body.iter(qn("w:fldChar"))
    ]
    assert fld_chars == ["begin", "separate", "end"]


def test_toc_finalize_noop_when_no_headings() -> None:
    """If somehow no real headings exist after the TOC, the placeholder
    should stay rather than the field collapsing to empty."""
    gen = _new_generator()
    doc = Document()
    field_start_p, end_p = gen._add_table_of_contents(doc)
    gen._finalize_table_of_contents(doc, field_start_p, end_p)

    cached = _cached_body_paragraphs(doc, field_start_p, end_p)
    assert len(cached) == 1
    assert "Ctrl+A" in cached[0]
