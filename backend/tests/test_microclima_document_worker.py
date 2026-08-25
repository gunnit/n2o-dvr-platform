"""Regression coverage for thermal calculations in document workers."""

import os
import subprocess
import sys
import uuid
from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace

import pytest
import yaml
from docx import Document
from docx.shared import Cm

BACKEND_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BACKEND_ROOT))

from app.services.document_generator.allegato_microclima import (  # noqa: E402
    _compute_pmv_ppd,
)
from app.services.document_generator import (  # noqa: E402
    allegato_microclima_severo as severe_module,
)
from app.services.document_generator.allegato_microclima_severo import (  # noqa: E402
    _compute_phs,
)
from app.services.document_generator.docx_utils import add_data_table  # noqa: E402


def test_moderate_document_uses_current_iso_pmv_model():
    """The document must expose the same ISO result as the calculator API."""
    pmv, ppd = _compute_pmv_ppd(22.0, 22.0, 0.1, 50.0, 1.2, 0.5)

    assert pmv == pytest.approx(-0.81)
    assert ppd == pytest.approx(18.9)


def test_severe_heat_document_uses_current_phs_model():
    """PHS output must use current units and the binding exposure limit."""
    sweat_loss_g, rectal_temp, exposure_limit_min = _compute_phs(
        35.0, 38.0, 0.5, 50.0, 2.0, 0.5
    )

    assert sweat_loss_g == pytest.approx(3496.7)
    assert rectal_temp == pytest.approx(37.4)
    assert exposure_limit_min == pytest.approx(480.0)


def test_severe_heat_document_falls_back_for_non_finite_phs_results():
    """Persisted rows outside ISO applicability must never render NaN."""
    with pytest.warns(UserWarning, match="outside the applicability limits"):
        result = _compute_phs(10.0, 10.0, 0.5, 50.0, 2.0, 0.5)

    assert result == pytest.approx((1500.0, 37.0, 480.0))


def test_celery_parent_keeps_thermal_models_lazy_on_starter_worker():
    """The 512 MB worker parent must not retain the thermal model before fork."""
    script = """
import sys
from app.celery_app import celery_app

celery_app.loader.import_default_modules()
raise SystemExit(
    0 if "app.services.microclima_calculator" not in sys.modules else 1
)
"""
    env = os.environ.copy()
    env["PYTHONPATH"] = str(BACKEND_ROOT)

    completed = subprocess.run(
        [sys.executable, "-c", script],
        cwd=BACKEND_ROOT,
        env=env,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr


def test_starter_worker_serializes_memory_heavy_document_tasks():
    """One child keeps lazy thermal imports within the Starter memory budget."""
    script = """
from app.celery_app import celery_app

raise SystemExit(0 if celery_app.conf.worker_concurrency == 1 else 1)
"""
    env = os.environ.copy()
    env["PYTHONPATH"] = str(BACKEND_ROOT)

    completed = subprocess.run(
        [sys.executable, "-c", script],
        cwd=BACKEND_ROOT,
        env=env,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr


def test_render_worker_does_not_override_serial_concurrency():
    """The production command must preserve the one-child memory budget."""
    blueprint = yaml.safe_load((BACKEND_ROOT / "render.yaml").read_text())
    worker = next(
        service
        for service in blueprint["services"]
        if service["name"] == "n2o-dvr-worker"
    )

    assert worker["startCommand"].endswith("--concurrency=1")


def test_worker_startup_preserves_document_fallback_if_dependency_import_fails():
    """A broken optional model import must not disable every Celery task."""
    script = """
import builtins

real_import = builtins.__import__

def guarded_import(name, *args, **kwargs):
    if name.startswith("pythermalcomfort"):
        raise ImportError("simulated thermal-model import failure")
    return real_import(name, *args, **kwargs)

builtins.__import__ = guarded_import

from app.celery_app import celery_app
celery_app.loader.import_default_modules()

from app.services.document_generator.allegato_microclima import _compute_pmv_ppd
raise SystemExit(0 if _compute_pmv_ppd(22, 22, 0.1, 50, 1.2, 0.5) == (0.0, 5.0) else 2)
"""
    env = os.environ.copy()
    env["PYTHONPATH"] = str(BACKEND_ROOT)

    completed = subprocess.run(
        [sys.executable, "-c", script],
        cwd=BACKEND_ROOT,
        env=env,
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr


def test_data_table_honors_explicit_column_widths():
    """Dense document tables need deterministic room for narrative cells."""
    doc = Document()
    widths_cm = [2.0, 1.0, 3.0]

    table = add_data_table(
        doc,
        ["Ambiente", "PMV", "Classificazione"],
        [["Reparto caldo", "0.2", "Accettabile per l'intera giornata"]],
        column_widths_cm=widths_cm,
    )

    assert table.autofit is False
    assert all(
        column.width / Cm(1) == pytest.approx(width_cm, abs=0.002)
        for column, width_cm in zip(table.columns, widths_cm)
    )


@pytest.mark.asyncio
async def test_severe_heat_document_names_current_standard_and_fits_page(
    monkeypatch, tmp_path
):
    """Bind the real nine-column heat table to its standard and page width."""
    ambiente_id = uuid.uuid4()
    azienda = SimpleNamespace(ragione_sociale="Azienda Verifica Microclima")
    ambiente = SimpleNamespace(id=ambiente_id, nome="Reparto caldo")
    row = SimpleNamespace(
        ambiente_id=ambiente_id,
        nome_area="Reparto caldo",
        tipo_ambiente="severo_caldo",
        temperatura_aria=35.0,
        temperatura_radiante=38.0,
        velocita_aria=0.5,
        umidita_relativa=50.0,
        metabolismo=2.0,
        isolamento_vestiario=0.5,
    )

    async def fake_load_microclima(db, azienda_id):
        return [row]

    async def fake_load_data():
        return {
            "azienda": azienda,
            "generated_at": datetime(2026, 8, 24, tzinfo=timezone.utc),
            "ambienti": [ambiente],
        }

    async def fake_next_version():
        return 1

    monkeypatch.setattr(severe_module, "load_microclima", fake_load_microclima)
    generator = severe_module.AllegatoMicroclimaSeveroGenerator(uuid.uuid4(), None)
    monkeypatch.setattr(generator, "load_data", fake_load_data)
    monkeypatch.setattr(generator, "_next_version", fake_next_version)
    monkeypatch.setattr(generator, "_get_output_dir", lambda: str(tmp_path))

    path = await generator.generate()
    doc = Document(path)
    document_text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    heat_table = next(
        table
        for table in doc.tables
        if [cell.text for cell in table.rows[0].cells][:2] == ["Ambiente", "t_aria"]
    )
    expected_widths_cm = [1.8, 1.1, 1.1, 0.9, 0.9, 1.7, 1.1, 1.1, 5.3]
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    assert "UNI EN ISO 7933:2023" in document_text
    assert heat_table.autofit is False
    assert len(heat_table.columns) == len(expected_widths_cm)
    assert all(
        column.width / Cm(1) == pytest.approx(width_cm, abs=0.002)
        for column, width_cm in zip(heat_table.columns, expected_widths_cm)
    )
    assert sum(column.width for column in heat_table.columns) <= usable_width
