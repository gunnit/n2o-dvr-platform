"""Unit tests for the document branding resolver (per-organization letterhead).

Pure logic — no DB. Guards the fallback behaviour that keeps document
generation working for a sparse org and keeps the committed N2O logo as the
default so existing output never regresses. The logo now lives in the DB as
bytes (not a disk path), so the worker that generates documents can read it.
"""

from __future__ import annotations

import io
import types

from app.services.document_generator.branding import (
    DEFAULT_LOGO_PATH,
    Branding,
    resolve_logo_source,
)


def _fake_org(**overrides):
    """A duck-typed Organization-like object."""
    base = {
        "name": "Acme Safety SRL",
        "logo_bytes": None,
        "logo_content_type": None,
        "indirizzo": None,
        "cap": None,
        "citta": None,
        "provincia": None,
        "partita_iva": None,
        "codice_fiscale": None,
        "telefono": None,
        "email": None,
        "sito_web": None,
        "rspp_nome": None,
    }
    base.update(overrides)
    return types.SimpleNamespace(**base)


def test_default_uses_n2o_firm_name():
    b = Branding.default()
    assert b.firm_name == "N2O SRL"
    assert b.logo_bytes is None


def test_from_organization_reads_fields():
    org = _fake_org(
        name="Acme Safety SRL",
        indirizzo="Via Roma 1",
        partita_iva="01234567890",
        rspp_nome="Mario Rossi",
        logo_bytes=b"\x89PNG\r\n\x1a\n",
        logo_content_type="image/png",
    )
    b = Branding.from_organization(org)
    assert b.firm_name == "Acme Safety SRL"
    assert b.indirizzo == "Via Roma 1"
    assert b.partita_iva == "01234567890"
    assert b.rspp_nome == "Mario Rossi"
    assert b.logo_bytes == b"\x89PNG\r\n\x1a\n"
    assert b.logo_content_type == "image/png"


def test_from_organization_blank_name_falls_back_to_default_firm_name():
    org = _fake_org(name="   ")
    b = Branding.from_organization(org)
    assert b.firm_name == "N2O SRL"


def test_from_organization_handles_none_org():
    b = Branding.from_organization(None)
    assert b.firm_name == "N2O SRL"


def test_resolve_logo_source_falls_back_to_default_path_when_no_bytes():
    b = Branding.default()
    src = resolve_logo_source(b)
    # The committed asset exists, so we get its path as a string.
    assert src == str(DEFAULT_LOGO_PATH)


def test_resolve_logo_source_uses_custom_bytes_when_present():
    payload = b"\x89PNG\r\n\x1a\nDATA"
    b = Branding.default()
    b.logo_bytes = payload
    src = resolve_logo_source(b)
    assert isinstance(src, io.BytesIO)
    assert src.getvalue() == payload


def test_resolve_logo_source_returns_fresh_stream_each_call():
    b = Branding(logo_bytes=b"abc")
    a, c = resolve_logo_source(b), resolve_logo_source(b)
    assert a is not c  # fresh BytesIO so multiple embeds don't share a cursor


def test_has_letterhead_detail():
    b = Branding.default()
    assert b.has_letterhead_detail() is False
    b.indirizzo = "Via Roma 1"
    assert b.has_letterhead_detail() is True


def test_is_configured():
    assert Branding.default().is_configured() is False
    assert Branding(firm_name="Acme SRL").is_configured() is True  # custom name
    assert Branding(indirizzo="Via Roma 1").is_configured() is True  # detail
    assert Branding(logo_bytes=b"x").is_configured() is True  # custom logo


def test_address_line_full():
    b = Branding(indirizzo="Via Roma 1", cap="20100", citta="Milano", provincia="MI")
    assert b.address_line() == "Via Roma 1, 20100 Milano (MI)"


def test_address_line_province_only_has_no_stray_parens():
    # Province set, no cap/citta — must not emit a dangling "(MI)" fragment.
    b = Branding(provincia="MI")
    assert b.address_line() == "MI"


def test_address_line_empty_is_none():
    assert Branding.default().address_line() is None


def _valid_png(rgb=(0, 61, 116)) -> bytes:
    """Minimal valid 2x2 PNG so python-docx will actually embed it."""
    import struct
    import zlib

    def chunk(t, d):
        c = t + d
        return struct.pack(">I", len(d)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 2, 2, 8, 2, 0, 0, 0)
    px = bytes(rgb)
    raw = b"".join(b"\x00" + px * 2 for _ in range(2))
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


def test_custom_logo_bytes_embed_into_docx():
    """Regression for review finding #1: the per-org logo must end up inside the
    generated .docx (it now travels as DB bytes, not a worker-unreachable file)."""
    import zipfile

    from docx import Document
    from docx.shared import Inches

    png = _valid_png()
    branding = Branding(logo_bytes=png, logo_content_type="image/png")

    doc = Document()
    doc.add_paragraph().add_run().add_picture(resolve_logo_source(branding), width=Inches(2.0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with zipfile.ZipFile(buf) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        assert media, "expected an embedded image in word/media/"
        embedded = zf.read(media[0])
    assert embedded == png  # the customer's exact logo bytes, not the default

