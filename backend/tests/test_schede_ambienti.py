"""Scheda ambiente table shared by the incendio and PEE allegati."""

from decimal import Decimal
from types import SimpleNamespace

from docx import Document

from app.services.document_generator.schede_ambienti import (
    SCHEDA_HEADERS,
    add_schede_ambienti,
    scheda_rows,
)


def _amb(**kw):
    base = dict(
        nome="Magazzino",
        superficie_mq=None,
        descrizione_locale=None,
        materiali_presenti=None,
        max_persone=None,
        sorgenti_innesco=None,
    )
    base.update(kw)
    return SimpleNamespace(**base)


def test_rows_skip_ambienti_without_any_scheda_field():
    rows = scheda_rows([_amb(), _amb(nome="Ufficio", max_persone=4)])
    assert rows == [["Ufficio", "—", "—", "—", "4", "—"]]


def test_rows_format_every_field():
    amb = _amb(
        superficie_mq=Decimal("120.50"),
        descrizione_locale="Capannone in muratura",
        materiali_presenti="Cartone; bancali in legno",
        max_persone=12,
        sorgenti_innesco="Quadro elettrico",
    )
    assert scheda_rows([amb]) == [
        [
            "Magazzino",
            "Capannone in muratura",
            "120.5",
            "Cartone; bancali in legno",
            "12",
            "Quadro elettrico",
        ]
    ]


def test_zero_persone_is_a_value_not_a_gap():
    assert scheda_rows([_amb(max_persone=0)])[0][4] == "0"


def test_add_schede_prints_table_or_note():
    doc = Document()
    printed = add_schede_ambienti(doc, [_amb(sorgenti_innesco="Forno")])
    assert printed == 1
    table = doc.tables[-1]
    assert [c.text for c in table.rows[0].cells] == SCHEDA_HEADERS
    assert table.rows[1].cells[5].text == "Forno"

    empty = Document()
    assert add_schede_ambienti(empty, [_amb()]) == 0
    assert not empty.tables
    assert any("non compilate" in p.text for p in empty.paragraphs)
