"""Unit tests for POS phase validation + generator rendering (US-4.7).

Covers the three pieces the phase-builder UI depends on:

1. ``PosPhase`` Pydantic contract — shape, bounds, list dedup, sub-schemas.
2. ``validate_phases`` / ``normalize_ordering`` / ``dependency_violations_after_ordering``
   — the pure graph helpers the endpoint and the .docx generator share.
3. Route registration on the POS router + generator rendering of the
   new structured shape against a synthetic POS row (no DB required —
   the generator helper is a pure function on a list of dicts).

Mirrors the schema+route pattern used by ``test_rischi_misure_libreria.py``
so a reviewer only has to know one style.
"""

from __future__ import annotations

import pytest
from pydantic import ValidationError

from app.schemas.pos_phase import (
    PhaseNiosh,
    PhaseRumore,
    PhaseVibrazioni,
    PosPhase,
    PosPhasesUpdate,
)
from app.services.pos_phases import (
    PosPhaseError,
    dependency_violations_after_ordering,
    normalize_ordering,
    validate_phases,
)


# ---------------------------------------------------------------------------
# PosPhase schema
# ---------------------------------------------------------------------------


def _phase(**overrides) -> PosPhase:
    base = dict(id="p1", ordine=0, nome="Scavo")
    base.update(overrides)
    return PosPhase(**base)


def test_phase_minimum_fields():
    p = _phase()
    assert p.id == "p1"
    assert p.nome == "Scavo"
    assert p.descrizione is None
    assert p.rischi == []
    assert p.dipende_da == []


def test_phase_rejects_empty_nome():
    with pytest.raises(ValidationError):
        PosPhase(id="p1", ordine=0, nome="")


def test_phase_rejects_negative_ordine():
    with pytest.raises(ValidationError):
        PosPhase(id="p1", ordine=-1, nome="x")


def test_phase_strips_and_dedups_list_fields():
    p = _phase(
        rischi=["Caduta", "  Caduta  ", "Crollo", ""],
        dpi=["Casco", "Casco"],
        dipende_da=["p0", "p0", " p1 "],  # back-ref to self is fine for this level
    )
    assert p.rischi == ["Caduta", "Crollo"]
    assert p.dpi == ["Casco"]
    # dedup + strip applied; self-dep here is a string hygiene concern only,
    # the graph validator separately rejects self-dependencies.
    assert p.dipende_da == ["p0", "p1"]


def test_phase_extra_fields_forbidden():
    """Guards against typos (e.g. "rischi_extra") silently persisting."""
    with pytest.raises(ValidationError):
        PosPhase(id="p1", ordine=0, nome="x", rischi_extra=["y"])


# ---------------------------------------------------------------------------
# Sub-schemas
# ---------------------------------------------------------------------------


def test_niosh_rejects_missing_required_factors():
    with pytest.raises(ValidationError):
        PhaseNiosh(
            peso_sollevato=15,
            cp=25,
            fattore_a=1.0,
            fattore_b=1.0,
            fattore_c=1.0,
            fattore_d=1.0,
            fattore_e=1.0,
            # fattore_f intentionally omitted
        )


def test_niosh_accepts_valid_payload_with_computed_result():
    n = PhaseNiosh(
        peso_sollevato=20,
        cp=25,
        fattore_a=0.93,
        fattore_b=1.0,
        fattore_c=1.0,
        fattore_d=1.0,
        fattore_e=1.0,
        fattore_f=0.88,
        plr=20.46,
        ir=0.98,
        livello="GIALLA",
    )
    assert n.livello == "GIALLA"


def test_rumore_fascia_is_a_closed_set():
    with pytest.raises(ValidationError):
        PhaseRumore(lex_8h_dba=83.2, fascia="media")  # not in literal
    ok = PhaseRumore(lex_8h_dba=83.2, fascia="80-85", dpi_obbligatori=False)
    assert ok.fascia == "80-85"


def test_vibrazioni_all_fields_optional():
    v = PhaseVibrazioni()
    assert v.a8_mano_braccio is None
    assert v.entro_limiti is True  # safe default


# ---------------------------------------------------------------------------
# Graph validation
# ---------------------------------------------------------------------------


def test_validate_phases_accepts_happy_chain():
    phases = [
        _phase(id="a", ordine=0, nome="Scavo"),
        _phase(id="b", ordine=1, nome="Fondazioni", dipende_da=["a"]),
        _phase(id="c", ordine=2, nome="Getto", dipende_da=["b"]),
    ]
    validate_phases(phases)  # does not raise


def test_validate_phases_rejects_duplicate_ids():
    phases = [
        _phase(id="a", ordine=0, nome="Scavo"),
        _phase(id="a", ordine=1, nome="Altro scavo"),
    ]
    with pytest.raises(PosPhaseError, match="duplicato"):
        validate_phases(phases)


def test_validate_phases_rejects_unknown_dependency():
    phases = [
        _phase(id="a", ordine=0, nome="Scavo", dipende_da=["nonesiste"]),
    ]
    with pytest.raises(PosPhaseError, match="inesistente"):
        validate_phases(phases)


def test_validate_phases_rejects_self_dependency():
    phases = [
        _phase(id="a", ordine=0, nome="Scavo", dipende_da=["a"]),
    ]
    with pytest.raises(PosPhaseError, match="se stessa"):
        validate_phases(phases)


def test_validate_phases_rejects_cycle():
    phases = [
        _phase(id="a", ordine=0, nome="Scavo", dipende_da=["c"]),
        _phase(id="b", ordine=1, nome="Fondazioni", dipende_da=["a"]),
        _phase(id="c", ordine=2, nome="Getto", dipende_da=["b"]),
    ]
    with pytest.raises(PosPhaseError, match="cicliche"):
        validate_phases(phases)


def test_validate_phases_accepts_diamond():
    """a → b, a → c, b → d, c → d should be fine."""
    phases = [
        _phase(id="a", ordine=0, nome="A"),
        _phase(id="b", ordine=1, nome="B", dipende_da=["a"]),
        _phase(id="c", ordine=2, nome="C", dipende_da=["a"]),
        _phase(id="d", ordine=3, nome="D", dipende_da=["b", "c"]),
    ]
    validate_phases(phases)


# ---------------------------------------------------------------------------
# Ordering
# ---------------------------------------------------------------------------


def test_normalize_ordering_renumbers_0_to_n_minus_1():
    phases = [
        _phase(id="a", ordine=10, nome="A"),
        _phase(id="b", ordine=5, nome="B"),
        _phase(id="c", ordine=7, nome="C"),
    ]
    result = normalize_ordering(phases)
    assert [p.id for p in result] == ["b", "c", "a"]
    assert [p.ordine for p in result] == [0, 1, 2]


def test_normalize_ordering_stable_on_ties():
    """Equal ``ordine`` values preserve insertion order (stable sort)."""
    phases = [
        _phase(id="a", ordine=0, nome="A"),
        _phase(id="b", ordine=0, nome="B"),
        _phase(id="c", ordine=0, nome="C"),
    ]
    result = normalize_ordering(phases)
    assert [p.id for p in result] == ["a", "b", "c"]


def test_dependency_violations_reports_phase_before_its_predecessor():
    """If the operator drags a dependent phase before its predecessor,
    the generator footnote must name the pair."""
    phases = [
        _phase(id="a", ordine=1, nome="Getto", dipende_da=["b"]),
        _phase(id="b", ordine=0, nome="Fondazioni"),
    ]
    # Note: a depends on b, and b is ordine 0 while a is ordine 1 — that's
    # correct, b DOES precede a, so no violation.
    assert dependency_violations_after_ordering(phases) == []

    # Now flip: a before b.
    phases = [
        _phase(id="a", ordine=0, nome="Getto", dipende_da=["b"]),
        _phase(id="b", ordine=1, nome="Fondazioni"),
    ]
    violations = dependency_violations_after_ordering(phases)
    assert violations == [("Getto", "Fondazioni")]


# ---------------------------------------------------------------------------
# PosPhasesUpdate body
# ---------------------------------------------------------------------------


def test_phases_update_rejects_extra_keys():
    with pytest.raises(ValidationError):
        PosPhasesUpdate(fasi=[], note="typo_goes_here")


def test_phases_update_accepts_empty_list():
    """An operator deleting every phase must be able to save the empty list."""
    body = PosPhasesUpdate(fasi=[])
    assert body.fasi == []


# ---------------------------------------------------------------------------
# Route registration
# ---------------------------------------------------------------------------


def test_router_registers_phases_endpoint():
    """Fail loudly if anyone renames the PUT /fasi endpoint
    the frontend phase-builder card depends on."""
    from app.api.v1.router import api_router

    paths = {
        (method, getattr(r, "path", ""))
        for r in api_router.routes
        for method in getattr(r, "methods", set()) or set()
    }
    prefix = "/api/v1/aziende/{azienda_id}/pos/{pos_id}"
    assert ("PUT", f"{prefix}/fasi") in paths, "phase-list endpoint missing"


# ---------------------------------------------------------------------------
# Generator rendering
# ---------------------------------------------------------------------------


def test_generator_renders_structured_phases():
    """Feeds the structured shape through ``_render_phase_sections`` and
    checks the resulting docx has the synoptic + per-phase + footnote
    sections the AC calls for."""
    from docx import Document

    from app.services.document_generator.pos import _render_phase_sections

    doc = Document()
    fasi = [
        {
            "id": "a",
            "ordine": 0,
            "nome": "Scavo",
            "descrizione": "Escavazione fondazioni",
            "rischi": ["Caduta nel vuoto"],
            "dpi": ["Casco", "Scarpe"],
            "mezzi": ["Escavatore"],
            "niosh": None,
            "rumore": {"lex_8h_dba": 82.5, "fascia": "80-85", "dpi_obbligatori": False},
            "vibrazioni": None,
            "dipende_da": [],
        },
        {
            "id": "b",
            "ordine": 1,
            "nome": "Fondazioni",
            "descrizione": "Getto del plinto",
            "rischi": ["Schiacciamento"],
            "dpi": ["Casco"],
            "mezzi": [],
            "niosh": {
                "peso_sollevato": 20,
                "cp": 25,
                "fattore_a": 0.93,
                "fattore_b": 1.0,
                "fattore_c": 1.0,
                "fattore_d": 1.0,
                "fattore_e": 1.0,
                "fattore_f": 0.88,
                "plr": 20.46,
                "ir": 0.98,
                "livello": "GIALLA",
            },
            "rumore": None,
            "vibrazioni": {
                "a8_mano_braccio": 2.8,
                "a8_corpo_intero": None,
                "entro_limiti": True,
                "note": None,
            },
            "dipende_da": ["a"],
        },
    ]
    _render_phase_sections(doc, fasi)

    # Synoptic + per-phase headings + kv tables should all be there.
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Quadro sinottico" in paragraph_text
    assert "1. Scavo" in paragraph_text
    assert "2. Fondazioni" in paragraph_text
    # NIOSH + rumore sub-sections emit recognisable labels.
    assert "Valutazione NIOSH" in paragraph_text
    assert "Esposizione al rumore" in paragraph_text
    assert "Esposizione a vibrazioni" in paragraph_text


def test_generator_fallback_for_legacy_rows():
    """Pre-US-4.7 POS rows use loose ``{"fase": "...", ...}`` dicts.
    The generator must still render them without raising."""
    from docx import Document

    from app.services.document_generator.pos import _render_phase_sections

    doc = Document()
    fasi = [
        {"fase": "Smobilizzo", "descrizione": "Pulizia cantiere", "rischi": []},
    ]
    _render_phase_sections(doc, fasi)
    # Legacy rows get promoted to structured phases automatically, so the
    # synoptic + per-phase layout still applies — just with a synthesised id.
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Smobilizzo" in paragraph_text
