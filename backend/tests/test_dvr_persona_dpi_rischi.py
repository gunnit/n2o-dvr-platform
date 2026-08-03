"""Tests for per-persona DPI and specific-risk rendering in DVR Master.

Two persone with the same mansione can carry divergent flags. The
generator must:
  - retain each person's identity in the specific-risk table
  - aggregate (union) their codes only for the per-mansione DPI tables
"""

from dataclasses import dataclass, field

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_generator import dvr_master
from app.services.document_generator.dvr_master import DVRMasterGenerator


@dataclass
class _FakePersona:
    id: str
    nominativo: str
    mansione: str | None
    dpi_codes: list[str] = field(default_factory=list)
    rischi_specifici_codes: list[str] = field(default_factory=list)
    attrezzature_speciali: list[str] = field(default_factory=list)


def _new_generator() -> DVRMasterGenerator:
    """Construct a generator without DB context — we only call pure render
    helpers that take ``doc + persone + extras``."""
    gen = DVRMasterGenerator.__new__(DVRMasterGenerator)
    return gen


def _all_text(doc: Document) -> str:
    parts: list[str] = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_dpi_per_mansione_unions_codes_across_persone():
    """Two saldatori with overlapping but different DPI codes should both
    contribute to the mansione's DPI grid via union."""
    persone = [
        _FakePersona(
            id="p1",
            nominativo="Mario Rossi",
            mansione="Saldatore",
            dpi_codes=["caschi_industria", "guanti_meccanici"],
        ),
        _FakePersona(
            id="p2",
            nominativo="Anna Bianchi",
            mansione="Saldatore",
            dpi_codes=["guanti_meccanici", "occhiali_stanghette"],
        ),
    ]

    gen = _new_generator()
    doc = Document()
    gen._add_dpi_per_mansione_section(doc, persone, extras={})

    text = _all_text(doc).lower()
    # Union: all three DPI labels (any per-code keyword) should appear
    # under the SALDATORE heading.
    assert "caschi" in text
    assert "guanti" in text
    assert "occhiali" in text
    # Divergent flags → varia note must appear
    assert "varia tra i lavoratori" in text


def test_dpi_per_mansione_no_varia_note_when_aligned():
    """When all persone with the mansione carry identical DPI flags,
    the 'varia per lavoratore' note must NOT appear."""
    persone = [
        _FakePersona(
            id="p1",
            nominativo="Mario Rossi",
            mansione="Saldatore",
            dpi_codes=["caschi_industria", "guanti_meccanici"],
        ),
        _FakePersona(
            id="p2",
            nominativo="Anna Bianchi",
            mansione="Saldatore",
            dpi_codes=["caschi_industria", "guanti_meccanici"],
        ),
    ]

    gen = _new_generator()
    doc = Document()
    gen._add_dpi_per_mansione_section(doc, persone, extras={})

    text = _all_text(doc).lower()
    assert "varia tra i lavoratori" not in text


def test_specific_risks_render_one_row_per_exposed_person_without_cross_inheritance():
    mario = _FakePersona(
        "p1", "Mario Rossi", "Operaio", rischi_specifici_codes=["af_rumore"]
    )
    anna = _FakePersona(
        "p2", "Anna Bianchi", "Operaio", rischi_specifici_codes=["mmc"]
    )
    doc = Document()
    _new_generator()._add_mansioni_rischi_specifici_section(
        doc, [mario, anna], {"vdt_esposti_persona_ids": set()}
    )
    table = doc.tables[-1]
    assert [cell.text for cell in table.rows[0].cells] == [
        "Nominativo",
        "Mansione",
        "Rischio specifico",
    ]
    assert [cell.text for cell in table.rows[1].cells][:2] == [
        "MARIO ROSSI",
        "OPERAIO",
    ]
    assert "Rumore" in table.rows[1].cells[2].text
    assert "Movimentazione" not in table.rows[1].cells[2].text
    assert [cell.text for cell in table.rows[2].cells][:2] == [
        "ANNA BIANCHI",
        "OPERAIO",
    ]
    assert "Movimentazione" in table.rows[2].cells[2].text
    assert "Rumore" not in table.rows[2].cells[2].text


def test_person_specific_risks_keep_saved_order_and_deduplicate_labels():
    person = _FakePersona(
        "p1",
        "Mario Rossi",
        "Operaio",
        rischi_specifici_codes=["af_rumore", "mmc", "af_rumore"],
        attrezzature_speciali=[
            "lavori_in_quota",
            "trabattelli",
            "ponteggi",
            "carrello_elevatore",
            "ple",
            "gru",
            "ruspa_escavatore",
            "patente_cde",
            "adr",
            "adr",
        ],
    )

    assert dvr_master._person_specific_risk_labels(person, {"p1"}) == [
        "Videoterminali",
        "Lavori in quota",
        "Utilizzo di trabattelli",
        "Utilizzo di ponteggi",
        "Utilizzo di carrelli elevatori",
        "Utilizzo di piattaforme di lavoro elevabili (PLE)",
        "Utilizzo di gru",
        "Utilizzo di ruspe ed escavatori",
        "Guida professionale (patente C/D/E)",
        "Trasporto merci pericolose (ADR)",
        "Agenti fisici - Rumore",
        "Movimentazione manuale dei carichi (MMC)",
    ]


def test_only_dpi_marca_modello_dash_is_centered_both_ways():
    doc = Document()
    _new_generator()._add_dpi_per_mansione_section(
        doc, [_FakePersona("p", "Mario", "Operaio")], {}
    )
    table = next(
        table
        for table in doc.tables
        if table.rows
        and "Marca / Modello" in [
            cell.text.strip() for cell in table.rows[0].cells
        ]
    )
    column = [cell.text.strip() for cell in table.rows[0].cells].index(
        "Marca / Modello"
    )
    dash_cell = table.rows[1].cells[column]
    assert dash_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert dash_cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert (
        table.rows[1].cells[0].vertical_alignment
        != WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def test_dpi_per_mansione_documents_mansione_with_no_flags():
    """A mansione whose persone carry no DPI flags still gets its own grid
    with an explicit "nessun DPI" row (b780497).

    Silently omitting the mansione is the failure mode this guards: an
    inspector reading the DVR must be able to tell "no PPE is prescribed
    for this role, and here is why" apart from "this role was forgotten".
    """
    persone = [
        _FakePersona(id="p1", nominativo="Mario Rossi", mansione="Operaio"),
    ]
    gen = _new_generator()
    doc = Document()
    gen._add_dpi_per_mansione_section(doc, persone, extras={})

    text = _all_text(doc).lower()
    assert "operaio" in text
    assert "nessun dpi specifico richiesto" in text
    # The justification is the point of the row — a bare "nessuno" would
    # not be defensible at audit.
    assert "sorveglianza sanitaria" in text


def test_dpi_per_mansione_gestante_row_cites_dlgs_151():
    """art. 12 D.Lgs. 151/2001 requires documented protective measures for
    a gestante worker even when no PPE is prescribed."""
    persone = [
        _FakePersona(id="p1", nominativo="Anna Bianchi", mansione="Impiegata gestante"),
    ]
    gen = _new_generator()
    doc = Document()
    gen._add_dpi_per_mansione_section(doc, persone, extras={})

    text = _all_text(doc).lower()
    assert "151/2001" in text


def test_dpi_per_mansione_falls_back_when_no_mansioni_at_all():
    """With nothing to tabulate (no persona carries a mansione), the section
    must emit the 'in fase di compilazione' paragraph rather than a bare
    heading followed by nothing."""
    persone = [
        _FakePersona(id="p1", nominativo="Mario Rossi", mansione=None),
        _FakePersona(id="p2", nominativo="Anna Bianchi", mansione="   "),
    ]
    gen = _new_generator()
    doc = Document()
    gen._add_dpi_per_mansione_section(doc, persone, extras={})

    text = _all_text(doc).lower()
    assert "in fase di compilazione" in text
    assert not doc.tables


def test_sorveglianza_protocol_table_aggregates_per_mansione():
    """§4.3 protocol table must list each mansione once, aggregating the
    union of DPI + rischi codes from all persone with that mansione."""
    persone = [
        _FakePersona(
            id="p1",
            nominativo="Mario Rossi",
            mansione="Saldatore",
            dpi_codes=["caschi_industria"],
            rischi_specifici_codes=["af_rumore"],
        ),
        _FakePersona(
            id="p2",
            nominativo="Anna Bianchi",
            mansione="Saldatore",
            dpi_codes=["guanti_meccanici"],
            rischi_specifici_codes=["mmc"],
        ),
        _FakePersona(
            id="p3",
            nominativo="Lia Verdi",
            mansione="Impiegata",
            dpi_codes=["occhiali_stanghette"],
        ),
    ]
    gen = _new_generator()
    doc = Document()
    gen._add_sorveglianza_protocol_table(doc, persone)

    # Find the table; it must have headers + 2 mansione rows (Saldatore +
    # Impiegata), not 3 (one per persona).
    assert doc.tables, "protocol table should be rendered"
    table = doc.tables[-1]
    headers = [c.text.strip() for c in table.rows[0].cells]
    assert headers[0] == "Mansione"
    mansione_cells = [row.cells[0].text.strip() for row in table.rows[1:]]
    assert sorted(mansione_cells) == ["IMPIEGATA", "SALDATORE"]
